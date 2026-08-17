import ast
import json
import queue
import re
import shutil
import sys
import time
import unittest
import base64
from io import BytesIO
from pathlib import Path
from types import SimpleNamespace
from urllib.parse import quote
from unittest.mock import patch

from openpyxl import Workbook, load_workbook


WEB_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(WEB_ROOT))

import server
import estimate
import file_cache
import milp_solver
import plan_optimizer


class FakeCursor:
    def __init__(self, rows_by_sql):
        self.rows_by_sql = rows_by_sql
        self.last_sql = ""
        self.last_params = ()
        self.executed = []

    def execute(self, sql, params=()):
        self.last_sql = " ".join(sql.split())
        self.last_params = params
        self.executed.append((self.last_sql, params))

    def fetchall(self):
        return self.rows_by_sql.get(self.last_sql, [])

    def fetchone(self):
        rows = self.fetchall()
        if isinstance(rows, dict):
            return rows
        return rows[0] if rows else None

    def close(self):
        pass


class FakeConnection:
    def __init__(self, rows_by_sql):
        self.cursor_obj = FakeCursor(rows_by_sql)
        self.committed = False

    def cursor(self, *args, **kwargs):
        return self.cursor_obj

    def commit(self):
        self.committed = True

    def close(self):
        pass


class PowerPlanServerTest(unittest.TestCase):
    def setUp(self):
        self._original_data_source = server.DATA_SOURCE
        self._original_simu_runtime = server.SIMU_RUNTIME
        server.DATA_SOURCE = server.CsvDataSource()
        server.SIMU_RUNTIME = server.SimuRuntime()

    def tearDown(self):
        server.DATA_SOURCE = self._original_data_source
        server.SIMU_RUNTIME = self._original_simu_runtime

    def wait_optimization_runtime(self, runtime, timeout=30):
        deadline = time.time() + timeout
        payload = runtime.snapshot()
        while payload["status"] == "运行中" and time.time() < deadline:
            time.sleep(0.05)
            payload = runtime.snapshot()
        return payload

    def test_api_payload_excludes_removed_monitor_sections(self):
        payload = server.build_snapshot()

        self.assertEqual(payload["system"], "考察站风-光-氢-储-柴联合规划系统")
        self.assertIn("timestamp", payload)
        self.assertIn("summary", payload)
        self.assertNotIn("simu", payload)
        self.assertNotIn("scada", payload)
        self.assertNotIn("agc", payload)

    def test_index_page_has_visual_planning_entry_buttons(self):
        html = (WEB_ROOT / "index.html").read_text(encoding="utf-8")

        self.assertIn("assets/i18n.js", html)
        self.assertIn(".language-switch", html)
        self.assertIn(".home-theme-switch", html)
        self.assertIn('id="homeThemeSelect"', html)
        self.assertIn('aria-label="首页显示主题"', html)
        self.assertIn('data-home-theme="default"', html)
        self.assertIn('value="default">默认样式</option>', html)
        self.assertIn('value="fresh">轻快</option>', html)
        self.assertIn('value="bright">明亮</option>', html)
        self.assertIn('value="sci-fi">科幻</option>', html)
        self.assertIn('value="solemn">庄重</option>', html)
        additional_home_themes = [
            ("minimal", "极简"),
            ("dark", "黑暗"),
            ("illustration", "插画"),
            ("flat", "扁平"),
            ("neon-future", "霓虹未来"),
            ("glassmorphism", "玻璃拟态"),
            ("material", "材料设计"),
            ("magazine", "杂志排版"),
            ("dynamic", "动态交互"),
        ]
        for theme_value, theme_label in additional_home_themes:
            self.assertIn(f'value="{theme_value}">{theme_label}</option>', html)
            self.assertIn(f'.screen[data-home-theme="{theme_value}"]', html)
        i18n_script = (WEB_ROOT / "assets" / "i18n.js").read_text(encoding="utf-8")
        for _, theme_label in additional_home_themes:
            self.assertIn(theme_label, i18n_script)
        self.assertIn("powerPlanHomeTheme", html)
        self.assertIn("applyHomeTheme", html)
        self.assertIn("document.documentElement.dataset.homeTheme = theme", html)
        self.assertIn("document.body.dataset.homeTheme = theme", html)
        self.assertIn("homeDynamicPulse", html)
        self.assertIn(".screen[data-home-theme=\"fresh\"]", html)
        self.assertIn(".screen[data-home-theme=\"bright\"]", html)
        self.assertIn(".screen[data-home-theme=\"sci-fi\"]", html)
        self.assertIn(".screen[data-home-theme=\"solemn\"]", html)
        bright_entry_css = html.split('.screen[data-home-theme="bright"] .feature-entry-grid::before,', 1)[1].split("}", 1)[0]
        self.assertIn("background-color: #dceefa", bright_entry_css)
        illustration_entry_css = html.split('.screen[data-home-theme="illustration"] .feature-entry-grid::before,', 1)[1].split("}", 1)[0]
        self.assertIn("background-color: #cceaf6", illustration_entry_css)
        self.assertIn('background-image: url("assets/main-dashboard-bg.png?v=20260513-bg-refresh")', html)
        self.assertIn("background-size: contain", html)
        self.assertIn('<link rel="icon" href="data:,">', html)
        self.assertIn(".screen::before", html)
        self.assertIn("filter: saturate(1.08) brightness(0.74) contrast(1.08)", html)
        self.assertIn('class="home-title"', html)
        self.assertIn("考察站风-光-氢-储-柴联合规划系统</h1>", html)
        home_title_css = html.split(".home-title {", 1)[1].split("}", 1)[0]
        self.assertIn("z-index: 6", home_title_css)
        self.assertIn("text-shadow: none", home_title_css)
        self.assertIn("color: #ffffff", home_title_css)
        home_user_status_css = html.split(".home-user-status {", 1)[1].split("}", 1)[0]
        self.assertIn("border: 0", home_user_status_css)
        self.assertIn("background:", home_user_status_css)
        self.assertIn(".energy-side", html)
        self.assertIn(".energy-left", html)
        self.assertIn(".energy-right", html)
        energy_side_css = html.split(".energy-side {", 1)[1].split("}", 1)[0]
        self.assertIn("top: 50%", energy_side_css)
        self.assertIn("min-height: clamp(320px, 41vh, 400px)", energy_side_css)
        self.assertIn("align-content: space-around", energy_side_css)
        self.assertIn("transform: translateY(-29%)", energy_side_css)
        self.assertIn("clip-path: polygon", html)
        self.assertIn("box-shadow:", html)
        self.assertIn("color: #21d5ff", html)
        self.assertIn('class="feature-entry-grid"', html)
        self.assertIn('aria-label="规划功能快捷入口"', html)
        self.assertEqual(html.count('class="feature-entry"'), 5)
        self.assertEqual(html.count('class="feature-icon"'), 5)
        self.assertIn('class="energy-side energy-left"', html)
        self.assertIn('class="energy-side energy-right"', html)
        self.assertIn('<strong>参数维护</strong>', html)
        self.assertIn('<strong>规划求解</strong>', html)
        self.assertIn('<strong>方案评估</strong>', html)
        self.assertIn('<strong>结果对比</strong>', html)
        self.assertIn('<strong>任务并发</strong>', html)
        self.assertNotIn("规划参数维护", html)
        self.assertNotIn("规划算法", html)
        self.assertNotIn("规划方案评估", html)
        self.assertIn('href="planning.html"', html)
        self.assertIn('href="optimize.html"', html)
        self.assertIn('href="evaluation.html"', html)
        self.assertIn('href="comparison.html"', html)
        self.assertIn('href="tasks.html"', html)
        self.assertIn(".feature-entry-grid", html)
        self.assertIn("grid-template-columns: repeat(5, minmax(0, 1fr))", html)
        self.assertIn("top: 50%", html)
        self.assertIn("transform: translate(-50%, -29%)", html)
        feature_text_css = html.split(".feature-entry strong {", 1)[1].split("}", 1)[0]
        self.assertIn("white-space: nowrap", feature_text_css)
        self.assertIn("font-size: clamp(15px, min(1.32vw, 3.8vh), 28px)", feature_text_css)
        self.assertIn("max-width: 100%", feature_text_css)
        self.assertNotIn("text-overflow: ellipsis", feature_text_css)
        self.assertNotIn("overflow: hidden", feature_text_css)
        self.assertIn(".feature-icon svg", html)
        self.assertNotIn("hot-nav", html)
        self.assertNotIn("quick-links", html)
        self.assertNotIn("系统主导航", html)
        self.assertNotIn("在线监视快捷入口", html)
        self.assertNotIn("SIMU在线监视", html)

        self.assertIn("HOME_THEME_STORAGE_KEY", i18n_script)
        self.assertIn("HOME_THEMES", i18n_script)
        self.assertIn("applyStoredHomeTheme", i18n_script)
        self.assertIn("document.documentElement.dataset.homeTheme = theme", i18n_script)
        self.assertIn("document.body.dataset.homeTheme = theme", i18n_script)
        self.assertIn('"玻璃拟态": "Glassmorphism"', i18n_script)
        planning_css = (WEB_ROOT / "assets" / "planning.css").read_text(encoding="utf-8")
        self.assertIn('body[data-home-theme="glassmorphism"]', planning_css)
        self.assertIn('body[data-home-theme="neon-future"]', planning_css)
        self.assertIn("--theme-accent", planning_css)
        self.assertIn("--theme-control-text", planning_css)
        self.assertIn("--theme-active-bg", planning_css)
        self.assertIn("--theme-active-text", planning_css)
        self.assertIn("--theme-control-text: #ffffff", planning_css)
        self.assertIn("--theme-text: #183247", planning_css)
        self.assertIn("--theme-active-bg: #1f6f94", planning_css)
        self.assertIn("color: var(--theme-control-text)", planning_css)
        self.assertIn("background: var(--theme-active-bg)", planning_css)
        self.assertIn("color: var(--theme-active-text)", planning_css)
        self.assertIn('body[data-home-theme]:not([data-home-theme="default"]) .scheme-item,', planning_css)
        self.assertIn('body[data-home-theme]:not([data-home-theme="default"]) .editor-header,', planning_css)
        self.assertIn("background: color-mix(in srgb, var(--theme-control-bg) 36%, transparent)", planning_css)
        self.assertIn("background: color-mix(in srgb, var(--theme-control-bg) 78%, rgba(0, 0, 0, 0.38))", planning_css)
        self.assertIn("background: color-mix(in srgb, var(--theme-control-bg) 56%, transparent)", planning_css)
        self.assertIn('.screen[data-home-theme="bright"] .home-title', html)
        self.assertIn("0 -1px 0 rgba(3, 18, 32, 0.74)", html)

    def test_power_plan_pages_share_dark_hud_visual_theme(self):
        css = (WEB_ROOT / "assets" / "planning.css").read_text(encoding="utf-8")
        planning_html = (WEB_ROOT / "planning.html").read_text(encoding="utf-8")
        optimize_html = (WEB_ROOT / "optimize.html").read_text(encoding="utf-8")
        page_names = ("planning.html", "optimize.html", "evaluation.html", "comparison.html", "tasks.html", "users.html", "login.html", "register.html")

        for page_name in page_names:
            page_html = (WEB_ROOT / page_name).read_text(encoding="utf-8")
            self.assertIn("assets/planning.css?v=", page_html)
        self.assertIn('url("main-dashboard-bg.png?v=20260513-bg-refresh")', css)
        self.assertIn("--hud-cyan: #21d5ff", css)
        self.assertIn("--hud-panel:", css)
        self.assertIn("--hud-inner-surface:", css)
        self.assertIn("--hud-shadow-soft:", css)
        self.assertIn("rgba(20, 190, 255, 0.64)", css)
        self.assertIn(".scheme-rail,", css)
        self.assertIn(".optimization-command-card,", css)
        self.assertIn("background: var(--hud-panel)", css)
        self.assertIn("color: var(--hud-text)", css)
        self.assertIn(".composition-bar-summary span,", css)
        self.assertIn(".composition-bar-legend div {", css)
        self.assertIn("color: var(--hud-muted-strong)", css)
        self.assertIn("--chart-tooltip-bg:", css)
        self.assertIn("--chart-tooltip-text:", css)
        self.assertIn("--chart-tooltip-muted:", css)
        self.assertIn("--chart-tooltip-border:", css)
        self.assertIn("background: var(--chart-tooltip-bg", css)
        self.assertIn("color: var(--chart-tooltip-text", css)
        self.assertIn('body[data-home-theme]:not([data-home-theme="default"]) .chart-tip,', css)
        self.assertIn('body[data-home-theme]:not([data-home-theme="default"]) .chart-hover-tooltip,', css)
        self.assertIn('body[data-home-theme]:not([data-home-theme="default"]) .comparison-chart-tooltip,', css)
        self.assertIn('body[data-home-theme]:not([data-home-theme="default"]) .annual-chart-tooltip {', css)
        self.assertIn("color: var(--chart-tooltip-strong) !important", css)
        self.assertIn(".time-series-import-chart text,", css)
        self.assertIn(".load-generator-preview text,", css)
        self.assertIn("#timeChart text,", css)
        self.assertIn(".histogram-svg text {", css)
        self.assertIn("fill: var(--hud-muted-strong)", css)
        self.assertIn('body[data-home-theme]:not([data-home-theme="default"]) .time-series-import-chart text,', css)
        self.assertIn("fill: var(--theme-control-text)", css)
        self.assertIn("--hud-muted-strong: var(--theme-text)", css)
        self.assertIn("--control-bg: var(--theme-control-bg)", css)
        self.assertIn('body[data-home-theme]:not([data-home-theme="default"]) .tab,', css)
        self.assertIn('body[data-home-theme]:not([data-home-theme="default"]) .curve-button,', css)
        self.assertIn('body[data-home-theme]:not([data-home-theme="default"]) .optimization-actions .queue-action', css)
        self.assertIn("Theme contrast guard", css)
        self.assertIn("body[data-home-theme]:not([data-home-theme=\"default\"]) .main-nav a", css)
        self.assertIn("background: var(--theme-control-bg) !important", css)
        self.assertIn("color: var(--theme-control-text) !important", css)
        theme_queue_css = css.rsplit('body[data-home-theme]:not([data-home-theme="default"]) .secondary,\nbody[data-home-theme]:not([data-home-theme="default"]) .optimization-actions .queue-action {', 1)[1].split("}", 1)[0]
        self.assertIn("background: var(--theme-panel-bg) !important", theme_queue_css)
        self.assertIn("color: var(--theme-text) !important", theme_queue_css)
        self.assertIn('body[data-home-theme]:not([data-home-theme="default"]) .curve-range-filter label,', css)
        curve_filter_label_css = css.rsplit('body[data-home-theme]:not([data-home-theme="default"]) .curve-range-filter label,', 1)[1].split("}", 1)[0]
        self.assertIn("color: var(--theme-control-text) !important", curve_filter_label_css)
        self.assertIn('body[data-home-theme]:not([data-home-theme="default"]) .optimization-logs {', css)
        theme_log_css = css.rsplit('body[data-home-theme]:not([data-home-theme="default"]) .optimization-logs {', 1)[1].split("}", 1)[0]
        self.assertIn("background: var(--theme-log-bg)", theme_log_css)
        self.assertIn("color: var(--theme-log-text)", theme_log_css)
        self.assertIn("--theme-log-bg: var(--theme-control-bg)", css)
        self.assertIn("--theme-log-text: var(--theme-control-text)", css)
        for light_theme in ('body[data-home-theme="bright"]', 'body[data-home-theme="illustration"]', 'body[data-home-theme="glassmorphism"]', 'body[data-home-theme="magazine"]'):
            light_theme_css = css.split(light_theme, 1)[1].split("}", 1)[0]
            self.assertIn("--theme-log-bg: color-mix(in srgb, var(--theme-panel-bg) 92%, #ffffff)", light_theme_css)
            self.assertIn("--theme-log-text:", light_theme_css)
        self.assertIn('body[data-home-theme]:not([data-home-theme="default"]) .current-scheme strong,', css)
        self.assertIn('body[data-home-theme]:not([data-home-theme="default"]) .curve-range-scope button.active,', css)
        curve_scope_active_css = css.rsplit('body[data-home-theme]:not([data-home-theme="default"]) .curve-range-scope button.active,', 1)[1].split("}", 1)[0]
        self.assertIn("background: var(--theme-active-bg) !important", curve_scope_active_css)
        self.assertIn("color: var(--theme-active-text) !important", curve_scope_active_css)
        theme_language_css = css.rsplit('body[data-home-theme]:not([data-home-theme="default"]) .language-switch,', 1)[1].split("}", 1)[0]
        self.assertIn("background: var(--theme-control-bg) !important", theme_language_css)
        self.assertIn("color: var(--theme-control-text) !important", theme_language_css)
        theme_status_css = css.rsplit('body[data-home-theme]:not([data-home-theme="default"]) .task-status-pill.completed {', 1)[1].split("}", 1)[0]
        self.assertIn("background: #c9f4dc !important", theme_status_css)
        self.assertIn("color: #0d5e3f !important", theme_status_css)
        self.assertIn('body[data-home-theme]:not([data-home-theme="default"]) .auth-card label,', css)
        self.assertIn('body[data-home-theme]:not([data-home-theme="default"]) .task-table,', css)
        self.assertIn('body[data-home-theme]:not([data-home-theme="default"]) .comparison-curve-chart text,', css)
        self.assertIn('body[data-home-theme]:not([data-home-theme="default"]) .optimization-curve-name-list', css)
        theme_curve_name_css = css.rsplit('body[data-home-theme]:not([data-home-theme="default"]) .optimization-curve-name-list,', 1)[1].split("}", 1)[0]
        self.assertIn("background: var(--theme-panel-bg) !important", theme_curve_name_css)
        self.assertIn("color: var(--theme-text) !important", theme_curve_name_css)
        self.assertIn("--control-bg:", css)
        self.assertIn("--control-primary-bg:", css)
        self.assertIn("--control-danger-bg:", css)
        self.assertIn("--control-muted-bg:", css)
        self.assertIn(".optimization-actions button,", css)
        self.assertIn(".map-provider-tab,", css)
        self.assertIn(".time-series-import-curve-toggle,", css)
        self.assertIn(".task-actions button", css)
        self.assertIn(".optimization-actions .queue-action", css)
        inner_card_css = css.split(".evaluation-result-rail,\n.evaluation-planning-result-panel,", 1)[1].split("}", 1)[0]
        self.assertIn("background: var(--hud-inner-surface)", inner_card_css)
        self.assertIn("box-shadow: none", inner_card_css)
        self.assertIn(".task-scheme-filter select", css)
        self.assertIn(".language-switch select", css)
        self.assertIn(".users-table select", css)
        self.assertIn(".auth-card input", css)

    def test_monitor_static_pages_are_removed(self):
        for filename in ("simu.html", "scada.html", "agc.html"):
            self.assertFalse((WEB_ROOT / filename).exists())

    def test_auth_pages_and_topbars_include_user_controls(self):
        login_html = (WEB_ROOT / "login.html").read_text(encoding="utf-8")
        register_html = (WEB_ROOT / "register.html").read_text(encoding="utf-8")
        users_html = (WEB_ROOT / "users.html").read_text(encoding="utf-8")
        planning_html = (WEB_ROOT / "planning.html").read_text(encoding="utf-8")
        optimize_html = (WEB_ROOT / "optimize.html").read_text(encoding="utf-8")
        tasks_html = (WEB_ROOT / "tasks.html").read_text(encoding="utf-8")
        index_html = (WEB_ROOT / "index.html").read_text(encoding="utf-8")
        css = (WEB_ROOT / "assets" / "planning.css").read_text(encoding="utf-8")
        i18n_script = (WEB_ROOT / "assets" / "i18n.js").read_text(encoding="utf-8")

        self.assertIn('id="loginForm"', login_html)
        self.assertIn('id="registerForm"', register_html)
        self.assertIn('body data-admin-page="true"', users_html)
        self.assertIn('id="usersTable"', users_html)
        for html in (planning_html, optimize_html, index_html, users_html, tasks_html):
            self.assertIn("data-auth-user", html)
            self.assertIn("data-auth-username", html)
            self.assertIn("data-logout", html)
            self.assertIn("assets/auth.js", html)
            self.assertIn("assets/i18n.js", html)
        self.assertIn("assets/tasks.js", tasks_html)
        self.assertIn('<a class="active" href="tasks.html">任务并发</a>', tasks_html)
        self.assertIn('id="optimizationTaskTable"', tasks_html)
        self.assertIn('id="evaluationTaskTable"', tasks_html)
        self.assertIn('id="evaluationSchemeFilter"', tasks_html)
        self.assertIn("全部方案</option>", tasks_html)
        self.assertIn('id="taskTableResizeHandle"', tasks_html)
        self.assertNotIn('id="refreshTasks"', tasks_html)
        self.assertNotIn("<h1>任务并发</h1>", tasks_html)
        self.assertNotIn("展示所有规划计算任务和方案评估任务", tasks_html)
        self.assertIn("assets/i18n.js", login_html)
        self.assertIn("assets/i18n.js", register_html)
        self.assertIn("powerPlanLanguage", i18n_script)
        self.assertIn("languageSelect", i18n_script)
        self.assertIn("PowerPlanI18n", i18n_script)
        self.assertIn("Station Wind-Solar-Hydrogen-Storage-Diesel Planning System", i18n_script)
        self.assertIn("Scenario Evaluation", i18n_script)
        self.assertIn("Result Comparison", i18n_script)
        self.assertIn("Task Concurrency", i18n_script)
        self.assertIn("Refresh Status", i18n_script)
        self.assertIn("Fetch Weather", i18n_script)
        self.assertIn('"年份": "Year"', i18n_script)
        self.assertIn("Calculation End Time", i18n_script)
        self.assertIn("Start Now", i18n_script)
        self.assertIn("Queued", i18n_script)
        self.assertIn("Calculating", i18n_script)
        self.assertIn("Completed", i18n_script)
        self.assertIn("Planning Calculation", i18n_script)
        self.assertIn("Elapsed Time (s)", i18n_script)
        self.assertIn("Load Up Disturbance Factor", i18n_script)
        self.assertIn("Load Down Disturbance Factor", i18n_script)
        self.assertIn("Renewable Down Disturbance Factor", i18n_script)
        self.assertIn("Grid Up Regulation Requirement", i18n_script)
        self.assertIn("Grid Down Regulation Requirement", i18n_script)
        self.assertIn("Import File", i18n_script)
        self.assertIn("Importing load file", i18n_script)
        self.assertIn("Load file imported", i18n_script)
        self.assertIn("8760-point load curve", i18n_script)
        self.assertIn("Hourly curves are loading in the background", i18n_script)
        self.assertIn("Save Template", i18n_script)
        self.assertIn("Template name already exists", i18n_script)
        for translated_label in (
            "Solar Irradiance",
            "Ambient Temperature",
            "Load",
            "Capacity (kW)",
            "Capacity (kWh)",
            "Power Upper Limit (kW)",
            "Power Lower Limit (kW)",
            "Fuel Consumption Rate (kg/kWh)",
            "Cut-in Wind Speed (m/s)",
            "Rated Wind Speed (m/s)",
            "Cut-out Wind Speed (m/s)",
            "Quantity Lower Bound (units)",
            "Quantity Upper Bound (units)",
            "Design Life (years)",
            "Self-discharge Rate (0-1%/day)",
            "Hydrogen-to-Electric Efficiency (kWh/Nm3)",
            "Electric-to-Hydrogen Efficiency (Nm3/kWh)",
            "Statistical Histograms",
            "Candidate Device List",
            "Validation and Summary",
            "Time Series Rows",
            "Device Entries",
            "No device entries",
            "Name",
            "Device Type",
            "Cost Composition",
            "Annual Diesel Cost",
            "Annualized Construction Cost",
            "Diesel Generation",
            "Energy Composition",
            "Curve Preview",
            "Imported Curve Preview",
            "Imported Curve Display Toggle",
            "Resize imported curve preview height",
        ):
            self.assertIn(translated_label, i18n_script)
        for exact_mapping in (
            '"名称": "Name"',
            '"设备类型": "Device Type"',
            '"成本构成": "Cost Composition"',
            '"成本构成(单位: 万元)": "Cost Composition (Unit: 10k CNY)"',
            '"容量构成": "Capacity Composition"',
            '"容量构成(单位: kW)": "Capacity Composition (Unit: kW)"',
            '"柴发容量": "Diesel Capacity"',
            '"风电容量": "Wind Capacity"',
            '"光伏容量": "PV Capacity"',
            '"电储能容量": "Battery Storage Capacity"',
            '"燃料电池容量": "Fuel Cell Capacity"',
            '"电储": "Battery Storage"',
            '"燃电": "Fuel Cell"',
            '"柴发总容量": "Total Diesel Capacity"',
            '"风电总容量": "Total Wind Capacity"',
            '"光伏总容量": "Total PV Capacity"',
            '"氢能总容量": "Total Hydrogen Capacity"',
            '"储能总容量": "Total Storage Capacity"',
            '"绿电年发电量": "Annual Green Energy Generation"',
            '"总发电量": "Total Generation"',
            '"柴油消耗": "Diesel Consumption"',
            '"频率风险点": "Frequency Risk Points"',
            '"年总成本": "Annual Total Cost"',
            '"总成本": "Total Cost"',
            '"年柴油成本": "Annual Diesel Cost"',
            '"包含": "Contains"',
            '"文件": "File"',
            '"已解析": "Parsed"',
            '"共": "Total"',
            '"行": "Rows"',
            '"小时序号": "Hour Index"',
            '"时间": "Time"',
            '"风光柴": "Wind-Solar-Diesel"',
            '"电储能": "Battery Storage"',
            '"氢储能": "Hydrogen Storage"',
            '"柴发": "Diesel Generator"',
            '"风机": "Wind Turbine"',
            '"光伏": "PV"',
            '"储能PCS": "Storage PCS"',
            '"储能电池组": "Battery Pack"',
            '"电制氢": "Electrolyzer"',
            '"储氢罐": "Hydrogen Tank"',
            '"燃料电池": "Fuel Cell"',
            '"新增行": "Add Row"',
            '"是否考虑新能源扰动": "Renewable Disturbance Constraint"',
            '"是": "Yes"',
            '"否": "No"',
            '"取值范围": "Value Range"',
            '"参数": "Parameter"',
            '"参数名称": "Parameter Name"',
            '"当前": "Current"',
            '"行数正确": "Row Count OK"',
            '"分布": "Distribution"',
            '"正常": "Normal"',
            '"相关参数": "Related Parameters"',
            '"总容量": "Total Capacity"',
            '"规划年指标": "Planning Annual Metrics"',
            '"规划结果": "Planning Result"',
            '"计算中止": "Calculation Stopped"',
            '"计算失败": "Calculation Failed"',
            '"计算超时": "Calculation Timed Out"',
            '"请选择": "Please select"',
            '"暂无": "No data"',
            '"柴发总电量": "Total Diesel Energy"',
            '"制氢总量": "Total Hydrogen Production"',
            '"氢储总发电量": "Hydrogen Storage Generation"',
            '"电储总发电量": "Battery Storage Generation"',
            '"年均建设成本": "Annualized Construction Cost"',
            '"年柴油": "Annual Diesel"',
            '"年均建设": "Annualized Construction"',
            '"年均总成本": "Annualized Total Cost"',
            '"年运行成本": "Annual Operating Cost"',
            '"运行": "Operation"',
            '"建设": "Construction"',
            '"柴发电量": "Diesel Generation"',
            '"绿电电量": "Green Energy"',
            '"电量构成": "Energy Composition"',
            '"电量构成(单位: 万kWh)": "Energy Composition (Unit: 10k kWh)"',
            '"风电": "Wind Power"',
            '"新能源": "Renewable Energy"',
            '"绿电": "Green Energy"',
            '"表格显示": "Table View"',
            '"柱图对比": "Bar Comparison"',
            '"文件导入": "Import File"',
            '"气象获取": "Fetch Weather"',
            '"成本对比": "Cost Comparison"',
            '"电量对比": "Energy Comparison"',
            '"新能源利用对比": "Renewable Utilization Comparison"',
            '"储能氢能发电对比": "Storage and Hydrogen Generation Comparison"',
            '"负荷用电量": "Load Consumption"',
            '"柴油发电量": "Diesel Generation"',
            '"新能源总发电量": "Total Renewable Generation"',
            '"新能源实际电量": "Actual Renewable Energy"',
            '"储能发电量": "Storage Generation"',
            '"暂无柱图对比数据": "No bar comparison data"',
            '"调整年度柱图左右宽度": "Resize annual bar chart columns"',
            '"调整年度柱图上下高度": "Resize annual bar chart rows"',
            '"对比项": "Comparison Item"',
            '"负荷上扰动功率": "Load Up Disturbance Power"',
            '"负荷下扰动功率": "Load Down Disturbance Power"',
            '"新能源下扰动功率": "Renewable Down Disturbance Power"',
            '"风光单机功率最大值": "Max Single Wind/PV Unit Power"',
            '"电网向上调节能力": "Grid Up Regulation Capability"',
            '"电网向下调节能力": "Grid Down Regulation Capability"',
            '"电网向上调节需求": "Grid Up Regulation Requirement"',
            '"电网向下调节需求": "Grid Down Regulation Requirement"',
            '"曲线显示和统计信息": "Curve Legend and Statistics"',
            '"统计信息菜单": "Statistics Menu"',
            '"隐藏统计信息": "Hide Statistics"',
            '"显示统计信息": "Show Statistics"',
            '"恢复统计位置": "Reset Statistics Position"',
            '"拖动可移动统计信息，右键显示菜单": "Drag to move statistics; right-click for menu"',
        ):
            self.assertIn(exact_mapping, i18n_script)
        self.assertNotIn("新能源N-1功率缺口", i18n_script)
        for mixed_label in (
            '"切入风速(m/s)"',
            '"额定风速(m/s)"',
            '"切出风速(m/s)"',
        ):
            self.assertIn(mixed_label, i18n_script)
        self.assertIn("MutationObserver", i18n_script)
        self.assertIn("patchDialogs", i18n_script)
        self.assertIn("target.parentNode.insertBefore(wrap, target)", i18n_script)
        self.assertNotIn("target.insertBefore(wrap, target.firstElementChild)", i18n_script)
        self.assertIn("const translated = translateText(node.nodeValue, language);", i18n_script)
        self.assertIn("if (translated !== node.nodeValue) node.nodeValue = translated;", i18n_script)
        self.assertIn("filter(([key]) => key.length > 1)", i18n_script)
        self.assertIn("captureFitTextBaselines", i18n_script)
        self.assertIn("fitTranslatedText", i18n_script)
        self.assertIn("i18n-fit-text", i18n_script)
        self.assertIn("element.style.maxWidth", i18n_script)
        self.assertIn("shrinkElementFontToFit", i18n_script)
        self.assertNotIn('href="users.html" data-admin-only hidden>用户管理</a>', planning_html)
        self.assertNotIn('href="users.html" data-admin-only hidden>用户管理</a>', optimize_html)
        self.assertNotIn('href="users.html" data-admin-only hidden>用户管理</a>', tasks_html)
        self.assertIn("data-admin-only", index_html)
        self.assertIn(".user-status", css)
        self.assertIn(".language-switch", css)
        self.assertIn(".auth-shell > .language-switch", css)
        self.assertIn(".i18n-fit-text", css)
        user_status_css = css.split(".user-status {", 1)[1].split("}", 1)[0]
        self.assertIn("border: 0", user_status_css)
        self.assertIn(".auth-card", css)
        self.assertIn("font-size: 17px", css)

    def test_sqlite_user_store_registers_first_user_as_admin_and_authenticates(self):
        db_path = WEB_ROOT / "tests" / "tmp_users.sqlite3"
        db_path.unlink(missing_ok=True)
        try:
            store = server.UserStore(db_path)
            admin = store.create_user("adminA", "secret1")
            normal = store.create_user("userA", "secret2")

            self.assertEqual(admin["role"], "admin")
            self.assertEqual(normal["role"], "user")
            self.assertEqual(store.authenticate("adminA", "secret1")["id"], admin["id"])
            with self.assertRaises(ValueError):
                store.authenticate("adminA", "bad-password")
            token = store.create_session(admin["id"])
            self.assertEqual(store.user_for_session(token)["username"], "adminA")
            store.delete_session(token)
            self.assertIsNone(store.user_for_session(token))
        finally:
            db_path.unlink(missing_ok=True)

    def test_auth_and_user_management_api_use_sqlite_sessions(self):
        original_store = server.USER_STORE
        db_path = WEB_ROOT / "tests" / "tmp_auth_api.sqlite3"
        db_path.unlink(missing_ok=True)
        temp_store = server.UserStore(db_path)
        server.USER_STORE = temp_store
        try:
            status, headers, body = server.handle_auth_api_path(
                "/api/auth/register",
                "POST",
                json.dumps({"username": "adminA", "password": "secret1"}).encode("utf-8"),
            )
            data = json.loads(body.decode("utf-8"))
            self.assertEqual(status, 200)
            self.assertEqual(data["user"]["role"], "admin")
            self.assertIn("Set-Cookie", headers)
            token = headers["Set-Cookie"].split("=", 1)[1].split(";", 1)[0]

            status, headers, body = server.handle_auth_api_path("/api/auth/me", "GET", token=token)
            self.assertEqual(status, 200)
            self.assertEqual(json.loads(body.decode("utf-8"))["user"]["username"], "adminA")

            status, headers, body = server.handle_users_api_path("/api/users", "GET", b"", {"id": 1, "username": "adminA", "role": "admin"})
            self.assertEqual(status, 200)
            self.assertEqual(len(json.loads(body.decode("utf-8"))["users"]), 1)

            status, headers, body = server.handle_users_api_path("/api/users", "GET", b"", {"id": 2, "username": "userA", "role": "user"})
            self.assertEqual(status, 403)
        finally:
            server.USER_STORE = original_store
            del temp_store
            db_path.unlink(missing_ok=True)

    def test_optimization_runtime_can_clear_logs(self):
        runtime = server.OptimizationRuntime("测试方案")
        runtime._append_log_unlocked("info", "待清空日志")

        payload = runtime.apply("clear_logs", scheme="测试方案")

        self.assertEqual(payload["logs"], [])
        self.assertEqual(runtime.snapshot()["logs"], [])

    def test_evaluation_runtime_can_clear_logs(self):
        runtime = server.EvaluationRuntime("测试方案")
        runtime.result_filename = "case_results.xlsx"
        runtime._append_log_unlocked("info", "待清空日志")

        payload = runtime.apply("clear_logs", scheme="测试方案", filename="case_results.xlsx")

        self.assertEqual(payload["logs"], [])
        self.assertEqual(runtime.snapshot()["logs"], [])

    def test_api_response_is_json_for_known_endpoint(self):
        status, headers, body = server.handle_api_path("/api/overview")

        self.assertEqual(status, 200)
        self.assertEqual(headers["Content-Type"], "application/json; charset=utf-8")
        data = json.loads(body.decode("utf-8"))
        self.assertEqual(data["system"], "考察站风-光-氢-储-柴联合规划系统")
        self.assertIn("summary", data)
        self.assertNotIn("simu", data)
        self.assertNotIn("scada", data)
        self.assertNotIn("agc", data)

    def test_removed_monitor_api_endpoints_return_not_found(self):
        for path in ("/api/simu", "/api/scada", "/api/agc"):
            status, headers, body = server.handle_api_path(path)
            data = json.loads(body.decode("utf-8"))

            self.assertEqual(status, 404)
            self.assertEqual(headers["Content-Type"], "application/json; charset=utf-8")
            self.assertEqual(data["error"], "not_found")
            self.assertEqual(data["path"], path)

        status, headers, body = server.handle_control_path(
            "/api/simu/control",
            json.dumps({"action": "start"}).encode("utf-8"),
        )
        data = json.loads(body.decode("utf-8"))
        self.assertEqual(status, 404)
        self.assertEqual(data["error"], "not_found")

    def test_optimization_api_start_stop_and_logs(self):
        original_runtime = server.OPTIMIZATION_RUNTIME
        original_store = server.PLANNING_STORE
        planning_root = WEB_ROOT / "tests" / "tmp_optimization_api_start_stop"
        shutil.rmtree(planning_root, ignore_errors=True)
        planning_root.mkdir(parents=True)
        server.PLANNING_STORE = server.planning_store.PlanningStore(root=planning_root)
        server.OPTIMIZATION_RUNTIME = server.OptimizationRuntimeManager()

        def write_feasible_scheme(name: str) -> None:
            payload = server.planning_store.default_payload(name)
            for row in payload["time_series"]:
                row["wind_speed"] = 0
                row["solar_irradiance"] = 0
                row["load"] = 50
                row["temperature"] = 20
            payload["diesel_generators"][0].update(
                {
                    "capacity": 120,
                    "power_upper": 120,
                    "power_lower": 0,
                    "quantity_lower": 1,
                    "quantity_upper": 1,
                }
            )
            server.PLANNING_STORE.write_scheme(name, payload)

        try:
            write_feasible_scheme("方案A")
            write_feasible_scheme("方案B")

            status, headers, body = server.handle_api_path("/api/optimization/status?scheme=方案A")
            initial = json.loads(body.decode("utf-8"))
            self.assertEqual(status, 200)
            self.assertEqual(headers["Content-Type"], "application/json; charset=utf-8")
            self.assertEqual(initial["status"], "待启动")
            self.assertEqual(initial["scheme"], "方案A")
            self.assertIn("metrics", initial)
            self.assertIn("results", initial)
            self.assertIn("logs", initial)

            status, headers, body = server.handle_control_path(
                "/api/optimization/control",
                json.dumps({"action": "start", "scheme": "方案A"}, ensure_ascii=False).encode("utf-8"),
            )
            started = json.loads(body.decode("utf-8"))
            self.assertEqual(status, 200)
            self.assertEqual(started["state"]["status"], "运行中")
            self.assertEqual(started["state"]["scheme"], "方案A")
            self.assertTrue(started["state"]["start_time"])
            self.assertFalse(started["state"]["end_time"])
            self.assertTrue(any("启动规划求解" in item["message"] for item in started["state"]["logs"]))

            status, headers, body = server.handle_control_path(
                "/api/optimization/control",
                json.dumps({"action": "start", "scheme": "方案A"}, ensure_ascii=False).encode("utf-8"),
            )
            duplicate = json.loads(body.decode("utf-8"))
            self.assertEqual(status, 409)
            self.assertEqual(duplicate["error"], "running")
            self.assertIn("正在运行，无法再次启动", duplicate["message"])

            status, headers, body = server.handle_control_path(
                "/api/optimization/control",
                json.dumps({"action": "start", "scheme": "方案B"}, ensure_ascii=False).encode("utf-8"),
            )
            second_started = json.loads(body.decode("utf-8"))
            self.assertEqual(status, 200)
            self.assertEqual(second_started["state"]["status"], "运行中")
            self.assertEqual(second_started["state"]["scheme"], "方案B")

            status, headers, body = server.handle_api_path("/api/optimization/status?scheme=方案A")
            scheme_a = json.loads(body.decode("utf-8"))
            self.assertEqual(status, 200)
            self.assertEqual(scheme_a["status"], "运行中")
            self.assertEqual(scheme_a["scheme"], "方案A")
            self.assertIn("方案A", scheme_a["running_schemes"])
            self.assertIn("方案B", scheme_a["running_schemes"])

            status, headers, body = server.handle_api_path("/api/optimization/status?scheme=方案B")
            scheme_b = json.loads(body.decode("utf-8"))
            self.assertEqual(status, 200)
            self.assertEqual(scheme_b["status"], "运行中")
            self.assertEqual(scheme_b["scheme"], "方案B")
            self.assertTrue(any("方案：方案A" in item["message"] for item in scheme_a["logs"]))
            self.assertFalse(any("方案：方案B" in item["message"] for item in scheme_a["logs"]))
            self.assertTrue(any("方案：方案B" in item["message"] for item in scheme_b["logs"]))

            status, headers, body = server.handle_control_path(
                "/api/optimization/control",
                json.dumps({"action": "stop", "scheme": "方案A"}, ensure_ascii=False).encode("utf-8"),
            )
            stopped = json.loads(body.decode("utf-8"))
            self.assertEqual(status, 200)
            self.assertEqual(stopped["state"]["status"], "计算中止")
            self.assertTrue(stopped["state"]["end_time"])
            self.assertTrue(any("停止规划求解" in item["message"] for item in stopped["state"]["logs"]))

            status, headers, body = server.handle_control_path(
                "/api/optimization/control",
                json.dumps({"action": "stop", "scheme": "方案A"}, ensure_ascii=False).encode("utf-8"),
            )
            not_running = json.loads(body.decode("utf-8"))
            self.assertEqual(status, 409)
            self.assertEqual(not_running["error"], "not_running")
            self.assertIn("没有运行", not_running["message"])

            status, headers, body = server.handle_api_path("/api/optimization/status?scheme=方案B")
            scheme_b_after_a_stop = json.loads(body.decode("utf-8"))
            self.assertEqual(status, 200)
            self.assertEqual(scheme_b_after_a_stop["status"], "运行中")
            self.assertEqual(scheme_b_after_a_stop["scheme"], "方案B")

            status, headers, body = server.handle_control_path(
                "/api/optimization/control",
                json.dumps({"action": "bad"}).encode("utf-8"),
            )
            self.assertEqual(status, 400)
            self.assertEqual(json.loads(body.decode("utf-8"))["error"], "bad_request")
        finally:
            for runtime in server.OPTIMIZATION_RUNTIME.runtimes().values():
                if runtime.status == "运行中":
                    runtime.apply("stop", scheme=runtime.scheme)
            server.PLANNING_STORE = original_store
            server.OPTIMIZATION_RUNTIME = original_runtime
            shutil.rmtree(planning_root, ignore_errors=True)

    def test_optimization_start_rejects_fast_infeasible_planning_bounds_before_solving(self):
        planning_root = WEB_ROOT / "tests" / "tmp_optimization_fast_infeasible"
        shutil.rmtree(planning_root, ignore_errors=True)
        planning_root.mkdir(parents=True)
        original_store = server.PLANNING_STORE
        original_runtime = server.OPTIMIZATION_RUNTIME
        server.PLANNING_STORE = server.planning_store.PlanningStore(root=planning_root)
        server.OPTIMIZATION_RUNTIME = server.OptimizationRuntimeManager()

        def write_and_start(payload: dict) -> tuple[int, dict]:
            server.PLANNING_STORE.write_scheme("方案A", payload)
            status, headers, body = server.handle_control_path(
                "/api/optimization/control",
                json.dumps({"action": "start", "scheme": "方案A"}, ensure_ascii=False).encode("utf-8"),
            )
            return status, json.loads(body.decode("utf-8"))

        try:
            payload = server.planning_store.default_payload("方案A")
            payload["planning_parameters"][0]["green_power_ratio_lower"] = 0.2
            payload["wind_turbines"][0]["quantity_upper"] = 0
            payload["photovoltaics"][0]["quantity_upper"] = 0
            status, data = write_and_start(payload)
            self.assertEqual(status, 400)
            self.assertIn("风机数量上限和光伏数量上限均为0", data["message"])
            self.assertIn("绿色电量占比下限大于0", data["message"])
            status, headers, body = server.handle_api_path("/api/optimization/status?scheme=方案A&light=1")
            failed_task = json.loads(body.decode("utf-8"))
            self.assertEqual(status, 200)
            self.assertEqual(failed_task["status"], "失败")
            self.assertEqual(failed_task["task_status"], "计算失败")
            self.assertIn("绿色电量占比下限大于0", failed_task["logs"][-1]["message"])

            for row in payload["time_series"]:
                row["wind_speed"] = 0
                row["solar_irradiance"] = 0
                row["load"] = 100
            payload["wind_turbines"][0]["quantity_upper"] = 1
            payload["photovoltaics"][0]["quantity_upper"] = 1
            status, data = write_and_start(payload)
            self.assertEqual(status, 400)
            self.assertIn("风机和光伏最大可发电量", data["message"])
            self.assertIn("低于绿色电量占比要求", data["message"])

            payload["planning_parameters"][0]["green_power_ratio_lower"] = 0
            payload["diesel_generators"][0].update({"quantity_upper": 1, "power_upper": 20})
            payload["wind_turbines"][0]["quantity_upper"] = 0
            payload["photovoltaics"][0]["quantity_upper"] = 0
            payload["time_series"][0]["load"] = 100
            payload["diesel_generators"][0]["power_upper"] = 100
            payload["hydrogen_tanks"][0].update({"quantity_lower": 1, "quantity_upper": 1, "self_discharge_rate": 0.001})
            payload["hydrogen_electrolyzers"][0]["quantity_upper"] = 0
            status, data = write_and_start(payload)
            self.assertEqual(status, 400)
            self.assertIn("储氢罐数量下限大于0", data["message"])
            self.assertIn("电制氢数量上限为0", data["message"])
            self.assertIn("自损耗无法补偿", data["message"])

            payload["hydrogen_tanks"][0]["self_discharge_rate"] = 0
            payload["storage_battery_packs"][0].update({"quantity_lower": 1, "quantity_upper": 1, "self_discharge_rate": 0.01})
            payload["storage_pcs"][0]["quantity_upper"] = 0
            status, data = write_and_start(payload)
            self.assertEqual(status, 400)
            self.assertIn("储能电池数量下限大于0", data["message"])
            self.assertIn("储能PCS数量上限为0", data["message"])
            self.assertIn("自损耗无法补偿", data["message"])
        finally:
            for runtime in server.OPTIMIZATION_RUNTIME.runtimes().values():
                if runtime.status == "运行中":
                    runtime.apply("stop", scheme=runtime.scheme)
            server.PLANNING_STORE = original_store
            server.OPTIMIZATION_RUNTIME = original_runtime
            shutil.rmtree(planning_root, ignore_errors=True)

    def test_fast_feasibility_prechecks_are_shared_outside_server(self):
        server_source = (WEB_ROOT / "server.py").read_text(encoding="utf-8")
        precheck_source = (WEB_ROOT / "calculation_precheck.py").read_text(encoding="utf-8")
        docs = (WEB_ROOT / "docs" / "启动优化程序说明.md").read_text(encoding="utf-8")

        self.assertIn("import calculation_precheck", server_source)
        self.assertNotIn("def validate_optimization_fast_feasibility", server_source)
        self.assertNotIn("def validate_evaluation_fast_feasibility", server_source)
        self.assertIn("def validate_optimization_fast_feasibility", precheck_source)
        self.assertIn("def validate_evaluation_fast_feasibility", precheck_source)
        self.assertIn("启动前快速可行性预检查", docs)
        self.assertIn("年度风光最大可发电量不足", docs)
        self.assertNotIn("单小时最大供电功率不足", docs)
        self.assertNotIn("风机最大可发功率 + 光伏最大可发功率 + 柴发最大可发功率 < 负荷功率", docs)
        self.assertIn("储氢自损耗无法补偿", docs)
        self.assertIn("电储自损耗无法补偿", docs)

    def test_tasks_api_lists_and_controls_optimization_and_evaluation_jobs(self):
        original_optimization_runtime = server.OPTIMIZATION_RUNTIME
        original_evaluation_runtime = server.EVALUATION_RUNTIME
        planning_root = WEB_ROOT / "tests" / "tmp_tasks_api"
        shutil.rmtree(planning_root, ignore_errors=True)
        planning_root.mkdir(parents=True)
        original_store = server.PLANNING_STORE
        server.PLANNING_STORE = server.planning_store.PlanningStore(root=planning_root)
        server.OPTIMIZATION_RUNTIME = server.OptimizationRuntimeManager()
        server.EVALUATION_RUNTIME = server.EvaluationRuntimeManager()
        try:
            payload = server.planning_store.default_payload("方案A")
            for row in payload["time_series"]:
                row["wind_speed"] = 7
                row["solar_irradiance"] = 500
                row["load"] = 80
            payload["diesel_generators"][0]["quantity_upper"] = 2
            server.PLANNING_STORE.write_scheme("方案A", payload)
            workbook = Workbook()
            sheet = workbook.active
            sheet.title = "规划结果"
            sheet.append(["设备类型", "设计台数", "单台容量", "总容量", "单位"])
            sheet.append(["柴发", 2, 100, 200, "kW"])
            workbook.save(planning_root / "方案A" / "case_results.xlsx")
            workbook.close()
            default_result_workbook = Workbook()
            default_result_sheet = default_result_workbook.active
            default_result_sheet.title = "规划结果"
            default_result_sheet.append(["设备类型", "设计台数"])
            default_result_sheet.append(["柴发", 1])
            default_result_workbook.save(planning_root / "方案A" / "opt_results.xlsx")
            default_result_workbook.close()
            (planning_root / "方案A" / "broken_results.xlsx").write_text("not an xlsx workbook", encoding="utf-8")

            status, headers, body = server.handle_control_path(
                "/api/tasks/control",
                json.dumps({"action": "start", "task_type": "optimization", "scheme": "方案A"}, ensure_ascii=False).encode("utf-8"),
            )
            payload = json.loads(body.decode("utf-8"))
            self.assertEqual(status, 200)
            self.assertTrue(payload["ok"])
            self.assertEqual(payload["task"]["task_type"], "规划计算")
            self.assertEqual(payload["task"]["scheme"], "方案A")
            self.assertEqual(payload["task"]["status"], "计算中")
            self.assertIsInstance(payload["task"]["process_id"], int)
            self.assertNotEqual(payload["task"]["process_id"], server.os.getpid())
            self.assertEqual(payload["task"]["result"], "opt_results.xlsx")
            self.assertTrue(payload["task"]["start_time"])
            self.assertGreaterEqual(payload["task"]["elapsed_seconds"], 0)
            self.assertIn("后台规划求解程序已启动", payload["task"]["latest_log"])

            status, headers, body = server.handle_control_path(
                "/api/tasks/control",
                json.dumps({"action": "start", "task_type": "evaluation", "scheme": "方案A", "result": "case_results.xlsx"}, ensure_ascii=False).encode("utf-8"),
            )
            evaluation_payload = json.loads(body.decode("utf-8"))
            self.assertEqual(status, 200)
            self.assertEqual(evaluation_payload["task"]["task_type"], "方案评估")
            self.assertEqual(evaluation_payload["task"]["result"], "case_results.xlsx")
            self.assertEqual(evaluation_payload["task"]["status"], "计算中")
            self.assertIsInstance(evaluation_payload["task"]["process_id"], int)
            self.assertNotEqual(evaluation_payload["task"]["process_id"], server.os.getpid())

            status, headers, body = server.handle_api_path("/api/tasks")
            task_list = json.loads(body.decode("utf-8"))
            self.assertEqual(status, 200)
            self.assertIn("tasks", task_list)
            self.assertTrue(any(item["task_type"] == "规划计算" and item["scheme"] == "方案A" for item in task_list["tasks"]))
            self.assertTrue(any(item["task_type"] == "方案评估" and item["result"] == "case_results.xlsx" for item in task_list["tasks"]))
            self.assertFalse(any(item["task_type"] == "方案评估" and item["result"] == "opt_results.xlsx" for item in task_list["tasks"]))
            self.assertTrue(any(item["task_type"] == "方案评估" and item["result"] == "broken_results.xlsx" for item in task_list["tasks"]))
            type_order = [item["task_type_key"] for item in task_list["tasks"]]
            self.assertEqual(type_order, sorted(type_order, key=lambda value: 0 if value == "optimization" else 1))
            for task in task_list["tasks"]:
                for key in ("id", "task_key", "task_type", "scheme", "result", "status", "process_id", "start_time", "end_time", "elapsed_seconds", "latest_log", "can_start", "can_stop"):
                    self.assertIn(key, task)

            disabled_task = {
                "task_type_key": "evaluation",
                "scheme": "方案A",
                "result": "blocked_results.xlsx",
                "can_start": False,
                "can_queue": False,
                "can_stop": False,
                "queued": False,
            }
            self.assertFalse(server.task_list_item_is_visible(disabled_task))
            disabled_but_queued_task = {**disabled_task, "queued": True}
            self.assertTrue(server.task_list_item_is_visible(disabled_but_queued_task))
            disabled_but_running_task = {**disabled_task, "can_stop": True}
            self.assertTrue(server.task_list_item_is_visible(disabled_but_running_task))

            status, headers, body = server.handle_control_path(
                "/api/tasks/control",
                json.dumps({"action": "stop", "task_type": "optimization", "scheme": "方案A"}, ensure_ascii=False).encode("utf-8"),
            )
            stopped = json.loads(body.decode("utf-8"))
            self.assertEqual(status, 200)
            self.assertEqual(stopped["task"]["status"], "计算中止")
            self.assertIn("停止规划求解", stopped["task"]["latest_log"])

            status, headers, body = server.handle_control_path(
                "/api/tasks/control",
                json.dumps({"action": "bad", "task_type": "optimization", "scheme": "方案A"}, ensure_ascii=False).encode("utf-8"),
            )
            self.assertEqual(status, 400)
            self.assertEqual(json.loads(body.decode("utf-8"))["error"], "bad_request")
        finally:
            for runtime in server.OPTIMIZATION_RUNTIME.runtimes().values():
                if runtime.status == "运行中":
                    runtime.apply("stop", scheme=runtime.scheme)
            for runtime in server.EVALUATION_RUNTIME.runtimes().values():
                if runtime.status == "运行中":
                    runtime.apply("stop", scheme=runtime.scheme, filename=runtime.result_filename)
            server.OPTIMIZATION_RUNTIME = original_optimization_runtime
            server.EVALUATION_RUNTIME = original_evaluation_runtime
            server.PLANNING_STORE = original_store
            shutil.rmtree(planning_root, ignore_errors=True)

    def test_tasks_result_file_listing_checks_existence_without_opening_or_caching(self):
        planning_root = WEB_ROOT / "tests" / "tmp_tasks_result_cache"
        shutil.rmtree(planning_root, ignore_errors=True)
        planning_root.mkdir(parents=True)
        original_store = server.PLANNING_STORE
        original_health_check = server.result_workbook_error_message
        server.PLANNING_STORE = server.planning_store.PlanningStore(root=planning_root)

        def write_result(filename: str) -> None:
            workbook = Workbook()
            sheet = workbook.active
            sheet.title = "规划结果"
            sheet.append(["设备类型", "设计台数"])
            sheet.append(["柴发", 1])
            workbook.save(planning_root / "方案A" / filename)
            workbook.close()

        def forbidden_health_check(path: Path) -> str:
            raise AssertionError(f"任务并发刷新不应打开结果文件: {path.name}")

        try:
            server.PLANNING_STORE.create_scheme("方案A")
            write_result("case_results.xlsx")
            (planning_root / "方案A" / "broken_results.xlsx").write_text("not an xlsx workbook", encoding="utf-8")
            server.result_workbook_error_message = forbidden_health_check

            first = server.build_task_list(schedule=False)
            second = server.build_task_list(schedule=False)

            self.assertTrue(any(item["task_type_key"] == "evaluation" and item["result"] == "case_results.xlsx" for item in first))
            self.assertTrue(any(item["task_type_key"] == "evaluation" and item["result"] == "broken_results.xlsx" for item in first))
            self.assertEqual(first, second)

            write_result("case_2_results.xlsx")
            third = server.build_task_list(schedule=False)

            self.assertTrue(any(item["task_type_key"] == "evaluation" and item["result"] == "case_2_results.xlsx" for item in third))
        finally:
            server.result_workbook_error_message = original_health_check
            server.PLANNING_STORE = original_store
            shutil.rmtree(planning_root, ignore_errors=True)

    def test_status_apis_include_task_control_state_for_page_sync(self):
        original_optimization_runtime = server.OPTIMIZATION_RUNTIME
        original_evaluation_runtime = server.EVALUATION_RUNTIME
        original_scheduler = server.TASK_SCHEDULER
        planning_root = WEB_ROOT / "tests" / "tmp_task_status_sync"
        shutil.rmtree(planning_root, ignore_errors=True)
        planning_root.mkdir(parents=True)
        original_store = server.PLANNING_STORE
        server.PLANNING_STORE = server.planning_store.PlanningStore(root=planning_root)
        server.OPTIMIZATION_RUNTIME = server.OptimizationRuntimeManager()
        server.EVALUATION_RUNTIME = server.EvaluationRuntimeManager()
        server.TASK_SCHEDULER = server.TaskScheduler()
        try:
            for scheme in ("方案A", "方案B"):
                server.PLANNING_STORE.create_scheme(scheme)
            workbook = Workbook()
            sheet = workbook.active
            sheet.title = "规划结果"
            sheet.append(["设备类型", "设计台数", "单台容量", "总容量", "单位"])
            sheet.append(["柴发", 1, 100, 100, "kW"])
            workbook.save(planning_root / "方案A" / "case_results.xlsx")
            workbook.close()

            status, headers, body = server.handle_control_path(
                "/api/tasks/control",
                json.dumps({"action": "start", "task_type": "optimization", "scheme": "方案A"}, ensure_ascii=False).encode("utf-8"),
            )
            self.assertEqual(status, 200)
            status, headers, body = server.handle_control_path(
                "/api/tasks/control",
                json.dumps({"action": "queue", "task_type": "optimization", "scheme": "方案B"}, ensure_ascii=False).encode("utf-8"),
            )
            self.assertEqual(status, 200)

            status, headers, body = server.handle_api_path("/api/optimization/status?scheme=方案A&light=1")
            running = json.loads(body.decode("utf-8"))
            self.assertEqual(running["task_status"], "计算中")
            self.assertTrue(running["can_stop_task"])
            self.assertFalse(running["can_queue_task"])

            status, headers, body = server.handle_api_path("/api/optimization/status?scheme=方案B&light=1")
            queued = json.loads(body.decode("utf-8"))
            self.assertEqual(queued["task_status"], "排队中")
            self.assertEqual(queued["queue_position"], 1)
            self.assertTrue(queued["can_start_task"])
            self.assertFalse(queued["can_queue_task"])
            self.assertTrue(queued["can_cancel_queue_task"])

            status, headers, body = server.handle_control_path(
                "/api/tasks/control",
                json.dumps({"action": "cancel_queue", "task_type": "optimization", "scheme": "方案B"}, ensure_ascii=False).encode("utf-8"),
            )
            self.assertEqual(status, 200)
            cancelled = json.loads(body.decode("utf-8"))["task"]
            self.assertEqual(cancelled["status"], "退出队列")
            self.assertFalse(cancelled["queued"])
            self.assertTrue(cancelled["can_queue"])

            status, headers, body = server.handle_api_path("/api/optimization/status?scheme=方案B&light=1")
            exited = json.loads(body.decode("utf-8"))
            self.assertEqual(exited["task_status"], "退出队列")
            self.assertFalse(exited["queued"])
            self.assertTrue(exited["can_queue_task"])

            status, headers, body = server.handle_api_path("/api/evaluation/status?scheme=方案A&filename=case_results.xlsx&light=1")
            evaluation = json.loads(body.decode("utf-8"))
            self.assertEqual(evaluation["task_status"], "未计算")
            self.assertTrue(evaluation["can_start_task"])
            self.assertTrue(evaluation["can_queue_task"])

            status, headers, body = server.handle_control_path(
                "/api/tasks/control",
                json.dumps(
                    {"action": "queue", "task_type": "evaluation", "scheme": "方案A", "result": "case_results.xlsx"},
                    ensure_ascii=False,
                ).encode("utf-8"),
            )
            self.assertEqual(status, 200)
            queued_evaluation = json.loads(body.decode("utf-8"))["task"]
            self.assertEqual(queued_evaluation["status"], "排队中")

            status, headers, body = server.handle_control_path(
                "/api/tasks/control",
                json.dumps(
                    {"action": "cancel_queue", "task_type": "evaluation", "scheme": "方案A", "result": "case_results.xlsx"},
                    ensure_ascii=False,
                ).encode("utf-8"),
            )
            self.assertEqual(status, 200)
            exited_evaluation = json.loads(body.decode("utf-8"))["task"]
            self.assertEqual(exited_evaluation["status"], "退出队列")

            status, headers, body = server.handle_api_path("/api/evaluation/status?scheme=方案A&filename=case_results.xlsx&light=1")
            evaluation_after_exit = json.loads(body.decode("utf-8"))
            self.assertEqual(evaluation_after_exit["task_status"], "退出队列")
            self.assertFalse(evaluation_after_exit["queued"])
            self.assertTrue(evaluation_after_exit["can_queue_task"])
        finally:
            for runtime in server.OPTIMIZATION_RUNTIME.runtimes().values():
                if runtime.status == "运行中":
                    runtime.apply("stop", scheme=runtime.scheme)
            server.OPTIMIZATION_RUNTIME = original_optimization_runtime
            server.EVALUATION_RUNTIME = original_evaluation_runtime
            server.TASK_SCHEDULER = original_scheduler
            server.PLANNING_STORE = original_store
            shutil.rmtree(planning_root, ignore_errors=True)

    def test_task_page_splits_task_types_and_supports_resizing(self):
        html = (WEB_ROOT / "tasks.html").read_text(encoding="utf-8")
        script = (WEB_ROOT / "assets" / "tasks.js").read_text(encoding="utf-8")
        css = (WEB_ROOT / "assets" / "planning.css").read_text(encoding="utf-8")

        self.assertIn('id="optimizationTaskTable"', html)
        self.assertIn('id="evaluationTaskTable"', html)
        self.assertIn('id="evaluationSchemeFilter"', html)
        self.assertIn("renderTaskSection(\"optimization\"", script)
        self.assertIn("renderTaskSection(\"evaluation\"", script)
        self.assertIn("handleEvaluationSchemeFilterChange", script)
        self.assertIn("renderSchemeFilter", script)
        self.assertIn("filteredTasksForSection", script)
        self.assertIn('taskState.evaluationSchemeFilter', script)
        self.assertIn('taskState.frequencySchemeFilter', script)
        self.assertIn("lastRenderedTaskSignature", script)
        self.assertIn("function applyTasksPayload", script)
        self.assertIn("signature === taskState.lastRenderedTaskSignature", script)
        self.assertIn('renderSchemeFilter("evaluation", "evaluationSchemeFilter", "evaluationSchemeFilter")', script)
        self.assertIn('renderSchemeFilter("frequency", "frequencySchemeFilter", "frequencySchemeFilter")', script)
        self.assertIn('const selectedScheme = String(taskState[stateKey] || "").trim();', script)
        self.assertIn("if (selectedScheme && !schemeNames.includes(selectedScheme)) schemeNames.push(selectedScheme);", script)
        self.assertIn("select.value = selectedScheme;", script)
        self.assertNotIn('taskState[stateKey] = "";', script)
        self.assertNotIn('rememberTasksPageState({ [stateKey]: "" });', script)
        self.assertIn("task.task_type_key === taskTypeKey", script)
        self.assertIn('if (taskTypeKey === "evaluation" && taskState.evaluationSchemeFilter)', script)
        self.assertIn('if (taskTypeKey === "frequency" && taskState.frequencySchemeFilter)', script)
        self.assertIn("方案名称", (WEB_ROOT / "assets" / "i18n.js").read_text(encoding="utf-8"))
        self.assertIn("全部方案", (WEB_ROOT / "assets" / "i18n.js").read_text(encoding="utf-8"))
        self.assertIn("<colgroup>", script)
        self.assertIn("task-col-scheme", script)
        self.assertIn("task-col-actions", script)
        self.assertIn('class="task-actions-cell"', script)
        self.assertLess(script.index('<col class="task-col-actions">'), script.index('<col class="task-col-status">'))
        self.assertLess(script.index("<th>操作</th>"), script.index("<th>任务状态</th>"))
        self.assertLess(script.index('<div class="task-actions">'), script.index('<td><span class="task-status-pill'))
        self.assertIn(">启动</button>", script)
        self.assertIn(">排队</button>", script)
        self.assertIn('label: "离队"', script)
        self.assertIn('label: "停止"', script)
        self.assertNotIn(">立刻启动</button>", script)
        self.assertNotIn(">加入排队</button>", script)
        self.assertNotIn('label: "退出队列"', script)
        self.assertNotIn('label: "停止计算"', script)
        self.assertNotIn("refreshTasks", script)
        self.assertIn(".tasks-workspace > .tasks-panel", css)
        tasks_workspace_css = css.split(".tasks-workspace {", 1)[1].split("}", 1)[0]
        self.assertIn("grid-template-rows: minmax(0, 1fr)", tasks_workspace_css)
        self.assertIn(".task-section-evaluation", css)
        self.assertIn(".task-section-frequency", css)
        self.assertIn(".task-section.active", css)
        self.assertIn("grid-row: 2", css)
        self.assertIn(".task-section-evaluation .task-section-head", css)
        evaluation_head_css = css.split(".task-section-evaluation .task-section-head,", 1)[1].split("}", 1)[0]
        self.assertIn("justify-content: flex-start", evaluation_head_css)
        self.assertIn(".tasks-panel .task-table", css)
        task_table_panel_css = css.split(".tasks-panel .task-table {", 1)[1].split("}", 1)[0]
        self.assertIn("overflow-x: hidden", task_table_panel_css)
        self.assertIn("table-layout: fixed", css)
        task_table_css = css.split(".task-table table {", 1)[1].split("}", 1)[0]
        self.assertIn("min-width: 0", task_table_css)
        self.assertNotIn("min-width: 1120px", task_table_css)
        self.assertIn(".task-col-scheme", css)
        self.assertIn(".task-col-actions", css)
        self.assertIn(".task-actions-cell", css)
        self.assertNotIn(".task-table td:last-child", css)
        task_log_cell_css = css.split(".task-log-cell {", 1)[1].split("}", 1)[0]
        self.assertIn("max-width: 0", task_log_cell_css)
        self.assertIn("overflow: hidden", task_log_cell_css)
        self.assertIn("text-overflow: ellipsis", task_log_cell_css)
        self.assertIn('if (status === "计算中止") return "interrupted";', script)
        self.assertIn('if (status === "计算失败") return "failed";', script)
        self.assertIn('if (status === "计算超时") return "timeout";', script)
        self.assertIn(".task-status-pill.interrupted", css)
        self.assertIn(".task-status-pill.failed", css)
        self.assertIn(".task-status-pill.timeout", css)
        self.assertNotIn("<th>任务类型</th>", script)

    def test_timeout_solver_result_maps_to_timeout_task_status(self):
        self.assertTrue(milp_solver.is_timeout_result(SimpleNamespace(message="Gurobi status TIME_LIMIT", status=9)))
        self.assertTrue(milp_solver.is_timeout_result(SimpleNamespace(message="Time limit reached", status=1)))
        self.assertFalse(milp_solver.is_timeout_result(SimpleNamespace(message="Gurobi status INFEASIBLE", status=3)))

        with self.assertRaises(milp_solver.CalculationTimeoutError):
            plan_optimizer.raise_if_solver_timed_out(
                SimpleNamespace(message="Gurobi status TIME_LIMIT", status=9),
                "规划求解",
            )

        optimization_runtime = server.OptimizationRuntime("方案A")
        with optimization_runtime._lock:
            optimization_runtime.status = "运行中"
            optimization_runtime.scheme = "方案A"
            optimization_runtime.start_time = server._now_text()
            optimization_runtime._handle_process_event_unlocked({"type": "timeout", "message": "规划求解达到最大用时"})
            optimization_payload = optimization_runtime._payload_unlocked()
        self.assertEqual(optimization_payload["status"], "超时")
        self.assertEqual(server.task_display_status(optimization_payload["status"]), "计算超时")
        self.assertIn("最大用时", optimization_payload["logs"][-1]["message"])

        evaluation_runtime = server.EvaluationRuntime("方案A")
        with evaluation_runtime._lock:
            evaluation_runtime.status = "运行中"
            evaluation_runtime.scheme = "方案A"
            evaluation_runtime.result_filename = "aaa_results.xlsx"
            evaluation_runtime.start_time = server._now_text()
            evaluation_runtime._handle_process_event_unlocked({"type": "timeout", "message": "方案评估达到最大用时"})
            evaluation_payload = evaluation_runtime._payload_unlocked()
        self.assertEqual(evaluation_payload["status"], "超时")
        self.assertEqual(server.task_display_status(evaluation_payload["status"]), "计算超时")
        self.assertIn("最大用时", evaluation_payload["logs"][-1]["message"])

    def test_optimization_done_export_uses_runtime_results_not_stale_workbook(self):
        runtime = server.OptimizationRuntime("方案A")
        new_results = {"curves": {"green_hourly": [{"load_down_disturbance_power": -97.35}]}}
        captured = {}

        def fake_export(payload):
            captured["runtime_status_during_export"] = runtime.status
            captured["payload"] = payload
            return WEB_ROOT / "tests" / "tmp_opt_results.xlsx"

        with runtime._lock:
            runtime.status = "运行中"
            runtime.scheme = "方案A"
            runtime.start_time = server._now_text()
            with patch.object(
                server,
                "read_result_workbook_display_payload_for_response",
                side_effect=AssertionError("导出本次规划结果时不应读取旧结果工作簿"),
            ), patch.object(server, "export_optimization_results_workbook", side_effect=fake_export):
                runtime._handle_process_event_unlocked(
                    {
                        "type": "done",
                        "metrics": [{"label": "度电成本", "value": 1.23, "unit": "元"}],
                        "results": new_results,
                    }
                )

        self.assertEqual(runtime.status, "已完成")
        self.assertEqual(captured["runtime_status_during_export"], "运行中")
        self.assertEqual(captured["payload"]["status"], "已完成")
        self.assertTrue(captured["payload"]["end_time"])
        self.assertEqual(captured["payload"]["results"], new_results)
        self.assertEqual(
            captured["payload"]["results"]["curves"]["green_hourly"][0]["load_down_disturbance_power"],
            -97.35,
        )

    def test_evaluation_done_export_uses_runtime_results_not_stale_workbook(self):
        runtime = server.EvaluationRuntime("方案A")
        new_results = {"curves": {"green_hourly": [{"load_down_disturbance_power": -97.35}]}}
        dispatch_rows = [{"hour": 1, "load_down_disturbance_power": -97.35}]
        captured = {}

        def fake_export(payload, rows):
            captured["runtime_status_during_export"] = runtime.status
            captured["payload"] = payload
            captured["dispatch_rows"] = rows
            return WEB_ROOT / "tests" / "case_results.xlsx"

        with runtime._lock:
            runtime.status = "运行中"
            runtime.scheme = "方案A"
            runtime.result_filename = "case_results.xlsx"
            runtime.start_time = server._now_text()
            with patch.object(
                server,
                "read_result_workbook_display_payload_for_response",
                side_effect=AssertionError("导出本次评估结果时不应读取旧结果工作簿"),
            ), patch.object(server, "export_evaluation_results_workbook", side_effect=fake_export):
                runtime._handle_process_event_unlocked(
                    {
                        "type": "done",
                        "metrics": [{"label": "度电成本", "value": 1.23, "unit": "元"}],
                        "results": new_results,
                        "dispatch_rows": dispatch_rows,
                    }
                )

        self.assertEqual(runtime.status, "已完成")
        self.assertEqual(captured["runtime_status_during_export"], "运行中")
        self.assertEqual(captured["payload"]["status"], "已完成")
        self.assertTrue(captured["payload"]["end_time"])
        self.assertEqual(captured["payload"]["results"], new_results)
        self.assertEqual(captured["dispatch_rows"], dispatch_rows)
        self.assertEqual(
            captured["payload"]["results"]["curves"]["green_hourly"][0]["load_down_disturbance_power"],
            -97.35,
        )

    def test_tasks_api_queues_jobs_and_starts_next_after_current_finishes(self):
        original_optimization_runtime = server.OPTIMIZATION_RUNTIME
        original_evaluation_runtime = server.EVALUATION_RUNTIME
        original_scheduler = server.TASK_SCHEDULER
        planning_root = WEB_ROOT / "tests" / "tmp_tasks_queue"
        shutil.rmtree(planning_root, ignore_errors=True)
        planning_root.mkdir(parents=True)
        original_store = server.PLANNING_STORE
        server.PLANNING_STORE = server.planning_store.PlanningStore(root=planning_root)
        server.OPTIMIZATION_RUNTIME = server.OptimizationRuntimeManager()
        server.EVALUATION_RUNTIME = server.EvaluationRuntimeManager()
        server.TASK_SCHEDULER = server.TaskScheduler()
        try:
            for scheme in ("方案A", "方案B"):
                server.PLANNING_STORE.create_scheme(scheme)

            status, headers, body = server.handle_control_path(
                "/api/tasks/control",
                json.dumps({"action": "start", "task_type": "optimization", "scheme": "方案A"}, ensure_ascii=False).encode("utf-8"),
            )
            first = json.loads(body.decode("utf-8"))
            self.assertEqual(status, 200)
            self.assertEqual(first["task"]["status"], "计算中")

            status, headers, body = server.handle_control_path(
                "/api/tasks/control",
                json.dumps({"action": "queue", "task_type": "optimization", "scheme": "方案B"}, ensure_ascii=False).encode("utf-8"),
            )
            queued = json.loads(body.decode("utf-8"))
            self.assertEqual(status, 200)
            self.assertEqual(queued["task"]["status"], "排队中")
            self.assertEqual(queued["task"]["queue_position"], 1)
            self.assertFalse(queued["task"]["can_queue"])

            runtime_a = server.OPTIMIZATION_RUNTIME._runtime_for_scheme("方案A")
            runtime_a.apply("stop", scheme="方案A")

            status, headers, body = server.handle_api_path("/api/tasks")
            task_list = json.loads(body.decode("utf-8"))["tasks"]
            task_b = next(item for item in task_list if item["task_type_key"] == "optimization" and item["scheme"] == "方案B")
            self.assertEqual(task_b["status"], "计算中")
            self.assertIsInstance(task_b["process_id"], int)

            status, headers, body = server.handle_control_path(
                "/api/tasks/control",
                json.dumps({"action": "stop", "task_type": "optimization", "scheme": "方案B"}, ensure_ascii=False).encode("utf-8"),
            )
            stopped = json.loads(body.decode("utf-8"))["task"]
            self.assertEqual(stopped["status"], "计算中止")
        finally:
            for runtime in server.OPTIMIZATION_RUNTIME.runtimes().values():
                if runtime.status == "运行中":
                    runtime.apply("stop", scheme=runtime.scheme)
            server.OPTIMIZATION_RUNTIME = original_optimization_runtime
            server.EVALUATION_RUNTIME = original_evaluation_runtime
            server.TASK_SCHEDULER = original_scheduler
            server.PLANNING_STORE = original_store
            shutil.rmtree(planning_root, ignore_errors=True)

    def test_runtime_starts_child_process_for_optimization_tasks(self):
        planning_root = WEB_ROOT / "tests" / "tmp_process_runtime"
        shutil.rmtree(planning_root, ignore_errors=True)
        planning_root.mkdir(parents=True)
        original_store = server.PLANNING_STORE
        server.PLANNING_STORE = server.planning_store.PlanningStore(root=planning_root)
        try:
            server.PLANNING_STORE.create_scheme("方案A")
            runtime = server.OptimizationRuntime()
            payload = runtime.apply("start", scheme="方案A")

            self.assertEqual(payload["status"], "运行中")
            self.assertIsNotNone(runtime._process)
            self.assertFalse(hasattr(runtime, "_thread"))
            self.assertIsInstance(payload["process_id"], int)
            self.assertNotEqual(payload["process_id"], server.os.getpid())

            stopped = runtime.apply("stop", scheme="方案A")
            self.assertEqual(stopped["status"], "计算中止")
            self.assertFalse(runtime._process and runtime._process.is_alive())
        finally:
            server.PLANNING_STORE = original_store
            shutil.rmtree(planning_root, ignore_errors=True)

    def test_estimate_dispatch_minimizes_diesel_for_8760_hours(self):
        payload = server.planning_store.default_payload("方案A")
        for index, row in enumerate(payload["time_series"]):
            row["wind_speed"] = 8 if index % 2 == 0 else 0
            row["solar_irradiance"] = 800 if 8 <= index % 24 <= 16 else 0
            row["load"] = 100
            row["temperature"] = 20
        result_rows = [
            {"设备类型": "柴发", "设计台数": 2, "单台容量": 100, "总容量": 200, "单位": "kW"},
            {"设备类型": "风机", "设计台数": 1, "单台容量": 100, "总容量": 100, "单位": "kW"},
            {"设备类型": "光伏", "设计台数": 1, "单台容量": 100, "总容量": 100, "单位": "kW"},
            {"设备类型": "储能", "设计台数": 1, "单台容量": 100, "总容量": 100, "单位": "kWh"},
        ]
        events = []

        result = estimate.run_estimation(payload, result_rows, log=events.append)

        self.assertEqual(result["status"], "已完成")
        self.assertEqual(result["progress"], 100)
        self.assertEqual(len(result["dispatch_rows"]), 8760)
        self.assertEqual(result["dispatch_rows"][0]["hour_index"], 1)
        self.assertEqual(result["dispatch_rows"][-1]["hour_index"], 8760)
        self.assertLess(result["totals"]["diesel_energy"], result["totals"]["load_energy"])
        self.assertTrue(any("8760点优化调度" in item["message"] for item in events))
        self.assertTrue(any(item.get("progress") == 100 for item in events))

    def test_estimate_dispatch_reports_levelized_cost_and_green_ratio_metrics(self):
        payload = server.planning_store.default_payload("方案A")
        for row in payload["time_series"]:
            row["wind_speed"] = 0
            row["solar_irradiance"] = 0
            row["load"] = 50
            row["temperature"] = 20
        payload["planning_parameters"][0]["diesel_price"] = 2
        payload["diesel_generators"][0].update(
            {
                "name": "评估柴发",
                "capacity": 100,
                "cost": 100,
                "design_life_years": 20,
                "fuel_rate": 0.5,
                "power_lower": 0,
                "power_upper": 100,
            }
        )
        result_rows = [
            {"设备类型": "柴发", "名称": "评估柴发", "设计台数": 1, "单台容量": 100, "总容量": 100, "单位": "kW"},
        ]

        result = estimate.run_estimation(payload, result_rows)
        metrics = {item["label"]: item for item in result["metrics"]}
        annual_rows = {
            row["指标"]: row
            for row in result["results"]["overview_tables"][1]["rows"]
            if isinstance(row, dict) and row.get("指标")
        }
        expected_diesel_consumption = 8760 * 50 * 0.5 / 1000
        expected_diesel_cost = expected_diesel_consumption * 2
        expected_construction_cost = 100 / 20
        expected_lcoe = (expected_construction_cost + expected_diesel_cost) * 10000 / (8760 * 50)

        self.assertAlmostEqual(metrics["度电成本"]["value"], expected_lcoe, places=6)
        self.assertEqual(metrics["度电成本"]["unit"], "元")
        self.assertEqual(metrics["绿电占比"]["value"], 0.0)
        self.assertNotIn("综合评分", metrics)
        self.assertNotIn("风险等级", metrics)
        self.assertAlmostEqual(annual_rows["年均建设成本"]["数值"], expected_construction_cost, places=4)
        self.assertAlmostEqual(annual_rows["年柴油成本"]["数值"], expected_diesel_cost, places=4)
        self.assertAlmostEqual(annual_rows["度电成本"]["数值"], expected_lcoe, places=6)

    def test_estimate_dispatch_uses_joint_milp_with_curtailment_and_hydrogen(self):
        payload = server.planning_store.default_payload("方案A")
        for row in payload["time_series"]:
            row["wind_speed"] = 12
            row["solar_irradiance"] = 1000
            row["load"] = 40
            row["temperature"] = 20
        payload["diesel_generators"][0]["capacity"] = 100
        payload["diesel_generators"][0]["power_lower"] = 20
        payload["diesel_generators"][0]["power_upper"] = 100
        result_rows = [
            {"设备类型": "柴发", "设计台数": 1, "单台容量": 100, "总容量": 100, "单位": "kW"},
            {"设备类型": "风机", "设计台数": 1, "单台容量": 100, "总容量": 100, "单位": "kW"},
            {"设备类型": "光伏", "设计台数": 1, "单台容量": 100, "总容量": 100, "单位": "kW"},
            {"设备类型": "储能PCS", "设计台数": 1, "单台容量": 50, "总容量": 50, "单位": "kW"},
            {"设备类型": "储能电池组", "设计台数": 1, "单台容量": 100, "总容量": 100, "单位": "kWh"},
            {"设备类型": "电制氢", "设计台数": 1, "单台容量": 30, "总容量": 30, "单位": "kW"},
            {"设备类型": "储氢罐", "设计台数": 1, "单台容量": 100, "总容量": 100, "单位": "Nm3"},
            {"设备类型": "燃料电池", "设计台数": 1, "单台容量": 20, "总容量": 20, "单位": "kW"},
        ]
        events = []

        result = estimate.run_estimation(payload, result_rows, log=events.append)
        first_row = result["dispatch_rows"][0]

        self.assertEqual(result["status"], "已完成")
        self.assertEqual(len(result["dispatch_rows"]), 8760)
        self.assertIn("diesel_on", first_row)
        self.assertIn("hydrogen_production_power", first_row)
        self.assertIn("fuel_cell_power", first_row)
        self.assertIn("hydrogen_storage", first_row)
        self.assertGreater(result["totals"]["curtailed_energy"], 0)
        self.assertEqual(result["totals"]["unmet_load_energy"], 0)
        self.assertTrue(any("混合整数线性优化" in item["message"] for item in events))

    def test_estimate_dispatch_uses_surplus_renewable_hydrogen_to_reduce_later_diesel(self):
        payload = server.planning_store.default_payload("方案A")
        payload["time_series"] = payload["time_series"][:2]
        payload["time_series"][0].update({"wind_speed": 12, "solar_irradiance": 0, "load": 0, "temperature": 20})
        payload["time_series"][1].update({"wind_speed": 0, "solar_irradiance": 0, "load": 40, "temperature": 20})
        payload["diesel_generators"][0].update({"capacity": 100, "power_lower": 0, "power_upper": 100})
        payload["hydrogen_electrolyzers"][0].update(
            {"power_capacity": 100, "power_lower": 0, "electric_to_hydrogen_efficiency": 1}
        )
        payload["fuel_cells"][0].update({"power_capacity": 100, "hydrogen_to_electric_efficiency": 1})
        payload["hydrogen_tanks"][0].update({"hydrogen_tank_capacity": 200, "self_discharge_rate": 0})
        result_rows = [
            {"设备类型": "柴发", "设计台数": 1, "单台容量": 100, "总容量": 100, "单位": "kW"},
            {"设备类型": "风机", "设计台数": 1, "单台容量": 100, "总容量": 100, "单位": "kW"},
            {"设备类型": "电制氢", "设计台数": 1, "单台容量": 100, "总容量": 100, "单位": "kW"},
            {"设备类型": "储氢罐", "设计台数": 1, "单台容量": 200, "总容量": 200, "单位": "Nm3"},
            {"设备类型": "燃料电池", "设计台数": 1, "单台容量": 100, "总容量": 100, "单位": "kW"},
        ]

        dispatch_rows = estimate.solve_dispatch_model(estimate.build_dispatch_model(payload, result_rows))

        self.assertGreater(dispatch_rows[0]["hydrogen_production_power"], 0)
        self.assertGreater(dispatch_rows[0]["hydrogen_storage"], 100)
        self.assertGreater(dispatch_rows[1]["fuel_cell_power"], 0)
        self.assertLess(dispatch_rows[1]["diesel_power"], 40)

    def test_estimate_dispatch_uses_time_limit_from_planning_parameters(self):
        payload = server.planning_store.default_payload("方案A")
        payload["time_series"] = payload["time_series"][:1]
        payload["planning_parameters"][0]["optimization_time_limit_minutes"] = 45
        result_rows = [
            {"设备类型": "柴发", "设计台数": 1, "单台容量": 100, "总容量": 100, "单位": "kW"},
            {"设备类型": "电制氢", "设计台数": 1, "单台容量": 10, "总容量": 10, "单位": "kW"},
            {"设备类型": "储氢罐", "设计台数": 1, "单台容量": 20, "总容量": 20, "单位": "Nm3"},
            {"设备类型": "燃料电池", "设计台数": 1, "单台容量": 10, "总容量": 10, "单位": "kW"},
        ]
        model = estimate.build_dispatch_model(payload, result_rows)
        seen_options = {}

        def fake_solve_milp(c, integrality, lower_bounds, upper_bounds, constraints, constraint_lower, constraint_upper, options, log, problem_name):
            seen_options.update(options)
            return SimpleNamespace(success=True, x=lower_bounds.copy(), fun=0.0, message="ok")

        with patch.object(estimate, "solve_milp", side_effect=fake_solve_milp):
            estimate.solve_dispatch_model(model)

        self.assertEqual(seen_options["time_limit"], 2700)

    def test_estimate_dispatch_enforces_storage_daily_and_hydrogen_annual_cycle(self):
        payload = server.planning_store.default_payload("方案A")
        for row in payload["time_series"]:
            row["wind_speed"] = 12
            row["solar_irradiance"] = 1000
            row["load"] = 40
            row["temperature"] = 20
        payload["diesel_generators"][0]["capacity"] = 100
        payload["diesel_generators"][0]["power_lower"] = 20
        payload["diesel_generators"][0]["power_upper"] = 100
        result_rows = [
            {"设备类型": "柴发", "设计台数": 1, "单台容量": 100, "总容量": 100, "单位": "kW"},
            {"设备类型": "风机", "设计台数": 1, "单台容量": 100, "总容量": 100, "单位": "kW"},
            {"设备类型": "光伏", "设计台数": 1, "单台容量": 100, "总容量": 100, "单位": "kW"},
            {"设备类型": "储能PCS", "设计台数": 1, "单台容量": 50, "总容量": 50, "单位": "kW"},
            {"设备类型": "储能电池组", "设计台数": 1, "单台容量": 100, "总容量": 100, "单位": "kWh"},
            {"设备类型": "电制氢", "设计台数": 1, "单台容量": 30, "总容量": 30, "单位": "kW"},
            {"设备类型": "储氢罐", "设计台数": 1, "单台容量": 100, "总容量": 100, "单位": "Nm3"},
            {"设备类型": "燃料电池", "设计台数": 1, "单台容量": 20, "总容量": 20, "单位": "kW"},
        ]

        result = estimate.run_estimation(payload, result_rows)
        dispatch_rows = result["dispatch_rows"]
        model = estimate.build_dispatch_model(payload, result_rows)
        fixed_quantities = {
            (device["key"], device["index"]): device["quantity_lower"]
            for devices in model["device_rows"].values()
            for device in devices
        }
        expected_storage_start = sum(
            device["capacity"] * fixed_quantities[(device["key"], device["index"])]
            for device in model["device_rows"]["storage_battery_packs"]
        ) * 0.5
        expected_hydrogen_start = sum(
            device["capacity"] * fixed_quantities[(device["key"], device["index"])]
            for device in model["device_rows"]["hydrogen_tanks"]
        ) * 0.5

        for day_end_hour in range(23, 8760, 24):
            self.assertAlmostEqual(dispatch_rows[day_end_hour]["storage_soc"], expected_storage_start, places=3)
        self.assertAlmostEqual(dispatch_rows[-1]["hydrogen_storage"], expected_hydrogen_start, places=3)

    def test_estimate_capacity_mapping_keeps_fuel_cells_out_of_storage_energy(self):
        result_rows = [
            {"设备类型": "储能PCS", "设计台数": 1, "单台容量": 50, "总容量": 50, "单位": "kW"},
            {"设备类型": "储能电池组", "设计台数": 1, "单台容量": 100, "总容量": 100, "单位": "kWh"},
            {"设备类型": "燃料电池", "设计台数": 1, "单台容量": 20, "总容量": 20, "单位": "kW"},
        ]

        capacities = estimate.capacities_from_planning_rows(result_rows)

        self.assertEqual(capacities["storage_power_capacity"], 50)
        self.assertEqual(capacities["storage_energy_capacity"], 100)
        self.assertEqual(capacities["fuel_cell_power_capacity"], 20)

    def test_estimate_capacity_mapping_recomputes_total_capacity_from_count_and_unit_capacity(self):
        result_rows = [
            {"设备类型": "电制氢", "设计台数": 1, "单台容量": 80, "总容量": 0, "单位": "kW"},
            {"设备类型": "储氢罐", "设计台数": 1, "单台容量": 1000, "总容量": 0, "单位": "Nm3"},
            {"设备类型": "燃料电池", "设计台数": 1, "单台容量": 50, "总容量": 0, "单位": "kW"},
            {"设备类型": "柴发", "设计台数": 0, "单台容量": 300, "总容量": 300, "单位": "kW"},
        ]

        capacities = estimate.capacities_from_planning_rows(result_rows)

        self.assertEqual(capacities["electrolyzer_power_capacity"], 80)
        self.assertEqual(capacities["hydrogen_tank_capacity"], 1000)
        self.assertEqual(capacities["fuel_cell_power_capacity"], 50)
        self.assertEqual(capacities["diesel_capacity"], 0)

    def test_estimate_dispatch_uses_count_variables_and_initial_storage_ratios(self):
        payload = server.planning_store.default_payload("方案A")
        payload["time_series"] = payload["time_series"][:1]
        payload["time_series"][0]["load"] = 0
        payload["time_series"][0]["wind_speed"] = 0
        payload["time_series"][0]["solar_irradiance"] = 0
        payload["planning_parameters"][0]["initial_storage_soc_ratio"] = 0.2
        payload["planning_parameters"][0]["initial_hydrogen_storage_ratio"] = 0.8
        payload["storage_pcs"][0]["storage_charge_efficiency"] = 0.91
        payload["storage_pcs"][0]["storage_discharge_efficiency"] = 0.89
        payload["planning_parameters"][0]["post_disturbance_power_balance_enabled"] = 1
        payload["diesel_generators"][0]["capacity"] = 100
        payload["diesel_generators"][0]["power_lower"] = 20
        payload["diesel_generators"][0]["power_upper"] = 100
        payload["storage_pcs"][0]["is_grid_forming"] = 1
        payload["storage_battery_packs"][0]["soc_upper"] = 0.8
        payload["storage_battery_packs"][0]["soc_lower"] = 0.2
        payload["hydrogen_electrolyzers"][0]["power_capacity"] = 30
        payload["hydrogen_electrolyzers"][0]["power_lower"] = 5
        result_rows = [
            {"设备类型": "柴发", "设计台数": 2, "单台容量": 100, "总容量": 200, "单位": "kW"},
            {"设备类型": "储能PCS", "设计台数": 1, "单台容量": 50, "总容量": 50, "单位": "kW"},
            {"设备类型": "储能电池组", "设计台数": 1, "单台容量": 100, "总容量": 100, "单位": "kWh"},
            {"设备类型": "电制氢", "设计台数": 2, "单台容量": 30, "总容量": 60, "单位": "kW"},
            {"设备类型": "储氢罐", "设计台数": 1, "单台容量": 100, "总容量": 100, "单位": "Nm3"},
            {"设备类型": "燃料电池", "设计台数": 1, "单台容量": 20, "总容量": 20, "单位": "kW"},
        ]
        model = estimate.build_dispatch_model(payload, result_rows)
        captured = {}

        def fake_solve_milp(c, integrality, lower_bounds, upper_bounds, constraints, constraint_lower, constraint_upper, options, log, problem_name):
            captured["objective"] = c.copy()
            captured["upper_bounds"] = upper_bounds.copy()
            return SimpleNamespace(success=True, x=lower_bounds.copy(), fun=0.0, message="ok")

        with patch.object(estimate, "solve_milp", side_effect=fake_solve_milp):
            estimate.solve_dispatch_model(model)

        variables = model["variables"]
        objective = captured["objective"]

        def objective_cost(key):
            return objective[variables[key]]

        diesel_device = model["device_rows"]["diesel_generators"][0]
        expected_diesel_cost = diesel_device["fuel_rate"] * model["diesel_objective_price"] / 1000
        self.assertEqual(objective_cost(("diesel_power", 0, 0)), expected_diesel_cost)
        self.assertEqual(objective_cost(("unmet_load", 0)), 0.0)
        self.assertEqual(captured["upper_bounds"][variables[("unmet_load", 0)]], 0.0)
        self.assertEqual(objective_cost(("diesel_on_count", 0, 0)), plan_optimizer.DIESEL_ON_COUNT_PENALTY)
        self.assertEqual(objective_cost(("electrolyzer_on_count", 0, 0)), plan_optimizer.ELECTROLYZER_ON_COUNT_PENALTY)
        for key in (
            ("storage_charge", 0),
            ("storage_discharge", 0),
            ("storage_charge_mode", 0),
            ("electrolyzer_power", 0, 0),
            ("fuel_cell_power", 0, 0),
            ("wind_curtailed", 0),
            ("pv_curtailed", 0),
            ("grid_storage_on_count", 0, 0),
        ):
            self.assertEqual(objective_cost(key), 0.0)
        for unit in range(2):
            self.assertNotIn(("diesel_on_unit", 0, unit), variables)
            self.assertNotIn(("electrolyzer_on_unit", 0, unit), variables)
            self.assertNotIn(("grid_storage_on_unit", 0, unit), variables)
        self.assertIn(("diesel_on_count", 0, 0), variables)
        self.assertIn(("electrolyzer_on_count", 0, 0), variables)
        self.assertIn(("grid_storage_on_count", 0, 0), variables)
        self.assertIn(("grid_storage_up_available_count", 0, 0), variables)
        self.assertIn(("grid_storage_down_available_count", 0, 0), variables)
        self.assertIn(("storage_charge_mode", 0), variables)
        self.assertNotIn(("storage_charge_on", 0), variables)
        self.assertNotIn(("storage_discharge_on", 0), variables)
        self.assertEqual(model["initial_storage_soc_ratio"], 0.2)
        self.assertEqual(model["initial_hydrogen_storage_ratio"], 0.8)
        self.assertEqual(model["storage_charge_efficiency"], 0.91)
        self.assertEqual(model["storage_discharge_efficiency"], 0.89)
        self.assertTrue(model["post_disturbance_power_balance_enabled"])
        self.assertEqual(model["storage_soc_upper_ratio"], 0.8)
        self.assertEqual(model["storage_soc_lower_ratio"], 0.2)

    def test_pv_generation_uses_capacity_times_irradiance_without_efficiency_parameter(self):
        self.assertAlmostEqual(estimate.pv_generation(500, 100, {"generation_efficiency": 0.1}), 50.0)
        self.assertAlmostEqual(estimate.pv_generation(1200, 100, {}), 100.0)

    def test_wind_generation_uses_configured_rated_wind_speed(self):
        params = {"cut_in_wind_speed": 3, "rated_wind_speed": 9, "cut_out_wind_speed": 25}

        self.assertAlmostEqual(estimate.wind_generation(9, 100, params), 100.0)
        self.assertAlmostEqual(estimate.wind_generation(6, 100, params), 12.5)
        self.assertAlmostEqual(estimate.wind_generation(25, 100, params), 0.0)

    def test_estimate_dispatch_outputs_requested_8760_curve_columns(self):
        payload = server.planning_store.default_payload("方案A")
        payload["planning_parameters"][0]["post_disturbance_power_balance_enabled"] = 1
        payload["planning_parameters"][0]["load_disturbance_enabled"] = 1
        payload["planning_parameters"][0]["renewable_disturbance_enabled"] = 1
        payload["planning_parameters"][0]["load_up_disturbance_factor"] = 0.1
        payload["planning_parameters"][0]["load_down_disturbance_factor"] = 0.2
        payload["planning_parameters"][0]["renewable_down_disturbance_factor"] = 0.3
        payload["storage_pcs"][0]["is_grid_forming"] = 1
        for row in payload["time_series"]:
            row["wind_speed"] = 12
            row["solar_irradiance"] = 900
            row["temperature"] = -5
            row["load"] = 60
        result_rows = [
            {"设备类型": "柴发", "设计台数": 1, "单台容量": 100, "总容量": 100, "单位": "kW"},
            {"设备类型": "风机", "设计台数": 1, "单台容量": 100, "总容量": 100, "单位": "kW"},
            {"设备类型": "光伏", "设计台数": 1, "单台容量": 100, "总容量": 100, "单位": "kW"},
            {"设备类型": "储能PCS", "设计台数": 1, "单台容量": 50, "总容量": 50, "单位": "kW"},
            {"设备类型": "储能电池组", "设计台数": 1, "单台容量": 100, "总容量": 100, "单位": "kWh"},
            {"设备类型": "电制氢", "设计台数": 1, "单台容量": 30, "总容量": 30, "单位": "kW"},
            {"设备类型": "储氢罐", "设计台数": 1, "单台容量": 100, "总容量": 100, "单位": "Nm3"},
            {"设备类型": "燃料电池", "设计台数": 1, "单台容量": 20, "总容量": 20, "单位": "kW"},
        ]

        result = estimate.run_estimation(payload, result_rows)
        row = result["dispatch_rows"][0]

        for field in (
            "wind_speed",
            "solar_irradiance",
            "temperature",
            "load",
            "diesel_power",
            "wind_available",
            "wind_power",
            "pv_available",
            "pv_power",
            "renewable_available",
            "storage_power",
            "storage_soc",
            "hydrogen_production_power",
            "hydrogen_storage",
            "fuel_cell_power",
            "wind_curtailed_power",
            "pv_curtailed_power",
            "curtailed_power",
            "unmet_load",
            "renewable_ratio",
            "renewable_curtailed_rate",
            "load_up_disturbance_power",
            "load_down_disturbance_power",
            "renewable_down_disturbance_power",
            "renewable_single_unit_power_max",
            "grid_up_regulation_capacity",
            "grid_down_regulation_capacity",
            "grid_up_regulation_requirement",
            "grid_down_regulation_requirement",
        ):
            self.assertIn(field, row)
            self.assertIsInstance(row[field], (int, float))
        self.assertAlmostEqual(row["curtailed_power"], row["wind_curtailed_power"] + row["pv_curtailed_power"], places=3)
        self.assertAlmostEqual(row["load_up_disturbance_power"], row["load"] * 0.1, places=3)
        self.assertAlmostEqual(row["load_down_disturbance_power"], -row["load"] * 0.2, places=3)
        self.assertAlmostEqual(
            row["renewable_down_disturbance_power"],
            (row["wind_power"] + row["pv_power"]) * 0.3,
            places=3,
        )
        self.assertNotIn("renewable_n1_power_gap", row)
        self.assertAlmostEqual(
            row["grid_up_regulation_requirement"],
            row["load_up_disturbance_power"] + row["renewable_down_disturbance_power"],
            places=3,
        )
        self.assertAlmostEqual(row["grid_down_regulation_requirement"], row["load_down_disturbance_power"], places=3)

        requested_energy_fields = (
            "load_energy",
            "diesel_energy",
            "wind_energy",
            "pv_energy",
            "storage_charge_energy",
            "storage_discharge_energy",
            "hydrogen_production_energy",
            "hydrogen_storage_increase",
            "hydrogen_storage_decrease",
            "fuel_cell_energy",
            "wind_available_energy",
            "pv_available_energy",
            "renewable_available_energy",
            "renewable_energy",
            "wind_curtailed_energy",
            "pv_curtailed_energy",
            "curtailed_energy",
            "unmet_load_energy",
            "renewable_ratio",
            "renewable_curtailed_rate",
        )
        daily = result["results"]["curves"]["green_daily"]
        monthly = result["results"]["curves"]["green_monthly"]
        self.assertEqual(len(daily), 365)
        self.assertEqual(len(monthly), 12)
        for field in requested_energy_fields:
            self.assertIn(field, daily[0])
            self.assertIsInstance(daily[0][field], (int, float))
            self.assertIn(field, monthly[0])
            self.assertIsInstance(monthly[0][field], (int, float))

        annual_metric_names = {
            row["指标"]
            for table in result["results"]["overview_tables"]
            if table["title"] == "规划年指标"
            for row in table["rows"]
        }
        green_metric_names = {row["指标"] for row in result["results"]["green_table"]}
        for name in (
            "负荷总电量",
            "柴发总发电量",
            "风机总发电量",
            "光伏总发电量",
            "电储能总储电量",
            "电储能总放电量",
            "电制氢总用电量",
            "氢储总增加量",
            "氢储总消耗量",
            "燃料电池总发电量",
            "风力最大可发电量",
            "光伏最大可发电量",
            "新能源最大可发电量",
            "新能源实发电量",
            "弃风总电量",
            "弃光总电量",
            "新能源总弃电量",
            "切负荷总电量",
            "新能源占比",
            "新能源弃电率",
        ):
            self.assertIn(name, annual_metric_names)
            self.assertIn(name, green_metric_names)

    def test_evaluation_api_uses_independent_runtime_and_estimate_script(self):
        planning_root = WEB_ROOT / "tests" / "tmp_evaluation_runtime"
        shutil.rmtree(planning_root, ignore_errors=True)
        planning_root.mkdir(parents=True)
        original_store = server.PLANNING_STORE
        original_runtime = server.EVALUATION_RUNTIME
        original_scheduler = server.TASK_SCHEDULER
        server.PLANNING_STORE = server.planning_store.PlanningStore(root=planning_root)
        server.EVALUATION_RUNTIME = server.EvaluationRuntimeManager()
        server.TASK_SCHEDULER = server.TaskScheduler()
        try:
            payload = server.planning_store.default_payload("方案A")
            for row in payload["time_series"]:
                row["wind_speed"] = 7
                row["solar_irradiance"] = 500
                row["load"] = 80
            server.PLANNING_STORE.write_scheme("方案A", payload)
            workbook = Workbook()
            sheet = workbook.active
            sheet.title = "规划结果"
            sheet.append(["设备类型", "设计台数", "单台容量", "总容量", "单位"])
            sheet.append(["柴发", 2, 100, 200, "kW"])
            sheet.append(["风机", 1, 100, 100, "kW"])
            sheet.append(["光伏", 1, 100, 100, "kW"])
            workbook.save(planning_root / "方案A" / "case_results.xlsx")

            status, headers, body = server.handle_api_path("/api/evaluation/status?scheme=方案A")
            initial = json.loads(body.decode("utf-8"))
            self.assertEqual(status, 200)
            self.assertEqual(initial["status"], "待启动")

            status, headers, body = server.handle_control_path(
                "/api/evaluation/control",
                json.dumps(
                    {"action": "start", "scheme": "方案A", "filename": "case_results.xlsx"},
                    ensure_ascii=False,
                ).encode("utf-8"),
            )
            started = json.loads(body.decode("utf-8"))
            self.assertEqual(status, 200)
            self.assertEqual(started["state"]["status"], "运行中")
            self.assertEqual(started["state"]["result_filename"], "case_results.xlsx")
            self.assertTrue(any("启动方案评估" in item["message"] for item in started["state"]["logs"]))

            state = started["state"]
            for _ in range(20):
                status, headers, body = server.handle_api_path("/api/evaluation/status?scheme=方案A")
                state = json.loads(body.decode("utf-8"))
                if any("读取方案参数" in item["message"] for item in state["logs"]):
                    break
                time.sleep(0.05)

            self.assertEqual(state["status"], "运行中")
            self.assertEqual(state["result_filename"], "case_results.xlsx")
            self.assertIsInstance(state["process_id"], int)
            self.assertNotEqual(state["process_id"], server.os.getpid())
            self.assertTrue(any("后台评估程序已启动" in item["message"] for item in state["logs"]))

            status, headers, body = server.handle_control_path(
                "/api/evaluation/control",
                json.dumps(
                    {"action": "stop", "scheme": "方案A", "filename": "case_results.xlsx"},
                    ensure_ascii=False,
                ).encode("utf-8"),
            )
            stopped = json.loads(body.decode("utf-8"))
            self.assertEqual(status, 200)
            self.assertEqual(stopped["state"]["status"], "计算中止")
        finally:
            server.PLANNING_STORE = original_store
            server.EVALUATION_RUNTIME = original_runtime
            server.TASK_SCHEDULER = original_scheduler
            shutil.rmtree(planning_root, ignore_errors=True)

    def test_evaluation_start_rejects_fast_infeasible_fixed_results_before_solving(self):
        planning_root = WEB_ROOT / "tests" / "tmp_evaluation_fast_infeasible"
        shutil.rmtree(planning_root, ignore_errors=True)
        planning_root.mkdir(parents=True)
        original_store = server.PLANNING_STORE
        original_runtime = server.EVALUATION_RUNTIME
        original_scheduler = server.TASK_SCHEDULER
        server.PLANNING_STORE = server.planning_store.PlanningStore(root=planning_root)
        server.EVALUATION_RUNTIME = server.EvaluationRuntimeManager()
        server.TASK_SCHEDULER = server.TaskScheduler()

        def write_result(filename: str, rows: list[list[object]]) -> None:
            workbook = Workbook()
            sheet = workbook.active
            sheet.title = "规划结果"
            sheet.append(["设备类型", "名称", "设计台数", "单台容量", "总容量", "单位"])
            for row in rows:
                sheet.append(row)
            workbook.save(planning_root / "方案A" / filename)
            workbook.close()

        def start_result(filename: str) -> tuple[int, dict]:
            status, headers, body = server.handle_control_path(
                "/api/evaluation/control",
                json.dumps(
                    {"action": "start", "scheme": "方案A", "filename": filename},
                    ensure_ascii=False,
                ).encode("utf-8"),
            )
            return status, json.loads(body.decode("utf-8"))

        def queue_result(filename: str) -> tuple[int, dict]:
            status, headers, body = server.handle_control_path(
                "/api/tasks/control",
                json.dumps(
                    {"action": "queue", "task_type": "evaluation", "scheme": "方案A", "result": filename},
                    ensure_ascii=False,
                ).encode("utf-8"),
            )
            return status, json.loads(body.decode("utf-8"))

        try:
            payload = server.planning_store.default_payload("方案A")
            for row in payload["time_series"]:
                row["wind_speed"] = 7
                row["solar_irradiance"] = 500
                row["load"] = 80
            payload["planning_parameters"][0]["green_power_ratio_lower"] = 0
            payload["planning_parameters"][0]["initial_storage_soc_ratio"] = 0.5
            payload["planning_parameters"][0]["initial_hydrogen_storage_ratio"] = 0.5
            payload["storage_battery_packs"][0]["battery_capacity"] = 200
            payload["storage_battery_packs"][0]["self_discharge_rate"] = 0.01
            payload["storage_pcs"][0]["power_capacity"] = 50
            payload["hydrogen_tanks"][0]["hydrogen_tank_capacity"] = 1000
            payload["hydrogen_tanks"][0]["self_discharge_rate"] = 0.001
            payload["hydrogen_electrolyzers"][0]["power_capacity"] = 80
            server.PLANNING_STORE.write_scheme("方案A", payload)

            write_result(
                "hydrogen_no_makeup_results.xlsx",
                [
                    ["柴发", "柴发1", 2, 100, 200, "kW"],
                    ["风机", "风机1", 1, 100, 100, "kW"],
                    ["光伏", "光伏1", 1, 100, 100, "kW"],
                    ["电制氢", "电制氢1", 0, 80, 0, "kW"],
                    ["储氢罐", "储氢罐1", 1, 1000, 1000, "Nm3"],
                    ["燃料电池", "燃料电池1", 0, 50, 0, "kW"],
                ],
            )
            status, data = start_result("hydrogen_no_makeup_results.xlsx")
            self.assertEqual(status, 400)
            self.assertEqual(data["error"], "bad_request")
            self.assertIn("储氢罐", data["message"])
            self.assertIn("电制氢", data["message"])
            self.assertIn("自损耗无法补偿", data["message"])

            write_result(
                "green_ratio_without_renewables_results.xlsx",
                [
                    ["柴发", "柴发1", 2, 100, 200, "kW"],
                    ["风机", "风机1", 0, 100, 0, "kW"],
                    ["光伏", "光伏1", 0, 100, 0, "kW"],
                ],
            )
            payload["planning_parameters"][0]["green_power_ratio_lower"] = 0.2
            server.PLANNING_STORE.write_scheme("方案A", payload)
            status, data = start_result("green_ratio_without_renewables_results.xlsx")
            self.assertEqual(status, 400)
            self.assertIn("风机和光伏设计台数均为0", data["message"])
            self.assertIn("绿色电量占比下限大于0", data["message"])

            for row in payload["time_series"]:
                row["wind_speed"] = 0
                row["solar_irradiance"] = 0
                row["load"] = 80
            server.PLANNING_STORE.write_scheme("方案A", payload)
            write_result(
                "green_ratio_insufficient_renewables_results.xlsx",
                [
                    ["柴发", "柴发1", 2, 100, 200, "kW"],
                    ["风机", "风机1", 1, 100, 100, "kW"],
                    ["光伏", "光伏1", 1, 100, 100, "kW"],
                ],
            )
            status, data = start_result("green_ratio_insufficient_renewables_results.xlsx")
            self.assertEqual(status, 400)
            self.assertIn("风机和光伏最大可发电量", data["message"])
            self.assertIn("低于绿色电量占比要求", data["message"])

            write_result(
                "battery_no_pcs_results.xlsx",
                [
                    ["柴发", "柴发1", 2, 100, 200, "kW"],
                    ["风机", "风机1", 1, 100, 100, "kW"],
                    ["光伏", "光伏1", 1, 100, 100, "kW"],
                    ["储能PCS", "储能PCS1", 0, 50, 0, "kW"],
                    ["储能电池组", "储能电池组1", 1, 200, 200, "kWh"],
                ],
            )
            payload["planning_parameters"][0]["green_power_ratio_lower"] = 0
            payload["diesel_generators"][0]["power_upper"] = 100
            server.PLANNING_STORE.write_scheme("方案A", payload)
            status, data = start_result("battery_no_pcs_results.xlsx")
            self.assertEqual(status, 400)
            self.assertIn("储能电池", data["message"])
            self.assertIn("储能PCS", data["message"])
            self.assertIn("自损耗无法补偿", data["message"])

            status, data = queue_result("battery_no_pcs_results.xlsx")
            self.assertEqual(status, 200)
            self.assertEqual(data["task"]["status"], "排队中")
            self.assertTrue(server.TASK_SCHEDULER.is_queued("evaluation", "方案A", "battery_no_pcs_results.xlsx"))
            scheduled_tasks = server.build_task_list()
            battery_task = next(
                item
                for item in scheduled_tasks
                if item["task_type_key"] == "evaluation" and item["result"] == "battery_no_pcs_results.xlsx"
            )
            self.assertEqual(battery_task["status"], "计算失败")
            self.assertFalse(battery_task["queued"])
            self.assertIn("自损耗无法补偿", battery_task["latest_log"])
        finally:
            for runtime in server.EVALUATION_RUNTIME.runtimes().values():
                if runtime.status == "运行中":
                    runtime.apply("stop", scheme=runtime.scheme, filename=runtime.result_filename)
            server.PLANNING_STORE = original_store
            server.EVALUATION_RUNTIME = original_runtime
            server.TASK_SCHEDULER = original_scheduler
            shutil.rmtree(planning_root, ignore_errors=True)

    def test_evaluation_status_is_scoped_by_result_file(self):
        planning_root = WEB_ROOT / "tests" / "tmp_evaluation_result_switch"
        shutil.rmtree(planning_root, ignore_errors=True)
        planning_root.mkdir(parents=True)
        original_store = server.PLANNING_STORE
        original_runtime = server.EVALUATION_RUNTIME
        server.PLANNING_STORE = server.planning_store.PlanningStore(root=planning_root)
        server.EVALUATION_RUNTIME = server.EvaluationRuntimeManager()
        try:
            payload = server.planning_store.default_payload("方案A")
            for row in payload["time_series"]:
                row["wind_speed"] = 7
                row["solar_irradiance"] = 500
                row["load"] = 80
            server.PLANNING_STORE.write_scheme("方案A", payload)

            for filename, diesel_count in (("case_a_results.xlsx", 2), ("case_b_results.xlsx", 3)):
                workbook = Workbook()
                sheet = workbook.active
                sheet.title = "规划结果"
                sheet.append(["设备类型", "设计台数", "单台容量", "总容量", "单位"])
                sheet.append(["柴发", diesel_count, 100, diesel_count * 100, "kW"])
                sheet.append(["风机", 1, 100, 100, "kW"])
                sheet.append(["光伏", 1, 100, 100, "kW"])
                workbook.save(planning_root / "方案A" / filename)

            status, headers, body = server.handle_control_path(
                "/api/evaluation/control",
                json.dumps(
                    {"action": "start", "scheme": "方案A", "filename": "case_a_results.xlsx"},
                    ensure_ascii=False,
                ).encode("utf-8"),
            )
            self.assertEqual(status, 200)

            case_a = {}
            for _ in range(200):
                status, headers, body = server.handle_api_path(
                    "/api/evaluation/status?scheme=方案A&filename=case_a_results.xlsx"
                )
                case_a = json.loads(body.decode("utf-8"))
                if case_a["status"] == "已完成":
                    break
                time.sleep(0.05)

            status, headers, body = server.handle_api_path(
                "/api/evaluation/status?scheme=方案A&filename=case_b_results.xlsx"
            )
            case_b = json.loads(body.decode("utf-8"))

            self.assertEqual(case_a["status"], "已完成")
            self.assertEqual(case_a["result_filename"], "case_a_results.xlsx")
            self.assertTrue(any("8760点优化调度完成" in item["message"] for item in case_a["logs"]))
            self.assertEqual(len(case_a["results"]["curves"]["green_hourly"]), 8760)
            self.assertEqual(case_b["status"], "待启动")
            self.assertEqual(case_b["result_filename"], "case_b_results.xlsx")
            self.assertFalse(any("8760点优化调度完成" in item["message"] for item in case_b["logs"]))
            self.assertEqual(case_b["results"]["curves"]["green_hourly"], [])
        finally:
            server.PLANNING_STORE = original_store
            server.EVALUATION_RUNTIME = original_runtime
            shutil.rmtree(planning_root, ignore_errors=True)

    def test_optimization_overview_results_are_two_tables_with_composition_bars(self):
        planning_root = WEB_ROOT / "tests" / "tmp_optimization_overview"
        shutil.rmtree(planning_root, ignore_errors=True)
        planning_root.mkdir(parents=True)
        original_store = server.PLANNING_STORE
        server.PLANNING_STORE = server.planning_store.PlanningStore(root=planning_root)
        try:
            server.PLANNING_STORE.create_scheme("caseA")
            runtime = server.OptimizationRuntime()
            runtime.apply("start", scheme="caseA")
            payload = self.wait_optimization_runtime(runtime)

            self.assertEqual(payload["status"], "已完成")
            tables = payload["results"]["overview_tables"]
            self.assertEqual([table["title"] for table in tables], ["规划结果", "规划年指标"])
            self.assertEqual(len(tables), 2)
            self.assertTrue(any(row["设备类型"] == "柴发" and "设计台数" in row for row in tables[0]["rows"]))
            self.assertTrue(any(row["设备类型"] == "储能电池组" and "设计台数" in row for row in tables[0]["rows"]))
            annual_metric_names = {row["指标"] for row in tables[1]["rows"]}
            for name in (
                "柴发总容量",
                "风电总容量",
                "光伏总容量",
                "氢能总容量",
                "储能总容量",
                "负荷总电量",
                "柴发总发电量",
                "风机总发电量",
                "光伏总发电量",
                "风力最大可发电量",
                "光伏最大可发电量",
                "新能源最大可发电量",
                "新能源实发电量",
                "弃风总电量",
                "弃光总电量",
                "新能源总弃电量",
                "切负荷总电量",
                "新能源占比",
                "新能源弃电率",
                "年均建设成本",
                "年柴油成本",
                "年总成本",
                "总成本",
                "绿电占比",
                "频率风险点",
            ):
                self.assertIn(name, annual_metric_names)
            self.assertNotIn("规划年效益", [table["title"] for table in tables])

            disks = payload["results"]["overview_disks"]
            self.assertEqual([disk["title"] for disk in disks], ["成本构成", "容量构成", "电量构成"])
            self.assertEqual(disks[0]["left_label"], "年柴油成本")
            self.assertEqual(disks[0]["right_label"], "年均建设成本")
            capacity_segments = disks[1]["segments"]
            self.assertEqual(
                [segment["label"] for segment in capacity_segments],
                ["柴发容量", "风电容量", "光伏容量", "电储能容量", "燃料电池容量"],
            )
            self.assertEqual(capacity_segments[0]["unit"], "kW")
            self.assertEqual(capacity_segments[3]["unit"], "kWh")
            self.assertEqual(disks[2]["left_label"], "柴发电量")
            self.assertEqual(disks[2]["right_label"], "绿电电量")
        finally:
            server.PLANNING_STORE = original_store
            shutil.rmtree(planning_root, ignore_errors=True)

    def test_optimization_green_result_has_summary_table_and_daily_curve(self):
        planning_root = WEB_ROOT / "tests" / "tmp_optimization_green"
        shutil.rmtree(planning_root, ignore_errors=True)
        planning_root.mkdir(parents=True)
        original_store = server.PLANNING_STORE
        server.PLANNING_STORE = server.planning_store.PlanningStore(root=planning_root)
        try:
            server.PLANNING_STORE.create_scheme("caseA")
            runtime = server.OptimizationRuntime()
            runtime.apply("start", scheme="caseA")
            payload = self.wait_optimization_runtime(runtime)

            self.assertEqual(payload["status"], "已完成")
            green_rows = payload["results"]["green_table"]
            metric_names = {row["指标"] for row in green_rows}
            for name in (
                "负荷总电量",
                "柴发总电量",
                "风机总发电量",
                "光伏总发电量",
                "电储总发电量",
                "氢储总发电量",
                "新能源总弃电量",
                "柴油消耗",
                "制氢总量",
                "电储能总储电量",
                "电储能总放电量",
                "电制氢总用电量",
                "氢储总增加量",
                "氢储总消耗量",
                "燃料电池总发电量",
                "风力最大可发电量",
                "光伏最大可发电量",
                "新能源最大可发电量",
                "新能源实发电量",
                "弃风总电量",
                "弃光总电量",
                "新能源总弃电量",
                "切负荷总电量",
                "新能源占比",
                "新能源弃电率",
            ):
                self.assertIn(name, metric_names)
            self.assertTrue(all(set(row) == {"指标", "数值", "单位"} for row in green_rows))
            units = {row["指标"]: row["单位"] for row in green_rows}
            self.assertEqual(units["负荷总电量"], "kWh")
            self.assertEqual(units["柴发总电量"], "kWh")
            self.assertEqual(units["风机总发电量"], "kWh")
            self.assertEqual(units["光伏总发电量"], "kWh")
            self.assertEqual(units["电储总发电量"], "kWh")
            self.assertEqual(units["氢储总发电量"], "kWh")
            self.assertEqual(units["新能源总弃电量"], "kWh")
            self.assertEqual(units["新能源占比"], "%")
            self.assertEqual(units["新能源弃电率"], "%")
            self.assertEqual(units["风力最大可发电量"], "kWh")
            self.assertEqual(units["光伏最大可发电量"], "kWh")
            self.assertEqual(units["新能源最大可发电量"], "kWh")
            self.assertEqual(units["新能源实发电量"], "kWh")
            self.assertEqual(units["弃风总电量"], "kWh")
            self.assertEqual(units["弃光总电量"], "kWh")
            self.assertEqual(units["切负荷总电量"], "kWh")
            self.assertEqual(units["柴油消耗"], "吨")
            self.assertEqual(units["制氢总量"], "Nm3")

            daily = payload["results"]["curves"]["green_daily"]
            self.assertEqual(len(daily), 365)
            self.assertEqual(daily[0]["day"], 1)
            self.assertEqual(daily[-1]["day"], 365)
            for field in (
                "load_energy",
                "diesel_energy",
                "wind_energy",
                "pv_energy",
                "storage_charge_energy",
                "storage_discharge_energy",
                "hydrogen_production_energy",
                "hydrogen_storage_increase",
                "hydrogen_storage_decrease",
                "fuel_cell_energy",
                "wind_available_energy",
                "pv_available_energy",
                "renewable_available_energy",
                "renewable_energy",
                "wind_curtailed_energy",
                "pv_curtailed_energy",
                "curtailed_energy",
                "unmet_load_energy",
                "renewable_ratio",
                "renewable_curtailed_rate",
            ):
                self.assertIn(field, daily[0])
                self.assertIsInstance(daily[0][field], (int, float))

            monthly = payload["results"]["curves"]["green_monthly"]
            self.assertEqual(len(monthly), 12)
            self.assertEqual(monthly[0]["month"], 1)
            self.assertEqual(monthly[-1]["month"], 12)
            for field in (
                "load_energy",
                "diesel_energy",
                "wind_energy",
                "pv_energy",
                "storage_charge_energy",
                "storage_discharge_energy",
                "hydrogen_production_energy",
                "hydrogen_storage_increase",
                "hydrogen_storage_decrease",
                "fuel_cell_energy",
                "wind_available_energy",
                "pv_available_energy",
                "renewable_available_energy",
                "renewable_energy",
                "wind_curtailed_energy",
                "pv_curtailed_energy",
                "curtailed_energy",
                "unmet_load_energy",
                "renewable_ratio",
                "renewable_curtailed_rate",
            ):
                self.assertIn(field, monthly[0])
                self.assertIsInstance(monthly[0][field], (int, float))

            hourly = payload["results"]["curves"]["green_hourly"]
            self.assertEqual(len(hourly), 8760)
            self.assertEqual(hourly[0]["hour_index"], 1)
            self.assertEqual(hourly[-1]["hour_index"], 8760)
            for field in (
                "load",
                "wind_power",
                "pv_power",
                "storage_charge",
                "storage_discharge",
                "diesel_power",
                "curtailed_power",
                "unmet_load",
                "storage_soc",
            ):
                self.assertIn(field, hourly[0])
                self.assertIsInstance(hourly[0][field], (int, float))
        finally:
            server.PLANNING_STORE = original_store
            shutil.rmtree(planning_root, ignore_errors=True)

    def test_optimization_safety_result_has_summary_table_and_daily_frequency_curve(self):
        planning_root = WEB_ROOT / "tests" / "tmp_optimization_safety"
        shutil.rmtree(planning_root, ignore_errors=True)
        planning_root.mkdir(parents=True)
        original_store = server.PLANNING_STORE
        server.PLANNING_STORE = server.planning_store.PlanningStore(root=planning_root)
        try:
            server.PLANNING_STORE.create_scheme("caseA")
            runtime = server.OptimizationRuntime()
            runtime.apply("start", scheme="caseA")
            payload = self.wait_optimization_runtime(runtime)

            self.assertEqual(payload["status"], "已完成")
            safety_rows = payload["results"]["safety_table"]
            metric_names = {row["指标"] for row in safety_rows}
            for name in (
                "向上扰动最大量",
                "向下扰动最大量",
                "最高频率",
                "最低频率",
                "频率安全风险小时数",
            ):
                self.assertIn(name, metric_names)
            units = {row["指标"]: row["单位"] for row in safety_rows}
            self.assertEqual(units["向上扰动最大量"], "kW")
            self.assertEqual(units["向下扰动最大量"], "kW")
            self.assertEqual(units["最高频率"], "Hz")
            self.assertEqual(units["最低频率"], "Hz")
            self.assertEqual(units["频率安全风险小时数"], "h")

            daily = payload["results"]["curves"]["safety_daily"]
            self.assertEqual(len(daily), 365)
            self.assertEqual(daily[0]["day"], 1)
            self.assertEqual(daily[-1]["day"], 365)
            self.assertGreaterEqual(daily[0]["frequency_max"], 50)
            self.assertLessEqual(daily[0]["frequency_min"], 50)
            for field in ("frequency_max", "frequency_min"):
                self.assertIn(field, daily[0])
                self.assertIsInstance(daily[0][field], (int, float))
        finally:
            server.PLANNING_STORE = original_store
            shutil.rmtree(planning_root, ignore_errors=True)

    def test_completed_optimization_writes_result_workbook_and_overwrites_existing_file(self):
        planning_root = WEB_ROOT / "tests" / "tmp_optimization_results"
        shutil.rmtree(planning_root, ignore_errors=True)
        planning_root.mkdir(parents=True)
        original_store = server.PLANNING_STORE
        server.PLANNING_STORE = server.planning_store.PlanningStore(root=planning_root)
        try:
            server.PLANNING_STORE.create_scheme("方案A")
            result_path = planning_root / "方案A" / "opt_results.xlsx"
            result_path.write_text("old result", encoding="utf-8")

            runtime = server.OptimizationRuntime()
            runtime.apply("start", scheme="方案A")
            payload = self.wait_optimization_runtime(runtime)

            self.assertEqual(payload["status"], "已完成")
            self.assertEqual(payload["result_file"], str(result_path))
            self.assertTrue(result_path.exists())
            curve_path = server.result_curves_workbook_path(result_path)
            self.assertTrue(curve_path.exists())
            self.assertTrue(server.result_curves_sqlite_path(result_path).exists())
            workbook = load_workbook(result_path, data_only=True, read_only=True)
            try:
                self.assertEqual(
                    workbook.sheetnames,
                    ["总体指标", "规划结果", "规划年指标", "供能分析", "安全评估", "运行日志"],
                )
                self.assertEqual(workbook["总体指标"]["A1"].value, "指标")
                self.assertEqual(workbook["总体指标"]["B1"].value, "数值")
                self.assertEqual(workbook["规划结果"]["A1"].value, "设备类型")
                self.assertEqual(workbook["规划结果"]["A2"].value, "柴发")
                self.assertEqual(workbook["运行日志"]["A1"].value, "时间")
                log_messages = [row[2] for row in workbook["运行日志"].iter_rows(min_row=2, values_only=True)]
                self.assertIn("规划求解完成", log_messages)
            finally:
                workbook.close()
            curve_workbook = load_workbook(curve_path, data_only=True, read_only=True)
            try:
                self.assertEqual(curve_workbook.sheetnames, ["供能日曲线", "供能月曲线", "安全日曲线", "调度结果"])
                self.assertEqual(curve_workbook["供能日曲线"].max_row, 366)
                self.assertIn("供能月曲线", curve_workbook.sheetnames)
                self.assertEqual(curve_workbook["供能月曲线"].max_row, 13)
                self.assertEqual(curve_workbook["安全日曲线"].max_row, 366)
                self.assertEqual(curve_workbook["调度结果"].max_row, 8761)
                self.assertEqual(curve_workbook["调度结果"]["A1"].value, "小时")
                self.assertEqual(curve_workbook["调度结果"]["C1"].value, "风速")
                self.assertEqual(curve_workbook["调度结果"]["F1"].value, "负荷总功率")
                hourly_headers = [cell.value for cell in curve_workbook["调度结果"][1]]
                self.assertIn("新能源最大可发", hourly_headers)
                self.assertIn("新能源总出力", hourly_headers)
                self.assertIn("新能源占比", hourly_headers)
                self.assertIn("新能源弃电率", hourly_headers)
                self.assertIn("负荷上扰动功率", hourly_headers)
                self.assertIn("负荷下扰动功率", hourly_headers)
                self.assertIn("新能源下扰动功率", hourly_headers)
                self.assertIn("风光单机功率最大值", hourly_headers)
                self.assertNotIn("新能源N-1功率缺口", hourly_headers)
                self.assertIn("电网向上调节能力", hourly_headers)
                self.assertIn("电网向下调节能力", hourly_headers)
                self.assertIn("电网向上调节需求", hourly_headers)
                self.assertIn("电网向下调节需求", hourly_headers)
            finally:
                curve_workbook.close()
        finally:
            server.PLANNING_STORE = original_store
            shutil.rmtree(planning_root, ignore_errors=True)

    def test_optimization_status_reads_display_results_from_result_workbook(self):
        planning_root = WEB_ROOT / "tests" / "tmp_optimization_status_workbook"
        shutil.rmtree(planning_root, ignore_errors=True)
        planning_root.mkdir(parents=True)
        original_store = server.PLANNING_STORE
        original_runtime = server.OPTIMIZATION_RUNTIME
        server.PLANNING_STORE = server.planning_store.PlanningStore(root=planning_root)
        server.OPTIMIZATION_RUNTIME = server.OptimizationRuntimeManager()
        try:
            server.PLANNING_STORE.create_scheme("方案A")
            result_path = planning_root / "方案A" / "opt_results.xlsx"
            workbook = Workbook()
            workbook.active.title = "总体指标"
            workbook.active.append(["指标", "数值", "单位"])
            workbook.active.append(["度电成本", 9.99, "元/kWh"])
            planning_sheet = workbook.create_sheet("规划结果")
            planning_sheet.append(["设备类型", "设计台数", "单台容量", "总容量", "单位"])
            planning_sheet.append(["工作簿柴发", 3, 111, 333, "kW"])
            annual_sheet = workbook.create_sheet("规划年指标")
            annual_sheet.append(["指标", "数值", "单位"])
            annual_sheet.append(["工作簿年指标", 1234, "kWh"])
            green_sheet = workbook.create_sheet("供能分析")
            green_sheet.append(["指标", "数值", "单位"])
            green_sheet.append(["工作簿供能指标", 5678, "kWh"])
            daily_sheet = workbook.create_sheet("供能日曲线")
            daily_sheet.append(["day", "load_energy", "diesel_energy"])
            daily_sheet.append([1, 10, 2])
            monthly_sheet = workbook.create_sheet("供能月曲线")
            monthly_sheet.append(["month", "load_energy", "diesel_energy"])
            monthly_sheet.append([1, 310, 62])
            safety_sheet = workbook.create_sheet("安全评估")
            safety_sheet.append(["指标", "数值", "单位"])
            safety_sheet.append(["工作簿安全指标", 50.1, "Hz"])
            safety_daily_sheet = workbook.create_sheet("安全日曲线")
            safety_daily_sheet.append(["day", "frequency_max", "frequency_min"])
            safety_daily_sheet.append([1, 50.2, 49.8])
            dispatch_sheet = workbook.create_sheet("调度结果")
            dispatch_sheet.append(["小时", "负荷总功率", "柴发总功率", "新能源N-1功率缺口"])
            dispatch_sheet.append([1, 100, 30, 88])
            workbook.save(result_path)
            workbook.close()

            status, headers, body = server.handle_api_path("/api/optimization/status?scheme=方案A")
            payload = json.loads(body.decode("utf-8"))

            self.assertEqual(status, 200)
            self.assertEqual(payload["result_file"], str(result_path))
            self.assertEqual(payload["metrics"][3], {"label": "度电成本", "value": 9.99, "unit": "元"})
            self.assertEqual(payload["results"]["overview_tables"][0]["rows"][0]["设备类型"], "工作簿柴发")
            self.assertEqual(payload["results"]["overview_tables"][1]["rows"][0]["指标"], "工作簿年指标")
            self.assertEqual(payload["results"]["green_table"][0]["指标"], "工作簿供能指标")
            self.assertEqual(payload["results"]["safety_table"][0]["指标"], "工作簿安全指标")
            self.assertEqual(payload["results"]["curves"]["green_daily"][0]["load_energy"], 10)
            self.assertEqual(payload["results"]["curves"]["green_monthly"][0]["load_energy"], 310)
            self.assertEqual(payload["results"]["curves"]["safety_daily"][0]["frequency_max"], 50.2)
            self.assertEqual(payload["results"]["curves"]["green_hourly"][0]["load"], 100)
            self.assertNotIn("renewable_n1_power_gap", payload["results"]["curves"]["green_hourly"][0])
        finally:
            server.PLANNING_STORE = original_store
            server.OPTIMIZATION_RUNTIME = original_runtime
            shutil.rmtree(planning_root, ignore_errors=True)

    def test_evaluation_report_api_exports_word_document(self):
        from docx import Document

        planning_root = WEB_ROOT / "tests" / "tmp_evaluation_report_export"
        shutil.rmtree(planning_root, ignore_errors=True)
        planning_root.mkdir(parents=True)
        original_store = server.PLANNING_STORE
        server.PLANNING_STORE = server.planning_store.PlanningStore(root=planning_root)
        try:
            server.PLANNING_STORE.create_scheme("方案A", owner_username="alice")
            scheme_payload = server.planning_store.default_payload("方案A")
            scheme_payload["planning_parameters"][0]["diesel_price"] = 7.5
            scheme_payload["planning_parameters"][0]["storage_balance_mode"] = "weekly"
            scheme_payload["diesel_generators"][0]["name"] = "输入柴发"
            for index, row in enumerate(scheme_payload["time_series"]):
                hour = index % 24
                row["load"] = 100 + hour
                row["wind_speed"] = 5 + (hour / 24)
                row["solar_irradiance"] = max(0, 700 - abs(12 - hour) * 55)
                row["temperature"] = 18 + (hour / 12)
            server.PLANNING_STORE.write_scheme("方案A", scheme_payload)
            result_path = planning_root / "方案A" / "opt_results.xlsx"
            workbook = Workbook()
            workbook.active.title = "总体指标"
            workbook.active.append(["指标", "数值", "单位"])
            workbook.active.append(["度电成本", 9.99, "元/kWh"])
            workbook.active.append(["绿电占比", 0.62, "%"])
            planning_sheet = workbook.create_sheet("规划结果")
            planning_sheet.append(["设备类型", "设计台数", "单台容量", "总容量", "单位"])
            planning_sheet.append(["工作簿柴发", 3, 111, 333, "kW"])
            annual_sheet = workbook.create_sheet("规划年指标")
            annual_sheet.append(["指标", "数值", "单位"])
            annual_sheet.append(["年柴油成本", 12.5, "万元"])
            annual_sheet.append(["年均建设成本", 8.2, "万元"])
            green_sheet = workbook.create_sheet("供能分析")
            green_sheet.append(["指标", "数值", "单位"])
            green_sheet.append(["新能源实发电量", 5678, "kWh"])
            safety_sheet = workbook.create_sheet("安全评估")
            safety_sheet.append(["指标", "数值", "单位"])
            safety_sheet.append(["最低频率", 49.8, "Hz"])
            daily_sheet = workbook.create_sheet("供能日曲线")
            daily_sheet.append(["day", "load_energy", "diesel_energy", "renewable_energy"])
            daily_sheet.append([1, 100, 30, 70])
            daily_sheet.append([2, 120, 36, 84])
            safety_daily_sheet = workbook.create_sheet("安全日曲线")
            safety_daily_sheet.append(["day", "frequency_min", "frequency_max"])
            safety_daily_sheet.append([1, 49.8, 50.15])
            safety_daily_sheet.append([2, 49.82, 50.12])
            workbook.save(result_path)
            workbook.close()

            status, headers, body = server.handle_evaluation_results_api_path(
                "/api/evaluation/report",
                "GET",
                b"",
                "scheme=方案A&filename=opt_results.xlsx",
                current_user={"id": 1, "username": "alice", "role": "user"},
            )

            self.assertEqual(status, 200)
            self.assertEqual(headers["Content-Type"], server.EVALUATION_REPORT_CONTENT_TYPE)
            self.assertIn(".docx", headers["Content-Disposition"])
            self.assertTrue(body.startswith(b"PK"))
            document = Document(BytesIO(body))
            paragraph_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            table_text = "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
            self.assertIn("方案结果报告", paragraph_text)
            self.assertIn("方案A / opt", paragraph_text)
            self.assertIn("输入数据", paragraph_text)
            self.assertIn("规划参数", paragraph_text)
            self.assertIn("8760时序数据样表", paragraph_text)
            self.assertIn("输入负荷日曲线", paragraph_text)
            self.assertIn("方案名称", table_text)
            self.assertIn("柴油价格(万元/吨)", table_text)
            self.assertIn("7.5", table_text)
            self.assertIn("输入柴发", table_text)
            self.assertIn("负荷(kW)", table_text)
            self.assertIn("度电成本", table_text)
            self.assertIn("工作簿柴发", table_text)
            self.assertGreaterEqual(len(document.inline_shapes), 4)

            status, headers, body = server.handle_evaluation_results_api_path(
                "/api/evaluation/report",
                "GET",
                b"",
                "scheme=方案A&filename=missing_results.xlsx",
                current_user={"id": 1, "username": "alice", "role": "user"},
            )
            error_payload = json.loads(body.decode("utf-8"))
            self.assertEqual(status, 404)
            self.assertEqual(error_payload["error"], "not_found")
            self.assertIn("missing_results.xlsx", error_payload["message"])
        finally:
            server.PLANNING_STORE = original_store
            shutil.rmtree(planning_root, ignore_errors=True)

    def test_result_workbook_display_payload_uses_file_cache_until_file_changes(self):
        planning_root = WEB_ROOT / "tests" / "tmp_result_workbook_cache"
        shutil.rmtree(planning_root, ignore_errors=True)
        planning_root.mkdir(parents=True)
        result_path = planning_root / "case_results.xlsx"
        workbook = Workbook()
        workbook.active.title = "总体指标"
        workbook.active.append(["指标", "数值", "单位"])
        workbook.active.append(["度电成本", 1.23, "元/kWh"])
        planning_sheet = workbook.create_sheet("规划结果")
        planning_sheet.append(["设备类型", "设计台数", "单台容量", "总容量", "单位"])
        planning_sheet.append(["柴发", 1, 100, 100, "kW"])
        workbook.save(result_path)
        workbook.close()
        server.RESULT_DISPLAY_PAYLOAD_CACHE.clear()
        original_load_workbook = server.load_workbook
        load_count = 0

        def counting_load_workbook(*args, **kwargs):
            nonlocal load_count
            load_count += 1
            return original_load_workbook(*args, **kwargs)

        try:
            with patch.object(server, "load_workbook", side_effect=counting_load_workbook):
                first = server.read_result_workbook_display_payload(result_path, include_hourly_curves=False)
                first["results"]["overview_tables"][0]["rows"][0]["设备类型"] = "外部修改"
                second = server.read_result_workbook_display_payload(result_path, include_hourly_curves=False)
                self.assertEqual(load_count, 1)
                self.assertEqual(second["results"]["overview_tables"][0]["rows"][0]["设备类型"], "柴发")

                time.sleep(0.02)
                changed_workbook = load_workbook(result_path)
                changed_workbook["规划结果"]["A2"] = "风机"
                changed_workbook.save(result_path)
                changed_workbook.close()
                third = server.read_result_workbook_display_payload(result_path, include_hourly_curves=False)

            self.assertEqual(load_count, 2)
            self.assertEqual(third["results"]["overview_tables"][0]["rows"][0]["设备类型"], "风机")
        finally:
            shutil.rmtree(planning_root, ignore_errors=True)

    def test_result_workbook_display_payload_refreshes_when_split_curve_file_changes(self):
        planning_root = WEB_ROOT / "tests" / "tmp_result_workbook_curve_cache"
        shutil.rmtree(planning_root, ignore_errors=True)
        planning_root.mkdir(parents=True)
        result_path = planning_root / "case_results.xlsx"
        curve_path = server.result_curves_workbook_path(result_path)

        workbook = Workbook()
        workbook.active.title = "总体指标"
        workbook.active.append(["指标", "数值", "单位"])
        workbook.active.append(["度电成本", 1.23, "元/kWh"])
        planning_sheet = workbook.create_sheet("规划结果")
        planning_sheet.append(["设备类型", "设计台数", "单台容量", "总容量", "单位"])
        planning_sheet.append(["柴发", 1, 100, 100, "kW"])
        annual_sheet = workbook.create_sheet("规划年指标")
        annual_sheet.append(["指标", "数值", "单位"])
        green_sheet = workbook.create_sheet("供能分析")
        green_sheet.append(["指标", "数值", "单位"])
        safety_sheet = workbook.create_sheet("安全评估")
        safety_sheet.append(["指标", "数值", "单位"])
        workbook.save(result_path)
        workbook.close()

        curve_workbook = Workbook()
        curve_workbook.active.title = "供能日曲线"
        curve_workbook.active.append(["day", "load_energy"])
        curve_workbook.active.append([1, 56])
        monthly_sheet = curve_workbook.create_sheet("供能月曲线")
        monthly_sheet.append(["month", "load_energy"])
        monthly_sheet.append([1, 310])
        safety_daily_sheet = curve_workbook.create_sheet("安全日曲线")
        safety_daily_sheet.append(["day", "frequency_max", "frequency_min"])
        safety_daily_sheet.append([1, 50.1, 49.9])
        dispatch_sheet = curve_workbook.create_sheet("调度结果")
        dispatch_sheet.append(["小时", "时间", "负荷总功率"])
        dispatch_sheet.append([1, "2026-01-01 00:00", 100])
        curve_workbook.save(curve_path)
        curve_workbook.close()

        server.RESULT_DISPLAY_PAYLOAD_CACHE.clear()
        original_load_workbook = server.load_workbook
        curve_load_count = 0

        def counting_load_workbook(*args, **kwargs):
            nonlocal curve_load_count
            if Path(args[0]) == curve_path:
                curve_load_count += 1
            return original_load_workbook(*args, **kwargs)

        try:
            with patch.object(server, "load_workbook", side_effect=counting_load_workbook):
                first = server.read_result_workbook_display_payload(result_path)
                second = server.read_result_workbook_display_payload(result_path)
                self.assertEqual(curve_load_count, 1)
                self.assertEqual(first["results"]["curves"]["green_daily"][0]["load_energy"], 56)
                self.assertEqual(second["results"]["curves"]["green_daily"][0]["load_energy"], 56)

                time.sleep(0.02)
                changed_curve_workbook = load_workbook(curve_path)
                changed_curve_workbook["供能日曲线"]["B2"] = 78
                changed_curve_workbook.save(curve_path)
                changed_curve_workbook.close()
                third = server.read_result_workbook_display_payload(result_path)

            self.assertEqual(curve_load_count, 2)
            self.assertEqual(third["metrics"][0]["label"], "度电成本")
            self.assertEqual(third["results"]["overview_tables"][0]["rows"][0]["设备类型"], "柴发")
            self.assertEqual(third["results"]["curves"]["green_daily"][0]["load_energy"], 78)
            self.assertEqual(third["results"]["curves"]["green_monthly"][0]["load_energy"], 310)
            self.assertEqual(third["results"]["curves"]["safety_daily"][0]["frequency_min"], 49.9)
            self.assertEqual(third["results"]["curves"]["green_hourly"][0]["load"], 100)

            self.assertTrue(server.result_curves_sqlite_path(result_path).exists())
            server.RESULT_DISPLAY_PAYLOAD_CACHE.clear()
            server.COMPARISON_CURVE_SLICE_CACHE.clear()

            def reject_curve_workbook_load(path, *args, **kwargs):
                if Path(path) == curve_path:
                    raise AssertionError("curve workbook should not be loaded when sqlite mirror is fresh")
                return original_load_workbook(path, *args, **kwargs)

            with patch.object(server, "load_workbook", side_effect=reject_curve_workbook_load):
                sqlite_payload = server.read_result_workbook_display_payload(result_path)
                sqlite_group = server.read_comparison_curve_group(result_path, "hourly", ["负荷总功率"])
            self.assertEqual(sqlite_payload["results"]["curves"]["green_daily"][0]["load_energy"], 78)
            self.assertEqual(sqlite_payload["results"]["curves"]["green_hourly"][0]["load"], 100)
            self.assertEqual(sqlite_group["负荷总功率"][0]["y"], 100)
        finally:
            shutil.rmtree(planning_root, ignore_errors=True)

    def test_curve_sqlite_writer_closes_connection_before_replacing_unique_temp_file(self):
        planning_root = WEB_ROOT / "tests" / "tmp_curve_sqlite_writer"
        shutil.rmtree(planning_root, ignore_errors=True)
        planning_root.mkdir(parents=True)
        result_path = planning_root / "opt_results.xlsx"
        curve_path = server.result_curves_workbook_path(result_path)
        curve_path.write_bytes(b"curve source")
        db_path = server.result_curves_sqlite_path(result_path)
        real_connect = server.sqlite3.connect
        real_replace = server.file_ops.replace_file_with_retry
        connections = []
        temp_paths = []

        def tracking_connect(*args, **kwargs):
            connection = real_connect(*args, **kwargs)
            connections.append(connection)
            return connection

        def checked_replace(source, target, label):
            with self.assertRaises(server.sqlite3.ProgrammingError):
                connections[-1].execute("SELECT 1")
            temp_paths.append(Path(source))
            return real_replace(source, target, label)

        curves = {
            "green_daily": [{"day": 1, "load_energy": 10}],
            "green_monthly": [],
            "green_hourly": [],
            "safety_daily": [],
        }
        try:
            with patch.object(server.sqlite3, "connect", side_effect=tracking_connect), patch.object(
                server.file_ops,
                "replace_file_with_retry",
                side_effect=checked_replace,
            ):
                server.write_result_curve_sqlite(result_path, curves)
                server.write_result_curve_sqlite(result_path, curves)

            self.assertTrue(db_path.exists())
            self.assertEqual(len(temp_paths), 2)
            self.assertNotEqual(temp_paths[0], temp_paths[1])
            self.assertFalse(list(planning_root.glob(f"{db_path.name}.*.tmp")))
        finally:
            shutil.rmtree(planning_root, ignore_errors=True)

    def test_comparison_workbook_reads_curves_from_split_curve_workbook(self):
        result_path = WEB_ROOT / "tests" / "tmp_comparison_split_results.xlsx"
        curve_path = server.result_curves_workbook_path(result_path)
        for path in (result_path, curve_path):
            path.unlink(missing_ok=True)

        workbook = Workbook()
        try:
            workbook.active.title = "总体指标"
            planning_sheet = workbook.create_sheet("规划结果")
            planning_sheet.append(["设备类型", "设计台数", "单台容量", "总容量", "单位"])
            planning_sheet.append(["柴发", 1, 100, 100, "kW"])
            energy_sheet = workbook.create_sheet("供能分析")
            energy_sheet.append(["指标", "数值", "单位"])
            energy_sheet.append(["柴油消耗", 1, "吨"])
            safety_sheet = workbook.create_sheet("安全评估")
            safety_sheet.append(["指标", "数值", "单位"])
            safety_sheet.append(["最高频率", 50.1, "Hz"])
            annual_sheet = workbook.create_sheet("规划年指标")
            annual_sheet.append(["指标", "数值", "单位"])
            annual_sheet.append(["年供电量", 1000, "kWh"])
            workbook.save(result_path)
        finally:
            workbook.close()

        curve_workbook = Workbook()
        try:
            daily_sheet = curve_workbook.active
            daily_sheet.title = "供能日曲线"
            daily_sheet.append(["day", "load_energy"])
            daily_sheet.append([1, 10])
            monthly_sheet = curve_workbook.create_sheet("供能月曲线")
            monthly_sheet.append(["month", "load_energy"])
            monthly_sheet.append([1, 310])
            safety_daily_sheet = curve_workbook.create_sheet("安全日曲线")
            safety_daily_sheet.append(["day", "frequency_max", "frequency_min"])
            safety_daily_sheet.append([1, 50.2, 49.8])
            dispatch_sheet = curve_workbook.create_sheet("调度结果")
            dispatch_sheet.append(["小时", "负荷总功率"])
            dispatch_sheet.append([1, 123])
            curve_workbook.save(curve_path)

            payload = server.read_comparison_workbook(result_path)
        finally:
            curve_workbook.close()
            for path in (result_path, curve_path):
                path.unlink(missing_ok=True)

        self.assertEqual(payload["capacity"][0]["设备类型"], "柴发")
        self.assertEqual(payload["energy"][0]["指标"], "柴油消耗")
        self.assertIn("负荷总功率", payload["curve_groups"]["hourly"])
        self.assertEqual(payload["curve_groups"]["hourly"]["负荷总功率"][0]["y"], 123)
        self.assertEqual(payload["curve_groups"]["daily"]["负荷总电量"][0]["y"], 10)
        self.assertEqual(payload["curve_groups"]["monthly"]["负荷总电量"][0]["y"], 310)

    def test_light_optimization_status_skips_hourly_dispatch_sheet(self):
        planning_root = WEB_ROOT / "tests" / "tmp_optimization_status_light"
        shutil.rmtree(planning_root, ignore_errors=True)
        planning_root.mkdir(parents=True)
        original_store = server.PLANNING_STORE
        original_runtime = server.OPTIMIZATION_RUNTIME
        server.PLANNING_STORE = server.planning_store.PlanningStore(root=planning_root)
        server.OPTIMIZATION_RUNTIME = server.OptimizationRuntimeManager()
        try:
            server.PLANNING_STORE.create_scheme("方案A")
            result_path = planning_root / "方案A" / "opt_results.xlsx"
            workbook = Workbook()
            workbook.active.title = "总体指标"
            workbook.active.append(["指标", "数值", "单位"])
            planning_sheet = workbook.create_sheet("规划结果")
            planning_sheet.append(["设备类型", "设计台数", "单台容量", "总容量", "单位"])
            planning_sheet.append(["柴发", 1, 100, 100, "kW"])
            annual_sheet = workbook.create_sheet("规划年指标")
            annual_sheet.append(["指标", "数值", "单位"])
            green_sheet = workbook.create_sheet("供能分析")
            green_sheet.append(["指标", "数值", "单位"])
            safety_sheet = workbook.create_sheet("安全评估")
            safety_sheet.append(["指标", "数值", "单位"])
            daily_sheet = workbook.create_sheet("供能日曲线")
            daily_sheet.append(["day", "load_energy"])
            daily_sheet.append([1, 100])
            monthly_sheet = workbook.create_sheet("供能月曲线")
            monthly_sheet.append(["month", "load_energy"])
            monthly_sheet.append([1, 100])
            safety_daily_sheet = workbook.create_sheet("安全日曲线")
            safety_daily_sheet.append(["day", "frequency_max"])
            safety_daily_sheet.append([1, 50.1])
            dispatch_sheet = workbook.create_sheet("调度结果")
            dispatch_sheet.append(["小时", "负荷总功率"])
            dispatch_sheet.append([1, 100])
            workbook.save(result_path)
            workbook.close()

            read_sheets = []
            original_reader = server.read_workbook_rows_with_field_map

            def tracking_reader(workbook, sheet_name, limit=None):
                read_sheets.append(sheet_name)
                return original_reader(workbook, sheet_name, limit)

            with patch.object(server, "read_workbook_rows_with_field_map", side_effect=tracking_reader):
                status, headers, body = server.handle_api_path("/api/optimization/status?scheme=方案A&light=1")
            payload = json.loads(body.decode("utf-8"))

            self.assertEqual(status, 200)
            self.assertNotIn("调度结果", read_sheets)
            self.assertEqual(payload["results"]["curves"]["green_hourly"], [])
            self.assertEqual(payload["results"]["curves"]["green_daily"][0]["load_energy"], 100)
        finally:
            server.PLANNING_STORE = original_store
            server.OPTIMIZATION_RUNTIME = original_runtime
            shutil.rmtree(planning_root, ignore_errors=True)

    def test_runtime_running_schemes_does_not_parse_finished_workbooks(self):
        manager = server.OptimizationRuntimeManager()
        finished = manager._runtime_for_scheme("已完成方案")
        finished.status = "已完成"
        running = manager._runtime_for_scheme("运行中方案")
        running.status = "运行中"

        with patch.object(server, "read_result_workbook_display_payload_for_response") as workbook_reader:
            self.assertEqual(manager.running_schemes(), ["运行中方案"])
            workbook_reader.assert_not_called()

    def test_evaluation_status_reads_display_results_from_selected_workbook(self):
        planning_root = WEB_ROOT / "tests" / "tmp_evaluation_status_workbook"
        shutil.rmtree(planning_root, ignore_errors=True)
        planning_root.mkdir(parents=True)
        original_store = server.PLANNING_STORE
        original_runtime = server.EVALUATION_RUNTIME
        server.PLANNING_STORE = server.planning_store.PlanningStore(root=planning_root)
        server.EVALUATION_RUNTIME = server.EvaluationRuntimeManager()
        try:
            server.PLANNING_STORE.create_scheme("方案A")
            result_path = planning_root / "方案A" / "case_results.xlsx"
            workbook = Workbook()
            workbook.active.title = "总体指标"
            workbook.active.append(["指标", "数值", "单位"])
            workbook.active.append(["柴油消耗", 8.8, "吨"])
            planning_sheet = workbook.create_sheet("规划结果")
            planning_sheet.append(["设备类型", "设计台数", "单台容量", "总容量", "单位"])
            planning_sheet.append(["评估柴发", 1, 200, 200, "kW"])
            annual_sheet = workbook.create_sheet("规划年指标")
            annual_sheet.append(["指标", "数值", "单位"])
            annual_sheet.append(["评估年指标", 12, "kWh"])
            green_sheet = workbook.create_sheet("供能分析")
            green_sheet.append(["指标", "数值", "单位"])
            green_sheet.append(["评估供能指标", 34, "kWh"])
            daily_sheet = workbook.create_sheet("供能日曲线")
            daily_sheet.append(["day", "load_energy", "diesel_energy"])
            daily_sheet.append([1, 56, 7])
            safety_sheet = workbook.create_sheet("安全评估")
            safety_sheet.append(["指标", "数值", "单位"])
            safety_sheet.append(["评估安全指标", 49.9, "Hz"])
            safety_daily_sheet = workbook.create_sheet("安全日曲线")
            safety_daily_sheet.append(["day", "frequency_max", "frequency_min"])
            safety_daily_sheet.append([1, 50.1, 49.9])
            workbook.save(result_path)
            workbook.close()

            status, headers, body = server.handle_api_path("/api/evaluation/status?scheme=方案A&filename=case_results.xlsx")
            payload = json.loads(body.decode("utf-8"))

            self.assertEqual(status, 200)
            self.assertEqual(payload["result_filename"], "case_results.xlsx")
            self.assertEqual(payload["result_file"], str(result_path))
            metric_labels = [item["label"] for item in payload["metrics"]]
            self.assertIn("度电成本", metric_labels)
            self.assertIn("绿电占比", metric_labels)
            self.assertNotIn("综合评分", metric_labels)
            self.assertNotIn("风险等级", metric_labels)
            self.assertEqual(payload["metrics"][3], {"label": "柴油消耗", "value": 8.8, "unit": "吨"})
            self.assertEqual(payload["results"]["overview_tables"][0]["rows"][0]["设备类型"], "评估柴发")
            self.assertEqual(payload["results"]["overview_tables"][1]["rows"][0]["指标"], "评估年指标")
            self.assertEqual(payload["results"]["green_table"][0]["指标"], "评估供能指标")
            self.assertEqual(payload["results"]["safety_table"][0]["指标"], "评估安全指标")
            self.assertEqual(payload["results"]["curves"]["green_daily"][0]["load_energy"], 56)
            self.assertEqual(payload["results"]["curves"]["safety_daily"][0]["frequency_min"], 49.9)
        finally:
            server.PLANNING_STORE = original_store
            server.EVALUATION_RUNTIME = original_runtime
            shutil.rmtree(planning_root, ignore_errors=True)

    def test_evaluation_runtime_default_metrics_expose_cost_and_green_ratio(self):
        runtime = server.EvaluationRuntime(scheme="方案A")

        payload = runtime.snapshot()
        metric_labels = [item["label"] for item in payload["metrics"]]

        self.assertEqual(metric_labels[:5], ["状态", "开始", "完成", "度电成本", "绿电占比"])
        self.assertNotIn("综合评分", metric_labels)
        self.assertNotIn("风险等级", metric_labels)

    def test_light_evaluation_status_skips_hourly_dispatch_sheet(self):
        planning_root = WEB_ROOT / "tests" / "tmp_evaluation_status_light"
        shutil.rmtree(planning_root, ignore_errors=True)
        planning_root.mkdir(parents=True)
        original_store = server.PLANNING_STORE
        original_runtime = server.EVALUATION_RUNTIME
        server.PLANNING_STORE = server.planning_store.PlanningStore(root=planning_root)
        server.EVALUATION_RUNTIME = server.EvaluationRuntimeManager()
        try:
            server.PLANNING_STORE.create_scheme("方案A")
            result_path = planning_root / "方案A" / "case_results.xlsx"
            workbook = Workbook()
            workbook.active.title = "总体指标"
            workbook.active.append(["指标", "数值", "单位"])
            planning_sheet = workbook.create_sheet("规划结果")
            planning_sheet.append(["设备类型", "设计台数", "单台容量", "总容量", "单位"])
            planning_sheet.append(["柴发", 1, 100, 100, "kW"])
            annual_sheet = workbook.create_sheet("规划年指标")
            annual_sheet.append(["指标", "数值", "单位"])
            green_sheet = workbook.create_sheet("供能分析")
            green_sheet.append(["指标", "数值", "单位"])
            safety_sheet = workbook.create_sheet("安全评估")
            safety_sheet.append(["指标", "数值", "单位"])
            daily_sheet = workbook.create_sheet("供能日曲线")
            daily_sheet.append(["day", "load_energy"])
            daily_sheet.append([1, 200])
            monthly_sheet = workbook.create_sheet("供能月曲线")
            monthly_sheet.append(["month", "load_energy"])
            monthly_sheet.append([1, 200])
            safety_daily_sheet = workbook.create_sheet("安全日曲线")
            safety_daily_sheet.append(["day", "frequency_min"])
            safety_daily_sheet.append([1, 49.9])
            dispatch_sheet = workbook.create_sheet("调度结果")
            dispatch_sheet.append(["小时", "负荷总功率"])
            dispatch_sheet.append([1, 200])
            workbook.save(result_path)
            workbook.close()

            read_sheets = []
            original_reader = server.read_workbook_rows_with_field_map

            def tracking_reader(workbook, sheet_name, limit=None):
                read_sheets.append(sheet_name)
                return original_reader(workbook, sheet_name, limit)

            with patch.object(server, "read_workbook_rows_with_field_map", side_effect=tracking_reader):
                status, headers, body = server.handle_api_path(
                    "/api/evaluation/status?scheme=方案A&filename=case_results.xlsx&light=1"
                )
            payload = json.loads(body.decode("utf-8"))

            self.assertEqual(status, 200)
            self.assertNotIn("调度结果", read_sheets)
            self.assertEqual(payload["results"]["curves"]["green_hourly"], [])
            self.assertEqual(payload["results"]["curves"]["green_daily"][0]["load_energy"], 200)
        finally:
            server.PLANNING_STORE = original_store
            server.EVALUATION_RUNTIME = original_runtime
            shutil.rmtree(planning_root, ignore_errors=True)

    def test_evaluation_results_api_manages_scheme_result_workbooks(self):
        planning_root = WEB_ROOT / "tests" / "tmp_evaluation_results"
        shutil.rmtree(planning_root, ignore_errors=True)
        planning_root.mkdir(parents=True)
        original_store = server.PLANNING_STORE
        original_runtime = server.OPTIMIZATION_RUNTIME
        server.PLANNING_STORE = server.planning_store.PlanningStore(root=planning_root)
        server.OPTIMIZATION_RUNTIME = server.OptimizationRuntimeManager()
        try:
            server.PLANNING_STORE.create_scheme("方案A")

            status, headers, body = server.handle_api_path("/api/evaluation/results?scheme=方案A")
            self.assertEqual(status, 200)
            self.assertEqual(json.loads(body.decode("utf-8"))["results"], [])

            source_path = planning_root / "方案A" / "opt_results.xlsx"
            create_workbook = Workbook()
            create_workbook.active.title = "总体指标"
            create_workbook.active.append(["指标", "数值"])
            planning_sheet = create_workbook.create_sheet("规划结果")
            planning_sheet.append(["设备类型", "设计台数", "单台容量", "总容量", "单位"])
            planning_sheet.append(["柴发", 2, 320, 640, "kW"])
            planning_sheet.append(["储能", 4, 250, 1000, "kWh"])
            planning_sheet.append(["电制氢", 1, 80, 0, "kW"])
            planning_sheet.append(["燃料电池", 0, 50, 50, "kW"])
            embedded_curve_sheet = create_workbook.create_sheet("调度结果")
            embedded_curve_sheet.append(["小时", "负荷总功率"])
            embedded_curve_sheet.append([1, 100])
            create_workbook.save(source_path)
            source_curve_path = server.result_curves_workbook_path(source_path)
            source_curve_workbook = Workbook()
            source_curve_workbook.active.title = "调度结果"
            source_curve_workbook.active.append(["小时", "负荷总功率"])
            source_curve_workbook.active.append([1, 100])
            source_curve_workbook.save(source_curve_path)
            source_curve_workbook.close()

            broken_path = planning_root / "方案A" / "aaa_results.xlsx"
            broken_path.write_bytes(b"not a valid workbook")
            dead_path = planning_root / "方案A" / "dead_results.xlsx"
            dead_path.write_bytes(b"dead workbook")
            dead_curve_path = server.result_curves_workbook_path(dead_path)
            dead_curve_path.write_bytes(b"dead curve workbook")

            status, headers, body = server.handle_api_path(
                "/api/evaluation/results?scheme=方案A"
            )
            listed = json.loads(body.decode("utf-8"))
            self.assertEqual(status, 200)
            self.assertEqual(listed["selected"], "opt_results.xlsx")
            broken_item = next(item for item in listed["results"] if item["name"] == "aaa_results.xlsx")
            self.assertFalse(broken_item["readable"])
            self.assertIn("结果文件无法读取", broken_item["message"])
            self.assertEqual(
                listed["planning_result_rows"],
                [
                    {"设备类型": "柴发", "设计台数": 2, "单台容量": 320, "总容量": 640, "单位": "kW"},
                    {"设备类型": "储能", "设计台数": 4, "单台容量": 250, "总容量": 1000, "单位": "kWh"},
                    {"设备类型": "电制氢", "设计台数": 1, "单台容量": 80, "总容量": 80, "单位": "kW"},
                    {"设备类型": "燃料电池", "设计台数": 0, "单台容量": 50, "总容量": 0, "单位": "kW"},
                ],
            )
            status, headers, body = server.handle_api_path(
                "/api/evaluation/results?scheme=方案A&light=1"
            )
            light_listed = json.loads(body.decode("utf-8"))
            self.assertEqual(status, 200)
            self.assertEqual(light_listed["selected"], "opt_results.xlsx")
            self.assertIn("results", light_listed)
            self.assertNotIn("planning_result_rows", light_listed)

            status, headers, body = server.handle_api_path(
                "/api/evaluation/results?scheme=方案A&filename=aaa_results.xlsx"
            )
            selected_broken = json.loads(body.decode("utf-8"))
            self.assertEqual(status, 200)
            self.assertEqual(selected_broken["selected"], "aaa_results.xlsx")
            self.assertEqual(selected_broken["planning_result_rows"], [])

            status, headers, body = server.handle_evaluation_results_api_path(
                "/api/evaluation/results",
                "POST",
                json.dumps(
                    {
                        "scheme": "方案A",
                        "action": "copy",
                        "filename": "opt_results.xlsx",
                        "target_name": "custom",
                    },
                    ensure_ascii=False,
                ).encode("utf-8"),
            )
            copied = json.loads(body.decode("utf-8"))
            self.assertEqual(status, 200)
            self.assertEqual(copied["selected"], "custom_results.xlsx")
            self.assertTrue((planning_root / "方案A" / "custom_results.xlsx").exists())
            self.assertFalse((planning_root / "方案A" / "custom_curves.xlsx").exists())
            source_after_copy = load_workbook(source_path, read_only=True)
            try:
                self.assertNotIn("调度结果", source_after_copy.sheetnames)
            finally:
                source_after_copy.close()
            copied_workbook = load_workbook(planning_root / "方案A" / "custom_results.xlsx", read_only=True)
            try:
                self.assertNotIn("调度结果", copied_workbook.sheetnames)
            finally:
                copied_workbook.close()

            status, headers, body = server.handle_evaluation_results_api_path(
                "/api/evaluation/results",
                "POST",
                json.dumps(
                    {
                        "scheme": "方案A",
                        "action": "copy",
                        "filename": "opt_results.xlsx",
                        "target_name": "custom",
                    },
                    ensure_ascii=False,
                ).encode("utf-8"),
            )
            duplicate = json.loads(body.decode("utf-8"))
            self.assertEqual(status, 409)
            self.assertEqual(duplicate["error"], "exists")
            self.assertIn("复制失败", duplicate["message"])

            status, headers, body = server.handle_evaluation_results_api_path(
                "/api/evaluation/results",
                "POST",
                json.dumps(
                    {
                        "scheme": "方案A",
                        "action": "copy",
                        "filename": "opt_results.xlsx",
                        "target_name": "aaa",
                    },
                    ensure_ascii=False,
                ).encode("utf-8"),
            )
            overwritten = json.loads(body.decode("utf-8"))
            self.assertEqual(status, 200)
            self.assertEqual(overwritten["selected"], "aaa_results.xlsx")
            overwritten_item = next(item for item in overwritten["results"] if item["name"] == "aaa_results.xlsx")
            self.assertTrue(overwritten_item["readable"])
            workbook = load_workbook(planning_root / "方案A" / "aaa_results.xlsx", read_only=True)
            try:
                self.assertIn("规划结果", workbook.sheetnames)
            finally:
                workbook.close()

            status, headers, body = server.handle_evaluation_results_api_path(
                "/api/evaluation/results",
                "POST",
                json.dumps(
                    {"scheme": "方案A", "action": "delete", "filename": "dead_results.xlsx"},
                    ensure_ascii=False,
                ).encode("utf-8"),
            )
            deleted_broken = json.loads(body.decode("utf-8"))
            self.assertEqual(status, 200)
            self.assertFalse(dead_path.exists())
            self.assertFalse(dead_curve_path.exists())
            self.assertNotIn("dead_results.xlsx", [item["name"] for item in deleted_broken["results"]])

            status, headers, body = server.handle_evaluation_results_api_path(
                "/api/evaluation/results",
                "POST",
                json.dumps(
                    {"scheme": "方案A", "action": "delete", "filename": "opt_results.xlsx"},
                    ensure_ascii=False,
                ).encode("utf-8"),
            )
            protected_delete = json.loads(body.decode("utf-8"))
            self.assertEqual(status, 400)
            self.assertEqual(protected_delete["error"], "bad_request")
            self.assertIn("默认结果文件不允许删除", protected_delete["message"])
            self.assertTrue((planning_root / "方案A" / "opt_results.xlsx").exists())

            status, headers, body = server.handle_evaluation_results_api_path(
                "/api/evaluation/results",
                "POST",
                json.dumps(
                    {"scheme": "方案A", "action": "save", "filename": "opt_results.xlsx"},
                    ensure_ascii=False,
                ).encode("utf-8"),
            )
            protected_save = json.loads(body.decode("utf-8"))
            self.assertEqual(status, 400)
            self.assertEqual(protected_save["error"], "bad_request")
            self.assertIn("默认结果文件不允许修改", protected_save["message"])

            status, headers, body = server.handle_evaluation_results_api_path(
                "/api/evaluation/results",
                "POST",
                json.dumps(
                    {"scheme": "方案A", "action": "save", "filename": "custom_results.xlsx"},
                    ensure_ascii=False,
                ).encode("utf-8"),
            )
            saved = json.loads(body.decode("utf-8"))
            self.assertEqual(status, 200)
            self.assertEqual(saved["selected"], "custom_results.xlsx")
            workbook = load_workbook(planning_root / "方案A" / "custom_results.xlsx", read_only=True)
            try:
                self.assertIn("总体指标", workbook.sheetnames)
                self.assertIn("规划结果", workbook.sheetnames)
            finally:
                workbook.close()

            status, headers, body = server.handle_evaluation_results_api_path(
                "/api/evaluation/results",
                "POST",
                json.dumps(
                    {
                        "scheme": "方案A",
                        "action": "save",
                        "filename": "custom_results.xlsx",
                        "planning_result_rows": [
                            {"设备类型": "柴发", "设计台数": 5},
                            {"设备类型": "储能", "设计台数": 7},
                        ],
                    },
                    ensure_ascii=False,
                ).encode("utf-8"),
            )
            saved_with_counts = json.loads(body.decode("utf-8"))
            self.assertEqual(status, 200)
            self.assertEqual(saved_with_counts["selected"], "custom_results.xlsx")
            saved_counts = {row["设备类型"]: row["设计台数"] for row in saved_with_counts["planning_result_rows"]}
            saved_totals = {row["设备类型"]: row["总容量"] for row in saved_with_counts["planning_result_rows"]}
            self.assertEqual(saved_counts["柴发"], 5)
            self.assertEqual(saved_counts["储能"], 7)
            self.assertEqual(saved_totals["柴发"], 1600)
            self.assertEqual(saved_totals["储能"], 1750)
            workbook = load_workbook(planning_root / "方案A" / "custom_results.xlsx", read_only=True)
            try:
                workbook_counts = {
                    row[0]: row[1]
                    for row in workbook["规划结果"].iter_rows(min_row=2, values_only=True)
                    if row and row[0]
                }
                workbook_totals = {
                    row[0]: row[3]
                    for row in workbook["规划结果"].iter_rows(min_row=2, values_only=True)
                    if row and row[0]
                }
                self.assertEqual(workbook_counts["柴发"], 5)
                self.assertEqual(workbook_counts["储能"], 7)
                self.assertEqual(workbook_totals["柴发"], 1600)
                self.assertEqual(workbook_totals["储能"], 1750)
            finally:
                workbook.close()

            result_path_for_retry = planning_root / "方案A" / "custom_results.xlsx"
            replace_calls = {"count": 0}
            original_replace = Path.replace

            def flaky_replace(self, target):
                if Path(self).name == f".{result_path_for_retry.name}.tmp":
                    replace_calls["count"] += 1
                    if replace_calls["count"] == 1:
                        raise PermissionError("simulated workbook lock")
                return original_replace(self, target)

            with patch.object(Path, "replace", new=flaky_replace):
                status, headers, body = server.handle_evaluation_results_api_path(
                    "/api/evaluation/results",
                    "POST",
                    json.dumps(
                        {
                            "scheme": "方案A",
                            "action": "save",
                            "filename": "custom_results.xlsx",
                            "planning_result_rows": [{"设备类型": "柴发", "设计台数": 6}],
                        },
                        ensure_ascii=False,
                    ).encode("utf-8"),
                )
            retried_save = json.loads(body.decode("utf-8"))
            self.assertEqual(status, 200)
            self.assertEqual(retried_save["selected"], "custom_results.xlsx")
            self.assertGreaterEqual(replace_calls["count"], 2)
            retried_counts = {row["设备类型"]: row["设计台数"] for row in retried_save["planning_result_rows"]}
            self.assertEqual(retried_counts["柴发"], 6)

            def always_locked_replace(self, target):
                if Path(self).name == f".{result_path_for_retry.name}.tmp":
                    raise PermissionError("still locked")
                return original_replace(self, target)

            with patch.object(Path, "replace", new=always_locked_replace), patch.object(server.file_ops.time, "sleep", return_value=None):
                status, headers, body = server.handle_evaluation_results_api_path(
                    "/api/evaluation/results",
                    "POST",
                    json.dumps(
                        {
                            "scheme": "方案A",
                            "action": "save",
                            "filename": "custom_results.xlsx",
                            "planning_result_rows": [{"设备类型": "柴发", "设计台数": 6}],
                        },
                        ensure_ascii=False,
                    ).encode("utf-8"),
                )
            locked = json.loads(body.decode("utf-8"))
            self.assertEqual(status, 409)
            self.assertEqual(locked["error"], "file_locked")
            self.assertIn("结果文件被占用", locked["message"])
            self.assertIn("请关闭", locked["message"])

            original_copy2 = server.file_ops.shutil.copy2
            copy_calls = {"count": 0}

            def flaky_copy2(source, target):
                if Path(target).name == "copyretry_results.xlsx":
                    copy_calls["count"] += 1
                    if copy_calls["count"] == 1:
                        raise PermissionError("simulated copy lock")
                return original_copy2(source, target)

            with patch.object(server.file_ops.shutil, "copy2", side_effect=flaky_copy2):
                status, headers, body = server.handle_evaluation_results_api_path(
                    "/api/evaluation/results",
                    "POST",
                    json.dumps(
                        {
                            "scheme": "方案A",
                            "action": "copy",
                            "filename": "custom_results.xlsx",
                            "target_name": "copyretry",
                        },
                        ensure_ascii=False,
                    ).encode("utf-8"),
                )
            retried_copy = json.loads(body.decode("utf-8"))
            self.assertEqual(status, 200)
            self.assertEqual(retried_copy["selected"], "copyretry_results.xlsx")
            self.assertGreaterEqual(copy_calls["count"], 2)

            copyretry_curve_path = planning_root / "方案A" / "copyretry_curves.xlsx"
            curve_workbook = Workbook()
            curve_workbook.active.title = "调度结果"
            curve_workbook.save(copyretry_curve_path)
            curve_workbook.close()
            status, headers, body = server.handle_evaluation_results_api_path(
                "/api/evaluation/results",
                "POST",
                json.dumps(
                    {
                        "scheme": "方案A",
                        "action": "rename",
                        "filename": "copyretry_results.xlsx",
                        "target_name": "renamed",
                    },
                    ensure_ascii=False,
                ).encode("utf-8"),
            )
            renamed = json.loads(body.decode("utf-8"))
            self.assertEqual(status, 200)
            self.assertEqual(renamed["selected"], "renamed_results.xlsx")
            self.assertFalse((planning_root / "方案A" / "copyretry_results.xlsx").exists())
            self.assertFalse((planning_root / "方案A" / "copyretry_curves.xlsx").exists())
            self.assertTrue((planning_root / "方案A" / "renamed_results.xlsx").exists())
            self.assertTrue((planning_root / "方案A" / "renamed_curves.xlsx").exists())
            self.assertIn("renamed_results.xlsx", [item["name"] for item in renamed["results"]])
            self.assertNotIn("copyretry_results.xlsx", [item["name"] for item in renamed["results"]])

            status, headers, body = server.handle_evaluation_results_api_path(
                "/api/evaluation/results",
                "POST",
                json.dumps(
                    {
                        "scheme": "方案A",
                        "action": "rename",
                        "filename": "opt_results.xlsx",
                        "target_name": "opt2",
                    },
                    ensure_ascii=False,
                ).encode("utf-8"),
            )
            protected_rename = json.loads(body.decode("utf-8"))
            self.assertEqual(status, 400)
            self.assertEqual(protected_rename["error"], "bad_request")
            self.assertIn("默认结果文件不允许重命名", protected_rename["message"])

            def always_locked_copy2(source, target):
                if Path(target).name == "copylocked_results.xlsx":
                    raise PermissionError("copy target locked")
                return original_copy2(source, target)

            with patch.object(server.file_ops.shutil, "copy2", side_effect=always_locked_copy2), patch.object(server.file_ops.time, "sleep", return_value=None):
                status, headers, body = server.handle_evaluation_results_api_path(
                    "/api/evaluation/results",
                    "POST",
                    json.dumps(
                        {
                            "scheme": "方案A",
                            "action": "copy",
                            "filename": "custom_results.xlsx",
                            "target_name": "copylocked",
                        },
                        ensure_ascii=False,
                    ).encode("utf-8"),
                )
            locked_copy = json.loads(body.decode("utf-8"))
            self.assertEqual(status, 409)
            self.assertEqual(locked_copy["error"], "file_locked")
            self.assertIn("结果文件被占用", locked_copy["message"])
            self.assertIn("无法复制", locked_copy["message"])

            original_unlink = Path.unlink
            unlink_calls = {"count": 0}

            def flaky_unlink(self, *args, **kwargs):
                if Path(self).name == "renamed_results.xlsx":
                    unlink_calls["count"] += 1
                    if unlink_calls["count"] == 1:
                        raise PermissionError("simulated delete lock")
                return original_unlink(self, *args, **kwargs)

            with patch.object(Path, "unlink", new=flaky_unlink):
                status, headers, body = server.handle_evaluation_results_api_path(
                    "/api/evaluation/results",
                    "POST",
                    json.dumps(
                        {"scheme": "方案A", "action": "delete", "filename": "renamed_results.xlsx"},
                        ensure_ascii=False,
                    ).encode("utf-8"),
                )
            retried_delete = json.loads(body.decode("utf-8"))
            self.assertEqual(status, 200)
            self.assertGreaterEqual(unlink_calls["count"], 2)
            self.assertNotIn("renamed_results.xlsx", [item["name"] for item in retried_delete["results"]])
            self.assertFalse((planning_root / "方案A" / "renamed_curves.xlsx").exists())

            for invalid_count in (-1, 1.5, "2.2"):
                status, headers, body = server.handle_evaluation_results_api_path(
                    "/api/evaluation/results",
                    "POST",
                    json.dumps(
                        {
                            "scheme": "方案A",
                            "action": "save",
                            "filename": "custom_results.xlsx",
                            "planning_result_rows": [{"设备类型": "柴发", "设计台数": invalid_count}],
                        },
                        ensure_ascii=False,
                    ).encode("utf-8"),
                )
                invalid = json.loads(body.decode("utf-8"))
                self.assertEqual(status, 400)
                self.assertEqual(invalid["error"], "bad_request")
                self.assertIn("设计台数必须为非负整数", invalid["message"])

            status, headers, body = server.handle_evaluation_results_api_path(
                "/api/evaluation/results",
                "POST",
                json.dumps(
                    {"scheme": "方案A", "action": "delete", "filename": "custom_results.xlsx"},
                    ensure_ascii=False,
                ).encode("utf-8"),
            )
            deleted = json.loads(body.decode("utf-8"))
            self.assertEqual(status, 200)
            self.assertFalse((planning_root / "方案A" / "custom_results.xlsx").exists())
            self.assertEqual([item["name"] for item in deleted["results"]], ["aaa_results.xlsx", "opt_results.xlsx"])
            self.assertEqual(deleted["selected"], "aaa_results.xlsx")
        finally:
            server.PLANNING_STORE = original_store
            server.OPTIMIZATION_RUNTIME = original_runtime
            shutil.rmtree(planning_root, ignore_errors=True)

    def test_comparison_data_api_reads_selected_result_workbooks(self):
        planning_root = WEB_ROOT / "tests" / "tmp_comparison_data"
        shutil.rmtree(planning_root, ignore_errors=True)
        planning_root.mkdir(parents=True)
        original_store = server.PLANNING_STORE
        server.PLANNING_STORE = server.planning_store.PlanningStore(root=planning_root)
        try:
            server.PLANNING_STORE.create_scheme("方案A")
            result_path = planning_root / "方案A" / "case_results.xlsx"
            workbook = Workbook()
            planning_sheet = workbook.active
            planning_sheet.title = "规划结果"
            planning_sheet.append(["设备类型", "设计台数", "单台容量", "总容量", "单位"])
            planning_sheet.append(["柴发", 2, 100, 200, "kW"])
            planning_sheet.append(["风机", 1, 50, 50, "kW"])
            green_sheet = workbook.create_sheet("供能分析")
            green_sheet.append(["指标", "数值", "单位"])
            green_sheet.append(["柴油消耗", 12.5, "吨"])
            green_sheet.append(["绿电占比", 85, "%"])
            annual_sheet = workbook.create_sheet("规划年指标")
            annual_sheet.append(["指标", "数值", "单位"])
            annual_sheet.append(["年总成本", 123.4, "万元"])
            annual_sheet.append(["新能源弃电率", 1.2, "%"])
            safety_sheet = workbook.create_sheet("安全评估")
            safety_sheet.append(["指标", "数值", "单位"])
            safety_sheet.append(["最大未供负荷", 0, "kW"])
            daily_sheet = workbook.create_sheet("供能日曲线")
            daily_sheet.append(["day", "load_energy", "wind_energy"])
            daily_sheet.append([1, 1000, 220])
            daily_sheet.append([2, 1100, 240])
            monthly_sheet = workbook.create_sheet("供能月曲线")
            monthly_sheet.append(["month", "load_energy", "renewable_curtailed_rate"])
            monthly_sheet.append([1, 30000, 1.1])
            monthly_sheet.append([2, 28000, 1.4])
            safety_daily_sheet = workbook.create_sheet("安全日曲线")
            safety_daily_sheet.append(["day", "frequency_max", "frequency_min"])
            safety_daily_sheet.append([1, 50.1, 49.9])
            safety_daily_sheet.append([2, 50.2, 49.8])
            dispatch_sheet = workbook.create_sheet("调度结果")
            dispatch_sheet.append(["小时", "时间", "负荷", "风电出力", "光伏出力"])
            for hour in range(1, 8761):
                dispatch_sheet.append([hour, f"H{hour:04d}", 80 + hour % 3, 20 + hour % 5, 30 + hour % 7])
            workbook.save(result_path)

            status, headers, body = server.handle_api_path(
                "/api/comparison/data?items="
                + quote(json.dumps([{"scheme": "方案A", "filename": "case_results.xlsx"}], ensure_ascii=False))
            )
            payload = json.loads(body.decode("utf-8"))

            self.assertEqual(status, 200)
            self.assertEqual(payload["items"][0]["scheme"], "方案A")
            self.assertEqual(payload["items"][0]["result_display_name"], "case")
            self.assertEqual(payload["tables"]["capacity"][0]["设备类型"], "柴发")
            self.assertEqual(payload["tables"]["energy"][0]["指标"], "柴油消耗")
            self.assertEqual(payload["tables"]["safety"][0]["指标"], "最大未供负荷")
            self.assertIn("负荷", payload["curves"])
            self.assertIn("风电出力", payload["curves"])
            self.assertEqual(len(payload["series"]["负荷"][0]["points"]), 8760)
            self.assertEqual(payload["series"]["负荷"][0]["label"], "方案A / case")
            self.assertEqual(payload["curve_groups"]["hourly"]["title"], "小时级曲线")
            self.assertEqual(payload["curve_groups"]["daily"]["title"], "日级统计")
            self.assertEqual(payload["curve_groups"]["safety"]["title"], "安全日曲线")
            self.assertEqual(payload["curve_groups"]["monthly"]["title"], "月度统计")
            self.assertIn("负荷", payload["curve_groups"]["hourly"]["curves"])
            self.assertIn("负荷总电量", payload["curve_groups"]["daily"]["curves"])
            self.assertIn("风机总发电量", payload["curve_groups"]["daily"]["curves"])
            self.assertIn("最高频率", payload["curve_groups"]["safety"]["curves"])
            self.assertIn("最低频率", payload["curve_groups"]["safety"]["curves"])
            self.assertIn("新能源弃电率", payload["curve_groups"]["monthly"]["curves"])
            self.assertNotIn("load_energy", payload["curve_groups"]["daily"]["curves"])
            self.assertNotIn("renewable_curtailed_rate", payload["curve_groups"]["monthly"]["curves"])
            self.assertEqual(len(payload["curve_groups"]["daily"]["series"]["负荷总电量"][0]["points"]), 2)
            self.assertEqual(len(payload["curve_groups"]["safety"]["series"]["最高频率"][0]["points"]), 2)
            self.assertEqual(len(payload["curve_groups"]["monthly"]["series"]["负荷总电量"][0]["points"]), 2)
            self.assertEqual(payload["annual_table"][0]["指标"], "年总成本")
            self.assertEqual(payload["annual_table"][0]["方案A / case"], 123.4)
        finally:
            server.PLANNING_STORE = original_store
            shutil.rmtree(planning_root, ignore_errors=True)

    def test_comparison_data_api_accepts_up_to_eight_items(self):
        planning_root = WEB_ROOT / "tests" / "tmp_comparison_limit"
        shutil.rmtree(planning_root, ignore_errors=True)
        planning_root.mkdir(parents=True)
        original_store = server.PLANNING_STORE
        server.PLANNING_STORE = server.planning_store.PlanningStore(root=planning_root)
        try:
            server.PLANNING_STORE.create_scheme("方案A")
            result_path = planning_root / "方案A" / "case_results.xlsx"
            workbook = Workbook()
            planning_sheet = workbook.active
            planning_sheet.title = "规划结果"
            planning_sheet.append(["设备类型", "设计台数", "单台容量"])
            planning_sheet.append(["柴发", 2, 100])
            green_sheet = workbook.create_sheet("供能分析")
            green_sheet.append(["指标", "数值", "单位"])
            green_sheet.append(["柴油消耗", 12.5, "吨"])
            safety_sheet = workbook.create_sheet("安全评估")
            safety_sheet.append(["指标", "数值", "单位"])
            safety_sheet.append(["最大未供负荷", 0, "kW"])
            annual_sheet = workbook.create_sheet("规划年指标")
            annual_sheet.append(["指标", "数值", "单位"])
            annual_sheet.append(["年总成本", 123.4, "万元"])
            workbook.save(result_path)
            workbook.close()

            items = [{"scheme": "方案A", "filename": "case_results.xlsx"} for _ in range(9)]
            status, headers, body = server.handle_api_path(
                "/api/comparison/data?mode=summary&items=" + quote(json.dumps(items, ensure_ascii=False))
            )
            payload = json.loads(body.decode("utf-8"))

            self.assertEqual(status, 200)
            self.assertEqual(len(payload["items"]), 8)
            self.assertEqual(payload["items"][-1]["id"], "item-8")
        finally:
            server.PLANNING_STORE = original_store
            shutil.rmtree(planning_root, ignore_errors=True)

    def test_comparison_summary_mode_skips_hourly_dispatch_sheet(self):
        planning_root = WEB_ROOT / "tests" / "tmp_comparison_summary"
        shutil.rmtree(planning_root, ignore_errors=True)
        planning_root.mkdir(parents=True)
        original_store = server.PLANNING_STORE
        server.PLANNING_STORE = server.planning_store.PlanningStore(root=planning_root)
        try:
            server.PLANNING_STORE.create_scheme("方案A")
            result_path = planning_root / "方案A" / "case_results.xlsx"
            workbook = Workbook()
            planning_sheet = workbook.active
            planning_sheet.title = "规划结果"
            planning_sheet.append(["设备类型", "设计台数", "单台容量", "总容量", "单位"])
            planning_sheet.append(["柴发", 2, 100, 200, "kW"])
            green_sheet = workbook.create_sheet("供能分析")
            green_sheet.append(["指标", "数值", "单位"])
            green_sheet.append(["柴油消耗", 12.5, "吨"])
            annual_sheet = workbook.create_sheet("规划年指标")
            annual_sheet.append(["指标", "数值", "单位"])
            annual_sheet.append(["年总成本", 123.4, "万元"])
            safety_sheet = workbook.create_sheet("安全评估")
            safety_sheet.append(["指标", "数值", "单位"])
            daily_sheet = workbook.create_sheet("供能日曲线")
            daily_sheet.append(["day", "load_energy"])
            daily_sheet.append([1, 1000])
            monthly_sheet = workbook.create_sheet("供能月曲线")
            monthly_sheet.append(["month", "load_energy"])
            monthly_sheet.append([1, 30000])
            dispatch_sheet = workbook.create_sheet("调度结果")
            dispatch_sheet.append(["小时", "负荷"])
            dispatch_sheet.append([1, 80])
            workbook.save(result_path)
            workbook.close()

            read_sheets = []
            original_reader = server.read_curve_sheet

            def tracking_reader(workbook, sheet_name, limit=None):
                read_sheets.append(sheet_name)
                return original_reader(workbook, sheet_name, limit)

            with patch.object(server, "read_curve_sheet", side_effect=tracking_reader):
                status, headers, body = server.handle_api_path(
                    "/api/comparison/data?mode=summary&items="
                    + quote(json.dumps([{"scheme": "方案A", "filename": "case_results.xlsx"}], ensure_ascii=False))
                )
            payload = json.loads(body.decode("utf-8"))

            self.assertEqual(status, 200)
            self.assertNotIn("调度结果", read_sheets)
            self.assertIn("负荷", payload["curve_groups"]["hourly"]["curves"])
            self.assertEqual(payload["curve_groups"]["hourly"]["series"], {})
            self.assertIn("负荷", payload["curves"])
            self.assertEqual(payload["series"], {})
            self.assertIn("负荷总电量", payload["curve_groups"]["daily"]["curves"])
            self.assertEqual(payload["tables"]["capacity"][0]["设备类型"], "柴发")
        finally:
            server.PLANNING_STORE = original_store
            shutil.rmtree(planning_root, ignore_errors=True)

    def test_comparison_curve_mode_returns_only_requested_curves(self):
        planning_root = WEB_ROOT / "tests" / "tmp_comparison_curve_slice"
        shutil.rmtree(planning_root, ignore_errors=True)
        planning_root.mkdir(parents=True)
        original_store = server.PLANNING_STORE
        server.PLANNING_STORE = server.planning_store.PlanningStore(root=planning_root)
        file_cache.clear_all()
        try:
            server.PLANNING_STORE.create_scheme("方案A")
            result_path = planning_root / "方案A" / "case_results.xlsx"
            workbook = Workbook()
            planning_sheet = workbook.active
            planning_sheet.title = "规划结果"
            planning_sheet.append(["设备类型", "设计台数", "单台容量"])
            planning_sheet.append(["柴发", 2, 100])
            workbook.create_sheet("供能分析").append(["指标", "数值", "单位"])
            workbook.create_sheet("安全评估").append(["指标", "数值", "单位"])
            dispatch_sheet = workbook.create_sheet("调度结果")
            dispatch_sheet.append(["小时", "负荷", "风电出力", "光伏出力"])
            dispatch_sheet.append([1, 80, 20, 30])
            dispatch_sheet.append([2, 81, 21, 31])
            workbook.save(result_path)
            workbook.close()

            items = [{"scheme": "方案A", "filename": "case_results.xlsx"}]
            status, headers, body = server.handle_api_path(
                "/api/comparison/data?mode=curves&group=hourly&curves="
                + quote(json.dumps(["负荷"], ensure_ascii=False))
                + "&items="
                + quote(json.dumps(items, ensure_ascii=False))
            )
            payload = json.loads(body.decode("utf-8"))

            self.assertEqual(status, 200)
            self.assertEqual(payload["curve_groups"]["hourly"]["curves"], ["负荷"])
            self.assertEqual(set(payload["curve_groups"]["hourly"]["series"].keys()), {"负荷"})
            self.assertEqual(payload["curve_groups"]["hourly"]["series"]["负荷"][0]["points"][0]["y"], 80)
            self.assertNotIn("风电出力", payload["curve_groups"]["hourly"]["series"])
            self.assertEqual(payload["tables"], {"capacity": [], "energy": [], "safety": []})
        finally:
            server.PLANNING_STORE = original_store
            file_cache.clear_all()
            shutil.rmtree(planning_root, ignore_errors=True)

    def test_snapshot_reads_summary_from_csv_files(self):
        payload = server.build_snapshot(force_reload=True)

        self.assertEqual(payload["system"], "考察站风-光-氢-储-柴联合规划系统")
        self.assertIn("summary", payload)
        self.assertNotIn("simu", payload)
        self.assertNotIn("scada", payload)
        self.assertNotIn("agc", payload)

    def test_snapshot_reloads_after_configured_interval(self):
        data_dir = WEB_ROOT / "tests" / "tmp_data_source"
        if data_dir.exists():
            shutil.rmtree(data_dir)
        data_dir.mkdir(parents=True)
        try:
            (data_dir / "summary.csv").write_text("key,value,unit\nrunning_days,1,天\n", encoding="utf-8")

            reader = server.CsvDataSource(data_dir=data_dir, reload_interval=0)
            first = reader.snapshot()
            (data_dir / "summary.csv").write_text("key,value,unit\nrunning_days,2,天\n", encoding="utf-8")
            second = reader.snapshot()
        finally:
            shutil.rmtree(data_dir, ignore_errors=True)

        self.assertEqual(first["summary"]["running_days"], 1)
        self.assertEqual(second["summary"]["running_days"], 2)

    def test_mysql_data_source_builds_snapshot_from_database_rows(self):
        rows = {
            "SELECT `key`, value, unit FROM overview_summary ORDER BY display_order, id": [
                {"key": "running_days", "value": "449", "unit": "天"},
            ],
        }
        fake_connection = FakeConnection(rows)
        source = server.MySqlDataSource(connector_factory=lambda config: fake_connection, reload_interval=0)

        payload = source.snapshot()

        self.assertEqual(payload["summary"]["running_days"], 449)
        self.assertNotIn("simu", payload)
        self.assertNotIn("scada", payload)
        self.assertNotIn("agc", payload)
        self.assertEqual(
            fake_connection.cursor_obj.executed,
            [("SELECT `key`, value, unit FROM overview_summary ORDER BY display_order, id", ())],
        )

    def test_unknown_api_path_returns_404_json(self):
        status, headers, body = server.handle_api_path("/api/not-found")

        self.assertEqual(status, 404)
        self.assertEqual(headers["Content-Type"], "application/json; charset=utf-8")
        self.assertEqual(json.loads(body.decode("utf-8"))["error"], "not_found")

    def test_reliability_parameters_api_uses_result_counts_and_preserves_traceability_fields(self):
        planning_root = WEB_ROOT / "tests" / "tmp_reliability_parameters"
        shutil.rmtree(planning_root, ignore_errors=True)
        planning_root.mkdir(parents=True)
        original_store = server.PLANNING_STORE
        original_runtime = server.RELIABILITY_RUNTIME
        server.PLANNING_STORE = server.planning_store.PlanningStore(root=planning_root)
        server.RELIABILITY_RUNTIME = server.ReliabilityRuntimeManager()
        try:
            payload = server.planning_store.default_payload("方案A")
            payload["time_series"] = payload["time_series"][:24]
            payload["wind_turbines"][0].update({"capacity": 100, "quantity_lower": 0, "quantity_upper": 8, "cost": 120})
            payload["photovoltaics"][0].update({"capacity": 50, "quantity_lower": 0, "quantity_upper": 10, "cost": 22.5})
            payload["storage_pcs"][0].update({"power_capacity": 100, "quantity_lower": 0, "quantity_upper": 10, "cost": 70})
            payload["storage_battery_packs"][0].update(
                {"battery_capacity": 200, "quantity_lower": 0, "quantity_upper": 10, "cost": 300}
            )
            payload["diesel_generators"][0].update(
                {"capacity": 100, "power_upper": 100, "quantity_lower": 0, "quantity_upper": 5, "cost": 30}
            )
            server.PLANNING_STORE.write_scheme("方案A", payload)

            source_path = planning_root / "方案A" / "case_results.xlsx"
            workbook = Workbook()
            sheet = workbook.active
            sheet.title = server.PLANNING_RESULT_SHEET_NAME
            sheet.append(server.PLANNING_RESULT_HEADERS)
            sheet.append(["柴发", 3, 100, 300, "kW"])
            sheet.append(["风机", 2, 100, 200, "kW"])
            sheet.append(["光伏", 4, 50, 200, "kW"])
            sheet.append(["储能PCS", 3, 100, 300, "kW"])
            sheet.append(["储能电池组", 5, 200, 1000, "kWh"])
            workbook.save(source_path)
            workbook.close()

            status, headers, body = server.handle_api_path(
                "/api/reliability/parameters?scheme=方案A&filename=case_results.xlsx"
            )
            self.assertEqual(status, 200)
            initial = json.loads(body.decode("utf-8"))
            initial_by_type = {row["device_type"]: row for row in initial["parameters"]["devices"]}
            self.assertEqual(initial_by_type["diesel"]["unit_count"], 3)
            self.assertEqual(initial_by_type["wind"]["unit_count"], 2)
            self.assertEqual(initial_by_type["pv"]["unit_count"], 4)
            self.assertEqual(initial_by_type["storage"]["unit_count"], 3)
            self.assertEqual(initial_by_type["storage"]["unit_capacity_kwh"], 200)
            self.assertEqual(initial_by_type["storage"]["battery_forced_outage_rate"], 0.02)
            self.assertEqual(initial_by_type["storage"]["battery_mttr_hours"], 96)
            self.assertTrue(any("工作簿" in message for message in initial["assumption_warnings"]))

            parameters = initial["parameters"]
            parameters["simulation_years"] = 123
            parameters["devices"][0].update(
                {
                    "unit_count": 99,
                    "extreme_cold_capacity_factor": 0.63,
                    "capex_wan_per_unit": 188.5,
                    "fixed_om_rate": 0.031,
                    "design_life_years": 25,
                    "vendor_batch": "demo-wind-batch",
                }
            )
            diesel = next(row for row in parameters["devices"] if row["device_type"] == "diesel")
            diesel.update({"startup_failure_rate": 0.015, "variable_om_yuan_per_kwh": 0.21})
            storage = next(row for row in parameters["devices"] if row["device_type"] == "storage")
            storage.update({"forced_outage_rate": 0.035, "mttr_hours": 72, "battery_forced_outage_rate": 0.02, "battery_mttr_hours": 96})
            status, headers, body = server.handle_reliability_api_path(
                "/api/reliability/parameters",
                "PUT",
                json.dumps(
                    {
                        "scheme": "方案A",
                        "filename": "case_results.xlsx",
                        "parameters": parameters,
                    },
                    ensure_ascii=False,
                ).encode("utf-8"),
            )
            self.assertEqual(status, 200)
            saved = json.loads(body.decode("utf-8"))["parameters"]
            saved_by_type = {row["device_type"]: row for row in saved["devices"]}
            self.assertEqual(saved_by_type["wind"]["unit_count"], 2)
            self.assertEqual(saved_by_type["wind"]["vendor_batch"], "demo-wind-batch")
            self.assertEqual(saved_by_type["wind"]["extreme_cold_capacity_factor"], 0.63)
            self.assertEqual(saved_by_type["diesel"]["startup_failure_rate"], 0.015)
            self.assertEqual(saved_by_type["diesel"]["variable_om_yuan_per_kwh"], 0.21)
            prepared = server.prepare_reliability_scheme_payload(payload, saved, has_planning_result=True)
            self.assertEqual(prepared["storage_pcs"][0]["forced_outage_rate"], 0.035)
            self.assertEqual(prepared["storage_pcs"][0]["mttr_hours"], 72)
            self.assertEqual(prepared["storage_battery_packs"][0]["forced_outage_rate"], 0.02)
            self.assertEqual(prepared["storage_battery_packs"][0]["mttr_hours"], 96)
            planning_rows = server.read_evaluation_planning_result_rows("方案A", "case_results.xlsx")
            case = server.reliability.build_reliability_case(prepared, planning_rows, {"hours_per_year": 24})
            group_by_type = {group["device_type"]: group for group in case["groups"]}
            self.assertEqual(group_by_type["pcs"]["unit_count"], 3)
            self.assertEqual(group_by_type["battery"]["unit_count"], 5)
            parameter_path = planning_root / "方案A" / server.RELIABILITY_PARAMETERS_FILE_NAME
            self.assertTrue(parameter_path.exists())
            stored = json.loads(parameter_path.read_text(encoding="utf-8"))
            self.assertEqual(stored["simulation_years"], 123)

            export_path = planning_root / "方案A" / f"case_results{server.RELIABILITY_RESULT_WORKBOOK_SUFFIX}"
            export_book = Workbook()
            export_book.save(export_path)
            export_book.close()
            listed_names = [item["name"] for item in server.list_evaluation_result_files("方案A")]
            self.assertEqual(listed_names, ["case_results.xlsx"])
        finally:
            server.PLANNING_STORE = original_store
            server.RELIABILITY_RUNTIME = original_runtime
            shutil.rmtree(planning_root, ignore_errors=True)

    def test_reliability_control_runs_in_background_and_writes_independent_json_and_xlsx(self):
        planning_root = WEB_ROOT / "tests" / "tmp_reliability_control"
        shutil.rmtree(planning_root, ignore_errors=True)
        planning_root.mkdir(parents=True)
        original_store = server.PLANNING_STORE
        original_runtime = server.RELIABILITY_RUNTIME
        server.PLANNING_STORE = server.planning_store.PlanningStore(root=planning_root)
        server.RELIABILITY_RUNTIME = server.ReliabilityRuntimeManager()

        class InlineProcess:
            next_pid = 70000

            def __init__(self, target, args=(), daemon=None):
                self.target = target
                self.args = args
                self.daemon = daemon
                self.pid = None
                self.exitcode = None
                self._alive = False

            def start(self):
                type(self).next_pid += 1
                self.pid = type(self).next_pid
                self._alive = True
                try:
                    self.target(*self.args)
                    self.exitcode = 0
                finally:
                    self._alive = False

            def is_alive(self):
                return self._alive

            def join(self, timeout=None):
                return None

            def terminate(self):
                self._alive = False
                self.exitcode = -15

            def kill(self):
                self.terminate()

        try:
            payload = server.planning_store.default_payload("方案A")
            payload["time_series"] = payload["time_series"][:24]
            for hour, row in enumerate(payload["time_series"]):
                row.update({"wind_speed": 0, "solar_irradiance": 0, "load": 80, "temperature": -30})
            payload["diesel_generators"][0].update(
                {
                    "capacity": 100,
                    "power_upper": 100,
                    "power_lower": 0,
                    "quantity_lower": 2,
                    "quantity_upper": 2,
                }
            )
            for key in ("wind_turbines", "photovoltaics", "storage_pcs", "storage_battery_packs"):
                payload[key][0]["quantity_lower"] = 0
                payload[key][0]["quantity_upper"] = 0
            server.PLANNING_STORE.write_scheme("方案A", payload)

            source_path = planning_root / "方案A" / "case_results.xlsx"
            workbook = Workbook()
            sheet = workbook.active
            sheet.title = server.PLANNING_RESULT_SHEET_NAME
            sheet.append(server.PLANNING_RESULT_HEADERS)
            sheet.append(["柴发", 2, 100, 200, "kW"])
            workbook.save(source_path)
            workbook.close()
            source_bytes = source_path.read_bytes()

            parameters, _, _ = server.read_reliability_parameters("方案A", "case_results.xlsx")
            parameters["simulation_years"] = 2
            parameters["random_seed"] = 42
            diesel = next(row for row in parameters["devices"] if row["device_type"] == "diesel")
            diesel.update({"forced_outage_rate": 0.05, "mttr_hours": 12})

            with patch.object(server.multiprocessing, "Queue", queue.Queue), patch.object(
                server.multiprocessing,
                "Process",
                InlineProcess,
            ):
                status, headers, body = server.handle_control_path(
                    "/api/reliability/control",
                    json.dumps(
                        {
                            "action": "start",
                            "scheme": "方案A",
                            "filename": "case_results.xlsx",
                            "parameters": parameters,
                        },
                        ensure_ascii=False,
                    ).encode("utf-8"),
                )
                self.assertEqual(status, 200)
                self.assertEqual(json.loads(body.decode("utf-8"))["state"]["status"], "running")

                status, headers, body = server.handle_api_path(
                    "/api/reliability/status?scheme=方案A&filename=case_results.xlsx"
                )
                self.assertEqual(status, 200)
                completed = json.loads(body.decode("utf-8"))
                self.assertEqual(completed["status"], "completed")
                self.assertEqual(completed["progress"], 100)
                self.assertIsInstance(completed["result"], dict)
                self.assertIn("metrics", completed["result"])
                self.assertEqual(completed["result"]["parameters"]["simulation_years"], 2)

            json_path = planning_root / "方案A" / "case_results_reliability.json"
            xlsx_path = planning_root / "方案A" / "case_results_reliability_results.xlsx"
            self.assertTrue(json_path.exists())
            self.assertTrue(xlsx_path.exists())
            self.assertEqual(source_path.read_bytes(), source_bytes)
            reliability_payload = json.loads(json_path.read_text(encoding="utf-8"))
            self.assertEqual(reliability_payload["source_result_filename"], "case_results.xlsx")
            self.assertIn("annual_distribution", reliability_payload)
            self.assertIn("n1_scenarios", reliability_payload)

            export_book = load_workbook(xlsx_path, read_only=True, data_only=True)
            try:
                self.assertEqual(
                    export_book.sheetnames,
                    ["摘要", "N-1场景", "年度样本", "设备贡献", "假设与日志"],
                )
            finally:
                export_book.close()

            status, headers, body = server.handle_api_path(
                "/api/reliability/results?scheme=方案A&filename=case_results.xlsx"
            )
            self.assertEqual(status, 200)
            result_index = json.loads(body.decode("utf-8"))
            self.assertEqual([item["name"] for item in result_index["results"]], ["case_results.xlsx"])
            self.assertEqual(result_index["result_file"], "case_results_reliability.json")
            self.assertEqual(result_index["export_file"], "case_results_reliability_results.xlsx")
            self.assertIsInstance(result_index["result"], dict)
        finally:
            server.PLANNING_STORE = original_store
            server.RELIABILITY_RUNTIME = original_runtime
            shutil.rmtree(planning_root, ignore_errors=True)

    def test_planning_api_create_read_save_copy_rename(self):
        planning_root = WEB_ROOT / "tests" / "tmp_planning_api"
        shutil.rmtree(planning_root, ignore_errors=True)
        planning_root.mkdir(parents=True)
        original_store = server.PLANNING_STORE
        server.PLANNING_STORE = server.planning_store.PlanningStore(root=planning_root)
        try:
            status, headers, body = server.handle_planning_api_path(
                "/api/planning/schemes",
                "POST",
                json.dumps({"name": "方案A"}).encode("utf-8"),
            )
            self.assertEqual(status, 200)
            created = json.loads(body.decode("utf-8"))
            self.assertEqual(created["scheme"], "方案A")
            self.assertEqual(len(created["time_series"]), 8760)
            self.assertIn("planning_parameters", created)
            self.assertNotIn("design_life_years", created["planning_parameters"][0])
            self.assertEqual(created["storage_battery_packs"][0]["self_discharge_rate"], 0.01)
            self.assertEqual(created["hydrogen_tanks"][0]["self_discharge_rate"], 0.001)
            self.assertEqual(created["hydrogen_tanks"][0]["soc_upper"], 0.85)
            self.assertEqual(created["hydrogen_tanks"][0]["soc_lower"], 0.15)
            self.assertEqual(created["storage_pcs"][0]["storage_charge_efficiency"], 0.95)
            self.assertEqual(created["storage_pcs"][0]["storage_discharge_efficiency"], 0.95)

            created["time_series"][0]["load"] = 123.4
            created["planning_parameters"][0]["storage_frequency_regulation_enabled"] = 1
            status, headers, body = server.handle_planning_api_path(
                "/api/planning/schemes/方案A",
                "PUT",
                json.dumps(created, ensure_ascii=False).encode("utf-8"),
            )
            self.assertEqual(status, 200)

            status, headers, body = server.handle_planning_api_path("/api/planning/schemes/方案A", "GET", b"")
            loaded = json.loads(body.decode("utf-8"))
            self.assertEqual(loaded["time_series"][0]["load"], 123.4)
            self.assertNotIn("design_life_years", loaded["planning_parameters"][0])
            self.assertEqual(loaded["planning_parameters"][0]["storage_frequency_regulation_enabled"], 1)
            self.assertEqual(loaded["storage_battery_packs"][0]["self_discharge_rate"], 0.01)
            self.assertEqual(loaded["hydrogen_tanks"][0]["self_discharge_rate"], 0.001)
            self.assertEqual(loaded["hydrogen_tanks"][0]["soc_upper"], 0.85)
            self.assertEqual(loaded["hydrogen_tanks"][0]["soc_lower"], 0.15)
            self.assertEqual(loaded["storage_pcs"][0]["storage_charge_efficiency"], 0.95)
            self.assertEqual(loaded["storage_pcs"][0]["storage_discharge_efficiency"], 0.95)

            status, headers, body = server.handle_planning_api_path("/api/planning/schemes/方案A/overview", "GET", b"")
            overview = json.loads(body.decode("utf-8"))
            self.assertEqual(status, 200)
            self.assertNotIn("time_series", overview)
            self.assertFalse(overview["time_series_loaded"])
            self.assertEqual(overview["time_series_count"], 8760)
            self.assertIn("diesel_generators", overview)
            self.assertIn("planning_parameters", overview)
            self.assertNotIn("design_life_years", overview["planning_parameters"][0])

            save_labels = []
            original_save = server.planning_store.file_ops.save_workbook_with_retry

            def tracking_save(workbook, path, label):
                save_labels.append(label)
                return original_save(workbook, path, label)

            overview["diesel_generators"][0]["name"] = "只保存参数"
            with patch.object(server.planning_store.file_ops, "save_workbook_with_retry", side_effect=tracking_save):
                status, headers, body = server.handle_planning_api_path(
                    "/api/planning/schemes/方案A",
                    "PUT",
                    json.dumps(overview, ensure_ascii=False).encode("utf-8"),
                )
            saved_overview = json.loads(body.decode("utf-8"))
            self.assertEqual(status, 200)
            self.assertEqual(save_labels, ["参数文件"])
            self.assertNotIn("time_series", saved_overview)
            self.assertFalse(saved_overview["time_series_loaded"])
            self.assertEqual(saved_overview["diesel_generators"][0]["name"], "只保存参数")

            status, headers, body = server.handle_planning_api_path("/api/planning/schemes/方案A/time-series", "GET", b"")
            time_payload = json.loads(body.decode("utf-8"))
            self.assertEqual(status, 200)
            self.assertEqual(time_payload["time_series"][0]["load"], 123.4)
            self.assertNotIn("diesel_generators", time_payload)

            status, headers, body = server.handle_planning_api_path(
                "/api/planning/schemes/copy",
                "POST",
                json.dumps({"source": "方案A", "target": "方案B"}, ensure_ascii=False).encode("utf-8"),
            )
            self.assertEqual(status, 200)

            status, headers, body = server.handle_planning_api_path(
                "/api/planning/schemes/copy",
                "POST",
                json.dumps({"source": "方案A", "target": "方案B"}, ensure_ascii=False).encode("utf-8"),
            )
            duplicate_copy = json.loads(body.decode("utf-8"))
            self.assertEqual(status, 400)
            self.assertIn("目标方案已存在", duplicate_copy["message"])

            status, headers, body = server.handle_planning_api_path(
                "/api/planning/schemes/copy",
                "POST",
                json.dumps({"source": "方案A", "target": "方案B", "overwrite": True}, ensure_ascii=False).encode("utf-8"),
            )
            overwritten_copy = json.loads(body.decode("utf-8"))
            self.assertEqual(status, 200)
            self.assertEqual(overwritten_copy["scheme"], "方案B")

            status, headers, body = server.handle_planning_api_path(
                "/api/planning/schemes/rename",
                "POST",
                json.dumps({"source": "方案B", "target": "方案C"}, ensure_ascii=False).encode("utf-8"),
            )
            self.assertEqual(status, 200)

            status, headers, body = server.handle_planning_api_path("/api/planning/schemes", "GET", b"")
            names = [item["name"] for item in json.loads(body.decode("utf-8"))["schemes"]]
            self.assertEqual(names, ["方案A", "方案C"])
        finally:
            server.PLANNING_STORE = original_store
            shutil.rmtree(planning_root, ignore_errors=True)

    def test_planning_api_reports_locked_parameter_file_as_conflict(self):
        planning_root = WEB_ROOT / "tests" / "tmp_planning_api_locked"
        shutil.rmtree(planning_root, ignore_errors=True)
        planning_root.mkdir(parents=True)
        original_store = server.PLANNING_STORE
        server.PLANNING_STORE = server.planning_store.PlanningStore(root=planning_root)
        try:
            payload = server.PLANNING_STORE.create_scheme("方案A")
            payload["time_series"][0]["load"] = 789

            def always_locked_replace(self, target):
                if Path(self).name == f".{server.planning_store.WORKBOOK_NAME}.tmp":
                    raise PermissionError("still locked")
                return Path.replace(self, target)

            with patch.object(Path, "replace", new=always_locked_replace), patch.object(server.file_ops.time, "sleep", return_value=None):
                status, headers, body = server.handle_planning_api_path(
                    "/api/planning/schemes/方案A",
                    "PUT",
                    json.dumps(payload, ensure_ascii=False).encode("utf-8"),
                )

            locked = json.loads(body.decode("utf-8"))
            self.assertEqual(status, 409)
            self.assertEqual(locked["error"], "file_locked")
            self.assertIn("参数文件被占用", locked["message"])
            self.assertIn("请关闭", locked["message"])
        finally:
            server.PLANNING_STORE = original_store
            shutil.rmtree(planning_root, ignore_errors=True)

    def test_planning_store_backfills_key_specific_self_discharge_defaults_for_legacy_workbook(self):
        planning_root = WEB_ROOT / "tests" / "tmp_planning_legacy_defaults"
        shutil.rmtree(planning_root, ignore_errors=True)
        planning_root.mkdir(parents=True)
        path = planning_root / "parameters.xlsx"
        payload = server.planning_store.default_payload("方案A")
        workbook = server.planning_store.build_workbook(payload)
        try:
            for sheet_name in ("储能电池组参数", "储氢罐参数"):
                sheet = workbook[sheet_name]
                for cell in sheet[1]:
                    if cell.value == "self_discharge_rate":
                        sheet.delete_cols(cell.column)
                        break
            workbook.save(path)
        finally:
            workbook.close()

        try:
            loaded = server.planning_store.read_workbook(
                path,
                "方案A",
                include_keys=["storage_battery_packs", "hydrogen_tanks"],
            )

            self.assertEqual(loaded["storage_battery_packs"][0]["self_discharge_rate"], 0.01)
            self.assertEqual(loaded["hydrogen_tanks"][0]["self_discharge_rate"], 0.001)
        finally:
            shutil.rmtree(planning_root, ignore_errors=True)

    def test_planning_time_series_import_parses_csv_and_validates_required_columns(self):
        rows = ["风速,太阳辐射,室温,负荷"]
        rows.extend(f"{i % 20},{500 + i % 300},{-10 + i % 30},{100 + i % 50}" for i in range(8760))
        content = "\n".join(rows).encode("utf-8")

        status, headers, body = server.handle_planning_api_path(
            "/api/planning/time-series/import",
            "POST",
            json.dumps({"filename": "timeseries.csv", "content_base64": base64.b64encode(content).decode("ascii")}).encode("utf-8"),
        )

        self.assertEqual(status, 200)
        payload = json.loads(body.decode("utf-8"))
        self.assertEqual(payload["time_series_count"], 8760)
        self.assertEqual(payload["time_series"][0]["hour_index"], 1)
        self.assertEqual(payload["time_series"][0]["wind_speed"], 0)
        self.assertEqual(payload["time_series"][0]["solar_irradiance"], 500)
        self.assertEqual(payload["time_series"][0]["temperature"], -10)
        self.assertEqual(payload["time_series"][0]["load"], 100)

        bad_rows = ["风速,室温,负荷", "1,2,3"]
        status, headers, body = server.handle_planning_api_path(
            "/api/planning/time-series/import",
            "POST",
            json.dumps({"filename": "bad.csv", "content_base64": base64.b64encode("\n".join(bad_rows).encode("utf-8")).decode("ascii")}).encode("utf-8"),
        )

        self.assertEqual(status, 400)
        error_payload = json.loads(body.decode("utf-8"))
        self.assertIn("找不到对应的列", error_payload["message"])
        self.assertIn("太阳辐射", error_payload["message"])

    def test_planning_time_series_import_matches_fuzzy_headers_and_pads_short_files(self):
        rows = [
            "风速(m/s),单位面积太阳辐射(W/m^2),温度(摄氏度),用电功率(kW)",
            "3.5,600,18,120",
            "4.0,610,19,125",
            "4.5,620,20,130",
        ]
        content = "\n".join(rows).encode("utf-8")

        status, headers, body = server.handle_planning_api_path(
            "/api/planning/time-series/import",
            "POST",
            json.dumps({"filename": "short_timeseries.csv", "content_base64": base64.b64encode(content).decode("ascii")}).encode("utf-8"),
        )

        self.assertEqual(status, 200)
        payload = json.loads(body.decode("utf-8"))
        self.assertEqual(payload["time_series_count"], 8760)
        self.assertEqual(payload["time_series"][0]["solar_irradiance"], 600)
        self.assertEqual(payload["time_series"][0]["temperature"], 18)
        self.assertEqual(payload["time_series"][0]["load"], 120)
        self.assertEqual(payload["time_series"][8759]["hour_index"], 8760)
        self.assertEqual(payload["time_series"][8759]["datetime"], "H8760")
        self.assertEqual(payload["time_series"][8759]["wind_speed"], 4.5)
        self.assertEqual(payload["time_series"][8759]["solar_irradiance"], 620)
        self.assertIn("已按最后一行自动补齐8757行", payload["message"])

    def test_planning_time_series_import_fills_missing_middle_hours(self):
        rows = [
            "时间,风速,太阳辐照,环境温度,负荷功率",
            "H0001,3,500,10,100",
            "H0003,5,700,12,120",
        ]
        content = "\n".join(rows).encode("utf-8")

        status, headers, body = server.handle_planning_api_path(
            "/api/planning/time-series/import",
            "POST",
            json.dumps({"filename": "gap_timeseries.csv", "content_base64": base64.b64encode(content).decode("ascii")}).encode("utf-8"),
        )

        self.assertEqual(status, 200)
        payload = json.loads(body.decode("utf-8"))
        self.assertEqual(payload["time_series_count"], 8760)
        self.assertEqual(payload["time_series"][0]["datetime"], "H0001")
        self.assertEqual(payload["time_series"][0]["wind_speed"], 3)
        self.assertEqual(payload["time_series"][1]["datetime"], "H0002")
        self.assertEqual(payload["time_series"][1]["wind_speed"], 3)
        self.assertEqual(payload["time_series"][1]["solar_irradiance"], 500)
        self.assertEqual(payload["time_series"][2]["datetime"], "H0003")
        self.assertEqual(payload["time_series"][2]["wind_speed"], 5)
        self.assertEqual(payload["time_series"][8759]["datetime"], "H8760")
        self.assertEqual(payload["time_series"][8759]["load"], 120)
        self.assertIn("已补齐8758个缺失时点", payload["message"])

    def test_planning_time_series_import_repairs_invalid_numeric_values_with_neighbors(self):
        rows = [
            "时间,风速,太阳辐射,室温,负荷",
            "H0001,bad,500,10,100",
            "H0002,4,NaN,11,abc",
            "H0003,6,700,,120",
        ]
        content = "\n".join(rows).encode("utf-8")

        status, headers, body = server.handle_planning_api_path(
            "/api/planning/time-series/import",
            "POST",
            json.dumps({"filename": "dirty.csv", "content_base64": base64.b64encode(content).decode("ascii")}).encode("utf-8"),
        )

        self.assertEqual(status, 200)
        payload = json.loads(body.decode("utf-8"))
        self.assertEqual(payload["time_series_count"], 8760)
        self.assertEqual(payload["time_series"][0]["wind_speed"], 4)
        self.assertEqual(payload["time_series"][1]["solar_irradiance"], 500)
        self.assertEqual(payload["time_series"][1]["load"], 100)
        self.assertEqual(payload["time_series"][2]["temperature"], 11)
        self.assertEqual(payload["time_series"][2]["wind_speed"], 6)
        self.assertEqual(payload["time_series"][2]["solar_irradiance"], 700)
        self.assertEqual(payload["time_series"][2]["load"], 120)
        self.assertIn("已修复4个无效数值", payload["message"])

    def test_planning_time_series_import_parses_xlsx(self):
        workbook = Workbook()
        sheet = workbook.active
        sheet.append(["时间", "风速", "太阳辐照", "温度", "负荷"])
        for i in range(8760):
            sheet.append([f"H{i + 1:04d}", 3.5, 720, 15, 88])
        stream = BytesIO()
        workbook.save(stream)

        status, headers, body = server.handle_planning_api_path(
            "/api/planning/time-series/import",
            "POST",
            json.dumps({"filename": "timeseries.xlsx", "content_base64": base64.b64encode(stream.getvalue()).decode("ascii")}).encode("utf-8"),
        )

        self.assertEqual(status, 200)
        payload = json.loads(body.decode("utf-8"))
        self.assertEqual(payload["time_series_count"], 8760)
        self.assertEqual(payload["time_series"][8759]["datetime"], "H8760")
        self.assertEqual(payload["time_series"][8759]["wind_speed"], 3.5)

    def test_planning_load_curve_generation_matches_requested_statistics(self):
        for mode in ("random", "pattern1", "pattern2", "pattern3"):
            status, headers, body = server.handle_planning_api_path(
                "/api/planning/load-curve/generate",
                "POST",
                json.dumps({"mode": mode, "max": 180, "min": 40, "average": 95}).encode("utf-8"),
            )

            self.assertEqual(status, 200)
            payload = json.loads(body.decode("utf-8"))
            values = [row["load"] for row in payload["load_curve"]]
            self.assertEqual(payload["load_curve_count"], 8760)
            self.assertEqual(payload["mode"], mode)
            self.assertAlmostEqual(min(values), 40, places=3)
            self.assertAlmostEqual(max(values), 180, places=3)
            self.assertAlmostEqual(sum(values) / len(values), 95, places=3)

        status, headers, body = server.handle_planning_api_path(
            "/api/planning/load-curve/generate",
            "POST",
            json.dumps({"mode": "pattern1", "max": 180, "min": 40, "average": 200}).encode("utf-8"),
        )

        self.assertEqual(status, 400)
        self.assertIn("平均值必须介于最小值和最大值之间", json.loads(body.decode("utf-8"))["message"])

    def test_planning_wind_and_solar_curve_generation_matches_requested_statistics(self):
        cases = [
            ("wind_speed", 16, 0, 8, "wind_speed_curve"),
            ("solar_irradiance", 800, 0, 100, "solar_irradiance_curve"),
        ]
        for curve, maximum, minimum, average, row_key in cases:
            status, headers, body = server.handle_planning_api_path(
                "/api/planning/time-series-curve/generate",
                "POST",
                json.dumps({"curve": curve, "mode": "random", "max": maximum, "min": minimum, "average": average}).encode("utf-8"),
            )

            self.assertEqual(status, 200)
            payload = json.loads(body.decode("utf-8"))
            values = [row[curve] for row in payload[row_key]]
            self.assertEqual(payload["curve"], curve)
            self.assertEqual(payload["curve_count"], 8760)
            self.assertEqual(len(values), 8760)
            self.assertAlmostEqual(min(values), minimum, places=3)
            self.assertAlmostEqual(max(values), maximum, places=3)
            self.assertAlmostEqual(sum(values) / len(values), average, places=3)

    def test_planning_wind_and_solar_curve_import_can_return_raw_source(self):
        rows = ["时间,风速(m/s),太阳辐射(W/m^2)", "H0001,1,0", "H0002,2,100", "H0003,3,200"]
        content = "\n".join(rows).encode("utf-8")

        for curve, row_key, first, last in (
            ("wind_speed", "wind_speed_curve", 1, 3),
            ("solar_irradiance", "solar_irradiance_curve", 0, 200),
        ):
            status, headers, body = server.handle_planning_api_path(
                "/api/planning/time-series-curve/import",
                "POST",
                json.dumps(
                    {
                        "curve": curve,
                        "filename": "weather.csv",
                        "content_base64": base64.b64encode(content).decode("ascii"),
                        "raw": True,
                    }
                ).encode("utf-8"),
            )

            self.assertEqual(status, 200)
            payload = json.loads(body.decode("utf-8"))
            values = [row[curve] for row in payload[row_key]]
            self.assertEqual(payload["curve_count"], 8760)
            self.assertEqual(values[0], first)
            self.assertEqual(values[-1], last)
            self.assertIn("原始", payload["message"])

    def test_planning_load_curve_import_scales_and_pads_csv_to_8760(self):
        rows = ["时间,用电功率(kW)", "2024-01-01 00:00,10", "2024-01-01 02:00,30", "2024-01-01 05:00,50"]
        content = "\n".join(rows).encode("utf-8")

        status, headers, body = server.handle_planning_api_path(
            "/api/planning/load-curve/import",
            "POST",
            json.dumps(
                {
                    "filename": "load.csv",
                    "content_base64": base64.b64encode(content).decode("ascii"),
                    "max": 180,
                    "min": 40,
                    "average": 95,
                }
            ).encode("utf-8"),
        )

        self.assertEqual(status, 200)
        payload = json.loads(body.decode("utf-8"))
        values = [row["load"] for row in payload["load_curve"]]
        self.assertEqual(payload["load_curve_count"], 8760)
        self.assertAlmostEqual(min(values), 40, places=3)
        self.assertAlmostEqual(max(values), 180, places=3)
        self.assertAlmostEqual(sum(values) / len(values), 95, places=3)
        self.assertIn("已从load.csv导入8760点负荷曲线", payload["message"])
        self.assertIn("自动补齐", payload["message"])
        self.assertIn("缺失时点", payload["message"])

    def test_planning_load_curve_import_parses_xlsx_with_fuzzy_load_header(self):
        workbook = Workbook()
        sheet = workbook.active
        sheet.append(["小时", "负荷功率"])
        for i in range(24):
            sheet.append([i + 1, 50 + i])
        stream = BytesIO()
        workbook.save(stream)

        status, headers, body = server.handle_planning_api_path(
            "/api/planning/load-curve/import",
            "POST",
            json.dumps(
                {
                    "filename": "load.xlsx",
                    "content_base64": base64.b64encode(stream.getvalue()).decode("ascii"),
                    "max": 100,
                    "min": 10,
                    "average": 55,
                }
            ).encode("utf-8"),
        )

        self.assertEqual(status, 200)
        payload = json.loads(body.decode("utf-8"))
        self.assertEqual(payload["load_curve_count"], 8760)
        self.assertEqual(payload["source"], "file")
        self.assertAlmostEqual(payload["statistics"]["min"], 10, places=3)
        self.assertAlmostEqual(payload["statistics"]["max"], 100, places=3)

    def test_planning_load_curve_import_can_return_raw_source_and_generate_from_source(self):
        rows = ["时间,用电功率(kW)", "H0001,10", "H0002,20", "H0003,30"]
        content = "\n".join(rows).encode("utf-8")

        status, headers, body = server.handle_planning_api_path(
            "/api/planning/load-curve/import",
            "POST",
            json.dumps(
                {
                    "filename": "raw_load.csv",
                    "content_base64": base64.b64encode(content).decode("ascii"),
                    "raw": True,
                }
            ).encode("utf-8"),
        )

        self.assertEqual(status, 200)
        payload = json.loads(body.decode("utf-8"))
        raw_values = [row["load"] for row in payload["load_curve"]]
        self.assertEqual(payload["load_curve_count"], 8760)
        self.assertEqual(raw_values[0], 10)
        self.assertEqual(raw_values[-1], 30)
        self.assertAlmostEqual(min(raw_values), 10, places=3)
        self.assertAlmostEqual(max(raw_values), 30, places=3)
        self.assertEqual(payload["statistics"]["min"], 10)
        self.assertEqual(payload["statistics"]["max"], 30)
        self.assertIn("原始曲线", payload["message"])

        status, headers, body = server.handle_planning_api_path(
            "/api/planning/load-curve/generate",
            "POST",
            json.dumps(
                {
                    "mode": "file",
                    "source_load_curve": payload["load_curve"],
                    "max": 180,
                    "min": 40,
                    "average": 95,
                }
            ).encode("utf-8"),
        )

        self.assertEqual(status, 200)
        generated = json.loads(body.decode("utf-8"))
        values = [row["load"] for row in generated["load_curve"]]
        self.assertEqual(generated["mode"], "file")
        self.assertAlmostEqual(min(values), 40, places=3)
        self.assertAlmostEqual(max(values), 180, places=3)
        self.assertAlmostEqual(sum(values) / len(values), 95, places=3)

    def test_planning_load_curve_templates_save_conflict_overwrite_and_generate(self):
        template_path = WEB_ROOT / "tests" / "tmp_load_curve_templates.csv"
        if template_path.exists():
            template_path.unlink()
        original_path = server.LOAD_CURVE_TEMPLATE_PATH
        server.LOAD_CURVE_TEMPLATE_PATH = template_path
        try:
            first_curve = [{"hour_index": i + 1, "load": 10 + (i % 24)} for i in range(8760)]
            second_curve = [{"hour_index": i + 1, "load": 30 + (i % 48)} for i in range(8760)]

            status, headers, body = server.handle_planning_api_path(
                "/api/planning/load-curve/templates",
                "POST",
                json.dumps({"name": "模板A", "load_curve": first_curve}).encode("utf-8"),
            )

            self.assertEqual(status, 200)
            payload = json.loads(body.decode("utf-8"))
            self.assertEqual(payload["template"]["name"], "模板A")
            self.assertEqual(payload["template"]["load_curve_count"], 8760)
            self.assertIn("模板A", [item["name"] for item in payload["templates"]])
            self.assertTrue(template_path.exists())
            template_text = template_path.read_text(encoding="utf-8-sig")
            self.assertTrue(template_text.startswith("hour_index,模板A"))

            status, headers, body = server.handle_planning_api_path(
                "/api/planning/load-curve/templates",
                "POST",
                json.dumps({"name": "模板A", "load_curve": second_curve}).encode("utf-8"),
            )

            self.assertEqual(status, 409)
            self.assertIn("模板名称已存在", json.loads(body.decode("utf-8"))["message"])

            status, headers, body = server.handle_planning_api_path(
                "/api/planning/load-curve/templates",
                "POST",
                json.dumps({"name": "模板A", "load_curve": second_curve, "overwrite": True}).encode("utf-8"),
            )

            self.assertEqual(status, 200)
            self.assertIn("已覆盖负荷模板：模板A", json.loads(body.decode("utf-8"))["message"])

            status, headers, body = server.handle_planning_api_path(
                "/api/planning/load-curve/generate",
                "POST",
                json.dumps({"mode": "template:模板A", "max": 100, "min": 20, "average": 50}).encode("utf-8"),
            )

            self.assertEqual(status, 200)
            payload = json.loads(body.decode("utf-8"))
            values = [row["load"] for row in payload["load_curve"]]
            self.assertEqual(payload["mode"], "template:模板A")
            self.assertAlmostEqual(min(values), 20, places=3)
            self.assertAlmostEqual(max(values), 100, places=3)
            self.assertAlmostEqual(sum(values) / len(values), 50, places=3)
        finally:
            server.LOAD_CURVE_TEMPLATE_PATH = original_path
            if template_path.exists():
                template_path.unlink()

    def test_planning_load_curve_template_route_rejects_bad_post_body_not_not_found(self):
        status, headers, body = server.handle_planning_api_path(
            "/api/planning/load-curve/templates",
            "POST",
            json.dumps({"name": "模板X", "load_curve": []}, ensure_ascii=False).encode("utf-8"),
        )

        self.assertEqual(status, 400)
        payload = json.loads(body.decode("utf-8"))
        self.assertEqual(payload["error"], "bad_request")
        self.assertIn("负荷模板必须包含8760点曲线", payload["message"])

    def test_default_load_curve_templates_are_available(self):
        status, headers, body = server.handle_planning_api_path("/api/planning/load-curve/templates", "GET", b"")

        self.assertEqual(status, 200)
        payload = json.loads(body.decode("utf-8"))
        names = [item["name"] for item in payload["templates"]]
        self.assertGreaterEqual(names.index("模板1"), 0)
        self.assertIn("模板2", names)
        self.assertIn("模板3", names)

        for name in ("模板1", "模板2", "模板3"):
            status, headers, body = server.handle_planning_api_path(
                "/api/planning/load-curve/generate",
                "POST",
                json.dumps({"mode": f"template:{name}", "max": 120, "min": 20, "average": 65}, ensure_ascii=False).encode("utf-8"),
            )
            self.assertEqual(status, 200)
            payload = json.loads(body.decode("utf-8"))
            self.assertEqual(payload["load_curve_count"], 8760)
            self.assertAlmostEqual(payload["statistics"]["min"], 20, places=3)
            self.assertAlmostEqual(payload["statistics"]["max"], 120, places=3)
            self.assertAlmostEqual(payload["statistics"]["average"], 65, places=3)

    def test_planning_api_delete_scheme(self):
        planning_root = WEB_ROOT / "tests" / "tmp_planning_delete_api"
        shutil.rmtree(planning_root, ignore_errors=True)
        planning_root.mkdir(parents=True)
        original_store = server.PLANNING_STORE
        server.PLANNING_STORE = server.planning_store.PlanningStore(root=planning_root)
        try:
            server.PLANNING_STORE.create_scheme("方案A")
            server.PLANNING_STORE.create_scheme("方案B")

            status, headers, body = server.handle_planning_api_path("/api/planning/schemes/方案A", "DELETE", b"")

            self.assertEqual(status, 200)
            self.assertEqual(json.loads(body.decode("utf-8"))["deleted"], "方案A")
            self.assertFalse((planning_root / "方案A").exists())
            names = [item["name"] for item in server.PLANNING_STORE.list_schemes()]
            self.assertEqual(names, ["方案B"])
        finally:
            server.PLANNING_STORE = original_store
            shutil.rmtree(planning_root, ignore_errors=True)

    def test_planning_api_rejects_bad_scheme_name(self):
        status, headers, body = server.handle_planning_api_path(
            "/api/planning/schemes",
            "POST",
            json.dumps({"name": "../bad"}).encode("utf-8"),
        )

        self.assertEqual(status, 400)
        self.assertEqual(json.loads(body.decode("utf-8"))["error"], "bad_request")

    def test_planning_weather_history_endpoint_validates_year_before_current_year(self):
        status, headers, body = server.handle_planning_api_path(
            "/api/planning/weather-history",
            "POST",
            json.dumps({"latitude": 10, "longitude": 20, "year": 9999}).encode("utf-8"),
        )

        data = json.loads(body.decode("utf-8"))
        self.assertEqual(status, 400)
        self.assertEqual(data["error"], "bad_request")
        self.assertIn("历史数据年必须小于当前年", data["message"])

    def test_parse_nasa_power_hourly_response_requires_8760_rows(self):
        payload = {
            "header": {"fill_value": -999},
            "properties": {
                "parameter": {
                    "WS10M": {"2025010100": 7.1},
                    "ALLSKY_SFC_SW_DWN": {"2025010100": 0},
                    "T2M": {"2025010100": -12.5},
                }
            },
        }

        with self.assertRaises(server.WeatherHistoryError):
            server.parse_nasa_power_hourly_response(payload, 2025)

    def test_planning_geocode_endpoint_fills_coordinates_from_place_name(self):
        class FakeResponse:
            def __init__(self, payload):
                self.payload = payload

            def __enter__(self):
                return self

            def __exit__(self, exc_type, exc, tb):
                return False

            def read(self):
                return json.dumps(self.payload).encode("utf-8")

        def fake_urlopen(url, timeout):
            self.assertIn("geocoding-api.open-meteo.com", url)
            return FakeResponse(
                {
                    "results": [
                        {
                            "name": "北京",
                            "latitude": 39.9075,
                            "longitude": 116.39723,
                            "country": "中国",
                            "admin1": "北京市",
                        }
                    ]
                }
            )

        original_key = server.AMAP_WEB_SERVICE_KEY
        server.AMAP_WEB_SERVICE_KEY = ""
        try:
            with patch.object(server, "urlopen_with_user_agent", side_effect=fake_urlopen):
                status, headers, body = server.handle_planning_api_path(
                    "/api/planning/geocode",
                    "POST",
                    json.dumps({"place": "北京"}, ensure_ascii=False).encode("utf-8"),
                )
        finally:
            server.AMAP_WEB_SERVICE_KEY = original_key

        data = json.loads(body.decode("utf-8"))
        self.assertEqual(status, 200)
        self.assertEqual(data["latitude"], 39.9075)
        self.assertEqual(data["longitude"], 116.39723)
        self.assertEqual(data["source"], "Open-Meteo Geocoding API")

    def test_planning_geocode_prefers_amap_when_key_is_configured(self):
        class FakeResponse:
            def __enter__(self):
                return self

            def __exit__(self, exc_type, exc, tb):
                return False

            def read(self):
                return json.dumps(
                    {
                        "status": "1",
                        "geocodes": [
                            {
                                "formatted_address": "北京市",
                                "province": "北京市",
                                "city": "北京市",
                                "district": [],
                                "location": "116.407526,39.904030",
                            }
                        ],
                    },
                    ensure_ascii=False,
                ).encode("utf-8")

        original_key = server.AMAP_WEB_SERVICE_KEY
        server.AMAP_WEB_SERVICE_KEY = "test-key"
        try:
            def fake_urlopen(url, timeout):
                self.assertIn("restapi.amap.com", url)
                self.assertIn("key=test-key", url)
                return FakeResponse()

            with patch.object(server, "urlopen_with_user_agent", side_effect=fake_urlopen):
                status, headers, body = server.handle_planning_api_path(
                    "/api/planning/geocode",
                    "POST",
                    json.dumps({"place": "北京"}, ensure_ascii=False).encode("utf-8"),
                )
        finally:
            server.AMAP_WEB_SERVICE_KEY = original_key

        data = json.loads(body.decode("utf-8"))
        self.assertEqual(status, 200)
        self.assertEqual(data["latitude"], 39.90403)
        self.assertEqual(data["longitude"], 116.407526)
        self.assertEqual(data["source"], "高德地图 Web 服务地理编码 API")

    def test_planning_geocode_uses_global_provider_first_for_english_places(self):
        class FakeResponse:
            def __init__(self, payload):
                self.payload = payload

            def __enter__(self):
                return self

            def __exit__(self, exc_type, exc, tb):
                return False

            def read(self):
                return json.dumps(self.payload, ensure_ascii=False).encode("utf-8")

        requested_urls = []

        def fake_urlopen(url, timeout):
            requested_urls.append(url)
            if "geocoding-api.open-meteo.com" in url:
                return FakeResponse(
                    {
                        "results": [
                            {
                                "name": "New York",
                                "latitude": 40.71427,
                                "longitude": -74.00597,
                                "country": "美国",
                                "admin1": "纽约州",
                            }
                        ]
                    }
                )
            if "restapi.amap.com" in url:
                return FakeResponse(
                    {
                        "status": "1",
                        "geocodes": [
                            {
                                "formatted_address": "广东省惠州市惠东县New YorK(解放中路店)",
                                "province": "广东省",
                                "city": "惠州市",
                                "district": "惠东县",
                                "location": "114.721208,22.978660",
                            }
                        ],
                    }
                )
            raise AssertionError(f"unexpected url: {url}")

        original_key = server.AMAP_WEB_SERVICE_KEY
        server.AMAP_WEB_SERVICE_KEY = "test-key"
        try:
            with patch.object(server, "urlopen_with_user_agent", side_effect=fake_urlopen):
                status, headers, body = server.handle_planning_api_path(
                    "/api/planning/geocode",
                    "POST",
                    json.dumps({"place": "New York"}, ensure_ascii=False).encode("utf-8"),
                )
        finally:
            server.AMAP_WEB_SERVICE_KEY = original_key

        data = json.loads(body.decode("utf-8"))
        self.assertEqual(status, 200)
        self.assertEqual(data["latitude"], 40.71427)
        self.assertEqual(data["longitude"], -74.00597)
        self.assertEqual(data["source"], "Open-Meteo Geocoding API")
        self.assertIn("geocoding-api.open-meteo.com", requested_urls[0])
        self.assertTrue(all("restapi.amap.com" not in url for url in requested_urls))

    def test_planning_map_config_exposes_amap_key_when_configured(self):
        original_key = server.AMAP_WEB_SERVICE_KEY
        server.AMAP_WEB_SERVICE_KEY = "test-key"
        try:
            status, headers, body = server.handle_planning_api_path("/api/planning/map-config", "GET", b"")
        finally:
            server.AMAP_WEB_SERVICE_KEY = original_key

        data = json.loads(body.decode("utf-8"))
        self.assertEqual(status, 200)
        self.assertEqual(data["amap_key"], "test-key")
        self.assertEqual(data["preferred_provider"], "amap")

    def test_planning_map_config_has_default_amap_key(self):
        self.assertEqual(server.DEFAULT_AMAP_WEB_SERVICE_KEY, "21db26646aac8fed4620eaa36f210018")

    def test_planning_map_config_exposes_osm_provider_without_google_or_baidu(self):
        status, headers, body = server.handle_planning_api_path("/api/planning/map-config", "GET", b"")
        data = json.loads(body.decode("utf-8"))
        self.assertEqual(status, 200)
        self.assertNotIn("baidu_key", data)
        self.assertNotIn("google_key", data)
        self.assertEqual(data["osm_key"], "")
        self.assertIn({"key": "osm", "label": "OpenStreetMap", "enabled": True}, data["providers"])
        self.assertNotIn("google", {provider["key"] for provider in data["providers"]})

    def test_planning_geocode_endpoint_falls_back_to_nominatim(self):
        class FakeResponse:
            def __init__(self, payload):
                self.payload = payload

            def __enter__(self):
                return self

            def __exit__(self, exc_type, exc, tb):
                return False

            def read(self):
                return json.dumps(self.payload).encode("utf-8")

        def fake_urlopen(url, timeout):
            if "geocoding-api.open-meteo.com" in url:
                return FakeResponse({"results": []})
            if "photon.komoot.io" in url:
                return FakeResponse({"features": []})
            self.assertIn("nominatim.openstreetmap.org", url)
            return FakeResponse([{"lat": "39.9042", "lon": "116.4074", "display_name": "北京"}])

        original_key = server.AMAP_WEB_SERVICE_KEY
        server.AMAP_WEB_SERVICE_KEY = ""
        try:
            with patch.object(server, "urlopen_with_user_agent", side_effect=fake_urlopen):
                status, headers, body = server.handle_planning_api_path(
                    "/api/planning/geocode",
                    "POST",
                    json.dumps({"place": "北京"}, ensure_ascii=False).encode("utf-8"),
                )
        finally:
            server.AMAP_WEB_SERVICE_KEY = original_key

        data = json.loads(body.decode("utf-8"))
        self.assertEqual(status, 200)
        self.assertEqual(data["latitude"], 39.9042)
        self.assertEqual(data["longitude"], 116.4074)
        self.assertEqual(data["source"], "OpenStreetMap Nominatim")

    def test_planning_reverse_geocode_endpoint_returns_place_from_coordinates(self):
        class FakeResponse:
            def __init__(self, payload):
                self.payload = payload

            def __enter__(self):
                return self

            def __exit__(self, exc_type, exc, tb):
                return False

            def read(self):
                return json.dumps(self.payload, ensure_ascii=False).encode("utf-8")

        def fake_urlopen(url, timeout):
            self.assertIn("restapi.amap.com/v3/geocode/regeo", url)
            self.assertIn("location=116.407%2C39.904", url)
            return FakeResponse(
                {
                    "status": "1",
                    "regeocode": {
                        "formatted_address": "北京市东城区天安门",
                        "addressComponent": {
                            "province": "北京市",
                            "city": [],
                            "district": "东城区",
                            "township": "东华门街道",
                        },
                    },
                }
            )

        original_key = server.AMAP_WEB_SERVICE_KEY
        server.AMAP_WEB_SERVICE_KEY = "test-key"
        try:
            with patch.object(server, "urlopen_with_user_agent", side_effect=fake_urlopen):
                status, headers, body = server.handle_planning_api_path(
                    "/api/planning/reverse-geocode",
                    "POST",
                    json.dumps({"latitude": 39.904, "longitude": 116.407}).encode("utf-8"),
                )
        finally:
            server.AMAP_WEB_SERVICE_KEY = original_key

        data = json.loads(body.decode("utf-8"))
        self.assertEqual(status, 200)
        self.assertEqual(data["place"], "北京市东城区天安门")
        self.assertEqual(data["display_name"], "北京市东城区天安门")
        self.assertEqual(data["latitude"], 39.904)
        self.assertEqual(data["longitude"], 116.407)
        self.assertEqual(data["source"], "高德地图 Web 服务逆地理编码 API")

    def test_planning_geocode_uses_chinese_alias_when_original_query_fails(self):
        class FakeResponse:
            def __init__(self, payload):
                self.payload = payload

            def __enter__(self):
                return self

            def __exit__(self, exc_type, exc, tb):
                return False

            def read(self):
                return json.dumps(self.payload, ensure_ascii=False).encode("utf-8")

        requested_urls = []

        def fake_urlopen(url, timeout):
            requested_urls.append(url)
            if "restapi.amap.com" in url:
                return FakeResponse({"status": "0", "info": "ENGINE_RESPONSE_DATA_ERROR"})
            if "geocoding-api.open-meteo.com" in url and "shanghai" in url:
                return FakeResponse(
                    {
                        "results": [
                            {
                                "name": "上海",
                                "latitude": 31.22222,
                                "longitude": 121.45806,
                                "country": "中国",
                                "admin1": "上海市",
                            }
                        ]
                    }
                )
            if "geocoding-api.open-meteo.com" in url:
                return FakeResponse({"results": []})
            if "photon.komoot.io" in url or "nominatim.openstreetmap.org" in url:
                return FakeResponse({"features": []} if "photon.komoot.io" in url else [])
            raise AssertionError(f"unexpected url: {url}")

        original_key = server.AMAP_WEB_SERVICE_KEY
        server.AMAP_WEB_SERVICE_KEY = "test-key"
        try:
            with patch.object(server, "urlopen_with_user_agent", side_effect=fake_urlopen):
                status, headers, body = server.handle_planning_api_path(
                    "/api/planning/geocode",
                    "POST",
                    json.dumps({"place": "上海"}, ensure_ascii=False).encode("utf-8"),
                )
        finally:
            server.AMAP_WEB_SERVICE_KEY = original_key

        data = json.loads(body.decode("utf-8"))
        self.assertEqual(status, 200)
        self.assertEqual(data["place"], "上海")
        self.assertEqual(data["latitude"], 31.22222)
        self.assertEqual(data["longitude"], 121.45806)
        self.assertEqual(data["source"], "Open-Meteo Geocoding API")
        self.assertTrue(any("shanghai" in url for url in requested_urls))

    def test_planning_geocode_falls_back_to_photon(self):
        class FakeResponse:
            def __init__(self, payload):
                self.payload = payload

            def __enter__(self):
                return self

            def __exit__(self, exc_type, exc, tb):
                return False

            def read(self):
                return json.dumps(self.payload, ensure_ascii=False).encode("utf-8")

        def fake_urlopen(url, timeout):
            if "geocoding-api.open-meteo.com" in url:
                return FakeResponse({"results": []})
            if "photon.komoot.io" in url:
                return FakeResponse(
                    {
                        "features": [
                            {
                                "geometry": {"coordinates": [-0.1277653, 51.5074456]},
                                "properties": {
                                    "name": "London",
                                    "state": "England",
                                    "country": "United Kingdom",
                                },
                            }
                        ]
                    }
                )
            raise AssertionError(f"unexpected url: {url}")

        original_key = server.AMAP_WEB_SERVICE_KEY
        server.AMAP_WEB_SERVICE_KEY = ""
        try:
            with patch.object(server, "urlopen_with_user_agent", side_effect=fake_urlopen):
                status, headers, body = server.handle_planning_api_path(
                    "/api/planning/geocode",
                    "POST",
                    json.dumps({"place": "London"}, ensure_ascii=False).encode("utf-8"),
                )
        finally:
            server.AMAP_WEB_SERVICE_KEY = original_key

        data = json.loads(body.decode("utf-8"))
        self.assertEqual(status, 200)
        self.assertEqual(data["latitude"], 51.5074456)
        self.assertEqual(data["longitude"], -0.1277653)
        self.assertEqual(data["source"], "OpenStreetMap Photon API")

    def test_planning_page_has_current_scheme_display(self):
        html = (WEB_ROOT / "planning.html").read_text(encoding="utf-8")
        css = (WEB_ROOT / "assets" / "planning.css").read_text(encoding="utf-8")

        self.assertIn('id="currentSchemeName"', html)
        self.assertIn("当前:", html)
        self.assertIn(".current-scheme", css)
        self.assertIn("display: flex", css)
        self.assertIn("justify-content: flex-start", css)
        self.assertIn("white-space: nowrap", css)
        self.assertIn("text-overflow: ellipsis", css)
        self.assertLess(html.index('id="currentSchemeName"'), html.index(">时序数据<"))

    def test_planning_page_uses_requested_product_and_time_series_labels(self):
        html = (WEB_ROOT / "planning.html").read_text(encoding="utf-8")

        self.assertIn(">考察站风-光-氢-储-柴联合规划系统<", html)
        self.assertIn(">时序数据<", html)
        self.assertNotIn(">电网规划系统<", html)
        self.assertNotIn(">微电网风光氢储联合规划系统<", html)
        self.assertNotIn(">电网规划列表<", html)
        self.assertNotIn(">8760时序数据<", html)

    def test_scheme_lists_use_list_items_with_hover_response(self):
        planning_script = (WEB_ROOT / "assets" / "planning.js").read_text(encoding="utf-8")
        optimize_script = (WEB_ROOT / "assets" / "optimize.js").read_text(encoding="utf-8")
        evaluation_script = (WEB_ROOT / "assets" / "evaluation.js").read_text(encoding="utf-8")
        frequency_script = (WEB_ROOT / "assets" / "frequency.js").read_text(encoding="utf-8")
        css = (WEB_ROOT / "assets" / "planning.css").read_text(encoding="utf-8")

        for script in (planning_script, optimize_script):
            self.assertIn('<ul class="scheme-list-items" role="listbox">', script)
            self.assertIn('<li class="scheme-item', script)
            self.assertIn('role="option"', script)
            self.assertIn('tabindex="0"', script)
            self.assertNotIn("<button class=\"scheme-item", script)
        for script in (evaluation_script, frequency_script):
            self.assertIn('<ul class="scheme-list-items evaluation-scheme-tree" role="tree">', script)
            self.assertIn("scheme-tree-node", script)
            self.assertIn("scheme-tree-parent", script)
            self.assertIn("scheme-result-item", script)
            self.assertIn('role="treeitem"', script)
            self.assertIn('tabindex="0"', script)
            self.assertNotIn("<button class=\"scheme-item", script)
        for script in (planning_script, optimize_script):
            self.assertIn('aria-selected="${scheme.name === state.currentScheme ? "true" : "false"}"', script)
            self.assertIn("bindSchemeListItem", script)
            self.assertIn("event.key === \"Enter\" || event.key === \" \"", script)
        self.assertIn('aria-selected="${activeScheme ? "true" : "false"}"', evaluation_script)
        self.assertIn("bindSchemeListItem", evaluation_script)
        self.assertIn("event.key === \"Enter\" || event.key === \" \"", evaluation_script)
        self.assertIn('aria-selected="${activeScheme ? "true" : "false"}"', frequency_script)

        self.assertIn(".scheme-list-items", css)
        scheme_list_items_css = css.split(".scheme-list-items {", 1)[1].split("}", 1)[0]
        self.assertIn("position: relative", scheme_list_items_css)
        self.assertIn("grid-auto-rows: max-content", scheme_list_items_css)
        self.assertIn("align-content: start", scheme_list_items_css)
        self.assertIn("padding: 2px 0 2px 12px", scheme_list_items_css)
        self.assertIn(".scheme-list-items::before", css)
        self.assertIn(".scheme-item:hover", css)
        self.assertIn(".scheme-item:focus-visible", css)
        scheme_item_css = css.split(".scheme-item {", 1)[1].split("}", 1)[0]
        self.assertIn("position: relative", scheme_item_css)
        self.assertIn("cursor: pointer", scheme_item_css)
        self.assertIn("user-select: none", scheme_item_css)
        self.assertIn("min-height: 34px", scheme_item_css)
        self.assertIn("padding: 7px 10px 7px 30px", scheme_item_css)
        self.assertIn("border-bottom: 1px solid", scheme_item_css)
        self.assertIn("border-radius: 0", scheme_item_css)
        self.assertIn("line-height: 1.25", scheme_item_css)
        self.assertIn(".scheme-item::before", css)
        self.assertIn(".scheme-item::after", css)
        self.assertIn("clearPlanningDisplayForSchemeSwitch", planning_script)
        self.assertIn("renderPlanningSwitchingState", planning_script)
        planning_scheme_handler = planning_script.split("bindSchemeListItem(item, () =>", 1)[1].split("));", 1)[0]
        self.assertIn("selectSchemeWithSwitchFeedback(item.dataset.name)", planning_scheme_handler)
        clear_planning_script = planning_script.split("function clearPlanningDisplayForSchemeSwitch", 1)[1].split("function renderPlanningSwitchingState", 1)[0]
        for snippet in (
            "state.payload = renderPlanningSwitchingState(name)",
            "state.timeSeriesLoading = null",
            "renderAll()",
        ):
            self.assertIn(snippet, clear_planning_script)

    def test_optimization_page_has_requested_three_area_layout(self):
        html = (WEB_ROOT / "optimize.html").read_text(encoding="utf-8")
        planning_html = (WEB_ROOT / "planning.html").read_text(encoding="utf-8")
        css = (WEB_ROOT / "assets" / "planning.css").read_text(encoding="utf-8")

        self.assertIn(">考察站风-光-氢-储-柴联合规划系统<", html)
        self.assertIn('<a class="active" href="optimize.html">规划求解</a>', html)
        self.assertIn('<aside class="scheme-rail">', html)
        self.assertIn('<div class="scheme-list-title">方案列表</div>', html)
        self.assertIn('id="schemeList"', html)
        self.assertIn('class="optimization-panel"', html)
        self.assertIn('class="optimization-command-card"', html)
        self.assertIn('id="startOptimization"', html)
        self.assertIn('id="stopOptimization"', html)
        for label in ("状态", "开始", "完成", "度电成本", "绿电占比"):
            self.assertIn(label, html)
        for tab in ("结果概览", "经济性指标", "安全性指标"):
            self.assertIn(tab, html)
        self.assertNotIn("绿电结果", html)
        self.assertNotIn("安全结果", html)
        self.assertIn('id="overviewResult"', html)
        self.assertIn('id="greenResult"', html)
        self.assertIn('id="safetyResult"', html)
        self.assertIn('id="optimizationLogViewToggle"', html)
        self.assertIn('id="optimizationCurveViewToggle"', html)
        self.assertIn('id="optimizationLogPanel"', html)
        self.assertIn('id="optimizationCurvePanel"', html)
        self.assertIn('id="optimizationLogs"', html)
        self.assertIn('id="optimizationCurveNameList"', html)
        self.assertIn('id="optimizationCurveChart"', html)
        self.assertIn('id="queueOptimization" class="queue-action"', html)
        command_row = html.split('<div class="optimization-command-row">', 1)[1].split('</section>', 1)[0]
        self.assertLess(command_row.index('id="optimizationCurrentScheme"'), command_row.index('id="startOptimization"'))
        self.assertLess(command_row.index('id="startOptimization"'), command_row.index('class="optimization-status-grid"'))
        self.assertIn(">运行日志<", html)
        self.assertIn(">曲线展示<", html)
        self.assertIn("assets/result_curves.js", html)
        self.assertIn('assets/optimize.js', html)
        self.assertIn('href="optimize.html">规划求解</a>', planning_html)
        self.assertIn(".optimization-panel", css)
        self.assertIn("grid-template-rows: max-content minmax(220px, 1fr)", css)
        command_card_css = css.split(".optimization-command-card {", 1)[1].split("}", 1)[0]
        self.assertIn("display: block", command_card_css)
        command_row_css = css.split(".optimization-command-row {", 1)[1].split("}", 1)[0]
        self.assertIn("flex-wrap: nowrap", command_row_css)
        self.assertIn(".log-view-tabs", css)
        self.assertIn(".log-view-tab", css)
        self.assertIn(".log-view-panel", css)
        self.assertIn(".optimization-curve-panel", css)
        self.assertIn(".optimization-curve-name-list", css)
        self.assertIn(".optimization-curve-chart", css)
        self.assertIn(".optimization-actions .queue-action", css)
        self.assertIn(".optimization-actions .queue-action.is-active", css)
        status_tile_css = css.split(".status-tile {", 1)[1].split("}", 1)[0]
        self.assertIn("display: flex", status_tile_css)
        self.assertIn("align-items: center", status_tile_css)
        status_label_css = css.split(".status-tile span {", 1)[1].split("}", 1)[0]
        self.assertIn("white-space: nowrap", status_label_css)
        self.assertIn("font-size: 11px", status_label_css)
        status_value_css = css.split(".status-tile strong {", 1)[1].split("}", 1)[0]
        self.assertIn("white-space: nowrap", status_value_css)

    def test_optimization_frontend_polls_status_and_binds_controls(self):
        script = (WEB_ROOT / "assets" / "optimize.js").read_text(encoding="utf-8")
        html = (WEB_ROOT / "optimize.html").read_text(encoding="utf-8")

        self.assertIn("/api/planning/schemes", script)
        self.assertIn("/api/optimization/status", script)
        self.assertIn("/api/optimization/control", script)
        self.assertIn("/api/tasks/control", script)
        self.assertNotIn("/api/evaluation/status", script)
        self.assertNotIn("/api/evaluation/control", script)
        self.assertIn("startOptimization", script)
        self.assertIn("queueOptimization", script)
        self.assertIn("stopOptimization", script)
        self.assertIn("updateOptimizationActions", script)
        self.assertIn("data.can_start_task", script)
        self.assertIn("data.can_queue_task", script)
        self.assertIn("data.can_stop_task", script)
        self.assertIn("data.can_cancel_queue_task", script)
        self.assertIn("terminalOptimizationAction", script)
        self.assertIn("离队", script)
        self.assertIn("停止", script)
        self.assertIn("启动当前方案规划求解", script)
        self.assertIn("将当前方案排队", script)
        self.assertIn("从队列中移出当前方案", script)
        self.assertNotIn("退出队列", script)
        self.assertNotIn("退队", script)
        self.assertNotIn("停止计算", script)
        self.assertIn("classList.toggle(\"is-disabled\"", script)
        self.assertIn("classList.toggle(\"is-active\"", script)
        self.assertIn("正在运行，无法再次启动", script)
        self.assertIn("没有运行", script)
        self.assertIn("alert(data.message", script)
        self.assertIn("setInterval", script)
        self.assertIn("scheduleOptimizationPolling", script)
        self.assertIn('state.pollDelay = data.status === "运行中" || data.task_status === "排队中" ? 1000 : 4000', script)
        self.assertIn("renderOptimizationLogs", script)
        self.assertIn("isLogScrolledNearBottom", script)
        self.assertIn("const shouldStickToBottom = isLogScrolledNearBottom(box)", script)
        self.assertIn("const previousScrollTop = box.scrollTop", script)
        self.assertIn("box.scrollTop = shouldStickToBottom ? box.scrollHeight : previousScrollTop", script)
        self.assertIn("bindResultAxisRangeControls", script)
        self.assertIn("renderResultAxisRangeControls", script)
        self.assertIn("applyAxisRange", script)
        self.assertIn("data-result-axis-min", script)
        self.assertIn("data-result-axis-max", script)
        self.assertIn("bindLogContextMenu", script)
        self.assertIn("clearOptimizationLogs", script)
        self.assertIn("saveOptimizationLogs", script)
        self.assertIn("optimizationCurveViewer", script)
        self.assertIn("loadOptimizationCurveData", script)
        self.assertIn("scheduleOptimizationHourlyCurvePreload", script)
        self.assertIn("preloadOptimizationHourlyCurves", script)
        self.assertIn("syncOptimizationCurveViewerIfHourlyActive", script)
        self.assertIn("/api/comparison/data", script)
        self.assertIn("mode=summary", script)
        self.assertIn("mode=curves", script)
        self.assertIn("loadedCurveKeys", script)
        self.assertIn("hourlyCurvePreloadToken", script)
        self.assertIn("HOURLY_CURVE_PRELOAD_BATCH_SIZE = 8", script)
        self.assertIn("mergeCurvePayload", script)
        self.assertIn("onSelectionChange", script)
        self.assertIn("ResultCurveViewer.create", script)
        self.assertIn("暂无小时级曲线", script)
        self.assertIn("请选择小时级曲线", script)
        self.assertIn("正在加载小时级曲线", script)
        self.assertIn("setData", script)
        self.assertIn("scrollTop", script)
        self.assertIn("data-result-tab", script)
        self.assertIn("结果概览", script)
        self.assertIn("经济性指标", script)
        self.assertIn("安全性指标", script)
        self.assertNotIn("绿电结果", script)
        self.assertNotIn("安全结果", script)
        self.assertIn("optimizationStatusPath", script)
        self.assertIn("light=1", script)
        self.assertIn("defaultOptimizationState", script)
        self.assertIn("scheme=", script)
        self.assertIn("encodeURIComponent(scheme)", script)
        self.assertIn("refreshOptimizationStatus().catch(showError)", script)
        self.assertIn("clearOptimizationDisplayForSchemeSwitch", script)
        self.assertIn("renderOptimizationSwitchingState", script)
        self.assertIn("optimizationSchemeListResizeHandle", html)
        self.assertIn("bindOptimizationSchemeListResizeHandle", script)
        self.assertIn("setOptimizationSchemeRailHeight", script)
        self.assertIn("optimizationSchemeRailHeightBounds", script)
        self.assertIn("--optimization-scheme-rail-height", script)
        self.assertIn('const allowEmptyResult = data.status === "切换中"', script)
        self.assertIn("renderGreenResult(data.results?.green_table || [], data.results?.curves?.green_daily || [], { allowEmpty: allowEmptyResult })", script)
        self.assertIn("renderSafetyResult(data.results?.safety_table || [], data.results?.curves?.safety_daily || [], { allowEmpty: allowEmptyResult })", script)
        scheme_handler = script.split("bindSchemeListItem(item, () => {", 1)[1].split("});", 1)[0]
        self.assertIn("clearOptimizationDisplayForSchemeSwitch(state.currentScheme)", scheme_handler)
        self.assertIn("rememberOptimizationScheme()", scheme_handler)
        self.assertIn("OPTIMIZATION_SCHEME_STORAGE_KEY", script)
        self.assertIn("powerPlanLastOptimizationScheme", script)
        self.assertIn("readStoredText(OPTIMIZATION_SCHEME_STORAGE_KEY)", script)
        clear_scheme_script = script.split("function clearOptimizationDisplayForSchemeSwitch", 1)[1].split("function renderOptimizationSwitchingState", 1)[0]
        for snippet in (
            "window.clearInterval(state.pollTimer)",
            "state.curvePayload = null",
            "state.loadedCurveKeys = new Set()",
            "state.greenDailyPoints = []",
            "state.safetyDailyPoints = []",
            "state.optimization = defaultOptimizationState(scheme)",
            "renderOptimization(state.optimization)",
            'state.optimizationCurveViewer?.clear("正在加载小时级曲线")',
        ):
            self.assertIn(snippet, clear_scheme_script)
        self.assertIn(".optimization-actions button.is-disabled", css := (WEB_ROOT / "assets" / "planning.css").read_text(encoding="utf-8"))
        self.assertIn(".optimization-actions button.is-active", css)
        self.assertIn("--optimization-scheme-rail-height", css)
        self.assertIn(".optimization-workspace:not(.evaluation-workspace) > .scheme-rail", css)
        self.assertIn("grid-template-rows: auto minmax(0, 1fr)", css)
        self.assertIn(".optimization-workspace:not(.evaluation-workspace) > .scheme-rail .scheme-list", css)

    def test_evaluation_page_has_report_export_action(self):
        html = (WEB_ROOT / "evaluation.html").read_text(encoding="utf-8")
        script = (WEB_ROOT / "assets" / "evaluation.js").read_text(encoding="utf-8")

        self.assertIn('id="exportEvaluationReport"', html)
        self.assertIn(">导出报告</button>", html)
        self.assertIn("20260807-report-export", html)
        self.assertIn("exportEvaluationReport", script)
        self.assertIn("/api/evaluation/report", script)
        self.assertIn("filenameFromContentDisposition", script)
        self.assertIn("downloadBlob", script)
        self.assertIn("reportButton.disabled = !selectedResultIsReadable() || !hasScheme || !hasSelection", script)

    def test_evaluation_page_uses_optimization_layout_as_editable_base(self):
        html = (WEB_ROOT / "evaluation.html").read_text(encoding="utf-8")
        planning_html = (WEB_ROOT / "planning.html").read_text(encoding="utf-8")
        optimize_html = (WEB_ROOT / "optimize.html").read_text(encoding="utf-8")
        index_html = (WEB_ROOT / "index.html").read_text(encoding="utf-8")
        script = (WEB_ROOT / "assets" / "evaluation.js").read_text(encoding="utf-8")

        self.assertIn("assets/planning.css?v=", html)
        self.assertIn('<a class="active" href="evaluation.html">方案评估</a>', html)
        self.assertIn('href="evaluation.html">方案评估</a>', planning_html)
        self.assertIn('href="evaluation.html">方案评估</a>', optimize_html)
        self.assertIn('href="evaluation.html"', index_html)
        self.assertIn("<strong>方案评估</strong>", index_html)
        self.assertIn('<aside class="scheme-rail">', html)
        self.assertIn('id="schemeList"', html)
        self.assertIn('class="optimization-panel"', html)
        self.assertIn('id="startEvaluation"', html)
        self.assertIn('id="stopEvaluation"', html)
        command_row = html.split('<div class="optimization-command-row">', 1)[1].split('</section>', 1)[0]
        self.assertLess(command_row.index('id="optimizationCurrentScheme"'), command_row.index('id="startEvaluation"'))
        self.assertLess(command_row.index('id="startEvaluation"'), command_row.index('class="optimization-status-grid"'))
        self.assertNotIn('id="evaluationResultSelect"', html)
        self.assertIn('id="evaluationResultWarnings"', html)
        self.assertIn("result-file-warnings", html)
        self.assertIn('class="evaluation-result-rail"', html)
        self.assertIn('id="evaluationSchemeListResizeHandle"', html)
        self.assertNotIn('id="evaluationCurrentScheme"', html)
        self.assertIn('id="evaluationPlanningResultTitle"', html)
        self.assertIn('id="evaluationPlanningResultTable"', html)
        self.assertIn('id="evaluationMainResizeHandle"', html)
        self.assertIn("当前: 未选择方案", html)
        self.assertIn("当前: 未选择方案/未选择结果", html)
        self.assertNotIn("当前: 未选择方案，结果显示", html)
        self.assertNotIn("当前规划结果", html)
        self.assertIn("当前结果:未选择结果", html)
        self.assertNotIn(">结果文件<", html)
        self.assertNotIn('id="addEvaluationResult"', html)
        for control in ("deleteEvaluationResult", "copyEvaluationResult", "saveEvaluationResult", "renameEvaluationResult"):
            self.assertIn(f'id="{control}"', html)
        self.assertNotIn("增加结果", html)
        self.assertNotIn("删除结果", html)
        self.assertNotIn("复制结果", html)
        self.assertNotIn("保存结果", html)
        for label in ("删除", "复制", "保存", "重命名"):
            self.assertIn(label, html)
        for label in ("状态", "开始", "完成", "度电成本", "绿电占比"):
            self.assertIn(label, html)
        self.assertNotIn("综合评分", html)
        self.assertNotIn("风险等级", html)
        self.assertIn('setMetric("evaluationScore", byLabel.get("度电成本"))', script)
        self.assertIn('setMetric("evaluationRisk", byLabel.get("绿电占比"))', script)
        self.assertNotIn("evaluationScoreMetric", script)
        self.assertNotIn("evaluationRiskMetric", script)
        for tab in ("评估日志", "结果概览", "经济性指标", "安全性指标", "曲线展示"):
            self.assertIn(tab, html)
        self.assertNotIn("评估概览", html)
        self.assertIn('id="evaluationLogViewToggle" class="result-tab" data-result-tab="logs"', html)
        self.assertIn('id="evaluationCurveViewToggle" class="result-tab" data-result-tab="curves"', html)
        self.assertIn('id="evaluationLogPanel" class="result-panel optimization-log-result-panel" data-result-panel="logs"', html)
        self.assertIn('id="evaluationCurvePanel" class="result-panel optimization-curve-panel" data-result-panel="curves"', html)
        self.assertIn('id="evaluationLogs"', html)
        self.assertIn('id="evaluationCurveNameList"', html)
        self.assertIn('id="evaluationCurveChart"', html)
        self.assertIn('id="queueEvaluation" class="queue-action"', html)
        self.assertNotIn('class="evaluation-log-region optimization-log-card"', html)
        self.assertNotIn('id="optimizationLogResizeHandle"', html)
        self.assertIn(">评估日志<", html)
        self.assertIn(">曲线展示<", html)
        self.assertIn("assets/result_curves.js", html)
        self.assertIn("assets/evaluation.js", html)
        self.assertIn("/api/evaluation/status", script)
        self.assertIn("/api/evaluation/control", script)
        self.assertIn("/api/tasks/control", script)
        self.assertNotIn("/api/optimization/status", script)
        self.assertNotIn("/api/optimization/control", script)
        self.assertIn("queueEvaluation", script)
        self.assertIn("terminalEvaluationAction", script)
        self.assertIn("data.can_start_task", script)
        self.assertIn("data.can_queue_task", script)
        self.assertIn("data.can_stop_task", script)
        self.assertIn("data.can_cancel_queue_task", script)
        self.assertIn("离队", script)
        self.assertIn("停止", script)
        self.assertIn("启动当前方案评估", script)
        self.assertIn("将当前评估任务排队", script)
        self.assertIn("从队列中移出当前评估任务", script)
        self.assertNotIn("退出队列", script)
        self.assertNotIn("退队", script)
        self.assertNotIn("停止计算", script)
        self.assertIn("/api/evaluation/results", script)
        self.assertIn("loadEvaluationResults", script)
        self.assertIn("resultsByScheme", script)
        self.assertIn("collapsedSchemes", script)
        self.assertIn("loadEvaluationResultTree", script)
        self.assertIn("renderEvaluationSchemeTreeNode", script)
        self.assertIn("scheme-result-item", script)
        self.assertIn("data-scheme-tree-toggle", script)
        self.assertIn("bindEvaluationSchemeListResizeHandle", script)
        self.assertIn("setEvaluationSchemeRailHeight(evaluationSchemeRailHeightBounds().max)", script)
        self.assertIn("&light=1", script)
        self.assertIn("refreshOptimizationStatus(state.currentScheme, state.selectedResultFile).catch(showError)", script)
        self.assertIn("manageEvaluationResult", script)
        self.assertIn("resultDisplayName", script)
        self.assertNotIn("renderEvaluationResultOption", script)
        self.assertIn("renderEvaluationResultWarnings", script)
        self.assertIn("selectedResultIsReadable", script)
        self.assertIn("isLogScrolledNearBottom", script)
        self.assertIn("const shouldStickToBottom = isLogScrolledNearBottom(box)", script)
        self.assertIn("const previousScrollTop = box.scrollTop", script)
        self.assertIn("box.scrollTop = shouldStickToBottom ? box.scrollHeight : previousScrollTop", script)
        self.assertIn("bindResultAxisRangeControls", script)
        self.assertIn("renderResultAxisRangeControls", script)
        self.assertIn("data-result-axis-min", script)
        self.assertIn("data-result-axis-max", script)
        self.assertNotIn('unreadable ? " disabled" : ""', script)
        self.assertNotIn("暂无可读取结果文件", script)
        self.assertIn("无法读取", script)
        self.assertIn("请求后台失败，请检查 WEB 服务是否正常运行，或查看服务器错误日志。", script)
        self.assertIn("target_name", script)
        self.assertIn("filename=${encodeURIComponent(filename)}", script)
        self.assertIn("light=1", script)
        self.assertIn("planning_result_rows", script)
        self.assertIn("renderEvaluationPlanningResultTable", script)
        self.assertNotIn("renderEvaluationCurrentScheme", script)
        self.assertNotIn("bindLogViewTabs", script)
        self.assertIn('state.activeResultTab = target || "overview"', script)
        self.assertIn('if (target === "curves") loadEvaluationCurveData().catch(showError)', script)
        self.assertIn('state.activeResultTab === "curves"', script)
        self.assertIn("evaluationCurveViewer", script)
        self.assertIn("loadEvaluationCurveData", script)
        self.assertIn("scheduleEvaluationHourlyCurvePreload", script)
        self.assertIn("preloadEvaluationHourlyCurves", script)
        self.assertIn("syncEvaluationCurveViewerIfHourlyActive", script)
        self.assertIn("/api/comparison/data", script)
        self.assertIn("mode=summary", script)
        self.assertIn("mode=curves", script)
        self.assertIn("loadedCurveKeys", script)
        self.assertIn("hourlyCurvePreloadToken", script)
        self.assertIn("HOURLY_CURVE_PRELOAD_BATCH_SIZE = 8", script)
        self.assertIn("mergeCurvePayload", script)
        self.assertIn("onSelectionChange", script)
        self.assertIn("ResultCurveViewer.create", script)
        self.assertIn("暂无小时级曲线", script)
        self.assertIn("请选择小时级曲线", script)
        self.assertIn("正在加载小时级曲线", script)
        self.assertIn("clearEvaluationResultDisplayForSwitch", script)
        self.assertIn("clearEvaluationDisplayForSchemeSwitch", script)
        self.assertIn("renderEvaluationSwitchingState", script)
        self.assertIn("renderEvaluationSchemeSwitchingState", script)
        self.assertIn('const allowEmptyResult = data.status === "切换中"', script)
        self.assertIn("renderGreenResult(data.results?.green_table || [], data.results?.curves?.green_daily || [], { allowEmpty: allowEmptyResult })", script)
        self.assertIn("renderSafetyResult(data.results?.safety_table || [], data.results?.curves?.safety_daily || [], { allowEmpty: allowEmptyResult })", script)
        scheme_handler = script.split("bindSchemeListItem(item, () => {", 1)[1].split("});", 1)[0]
        self.assertIn("clearEvaluationDisplayForSchemeSwitch(state.currentScheme)", scheme_handler)
        self.assertNotIn('document.getElementById("evaluationResultSelect").addEventListener("change"', script)
        clear_switch_script = script.split("function clearEvaluationResultDisplayForSwitch", 1)[1].split("function renderEvaluationSwitchingState", 1)[0]
        for snippet in (
            "state.curvePayload = null",
            "state.loadedCurveKeys = new Set()",
            "state.planningResultRows = []",
            "state.greenDailyPoints = []",
            "state.safetyDailyPoints = []",
            "state.optimization = defaultOptimizationState(state.currentScheme)",
            "renderEvaluationPlanningResultTable()",
            "renderOptimization(state.optimization",
            'state.evaluationCurveViewer?.clear("正在加载小时级曲线")',
        ):
            self.assertIn(snippet, clear_switch_script)
        clear_scheme_script = script.split("function clearEvaluationDisplayForSchemeSwitch", 1)[1].split("function renderEvaluationSchemeSwitchingState", 1)[0]
        for snippet in (
            "window.clearInterval(state.pollTimer)",
            "state.resultFiles = []",
            "state.selectedResultFile = \"\"",
            "state.curvePayload = null",
            "state.loadedCurveKeys = new Set()",
            "state.planningResultRows = []",
            "state.greenDailyPoints = []",
            "state.safetyDailyPoints = []",
            "renderEvaluationResults()",
            "renderEvaluationPlanningResultTable()",
            "renderOptimization(state.optimization)",
            'state.evaluationCurveViewer?.clear("正在加载小时级曲线")',
        ):
            self.assertIn(snippet, clear_scheme_script)
        self.assertIn("bindEvaluationMainResizeHandle", script)
        self.assertIn("--evaluation-result-rail-width", script)
        self.assertIn("ArrowLeft", script)
        self.assertIn("ArrowRight", script)
        self.assertNotIn('document.getElementById("evaluationCurrentScheme")', script)
        self.assertIn("currentEvaluationResultLabel", script)
        self.assertIn("renderCurrentEvaluationResultTitle", script)
        self.assertIn('const schemeName = state.currentScheme || "未选择方案"', script)
        self.assertIn('const resultName = resultDisplayName(state.selectedResultFile) || "未选择结果"', script)
        self.assertIn("return `当前: ${schemeName}/${resultName}`", script)
        self.assertIn('title.textContent = `当前结果:${resultDisplayName(state.selectedResultFile) || "未选择结果"}`', script)
        self.assertNotIn('`当前: ${state.currentScheme || "未选择方案"}，结果显示`', script)
        self.assertIn('data-planning-count-index="${index}"', script)
        self.assertIn('pattern="[0-9]*"', script)
        self.assertIn('inputmode="numeric"', script)
        self.assertIn("validatePlanningCountInput", script)
        self.assertIn("collectPlanningResultRows", script)
        self.assertIn("设计台数", script)
        self.assertIn("prompt(", script)
        self.assertIn('`${resultDisplayName(state.selectedResultFile) || "当前结果"}_副本`', script)
        self.assertIn("复制失败", script)
        self.assertIn("renameEvaluationResult", script)
        self.assertIn('manageEvaluationResult("rename"', script)
        self.assertIn("默认结果文件不允许重命名", script)
        self.assertIn("重命名失败", script)
        self.assertIn("结果文件已重命名", script)
        self.assertIn("const message = messages[action]", script)
        self.assertIn("alert(message)", script)
        self.assertIn("EVALUATION_SELECTION_STORAGE_KEY", script)
        self.assertIn("powerPlanLastEvaluationSelection", script)
        self.assertIn("storedEvaluationSelection()", script)
        self.assertIn("rememberEvaluationSelection()", script)
        self.assertIn("selectedResultIsDefault", script)
        self.assertIn("deleteButton.disabled = selectedResultIsDefault() || !hasScheme || !hasSelection", script)
        self.assertIn("saveButton.disabled = !canEditWorkbook || !hasScheme || !hasSelection", script)
        self.assertIn("copyButton.disabled = !selectedResultIsReadable() || !hasScheme || !hasSelection", script)
        self.assertIn("renameButton.disabled = !canEditWorkbook || !hasScheme || !hasSelection", script)
        self.assertIn(">启动</button>", html)
        self.assertIn(">排队</button>", html)
        self.assertIn(">停止</button>", html)
        self.assertNotIn("立刻启动", html)
        self.assertNotIn("加入排队", html)
        self.assertNotIn("停止计算", html)

        css = (WEB_ROOT / "assets" / "planning.css").read_text(encoding="utf-8")
        self.assertIn(".evaluation-main-resize-handle", css)
        optimization_current_scheme_css = css.split(".optimization-current-scheme {", 1)[1].split("}", 1)[0]
        self.assertIn("flex: 0 1 220px", optimization_current_scheme_css)
        self.assertIn("text-align: left", optimization_current_scheme_css)
        self.assertNotIn("margin-left: auto", optimization_current_scheme_css)
        self.assertIn("grid-template-columns: minmax(0, var(--evaluation-result-rail-width, 340px)) 10px minmax(0, 1fr)", css)
        self.assertIn(".evaluation-workspace:not(.frequency-workspace) .evaluation-result-rail", css)
        self.assertIn("grid-template-rows: minmax(0, 1fr) auto", css)
        self.assertIn(".evaluation-workspace:not(.frequency-workspace) .evaluation-result-actions", css)
        self.assertIn("align-self: end", css)
        action_buttons_css = css.split(".evaluation-result-action-buttons {", 1)[1].split("}", 1)[0]
        self.assertIn("display: flex", action_buttons_css)
        self.assertIn("flex-wrap: nowrap", action_buttons_css)
        self.assertIn("overflow-x: auto", action_buttons_css)
        action_button_css = css.split(".evaluation-result-action-buttons button {", 1)[1].split("}", 1)[0]
        self.assertIn("flex: 1 0 50px", action_button_css)
        self.assertIn("white-space: nowrap", action_button_css)
        self.assertIn("grid-template-rows: minmax(150px, var(--evaluation-scheme-rail-height, 30vh)) 8px minmax(260px, 1fr)", css)
        self.assertIn(".evaluation-result-rail {\n  grid-column: 1;\n  grid-row: 3;", css)
        self.assertIn(".evaluation-workspace > .scheme-rail {\n  grid-column: 1;\n  grid-row: 1;", css)
        evaluation_scheme_rail_css = css.split(".evaluation-workspace > .scheme-rail {", 1)[1].split("}", 1)[0]
        self.assertIn("grid-template-rows: auto minmax(0, 1fr)", evaluation_scheme_rail_css)
        self.assertIn("overflow: hidden", evaluation_scheme_rail_css)
        evaluation_scheme_list_css = css.split(".evaluation-workspace > .scheme-rail .scheme-list {", 1)[1].split("}", 1)[0]
        self.assertIn("overflow: auto", evaluation_scheme_list_css)
        self.assertIn(".evaluation-workspace > .optimization-panel {\n  grid-column: 3;\n  grid-row: 1 / 4;", css)
        self.assertIn(".scheme-tree-caret", css)
        self.assertIn(".scheme-result-list", css)
        self.assertIn(".scheme-result-item", css)
        self.assertIn(".scheme-list-resize-handle", css)
        self.assertNotIn(".evaluation-log-resize-handle", css)
        self.assertNotIn(".evaluation-log-region", css)
        self.assertIn("grid-template-rows: max-content minmax(220px, 1fr)", css)
        self.assertIn(".optimization-curve-panel", css)
        self.assertIn(".optimization-curve-name-list", css)
        self.assertIn(".optimization-curve-chart", css)

    def test_comparison_page_has_tabs_tables_curves_and_result_selectors(self):
        html = (WEB_ROOT / "comparison.html").read_text(encoding="utf-8")
        planning_html = (WEB_ROOT / "planning.html").read_text(encoding="utf-8")
        optimize_html = (WEB_ROOT / "optimize.html").read_text(encoding="utf-8")
        evaluation_html = (WEB_ROOT / "evaluation.html").read_text(encoding="utf-8")
        index_html = (WEB_ROOT / "index.html").read_text(encoding="utf-8")
        script = (WEB_ROOT / "assets" / "comparison.js").read_text(encoding="utf-8")
        css = (WEB_ROOT / "assets" / "planning.css").read_text(encoding="utf-8")

        self.assertIn("assets/planning.css?v=", html)
        self.assertIn('<a class="active" href="comparison.html">结果对比</a>', html)
        self.assertIn('href="comparison.html">结果对比</a>', planning_html)
        self.assertIn('href="comparison.html">结果对比</a>', optimize_html)
        self.assertIn('href="comparison.html">结果对比</a>', evaluation_html)
        self.assertIn('href="comparison.html"', index_html)
        self.assertIn("comparison-tab-bar", html)
        self.assertIn("comparisonTabs", html)
        self.assertIn("comparisonResultWarnings", html)
        self.assertIn("result-file-warnings", html)
        self.assertIn("comparison-table-grid", html)
        for table_id in ("capacityComparisonTable", "energyComparisonTable", "safetyComparisonTable"):
            self.assertIn(table_id, html)
        for title in ("规划容量对比", "经济性指标对比", "安全性指标对比"):
            self.assertIn(title, html)
        self.assertIn("capacityEnergyResizeHandle", html)
        self.assertIn("energySafetyResizeHandle", html)
        self.assertIn('data-comparison-table-column-resize="capacity-energy"', html)
        self.assertIn('data-comparison-table-column-resize="energy-safety"', html)
        self.assertIn("comparisonTableCurveResizeHandle", html)
        self.assertIn("curveNameList", html)
        self.assertIn("comparisonCurveChart", html)
        self.assertIn("小时级曲线", html)
        self.assertIn("日级统计", html)
        self.assertIn("安全日曲线", html)
        self.assertIn("月度统计", html)
        self.assertIn("年度统计", html)
        self.assertNotIn("8760曲线", html)
        self.assertIn("assets/result_curves.js", html)
        self.assertIn("assets/comparison.js?v=20260530-page-state3", html)
        self.assertIn("assets/result_curves.js?v=20260530-page-state3", html)

        self.assertIn("/api/planning/schemes", script)
        self.assertIn("/api/evaluation/results", script)
        self.assertIn("loadResultFilesForTabs(state.tabs)", script)
        self.assertIn("resultsByScheme", script)
        self.assertIn("/api/comparison/data", script)
        self.assertIn("mode=summary", script)
        self.assertIn("mode=curves", script)
        self.assertIn("scheduleComparisonHourlyCurvePreload", script)
        self.assertIn("preloadComparisonHourlyCurves", script)
        self.assertIn("syncComparisonCurveViewerIfHourlyActive", script)
        self.assertIn("loadedCurveKeys", script)
        self.assertIn("hourlyCurvePreloadToken", script)
        self.assertIn("bindComparisonAxisRangeControls", script)
        self.assertIn("renderComparisonAxisRangeControls", script)
        self.assertIn("data-comparison-axis-min", script)
        self.assertIn("data-comparison-axis-max", script)
        self.assertIn("HOURLY_CURVE_PRELOAD_BATCH_SIZE = 8", script)
        self.assertIn("MAX_TABS = 8", script)
        self.assertIn("addComparisonTab", script)
        self.assertIn("renderAddComparisonTab", script)
        self.assertIn(".join(\"\") + renderAddComparisonTab()", script)
        self.assertIn('closest("#addComparisonTab")', script)
        self.assertIn("leftTabForNewComparison", script)
        self.assertIn('scheme: leftTabForNewComparison()?.scheme || state.schemes[0]?.name || ""', script)
        self.assertIn("COMPARISON_TABS_STORAGE_KEY", script)
        self.assertIn("powerPlanLastComparisonTabs", script)
        self.assertIn("restoreComparisonTabs()", script)
        self.assertIn("rememberComparisonTabs()", script)
        self.assertIn("draggable=\"true\"", script)
        self.assertIn("data-close-comparison-tab", script)
        self.assertIn("bindComparisonTableCurveResizeHandle", script)
        self.assertIn("bindComparisonTableColumnResizeHandles", script)
        self.assertIn("data-comparison-table-column-resize", script)
        self.assertIn("--comparison-capacity-table-width", script)
        self.assertIn("--comparison-energy-table-width", script)
        self.assertIn("--comparison-safety-table-width", script)
        self.assertIn("curveNameList", script)
        self.assertIn("comparisonCurveViewer", script)
        self.assertIn("ResultCurveViewer.create", script)
        self.assertIn("enableAnnualBarComparison: true", script)
        self.assertIn("curve_groups", script)
        self.assertIn("annual_table", script)
        self.assertIn("setData", script)
        self.assertIn("renderComparisonCurveChart", script)
        self.assertNotIn('<div class="comparison-curve-legend"', script)
        self.assertIn("selectedCurves", script)
        self.assertIn("selectedCurveNames", script)
        self.assertIn("selectedCurveSeries", script)
        self.assertIn("toggleSelectedCurve", script)
        self.assertIn("flatMap", script)
        self.assertIn('preserveAspectRatio="none"', script)
        self.assertIn("comparison-chart-hover-capture", script)
        self.assertIn("comparisonChartTooltip", script)
        self.assertIn("bindComparisonChartHover", script)
        self.assertIn("renderComparisonChartHover", script)
        self.assertIn("renderComparisonAxisLabels", script)
        self.assertIn("renderComparisonCurveStats", script)
        self.assertIn("comparison-chart-x-axis", script)
        self.assertIn("comparison-chart-y-axis", script)
        self.assertIn("comparison-curve-stats", script)
        self.assertIn("平均", script)
        self.assertIn("合计", script)
        self.assertIn("mouseleave", script)
        self.assertIn("loadResultFilesForTab", script)
        self.assertIn("schemeSelect", script)
        self.assertIn("resultSelect", script)
        self.assertIn("renderComparisonResultWarnings", script)
        self.assertIn("readable !== false", script)
        self.assertIn("无法读取", script)
        self.assertIn("请求后台失败，请检查 WEB 服务是否正常运行，或查看服务器错误日志。", script)
        self.assertIn('event.target.closest("select")', script)
        self.assertIn("event.stopPropagation()", script)
        self.assertIn('<ul aria-multiselectable="true">', script)
        self.assertIn("comparison-curve-name-item", script)
        self.assertIn("isMultiCurveSelectionEvent", script)
        self.assertIn("event?.ctrlKey || event?.shiftKey", script)
        self.assertIn('role="option"', script)
        self.assertIn('aria-multiselectable="true"', script)
        self.assertNotIn("<button type=\"button\" class=\"${name === state.selectedCurve", script)
        self.assertIn("comparison-table-curve-resize-handle", css)
        self.assertIn(".comparison-table-column-resize-handle", css)
        self.assertIn(".comparison-workspace > .comparison-panel", css)
        comparison_workspace_css = css.split(".comparison-workspace {", 1)[1].split("}", 1)[0]
        self.assertIn("grid-template-rows: minmax(0, 1fr)", comparison_workspace_css)
        comparison_panel_css = css.split(".comparison-panel {", 1)[1].split("}", 1)[0]
        self.assertIn("grid-template-rows: minmax(108px, auto) minmax(0, var(--comparison-table-height, 30vh)) 12px minmax(0, 1fr)", comparison_panel_css)
        comparison_tab_bar_css = css.split(".comparison-tab-bar {", 1)[1].split("}", 1)[0]
        self.assertIn("min-height: 108px", comparison_tab_bar_css)
        comparison_tabs_css = css.split(".comparison-tabs {", 1)[1].split("}", 1)[0]
        self.assertIn("min-height: 88px", comparison_tabs_css)
        comparison_curve_board_css = css.split(".comparison-curve-board {", 1)[1].split("}", 1)[0]
        self.assertIn("min-height: 0", comparison_curve_board_css)
        self.assertIn("height: 100%", comparison_curve_board_css)
        self.assertIn("grid-template-columns: minmax(0, var(--comparison-capacity-table-width, 1fr)) 10px minmax(0, var(--comparison-energy-table-width, 1fr)) 10px minmax(0, var(--comparison-safety-table-width, 1fr))", css)
        comparison_table_css = css.split(".comparison-table table {", 1)[1].split("}", 1)[0]
        self.assertIn("table-layout: fixed", comparison_table_css)
        self.assertIn("min-width: 0", comparison_table_css)
        self.assertIn(".comparison-curve-board", css)
        comparison_curve_chart_css = css.split(".comparison-curve-chart {", 1)[1].split("}", 1)[0]
        self.assertIn("width: 100%", comparison_curve_chart_css)
        comparison_curve_svg_css = css.split(".comparison-curve-chart svg {", 1)[1].split("}", 1)[0]
        self.assertIn("width: 100%", comparison_curve_svg_css)
        self.assertIn("height: 100%", comparison_curve_svg_css)
        self.assertIn("grid-template-rows: auto minmax(0, 1fr)", comparison_curve_chart_css)
        comparison_curve_chart_panel_css = css.split(".comparison-curve-chart-panel {", 1)[1].split("}", 1)[0]
        self.assertIn("grid-template-rows: minmax(0, 1fr)", comparison_curve_chart_panel_css)
        comparison_chart_frame_css = css.split(".comparison-chart-frame {", 1)[1].split("}", 1)[0]
        self.assertIn("height: 100%", comparison_chart_frame_css)
        self.assertIn(".comparison-curve-chart > .comparison-chart-frame:only-child", css)
        self.assertIn(".optimization-curve-chart > .comparison-chart-frame:only-child", css)
        monthly_chart_frame_fill_css = css.split(".optimization-curve-chart > .comparison-chart-frame:only-child {", 1)[1].split("}", 1)[0]
        self.assertIn("grid-row: 1 / -1", monthly_chart_frame_fill_css)
        self.assertIn("height: 100%", monthly_chart_frame_fill_css)
        optimization_curve_panel_active_css = css.split(".optimization-curve-panel.active {", 1)[1].split("}", 1)[0]
        self.assertIn("grid-template-rows: minmax(0, 1fr)", optimization_curve_panel_active_css)
        self.assertIn("height: 100%", optimization_curve_panel_active_css)
        self.assertIn(".comparison-curve-chart > .annual-stat-table", css)
        self.assertIn(".optimization-curve-chart > .annual-stat-table", css)
        annual_stat_table_fill_css = css.split(".optimization-curve-chart > .annual-stat-table {", 1)[1].split("}", 1)[0]
        self.assertIn("grid-row: 1 / -1", annual_stat_table_fill_css)
        self.assertIn("width: 100%", annual_stat_table_fill_css)
        self.assertIn("max-height: none", annual_stat_table_fill_css)
        self.assertIn(".comparison-curve-chart > .empty-summary", css)
        self.assertIn(".optimization-curve-chart > .empty-summary", css)
        curve_empty_summary_css = css.split(".optimization-curve-chart > .empty-summary {", 1)[1].split("}", 1)[0]
        self.assertIn("grid-row: 1 / -1", curve_empty_summary_css)
        self.assertIn("align-items: center", curve_empty_summary_css)
        self.assertIn(".comparison-chart-hover-line", css)
        self.assertIn(".comparison-chart-hover-capture", css)
        self.assertIn(".comparison-chart-tooltip", css)
        self.assertIn(".comparison-chart-x-axis", css)
        self.assertIn(".comparison-chart-y-axis", css)
        self.assertIn(".comparison-curve-stats", css)
        comparison_curve_stats_css = css.split(".comparison-curve-stats {", 1)[1].split("}", 1)[0]
        self.assertIn("position: absolute", comparison_curve_stats_css)
        self.assertIn("right:", comparison_curve_stats_css)
        self.assertIn("top:", comparison_curve_stats_css)
        self.assertIn("grid-template-columns: minmax(0, 1fr)", comparison_curve_stats_css)
        self.assertNotIn("<text class=\"comparison-chart-label\"", script)
        comparison_curve_stats_section_css = css.split(".comparison-curve-stats section {", 1)[1].split("}", 1)[0]
        self.assertIn("grid-template-columns: minmax(120px, 1fr) repeat(4, auto)", comparison_curve_stats_section_css)
        self.assertIn("min-width: 0", comparison_curve_stats_section_css)
        self.assertIn("white-space: nowrap", comparison_curve_stats_section_css)
        self.assertIn(".comparison-curve-name-item", css)
        self.assertIn(".curve-group-tabs", css)
        self.assertIn(".curve-group-tab", css)
        self.assertIn(".annual-stat-table", css)
        self.assertIn(".annual-view-switch", css)
        self.assertIn(".annual-view-toggle", css)
        self.assertIn(".annual-comparison-grid", css)
        self.assertIn(".annual-comparison-card", css)
        self.assertIn(".annual-comparison-chart", css)
        self.assertIn(".annual-chart-axis-label", css)
        self.assertIn(".annual-chart-x-axis", css)
        annual_comparison_card_css = css.split(".annual-comparison-card {", 1)[1].split("}", 1)[0]
        self.assertIn("grid-template-rows: auto minmax(0, 1fr)", annual_comparison_card_css)
        self.assertIn(".annual-head-actions", css)
        annual_comparison_head_css = css.split(".annual-comparison-head {", 1)[1].split("}", 1)[0]
        self.assertIn("display: grid", annual_comparison_head_css)
        self.assertIn("grid-template-columns: max-content minmax(0, 1fr) max-content", annual_comparison_head_css)
        annual_comparison_grid_css = css.split(".annual-comparison-grid {", 1)[1].split("}", 1)[0]
        self.assertIn("grid-template-columns: minmax(0, var(--annual-grid-left", annual_comparison_grid_css)
        self.assertIn("grid-template-rows: minmax(0, var(--annual-grid-top", annual_comparison_grid_css)
        self.assertIn("height: 100%", annual_comparison_grid_css)
        self.assertIn("overflow: hidden", annual_comparison_grid_css)
        self.assertIn(".annual-grid-resizer", css)
        self.assertIn(".annual-grid-resizer-col", css)
        self.assertIn(".annual-grid-resizer-row", css)
        self.assertIn(".annual-chart-tooltip", css)
        self.assertIn(".annual-comparison-legend button", css)
        self.assertIn(".annual-comparison-legend button.is-hidden", css)
        annual_comparison_legend_css = css.split(".annual-comparison-legend {", 1)[1].split("}", 1)[0]
        self.assertIn("flex-wrap: nowrap", annual_comparison_legend_css)
        self.assertIn("overflow-x: auto", annual_comparison_legend_css)
        self.assertIn(".annual-line-point", css)
        annual_line_point_css = css.split(".annual-line-point {", 1)[1].split("}", 1)[0]
        self.assertIn("position: absolute", annual_line_point_css)
        self.assertIn("border-radius: 50%", annual_line_point_css)
        self.assertIn("width: 8px", annual_line_point_css)
        self.assertIn("height: 8px", annual_line_point_css)
        annual_comparison_fill_css = css.split(".comparison-curve-chart > .annual-comparison-grid,", 1)[1].split("}", 1)[0]
        self.assertIn("grid-row: 1 / -1", annual_comparison_fill_css)
        self.assertIn("height: 100%", annual_comparison_fill_css)
        annual_stat_table_css = css.split(".annual-stat-table table {", 1)[1].split("}", 1)[0]
        self.assertIn("width: 100%", annual_stat_table_css)
        self.assertIn("height: 100%", annual_stat_table_css)
        curve_group_tabs_css = css.split(".curve-group-tabs {", 1)[1].split("}", 1)[0]
        self.assertIn("position: sticky", curve_group_tabs_css)
        self.assertIn("top: 0", curve_group_tabs_css)
        self.assertIn("z-index:", curve_group_tabs_css)
        self.assertIn("pointer-events: none", curve_group_tabs_css)
        curve_group_tab_css = css.split(".curve-group-tab {", 1)[1].split("}", 1)[0]
        self.assertIn("pointer-events: auto", curve_group_tab_css)
        self.assertNotIn(".comparison-curve-name-list button", css)
        comparison_curve_name_item_css = css.split(".comparison-curve-name-item {", 1)[1].split("}", 1)[0]
        self.assertIn("cursor: pointer", comparison_curve_name_item_css)
        self.assertIn("user-select: none", comparison_curve_name_item_css)
        comparison_curve_name_hover_css = css.split(".comparison-curve-name-item:hover,", 1)[1].split("}", 1)[0]
        self.assertIn("border-left-color: #0d5c59", comparison_curve_name_hover_css)
        self.assertIn("background: rgba(13, 92, 89, 0.12)", comparison_curve_name_hover_css)

        result_curve_script = (WEB_ROOT / "assets" / "result_curves.js").read_text(encoding="utf-8")
        self.assertIn("curve_groups", result_curve_script)
        self.assertIn("annual_table", result_curve_script)
        self.assertIn("小时级曲线", result_curve_script)
        self.assertIn("日级统计", result_curve_script)
        self.assertIn("月度统计", result_curve_script)
        self.assertIn("年度统计", result_curve_script)
        self.assertIn("data-curve-group", result_curve_script)
        self.assertIn("renderAnnualTable", result_curve_script)
        self.assertIn("enableAnnualBarComparison", result_curve_script)
        self.assertIn("annualViewMode", result_curve_script)
        self.assertIn('annualViewMode: options.enableAnnualBarComparison ? "bar" : "table"', result_curve_script)
        self.assertIn("data-annual-view-mode", result_curve_script)
        self.assertIn("表格显示", result_curve_script)
        self.assertIn("柱图对比", result_curve_script)
        self.assertIn("纵坐标配置", result_curve_script)
        self.assertIn("axis-range-toggle", result_curve_script)
        self.assertIn("axis-range-panel", result_curve_script)
        self.assertIn("renderAnnualBarComparison", result_curve_script)
        self.assertIn("renderAnnualComparisonHead", result_curve_script)
        self.assertIn("renderAnnualComparisonHead(definition, barMetrics, lineMetric)", result_curve_script)
        self.assertIn("annual-comparison-grid", result_curve_script)
        self.assertIn("renderAnnualChartLabels", result_curve_script)
        self.assertIn("annual-chart-axis-label", result_curve_script)
        self.assertIn("annual-chart-x-axis", result_curve_script)
        self.assertIn("renderAnnualAxisRangeControls", result_curve_script)
        self.assertIn("bindAnnualAxisRangeControls", result_curve_script)
        self.assertIn("annualAxisRangeKey", result_curve_script)
        self.assertIn("data-annual-axis-min", result_curve_script)
        self.assertIn("data-annual-axis-max", result_curve_script)
        self.assertIn("data-annual-axis-reset", result_curve_script)
        self.assertIn("bindAnnualGridResizers", result_curve_script)
        self.assertIn("data-annual-grid-resizer", result_curve_script)
        self.assertIn("setPointerCapture", result_curve_script)
        self.assertIn("renderAnnualLinePoints", result_curve_script)
        self.assertIn("bindAnnualChartHover", result_curve_script)
        self.assertIn("renderAnnualChartHover", result_curve_script)
        self.assertIn("data-annual-chart-hit", result_curve_script)
        self.assertIn("data-annual-chart-tooltip", result_curve_script)
        self.assertIn("annualHiddenSeries", result_curve_script)
        self.assertIn("bindAnnualLegendToggles", result_curve_script)
        self.assertIn("data-annual-series-toggle", result_curve_script)
        self.assertIn("toggleAnnualSeriesVisibility", result_curve_script)
        self.assertIn("formatAnnualMetricValue", result_curve_script)
        self.assertNotIn('<text class="annual-axis-label', result_curve_script)
        self.assertNotIn('<text class="annual-x-label', result_curve_script)
        self.assertNotIn('<circle class="annual-line-point', result_curve_script)
        self.assertIn("成本对比", result_curve_script)
        self.assertIn("年均总成本", result_curve_script)
        self.assertIn("年均建设成本", result_curve_script)
        self.assertIn("年运行成本", result_curve_script)
        self.assertIn("度电成本", result_curve_script)
        self.assertIn("负荷用电量", result_curve_script)
        self.assertIn("柴油发电量", result_curve_script)
        self.assertIn("新能源总发电量", result_curve_script)
        self.assertIn("新能源占比", result_curve_script)
        self.assertIn("新能源最大可发", result_curve_script)
        self.assertIn("新能源实际电量", result_curve_script)
        self.assertIn("新能源弃电率", result_curve_script)
        self.assertIn("储能发电量", result_curve_script)
        self.assertIn("燃料电池发电量", result_curve_script)
        self.assertIn("isMultiCurveSelectionEvent", result_curve_script)
        self.assertIn("event?.ctrlKey || event?.shiftKey", result_curve_script)
        self.assertIn("curveRangeFilter", result_curve_script)
        self.assertIn("renderRangeControls", result_curve_script)
        self.assertIn("renderAxisRangeControls", result_curve_script)
        self.assertIn("data-curve-axis-min", result_curve_script)
        self.assertIn("data-curve-axis-max", result_curve_script)
        self.assertIn("applyAxisRange", result_curve_script)
        self.assertIn("loadingText", result_curve_script)
        self.assertIn("小时级曲线正在后台加载", result_curve_script)
        self.assertIn("hiddenSeriesByGroup", result_curve_script)
        self.assertIn("renderCurveLegend", result_curve_script)
        self.assertIn("renderCurveLegend(allSeries, visibleSeries)", result_curve_script)
        self.assertIn("result-curve-legend-stat", result_curve_script)
        self.assertIn("curveStats(item)", result_curve_script)
        self.assertIn("bindCurveLegendToggles", result_curve_script)
        self.assertIn("resultLegendToggleBound", result_curve_script)
        self.assertIn('target.addEventListener("click"', result_curve_script)
        self.assertIn('event.target.closest("[data-result-series-toggle]")', result_curve_script)
        self.assertIn("data-result-stats-panel", result_curve_script)
        self.assertIn("data-result-stats-menu", result_curve_script)
        self.assertIn("startStatsPanelDrag", result_curve_script)
        self.assertIn("showStatsContextMenu", result_curve_script)
        self.assertIn("隐藏统计信息", result_curve_script)
        self.assertIn("显示统计信息", result_curve_script)
        self.assertIn("恢复统计位置", result_curve_script)
        self.assertIn("toggleSeriesVisibility", result_curve_script)
        self.assertIn("isSeriesHidden", result_curve_script)
        self.assertIn("data-result-series-toggle", result_curve_script)
        self.assertIn("data-result-series-swatch", result_curve_script)
        self.assertIn("stroke-dasharray", result_curve_script)
        self.assertIn("seriesLineStyle", result_curve_script)
        self.assertIn("filterSeriesByRange", result_curve_script)
        self.assertIn("availableDaysInMonth", result_curve_script)
        self.assertIn("data-curve-range-scope", result_curve_script)
        self.assertIn("relabelFilteredPoints", result_curve_script)
        self.assertIn("`第${dayNumber}日`", result_curve_script)
        self.assertIn("`${index + 1}时`", result_curve_script)
        self.assertIn("全年", result_curve_script)
        self.assertIn("指定月", result_curve_script)
        self.assertIn("指定日", result_curve_script)
        self.assertIn("日级统计只支持全年或指定月筛选", result_curve_script)
        self.assertNotIn('<div class="comparison-curve-legend"', result_curve_script)
        self.assertNotIn("${renderCurveLegend(allSeries)}", result_curve_script)
        self.assertNotIn("renderCurveStats(visibleSeries)", result_curve_script)
        self.assertIn("result-curve-legend", css)
        result_curve_legend_css = css.split(".result-curve-legend {", 1)[1].split("}", 1)[0]
        self.assertIn("position: absolute", result_curve_legend_css)
        self.assertIn("pointer-events: auto", result_curve_legend_css)
        self.assertIn("cursor: grab", result_curve_legend_css)
        self.assertIn(".result-curve-legend.dragging", css)
        self.assertIn(".result-curve-legend.stats-hidden .result-curve-legend-stat", css)
        self.assertIn(".result-curve-context-menu", css)
        self.assertIn(".result-curve-legend button", css)
        self.assertIn(".result-curve-legend button.is-hidden", css)
        self.assertIn(".result-curve-legend-label", css)
        self.assertIn(".result-curve-legend-stat", css)
        self.assertIn(".result-curve-legend-swatch", css)
        self.assertIn(".result-curve-empty-overlay", css)
        self.assertIn(".curve-range-filter", css)
        curve_range_active_css = css.split(".curve-range-scope button.active,", 1)[1].split("}", 1)[0]
        self.assertIn("rgba(33, 213, 255, 0.12)", curve_range_active_css)
        self.assertNotIn("background: #0d5c59", curve_range_active_css)
        time_chart_range_css = css.split(".time-chart-range {", 1)[1].split("}", 1)[0]
        self.assertIn("background: var(--hud-surface-soft)", time_chart_range_css)
        self.assertIn("box-shadow: inset 0 0 14px rgba(24, 175, 255, 0.08)", time_chart_range_css)

    def test_planning_time_series_chart_supports_year_month_day_filter(self):
        html = (WEB_ROOT / "planning.html").read_text(encoding="utf-8")
        script = (WEB_ROOT / "assets" / "planning.js").read_text(encoding="utf-8")
        css = (WEB_ROOT / "assets" / "planning.css").read_text(encoding="utf-8")
        i18n_script = (WEB_ROOT / "assets" / "i18n.js").read_text(encoding="utf-8")

        self.assertIn('id="timeChartRange"', html)
        self.assertIn('data-time-chart-scope="year"', html)
        self.assertIn('data-time-chart-scope="month"', html)
        self.assertIn('data-time-chart-scope="day"', html)
        self.assertIn('id="timeChartMonth"', html)
        self.assertIn('id="timeChartDay"', html)
        self.assertIn("timeChartRange", script)
        self.assertIn("state.timeChartRange", script)
        self.assertIn("bindTimeChartRangeControls", script)
        self.assertIn("filteredTimeChartRows", script)
        self.assertIn("availableDaysInMonth", script)
        self.assertIn("absoluteIndex", script)
        self.assertIn("updateVisibleTimeCell(absoluteIndex", script)
        self.assertIn("`第${Math.floor(pointIndex / 24) + 1}日`", script)
        self.assertIn(".time-chart-range", css)
        for text in ("全年", "指定月", "指定日", "月份", "日期"):
            self.assertIn(text, i18n_script)

    def test_optimization_page_removes_command_to_result_resize_handle(self):
        html = (WEB_ROOT / "optimize.html").read_text(encoding="utf-8")
        evaluation_html = (WEB_ROOT / "evaluation.html").read_text(encoding="utf-8")
        script = (WEB_ROOT / "assets" / "optimize.js").read_text(encoding="utf-8")
        evaluation_script = (WEB_ROOT / "assets" / "evaluation.js").read_text(encoding="utf-8")
        css = (WEB_ROOT / "assets" / "planning.css").read_text(encoding="utf-8")

        self.assertNotIn('id="optimizationResultResizeHandle"', html)
        self.assertNotIn('id="optimizationResultResizeHandle"', evaluation_html)
        self.assertNotIn('id="optimizationLogResizeHandle"', html)
        self.assertNotIn('aria-label="调整规划结果高度"', html)
        self.assertNotIn('aria-label="调整评估结果高度"', evaluation_html)
        self.assertNotIn("bindOptimizationResultResizeHandle", script)
        self.assertNotIn("bindOptimizationResultResizeHandle", evaluation_script)
        self.assertNotIn("bindOptimizationLogResizeHandle", script)
        self.assertNotIn("optimizationResultHeight", script)
        self.assertNotIn("--optimization-result-height", script)
        self.assertNotIn(".optimization-result-resize-handle", css)
        self.assertIn(".optimization-log-resize-handle", css)
        result_tabs_css = css.split(".result-tabs {", 1)[1].split("}", 1)[0]
        self.assertIn("flex: 0 0 auto", result_tabs_css)
        self.assertIn("min-height: 38px", result_tabs_css)
        result_tab_css = css.split(".result-tab {", 1)[1].split("}", 1)[0]
        self.assertIn("height: 36px", result_tab_css)
        self.assertIn("display: inline-flex", result_tab_css)
        self.assertIn("cursor: row-resize", css)
        self.assertIn("--optimization-log-height", css)

    def test_optimization_overview_frontend_renders_planning_table_and_composition_bars(self):
        script = (WEB_ROOT / "assets" / "optimize.js").read_text(encoding="utf-8")
        evaluation_script = (WEB_ROOT / "assets" / "evaluation.js").read_text(encoding="utf-8")
        css = (WEB_ROOT / "assets" / "planning.css").read_text(encoding="utf-8")

        self.assertIn("renderOverviewTables", script)
        self.assertIn("overview_tables", script)
        self.assertIn("renderOverviewCompositionBars", script)
        self.assertIn("renderOverviewCompositionBar", script)
        self.assertIn("normalizeOverviewCompositionDisplay", script)
        self.assertIn("normalizeOverviewCompositionSegments", script)
        self.assertIn("simplifyOverviewCompositionLabel", script)
        self.assertIn("normalizeOverviewCompositionValue", script)
        self.assertIn("formatOverviewCompositionNumber", script)
        self.assertIn('if (type === "capacity") return Math.round(number).toLocaleString("zh-CN")', script)
        self.assertIn('if (type === "energy" && text === "柴") text = "柴发"', script)
        self.assertIn('if (type === "energy") return "万kWh"', script)
        self.assertIn('if (normalizedUnit.includes("mwh")) return number / 10', script)
        self.assertIn("return number / 10000", script)
        self.assertIn("bindOverviewColumnResizeHandles", script)
        self.assertIn('data-overview-column-resize="left-middle"', script)
        self.assertNotIn('data-overview-column-resize="middle-right"', script)
        self.assertIn("--overview-left-column-width", script)
        self.assertNotIn("--overview-middle-column-width", script)
        self.assertIn("bindOverviewColumnResizeHandles", evaluation_script)
        self.assertIn('data-overview-column-resize="left-middle"', evaluation_script)
        self.assertNotIn('data-overview-column-resize="middle-right"', evaluation_script)
        self.assertIn("composition-bar-track", script)
        self.assertIn("composition-bar-segment", script)
        self.assertIn("composition-bar-percent", script)
        self.assertIn("composition-bar-summary-dot", script)
        self.assertIn("buildOverviewCompositionSummary(segments, positiveTotal)", script)
        self.assertNotIn('<div class="composition-bar-legend">', script)
        self.assertIn("overview_disks", script)
        self.assertIn("overview-composition-stack", script)
        self.assertIn("formatOverviewTableForDisplay", script)
        self.assertIn("formatOverviewPlanningRows", script)
        self.assertIn("formatOverviewTableForDisplay", evaluation_script)
        self.assertIn("formatOverviewPlanningRows", evaluation_script)
        self.assertIn("renderOverviewCompositionBars", evaluation_script)
        self.assertIn("normalizeOverviewCompositionDisplay", evaluation_script)
        self.assertIn("normalizeOverviewCompositionSegments", evaluation_script)
        self.assertIn("simplifyOverviewCompositionLabel", evaluation_script)
        self.assertIn('if (type === "capacity" && text === "电储能") text = "电储";', script)
        self.assertIn('if (type === "capacity" && text === "燃料电池") text = "燃电";', script)
        self.assertIn('if (type === "capacity" && text === "电储能") text = "电储";', evaluation_script)
        self.assertIn('if (type === "capacity" && text === "燃料电池") text = "燃电";', evaluation_script)
        self.assertIn("normalizeOverviewCompositionValue", evaluation_script)
        self.assertIn("formatOverviewCompositionNumber", evaluation_script)
        self.assertIn("composition-bar-track", evaluation_script)
        self.assertIn("composition-bar-percent", evaluation_script)
        self.assertIn("composition-bar-summary-dot", evaluation_script)
        self.assertIn("buildOverviewCompositionSummary(segments, positiveTotal)", evaluation_script)
        self.assertNotIn('<div class="composition-bar-legend">', evaluation_script)
        self.assertIn("optimization-overview-grid", script)
        self.assertIn("规划结果", script)
        self.assertNotIn("规划年指标", script)
        self.assertNotIn("规划年指标", evaluation_script)
        self.assertNotIn("规划年效益", script)
        for label in ("运行成本", "建设成本", "容量构成", "柴发容量", "风电容量", "光伏容量", "电储能容量", "燃料电池容量", "柴发电量", "新能源电量"):
            self.assertIn(label, script)
        self.assertIn('{ "名称": "-", "设计台数": "-", "单台容量": "-", "总容量": "-", "单位": "" }', script)
        self.assertIn('{ "名称": "-", "设计台数": "-", "单台容量": "-", "总容量": "-", "单位": "" }', evaluation_script)
        for field in ("名称", "设计台数", "指标", "数值", "单位"):
            self.assertIn(field, script)
        self.assertIn(".optimization-overview-grid", css)
        self.assertIn("grid-template-columns: minmax(0, var(--overview-left-column-width, 1fr)) 10px minmax(0, 1.35fr)", css)
        self.assertIn(".overview-column-resize-handle", css)
        self.assertIn("cursor: col-resize", css)
        self.assertIn(".overview-table-card", css)
        self.assertIn(".overview-composition-stack", css)
        composition_stack_css = css.split(".overview-composition-stack {", 1)[1].split("}", 1)[0]
        self.assertIn("display: grid", composition_stack_css)
        self.assertIn("grid-template-columns: minmax(0, 1fr)", composition_stack_css)
        self.assertIn("overflow: auto", composition_stack_css)
        composition_card_css = css.split(".composition-bar-card {", 1)[1].split("}", 1)[0]
        self.assertIn("grid-template-rows: auto auto", composition_card_css)
        self.assertIn("min-height: 128px", composition_card_css)
        self.assertIn("border: 1px solid #d7e4e0", composition_card_css)
        self.assertIn(".composition-bar-track", css)
        self.assertIn(".composition-bar-percent", css)
        self.assertIn(".composition-bar-summary-dot", css)
        composition_track_css = css.split(".composition-bar-track {", 1)[1].split("}", 1)[0]
        self.assertIn("display: flex", composition_track_css)
        self.assertIn("height: 22px", composition_track_css)
        composition_segment_css = css.split(".composition-bar-segment {", 1)[1].split("}", 1)[0]
        self.assertIn("height: 100%", composition_segment_css)
        self.assertIn("min-height: 0 !important", composition_segment_css)
        self.assertIn("max-height: 100%", composition_segment_css)
        self.assertIn("padding: 0 !important", composition_segment_css)
        self.assertIn("border: 0 !important", composition_segment_css)
        self.assertIn("box-shadow: none !important", composition_segment_css)
        composition_segment_label_css = css.split(".composition-bar-segment span {", 1)[1].split("}", 1)[0]
        self.assertIn("position: absolute", composition_segment_label_css)
        self.assertIn("top: 50%", composition_segment_label_css)
        self.assertIn("left: 50%", composition_segment_label_css)
        self.assertIn("display: inline-flex", composition_segment_label_css)
        self.assertIn("align-items: center", composition_segment_label_css)
        self.assertIn("justify-content: center", composition_segment_label_css)
        self.assertIn("line-height: 1", composition_segment_label_css)
        self.assertIn("transform: translate(-50%, -50%)", composition_segment_label_css)
        self.assertIn(".composition-bar-segment.primary", css)
        self.assertIn(".composition-bar-segment.secondary", css)
        composition_primary_css = css.split(".composition-bar-segment.primary {", 1)[1].split("}", 1)[0]
        self.assertIn("background: var(--composition-segment-color, #0d5c59) !important", composition_primary_css)
        composition_secondary_css = css.split(".composition-bar-segment.secondary {", 1)[1].split("}", 1)[0]
        self.assertIn("background: var(--composition-segment-color, #d8b35d) !important", composition_secondary_css)
        self.assertIn("--composition-segment-color", css)
        self.assertIn(".composition-bar-card.multi-segment", css)
        self.assertNotIn(".ratio-disk", css)
        self.assertNotIn("conic-gradient", css)

    def test_optimization_green_frontend_renders_daily_stacked_chart_and_table(self):
        script = (WEB_ROOT / "assets" / "optimize.js").read_text(encoding="utf-8")
        css = (WEB_ROOT / "assets" / "planning.css").read_text(encoding="utf-8")

        self.assertIn("renderGreenResult", script)
        self.assertIn("green_table", script)
        self.assertIn("green_daily", script)
        self.assertIn("renderGreenDailyChart", script)
        self.assertIn("bindResultColumnResizeHandles", script)
        self.assertIn("bindAdaptiveResultCharts", script)
        self.assertIn("bindChartHoverCursors", script)
        self.assertIn("updateChartHoverCursor", script)
        self.assertIn("renderGreenHoverTooltip", script)
        self.assertIn("ResizeObserver", script)
        self.assertIn("resultChartMargins", script)
        self.assertIn("chartTickIndexes", script)
        self.assertIn("height >= 200", script)
        self.assertIn("width < 620", script)
        self.assertNotIn("right: 22", script)
        self.assertIn('data-result-column-resize="green"', script)
        self.assertIn('data-result-chart-viewport="green"', script)
        self.assertIn("--green-result-table-width", script)
        self.assertIn("green-daily-chart", script)
        self.assertIn("green-zero-line", script)
        self.assertIn("green-zero-label", script)
        self.assertIn('data-chart-hover="green"', script)
        self.assertIn('data-chart-hover-line="green"', script)
        self.assertIn('data-chart-hover-tooltip="green"', script)
        self.assertIn('data-series-toggle', script)
        self.assertIn('aria-pressed', script)
        self.assertIn("greenSeriesVisibility", script)
        self.assertIn("toggleGreenSeriesVisibility", script)
        self.assertIn("isSeriesVisible", script)
        self.assertNotIn("green-axis-label", script)
        self.assertNotIn("green-y-axis-title", script)
        self.assertNotIn("green-x-axis-title", script)
        self.assertNotIn(">kWh<", script)
        self.assertNotIn("日序号", script)
        self.assertNotIn("renderGreenDailyDataTable", script)
        self.assertNotIn("green-daily-data-table", script)
        self.assertNotIn("日曲线测试数据", script)
        self.assertNotIn("第${day}日", script)
        green_render_script = script.split("function renderGreenResult", 1)[1].split("function renderGreenDailyChart", 1)[0]
        self.assertIn('"单位": row["单位"] || ""', green_render_script)
        self.assertLess(green_render_script.index("green-result-table"), green_render_script.index("green-chart-card"))
        for label in ("柴发日电量", "风电日电量", "光伏日电量", "氢能日电量", "储能放电量", "负荷电量", "制氢电量", "储能充电量"):
            self.assertIn(label, script)
        for metric in (
            '"指标": "负荷总电量", "数值": "-", "单位": "kWh"',
            '"指标": "柴发总电量", "数值": "-", "单位": "kWh"',
            '"指标": "风机总发电量", "数值": "-", "单位": "kWh"',
            '"指标": "光伏总发电量", "数值": "-", "单位": "kWh"',
            '"指标": "电储总发电量", "数值": "-", "单位": "kWh"',
            '"指标": "氢储总发电量", "数值": "-", "单位": "kWh"',
            '"指标": "新能源总弃电量", "数值": "-", "单位": "kWh"',
            '"指标": "新能源占比", "数值": "-", "单位": "%"',
            '"指标": "新能源弃电率", "数值": "-", "单位": "%"',
            '"指标": "柴油消耗", "数值": "-", "单位": "吨"',
            '"指标": "制氢总量", "数值": "-", "单位": "Nm3"',
        ):
            self.assertIn(metric, script)
        self.assertNotIn("负荷总电量(kWh)", script)

        self.assertIn(".green-result-layout", css)
        green_layout_css = css.split(".green-result-layout {", 1)[1].split("}", 1)[0]
        self.assertIn("display: grid", green_layout_css)
        self.assertIn("grid-template-columns: minmax(0, var(--green-result-table-width, 34%)) 10px minmax(0, 1fr)", green_layout_css)
        self.assertIn(".result-column-resize-handle", css)
        result_column_resize_css = css.split(".result-column-resize-handle {", 1)[1].split("}", 1)[0]
        self.assertIn("cursor: col-resize", result_column_resize_css)
        self.assertIn(".green-daily-chart", css)
        green_chart_card_css = css.split(".green-chart-card {", 1)[1].split("}", 1)[0]
        self.assertIn("display: grid", green_chart_card_css)
        self.assertIn("grid-template-rows: auto minmax(0, 1fr)", green_chart_card_css)
        self.assertIn(".green-chart-viewport", css)
        green_chart_viewport_css = css.split(".green-chart-viewport,", 1)[1].split("{", 1)[1].split("}", 1)[0]
        self.assertIn("height: 100%", green_chart_viewport_css)
        self.assertIn("min-height: 0", green_chart_viewport_css)
        self.assertIn(".green-chart-svg", css)
        self.assertIn(".chart-hover-line", css)
        self.assertIn(".chart-hover-tooltip", css)
        chart_tooltip_css = css.split(".chart-hover-tooltip {", 1)[1].split("}", 1)[0]
        self.assertIn("position: absolute", chart_tooltip_css)
        self.assertIn("pointer-events: none", chart_tooltip_css)
        self.assertIn('preserveAspectRatio="xMidYMid meet"', script)
        self.assertNotIn('preserveAspectRatio="none"', script)
        self.assertIn('resultChartSize("green"', script)
        green_chart_svg_css = css.split(".green-chart-svg {", 1)[1].split("}", 1)[0]
        self.assertIn("width: 100%", green_chart_svg_css)
        self.assertIn("height: 100%", green_chart_svg_css)
        self.assertIn("min-height: 0", green_chart_svg_css)
        self.assertIn(".green-chart-legend", css)
        self.assertIn(".green-chart-legend button", css)
        self.assertIn(".green-chart-legend button.is-hidden", css)
        self.assertIn(".green-result-table", css)
        self.assertNotIn(".green-daily-data-table", css)
        self.assertNotIn(".green-daily-data-section", css)

    def test_optimization_safety_frontend_renders_frequency_chart_and_table(self):
        script = (WEB_ROOT / "assets" / "optimize.js").read_text(encoding="utf-8")
        css = (WEB_ROOT / "assets" / "planning.css").read_text(encoding="utf-8")

        self.assertIn("renderSafetyResult", script)
        self.assertIn("safety_table", script)
        self.assertIn("safety_daily", script)
        self.assertIn("renderSafetyDailyChart", script)
        self.assertIn("bindResultColumnResizeHandles", script)
        self.assertIn("bindAdaptiveResultCharts", script)
        self.assertIn("bindChartHoverCursors", script)
        self.assertIn("updateChartHoverCursor", script)
        self.assertIn("renderSafetyHoverTooltip", script)
        self.assertIn("ResizeObserver", script)
        self.assertIn("resultChartMargins", script)
        self.assertIn("chartTickIndexes", script)
        self.assertIn("height >= 200", script)
        self.assertIn("width < 620", script)
        self.assertNotIn("right: 22", script)
        self.assertIn('data-result-column-resize="safety"', script)
        self.assertIn('data-result-chart-viewport="safety"', script)
        self.assertIn("--safety-result-table-width", script)
        self.assertIn("safety-frequency-chart", script)
        self.assertIn("safety-center-line", script)
        self.assertIn("safety-zero-label", script)
        self.assertIn('data-chart-hover="safety"', script)
        self.assertIn('data-chart-hover-line="safety"', script)
        self.assertIn('data-chart-hover-tooltip="safety"', script)
        self.assertIn('data-series-toggle', script)
        self.assertIn('aria-pressed', script)
        self.assertIn("safetySeriesVisibility", script)
        self.assertIn("toggleSafetySeriesVisibility", script)
        self.assertIn("isSeriesVisible", script)
        self.assertIn("formatFrequencyDeviation", script)
        self.assertNotIn("safety-axis-label", script)
        self.assertNotIn("safety-y-axis-title", script)
        self.assertNotIn("safety-x-axis-title", script)
        self.assertNotIn(">Hz<", script)
        self.assertNotIn("日序号", script)
        self.assertNotIn(">50Hz<", script)
        self.assertIn("向上频率最大值", script)
        self.assertIn("向下频率最小值", script)
        self.assertNotIn("第${day}日", script)
        safety_render_script = script.split("function renderSafetyResult", 1)[1].split("function renderSafetyDailyChart", 1)[0]
        self.assertLess(safety_render_script.index("safety-result-table"), safety_render_script.index("safety-chart-card"))
        for metric in (
            "向上扰动最大量",
            "向下扰动最大量",
            "最高频率",
            "最低频率",
            "频率安全风险小时数",
        ):
            self.assertIn(metric, script)

        self.assertIn(".safety-result-layout", css)
        safety_layout_css = css.split(".safety-result-layout {", 1)[1].split("}", 1)[0]
        self.assertIn("display: grid", safety_layout_css)
        self.assertIn("grid-template-columns: minmax(0, var(--safety-result-table-width, 34%)) 10px minmax(0, 1fr)", safety_layout_css)
        self.assertIn(".result-column-resize-handle", css)
        result_column_resize_css = css.split(".result-column-resize-handle {", 1)[1].split("}", 1)[0]
        self.assertIn("cursor: col-resize", result_column_resize_css)
        self.assertIn(".safety-chart-card", css)
        safety_chart_card_css = css.split(".safety-chart-card {", 1)[1].split("}", 1)[0]
        self.assertIn("display: grid", safety_chart_card_css)
        self.assertIn("grid-template-rows: auto minmax(0, 1fr)", safety_chart_card_css)
        self.assertIn(".safety-chart-viewport", css)
        safety_chart_viewport_css = css.split(".green-chart-viewport,", 1)[1].split("{", 1)[1].split("}", 1)[0]
        self.assertIn("height: 100%", safety_chart_viewport_css)
        self.assertIn("min-height: 0", safety_chart_viewport_css)
        self.assertIn(".safety-chart-svg", css)
        self.assertIn(".chart-hover-line", css)
        self.assertIn(".chart-hover-tooltip", css)
        chart_tooltip_css = css.split(".chart-hover-tooltip {", 1)[1].split("}", 1)[0]
        self.assertIn("position: absolute", chart_tooltip_css)
        self.assertIn("pointer-events: none", chart_tooltip_css)
        self.assertIn('preserveAspectRatio="xMidYMid meet"', script)
        self.assertNotIn('preserveAspectRatio="none"', script)
        self.assertIn('resultChartSize("safety"', script)
        safety_chart_svg_css = css.split(".safety-chart-svg {", 1)[1].split("}", 1)[0]
        self.assertIn("width: 100%", safety_chart_svg_css)
        self.assertIn("height: 100%", safety_chart_svg_css)
        self.assertIn("min-height: 0", safety_chart_svg_css)
        self.assertIn(".safety-chart-legend", css)
        self.assertIn(".safety-chart-legend button", css)
        self.assertIn(".safety-chart-legend button.is-hidden", css)
        self.assertIn(".safety-center-line", css)
        self.assertIn(".safety-result-table", css)

    def test_frequency_calculation_rows_include_source_and_calculated_extremes(self):
        scheme_payload = {"planning_parameters": [{"nominal_frequency_hz": 50, "frequency_governor_time_constant_s": 0.6}]}
        dispatch_rows = [
            {
                "hour_index": 1,
                "datetime": "2026-01-01 00:00",
                "load": 100,
                "wind_power": 20,
                "pv_power": 10,
                "diesel_capacity": 200,
                "diesel_power": 80,
                "renewable_power": 30,
                "grid_storage_capacity": 50,
                "storage_power": 1,
                "diesel_on": 2,
                "storage_charge": 3,
                "storage_discharge": 4,
                "equivalent_inertia_m": 10,
                "equivalent_primary_frequency_k": 1.2,
                "equivalent_damping_d": 0.8,
                "frequency_delta_p_mw": 0.2,
                "frequency_upper_delta_p_mw": -0.15,
                "frequency_max": 50.42,
                "frequency_min": 49.58,
            }
        ]

        outputs = list(server.iter_frequency_result_rows(scheme_payload, dispatch_rows))
        summary = outputs[0]["summary"]
        lower_curve, upper_curve = outputs[0]["curves"]
        time_headers = [server.frequency_curve_point_header(index) for index in range(server.FREQUENCY_CURVE_POINT_COUNT)]

        for header in (
            "max_up_disturbance_mw",
            "max_down_disturbance_mw",
            "source_frequency_max_hz",
            "source_frequency_min_hz",
            "calculated_max_frequency_hz",
            "calculated_min_frequency_hz",
            "diesel_capacity",
            "diesel_power",
            "load",
            "renewable_power",
            "grid_storage_capacity",
            "storage_power",
        ):
            self.assertIn(header, server.frequency_8760_result_headers())
            self.assertIn(header, server.frequency_curve_headers())

        self.assertEqual(summary["diesel_capacity"], 200)
        self.assertEqual(summary["diesel_power"], 80)
        self.assertEqual(summary["load"], 100)
        self.assertEqual(summary["renewable_power"], 30)
        self.assertEqual(summary["grid_storage_capacity"], 50)
        self.assertEqual(summary["storage_power"], 1)
        self.assertEqual(summary["max_up_disturbance_mw"], 0.2)
        self.assertEqual(summary["max_down_disturbance_mw"], -0.15)
        self.assertEqual(summary["source_frequency_max_hz"], 50.42)
        self.assertEqual(summary["source_frequency_min_hz"], 49.58)
        self.assertEqual(summary["calculated_min_frequency_hz"], min(lower_curve[header] for header in time_headers))
        self.assertEqual(summary["calculated_max_frequency_hz"], max(upper_curve[header] for header in time_headers))
        self.assertEqual(lower_curve["calculated_min_frequency_hz"], summary["calculated_min_frequency_hz"])
        self.assertEqual(upper_curve["calculated_max_frequency_hz"], summary["calculated_max_frequency_hz"])

    def test_frequency_8760_curve_board_exposes_requested_metric_series(self):
        frequency_script = (WEB_ROOT / "assets" / "frequency.js").read_text(encoding="utf-8")
        frequency_page = (WEB_ROOT / "frequency.html").read_text(encoding="utf-8")
        css = (WEB_ROOT / "assets" / "planning.css").read_text(encoding="utf-8")
        results = [
            {
                "summary": {
                    "hour_index": 1,
                    "datetime": "2026-01-01 00:00",
                    "diesel_capacity": 200,
                    "max_up_disturbance_mw": 0.2,
                    "max_down_disturbance_mw": -0.15,
                    "source_frequency_max_hz": 50.42,
                    "source_frequency_min_hz": 49.58,
                    "calculated_max_frequency_hz": 50.35,
                    "calculated_min_frequency_hz": 49.65,
                }
            }
        ]

        rows = server.frequency_8760_display_rows_from_results(results, {})
        self.assertEqual(rows[0]["柴发开机容量"], 200)
        for label in ("向上最大扰动", "向下最大扰动", "优化频率最大值", "优化频率最小值", "仿真频率最大值", "仿真频率最小值"):
            self.assertIn(label, rows[0])
            self.assertIn(label, frequency_script)
        self.assertIn("frequency8760CurveBoard", frequency_page)
        self.assertIn("function renderFrequency8760CurveBoard", frequency_script)
        self.assertIn("data.frequency_8760_table", frequency_script)
        self.assertIn("分时曲线", frequency_page)
        self.assertNotIn("frequencyCurveYear", frequency_page)
        self.assertIn("frequencyCurveDate", frequency_page)
        self.assertIn("frequencyTimeResultResizeHandle", frequency_page)
        self.assertIn("frequencyMetricsResizeHandle", frequency_page)
        self.assertIn("frequencySchemeListResizeHandle", frequency_page)
        self.assertNotIn("frequencyResultSelect", frequency_page)
        self.assertNotIn("frequencyResultSelect", frequency_script)
        self.assertNotIn("renderFrequencyResultOption", frequency_script)
        self.assertNotIn("频率校核摘要", frequency_page)
        self.assertIn("frequencySummaryTitle", frequency_page)
        self.assertIn("当前结果:未选择结果", frequency_page)
        self.assertIn("renderFrequencySummaryTitle", frequency_script)
        self.assertIn('setText("frequencySummaryTitle", `当前结果:${resultDisplayName(frequencyState.selectedResultFile) || "未选择结果"}`)', frequency_script)
        self.assertIn("resultsByScheme", frequency_script)
        self.assertIn("collapsedSchemes", frequency_script)
        self.assertIn("loadFrequencyResultTree", frequency_script)
        self.assertIn("renderFrequencySchemeTreeNode", frequency_script)
        self.assertIn("renderFrequencyResultTreeItem", frequency_script)
        self.assertIn("data-frequency-scheme-toggle", frequency_script)
        self.assertIn("function refreshFrequencyTimeCurve", frequency_script)
        self.assertIn("function bindFrequencySchemeListResize", frequency_script)
        self.assertIn("function setFrequencySchemeRailHeight", frequency_script)
        self.assertIn("function frequencySchemeRailHeightBounds", frequency_script)
        self.assertIn("setFrequencySchemeRailHeight(frequencySchemeRailHeightBounds().max", frequency_script)
        self.assertIn("--evaluation-scheme-rail-height", frequency_script)
        self.assertIn("function renderFrequencyLogs", frequency_script)
        self.assertIn("isLogScrolledNearBottom", frequency_script)
        self.assertIn("const shouldStickToBottom = isLogScrolledNearBottom(target)", frequency_script)
        self.assertIn("const previousScrollTop = target.scrollTop", frequency_script)
        self.assertIn("target.scrollTop = shouldStickToBottom ? target.scrollHeight : previousScrollTop", frequency_script)
        self.assertIn("bindFrequencyAxisRangeControls", frequency_script)
        self.assertIn("renderFrequencyAxisRangeControls", frequency_script)
        self.assertIn("data-frequency-axis-min", frequency_script)
        self.assertIn("data-frequency-axis-max", frequency_script)
        self.assertIn("function bindFrequencyMetricsResize", frequency_script)
        self.assertIn("function bindFrequencyTimeResultResize", frequency_script)
        self.assertIn("scheduleFrequencyPolling();", frequency_script)
        self.assertIn("window.setTimeout(() => refreshFrequencyStatus().catch(showFrequencyError), 250)", frequency_script)
        self.assertIn("/api/frequency/time-curve", frequency_script)
        self.assertIn('preserveAspectRatio="none"', frequency_script)
        self.assertIn("comparison-chart-frame frequency-8760-chart-frame", frequency_script)
        self.assertIn("function bindFrequency8760Hover", frequency_script)
        self.assertIn("function renderFrequency8760Stats", frequency_script)
        self.assertIn("frequencyDownsample(series.values, 720)", frequency_script)
        self.assertIn(".frequency-metrics-layout", css)
        self.assertIn("--frequency-metrics-table-width", css)
        self.assertIn(".frequency-8760-curve-board", css)
        self.assertIn(".frequency-8760-chart-frame", css)
        self.assertIn(".frequency-8760-line", css)
        self.assertIn(".frequency-8760-stats", css)
        self.assertIn(".frequency-8760-hover-capture", css)
        self.assertIn(".frequency-time-curve-panel", css)
        self.assertIn(".frequency-time-result-layout", css)
        self.assertIn(".frequency-workspace .evaluation-result-rail", css)
        self.assertIn("grid-template-rows: minmax(0, 1fr) auto", css)
        self.assertIn("--frequency-time-info-width", css)
        self.assertIn(".axis-range-controls", css)
        self.assertIn(".axis-range-panel", css)
        self.assertIn(".axis-range-controls:hover .axis-range-panel", css)
        self.assertIn(".axis-range-toggle", css)
        frequency_time_chart_css = css.split(".frequency-time-curve-chart {", 1)[1].split("}", 1)[0]
        self.assertIn("display: grid", frequency_time_chart_css)
        self.assertIn("grid-template-rows: minmax(0, 1fr) auto", frequency_time_chart_css)

    def test_frequency_time_curve_payload_reads_selected_hour_curves(self):
        path = WEB_ROOT / "tests" / "tmp_frequency_time_curve.xlsx"
        workbook = Workbook()
        summary = workbook.active
        summary.title = "频率8760结果"
        summary.append(server.frequency_8760_result_headers())
        summary.append([
            1,
            "2026-01-02 03:00",
            1.0,
            200,
            2,
            80,
            100,
            30,
            50,
            1,
            0,
            0,
            10,
            0.8,
            1.2,
            0.2,
            -0.15,
            50.42,
            49.58,
            50.35,
            49.65,
        ])
        curve = workbook.create_sheet("频率曲线")
        curve_headers = server.frequency_curve_headers()
        curve.append(curve_headers)
        low_row = {header: "" for header in curve_headers}
        high_row = {header: "" for header in curve_headers}
        low_row.update({"hour_index": 1, "datetime": "2026-01-02 03:00", "curve_type": "最低频率曲线"})
        high_row.update({"hour_index": 1, "datetime": "2026-01-02 03:00", "curve_type": "最高频率曲线"})
        for index in range(server.FREQUENCY_CURVE_POINT_COUNT):
            header = server.frequency_curve_point_header(index)
            low_row[header] = 50 - index * 0.001
            high_row[header] = 50 + index * 0.001
        curve.append([low_row.get(header, "") for header in curve_headers])
        curve.append([high_row.get(header, "") for header in curve_headers])
        try:
            workbook.save(path)
            payload = server.read_frequency_time_curve_payload(path, month="1", day="2", hour="3")
        finally:
            workbook.close()
            if path.exists():
                path.unlink()

        self.assertEqual(payload["selection"]["hour_index"], 1)
        metrics = {row["指标"]: row["数值"] for row in payload["summary_table"]}
        self.assertEqual(metrics["柴发开机总容量"], 200)
        self.assertEqual(metrics["柴发总功率"], 80)
        self.assertEqual(metrics["负荷总功率"], 100)
        self.assertEqual(metrics["新能源总出力"], 30)
        self.assertEqual(metrics["向上最大扰动"], 0.2)
        self.assertEqual(metrics["向下最大扰动"], -0.15)
        self.assertEqual(len(payload["curves"]["high"]), server.FREQUENCY_CURVE_POINT_COUNT)
        self.assertEqual(len(payload["curves"]["low"]), server.FREQUENCY_CURVE_POINT_COUNT)
        self.assertEqual(payload["curves"]["high"][0], {"time": 0.0, "frequency": 50.0})
        self.assertEqual(payload["curves"]["low"][1]["time"], 0.05)

    def test_frequency_curve_export_emits_progress_logs(self):
        scheme = "测试频率日志方案"
        filename = "logcase_results.xlsx"
        result_path = server.frequency_curve_result_path(scheme, filename)
        logs: list[tuple[str, str]] = []
        summary = {
            "hour_index": 1,
            "datetime": "2026-01-01 00:00",
            "grid_model": 1,
            "diesel_capacity": 100,
            "diesel_on": 1,
            "diesel_power": 50,
            "load": 80,
            "renewable_power": 30,
            "grid_storage_capacity": 20,
            "storage_power": 0,
            "storage_charge": 0,
            "storage_discharge": 0,
            "equivalent_inertia_m": 10,
            "load_response_d": 0.8,
            "equivalent_primary_frequency_k": 1.2,
            "max_up_disturbance_mw": 0.2,
            "max_down_disturbance_mw": -0.15,
            "source_frequency_max_hz": 50.2,
            "source_frequency_min_hz": 49.8,
            "calculated_max_frequency_hz": 50.1,
            "calculated_min_frequency_hz": 49.9,
        }
        curves = [
            {"hour_index": 1, "datetime": "2026-01-01 00:00", "curve_type": "最低频率曲线"},
            {"hour_index": 1, "datetime": "2026-01-01 00:00", "curve_type": "最高频率曲线"},
        ]
        try:
            server.export_frequency_curve_workbook(
                scheme,
                filename,
                {},
                [],
                frequency_results=[{"summary": summary, "curves": curves}],
                log_callback=lambda level, message: logs.append((level, message)),
            )
        finally:
            if result_path.exists():
                result_path.unlink()

        messages = [message for _, message in logs]
        self.assertTrue(any("频率结果写入进度" in message for message in messages))
        self.assertTrue(any("保存频率曲线Excel文件" in message for message in messages))
        self.assertTrue(any("替换频率曲线结果文件" in message for message in messages))

    def test_calculation_result_frontends_display_numeric_values_with_two_decimals(self):
        optimize_script = (WEB_ROOT / "assets" / "optimize.js").read_text(encoding="utf-8")
        evaluation_script = (WEB_ROOT / "assets" / "evaluation.js").read_text(encoding="utf-8")
        comparison_script = (WEB_ROOT / "assets" / "comparison.js").read_text(encoding="utf-8")
        result_curve_script = (WEB_ROOT / "assets" / "result_curves.js").read_text(encoding="utf-8")

        for script in (optimize_script, evaluation_script):
            self.assertIn("function formatDisplayValue(value)", script)
            self.assertIn("function formatMetricValue(item)", script)
            self.assertIn("function formatMetricTime(value)", script)
            self.assertIn("element.textContent = `${formatMetricValue(item)}${unit}`", script)
            self.assertIn('["开始", "完成"].includes(item.label)', script)
            self.assertIn('item.label === "度电成本"', script)
            self.assertIn("formatLevelizedCostValue(item.value)", script)
            self.assertIn('text.match(/(\\d{2}:\\d{2}:\\d{2})$/)', script)
            self.assertIn("formatDisplayValue(row[header], row, header)", script)
            self.assertIn("formatDisplayValue(value)", script)
            self.assertIn('row?.["指标"] === "度电成本"', script)
            self.assertIn("function formatLevelizedCostValue(value)", script)
            self.assertIn("minimumFractionDigits: 3", script)
            self.assertIn("maximumFractionDigits: 3", script)
            self.assertIn("minimumFractionDigits: 2", script)
            self.assertIn("maximumFractionDigits: 2", script)
            self.assertIn("return formatNumber(value);", script)
            self.assertNotIn("minimumSignificantDigits: 3", script)
            self.assertNotIn("maximumSignificantDigits: 3", script)
            self.assertNotIn("maximumFractionDigits: 1", script)

        for script in (comparison_script, result_curve_script):
            self.assertIn("function formatDisplayValue(value)", script)
            self.assertIn('formatDisplayValue(row[header] ?? "", row, header)', script)
            self.assertIn('row?.["指标"] === "度电成本"', script)
            self.assertIn("function formatLevelizedCostValue(value)", script)
            self.assertIn("minimumFractionDigits: 3", script)
            self.assertIn("maximumFractionDigits: 3", script)
            self.assertIn("minimumFractionDigits: 2", script)
            self.assertIn("maximumFractionDigits: 2", script)
            self.assertNotIn("minimumSignificantDigits: 3", script)
            self.assertNotIn("maximumSignificantDigits: 3", script)
            self.assertNotIn("maximumFractionDigits: 1", script)

    def test_drag_resize_controls_allow_full_panel_collapse(self):
        planning_script = (WEB_ROOT / "assets" / "planning.js").read_text(encoding="utf-8")
        optimize_script = (WEB_ROOT / "assets" / "optimize.js").read_text(encoding="utf-8")
        evaluation_script = (WEB_ROOT / "assets" / "evaluation.js").read_text(encoding="utf-8")
        comparison_script = (WEB_ROOT / "assets" / "comparison.js").read_text(encoding="utf-8")
        planning_html = (WEB_ROOT / "planning.html").read_text(encoding="utf-8")
        css = (WEB_ROOT / "assets" / "planning.css").read_text(encoding="utf-8")

        self.assertIn("const COLLAPSED_PANEL_SIZE = 0", planning_script)
        self.assertIn("Math.max(Number(height) || 240, COLLAPSED_PANEL_SIZE)", planning_script)
        self.assertIn('aria-valuemin="0"', planning_html)

        for script in (optimize_script, evaluation_script):
            self.assertIn("const COLLAPSED_PANEL_SIZE = 0", script)
            self.assertIn("return { min: COLLAPSED_PANEL_SIZE", script)
            self.assertIn("const maxTableWidth = layout.clientWidth - gap * 2 - handleWidth", script)
            self.assertNotIn("const minTableWidth = 260", script)
            self.assertNotIn("const minChartWidth = 320", script)
            self.assertNotIn("const max = grid.clientWidth - handleWidth * 2 - gap * 4 - 240 - 240", script)

        self.assertIn("const MIN_COMPARISON_TABLE_FR = 0", comparison_script)
        self.assertIn("Math.max(Number.isFinite(numericValue) ? numericValue : 1, MIN_COMPARISON_TABLE_FR)", comparison_script)
        self.assertIn("Math.max(COLLAPSED_PANEL_SIZE, Math.min", comparison_script)
        self.assertNotIn("const minWidth = 0.45", comparison_script)

        self.assertIn("grid-template-columns: minmax(0, var(--evaluation-result-rail-width, 340px)) 10px minmax(0, 1fr)", css)
        self.assertIn("grid-template-rows: minmax(108px, auto) minmax(0, var(--comparison-table-height, 30vh)) 12px minmax(0, 1fr)", css)
        self.assertIn("grid-template-columns: minmax(0, var(--overview-left-column-width, 1fr)) 10px minmax(0, 1.35fr)", css)
        self.assertIn("grid-template-columns: minmax(0, var(--green-result-table-width, 34%)) 10px minmax(0, 1fr)", css)
        self.assertIn("grid-template-columns: minmax(0, var(--safety-result-table-width, 34%)) 10px minmax(0, 1fr)", css)
        comparison_curve_board_css = css.split(".comparison-curve-board {", 1)[1].split("}", 1)[0]
        self.assertIn("min-height: 0", comparison_curve_board_css)

    def test_planning_scheme_rail_places_scheme_actions_around_list(self):
        html = (WEB_ROOT / "planning.html").read_text(encoding="utf-8")
        css = (WEB_ROOT / "assets" / "planning.css").read_text(encoding="utf-8")
        rail = html.split('<aside class="scheme-rail">', 1)[1].split("</aside>", 1)[0]

        self.assertNotIn("方案管理", rail)
        self.assertIn("方案列表", rail)
        self.assertIn('id="schemeList"', rail)
        for control in ("createScheme", "deleteScheme", "renameScheme", "copyScheme", "shareScheme"):
            self.assertIn(f'id="{control}"', rail)
        self.assertLess(rail.index('id="createScheme"'), rail.index("方案列表"))
        self.assertLess(rail.index('id="deleteScheme"'), rail.index("方案列表"))
        self.assertLess(rail.index('id="createScheme"'), rail.index('id="schemeList"'))
        self.assertLess(rail.index('id="deleteScheme"'), rail.index('id="schemeList"'))
        self.assertLess(rail.index('id="schemeList"'), rail.index('id="renameScheme"'))
        self.assertLess(rail.index('id="schemeList"'), rail.index('id="copyScheme"'))
        self.assertLess(rail.index('id="schemeList"'), rail.index('id="shareScheme"'))
        self.assertIn(".planning-scheme-rail-layout", css)
        self.assertIn("--planning-scheme-rail-height", css)
        planning_script = (WEB_ROOT / "assets" / "planning.js").read_text(encoding="utf-8")
        self.assertIn("applyAdaptiveSchemeRailLayout", planning_script)
        self.assertIn("scheduleSchemeRailLayout", planning_script)
        self.assertIn("function shareScheme()", planning_script)
        self.assertIn('document.getElementById("shareScheme").addEventListener("click", shareScheme)', planning_script)
        self.assertIn("currentSchemeCanManage", planning_script)
        self.assertIn("const workspaceContentHeight", planning_script)
        self.assertIn("summaryMinimumHeight = Math.max(280", planning_script)
        self.assertIn("workspaceContentHeight - rowGap * 2 - handleHeight - summaryMinimumHeight", planning_script)
        self.assertIn("const preferredHeight = Number.isFinite(state.schemeRailManualHeight) ? state.schemeRailManualHeight : maxRailHeight", planning_script)
        self.assertIn("grid-template-rows: minmax(150px, var(--planning-scheme-rail-height, auto)) 8px minmax(0, 1fr)", css)
        self.assertIn("grid-template-rows: auto auto minmax(0, 1fr) auto", css)
        self.assertIn(".scheme-rail.scheme-list-capped .planning-scheme-rail-layout", css)
        self.assertIn("grid-template-rows: auto auto minmax(0, 1fr) auto", css)
        self.assertIn("schemeListResizeHandle", html)
        self.assertGreater(html.index('id="schemeListResizeHandle"'), html.index("</aside>"))
        self.assertIn("bindSchemeListResizeHandle", planning_script)
        self.assertIn("schemeRailManualHeight", planning_script)
        self.assertIn("schemeRailHeightBounds", planning_script)
        self.assertIn(".scheme-list-resize-handle", css)
        self.assertIn(".scheme-rail.scheme-list-capped .scheme-list", css)
        self.assertIn("overflow: auto", css.split(".planning-scheme-rail-layout .scheme-list {", 1)[1].split("}", 1)[0])
        self.assertIn("overflow: auto", css.split(".scheme-rail.scheme-list-capped .scheme-list {", 1)[1].split("}", 1)[0])
        self.assertIn(".scheme-actions-rail", css)
        self.assertIn("grid-template-columns: repeat(2, minmax(0, 1fr))", css)
        self.assertIn(".scheme-actions-rail-bottom", css)
        self.assertIn("grid-template-columns: repeat(3, minmax(0, 1fr))", css)
        self.assertIn(".scheme-access-label", css)
        self.assertIn(".scheme-list-title", css)
        self.assertIn("color: #102b2a", css)
        self.assertIn("font-size: 18px", css)
        self.assertIn("font-weight: 900", css)
        self.assertIn("grid-template-columns: 280px minmax(0, 1fr)", css)
        self.assertIn(".workspace > .summary-rail {\n  grid-column: 1;\n  grid-row: 3;", css)
        self.assertIn(".workspace > .editor-panel {\n  grid-column: 2;\n  grid-row: 1 / 4;", css)

    def test_planning_page_save_button_is_in_scheme_actions(self):
        html = (WEB_ROOT / "planning.html").read_text(encoding="utf-8")
        current_scheme_panel = html.split('<div class="current-scheme-panel">', 1)[1].split('<div class="tabs"', 1)[0]
        editor_header = html.split('<div class="editor-header">', 1)[1].split("</div>\n\n        <section", 1)[0]
        topbar = html.split('<header class="topbar">', 1)[1].split("</header>", 1)[0]
        rail = html.split('<aside class="scheme-rail">', 1)[1].split("</aside>", 1)[0]

        self.assertNotIn('id="saveScheme"', current_scheme_panel)
        self.assertIn('class="scheme-actions"', editor_header)
        self.assertIn("margin-left: auto", (WEB_ROOT / "assets" / "planning.css").read_text(encoding="utf-8"))
        self.assertIn('id="importScheme"', editor_header)
        self.assertIn('id="exportScheme"', editor_header)
        self.assertIn('id="schemeImportFile"', editor_header)
        self.assertIn('id="saveScheme"', editor_header)
        self.assertNotIn('id="renameScheme"', editor_header)
        self.assertNotIn('id="copyScheme"', editor_header)
        self.assertNotIn('id="deleteScheme"', editor_header)
        self.assertNotIn('id="shareScheme"', editor_header)
        self.assertIn('id="renameScheme"', rail)
        self.assertIn('id="copyScheme"', rail)
        self.assertIn('id="deleteScheme"', rail)
        self.assertIn('id="shareScheme"', rail)
        self.assertNotIn('id="importScheme"', rail)
        self.assertNotIn('id="exportScheme"', rail)
        self.assertIn(">修改名称<", rail)
        self.assertNotIn("修改方案名称", editor_header)
        self.assertNotIn("修改方案名", editor_header)
        self.assertNotIn('id="saveScheme"', topbar)
        self.assertNotIn('id="saveScheme"', rail)
        self.assertLess(html.index('id="importScheme"'), html.index('id="exportScheme"'))
        self.assertLess(html.index('id="exportScheme"'), html.index('id="saveScheme"'))
        planning_script = (WEB_ROOT / "assets" / "planning.js").read_text(encoding="utf-8")
        self.assertIn('document.getElementById("importScheme").addEventListener("click", importScheme)', planning_script)
        self.assertIn('document.getElementById("exportScheme").addEventListener("click", exportScheme)', planning_script)
        self.assertIn('document.getElementById("schemeImportFile").addEventListener("change", onSchemeImportFileChange)', planning_script)
        self.assertIn("function importScheme()", planning_script)
        self.assertIn("async function exportScheme()", planning_script)
        self.assertIn("/api/planning/schemes/import", planning_script)
        self.assertIn("/export", planning_script)
        self.assertIn("downloadBlob", planning_script)
        self.assertIn("filenameFromContentDisposition", planning_script)
        self.assertLess(html.index('id="currentSchemeName"'), html.index(">时序数据<"))
        self.assertLess(html.index(">时序数据<"), html.index('id="saveScheme"'))

    def test_planning_scheme_actions_are_horizontal(self):
        css = (WEB_ROOT / "assets" / "planning.css").read_text(encoding="utf-8")

        self.assertIn("flex-wrap: nowrap", css)
        self.assertIn("overflow-x: auto", css)
        self.assertIn("white-space: nowrap", css)

    def test_planning_page_has_device_filter_tags(self):
        html = (WEB_ROOT / "planning.html").read_text(encoding="utf-8")
        script = (WEB_ROOT / "assets" / "planning.js").read_text(encoding="utf-8")
        css = (WEB_ROOT / "assets" / "planning.css").read_text(encoding="utf-8")

        self.assertIn('id="deviceFilters"', html)
        self.assertIn("renderDeviceFilters", script)
        self.assertIn("visibleDevices", script)
        self.assertIn("deviceGroups", script)
        self.assertIn("data-device-group", script)
        self.assertNotIn("<h2>设备类型显示</h2>", html)
        self.assertNotIn("默认全部显示，取消勾选则隐藏对应表格。", html)
        self.assertIn('class="device-filter-row"', html)
        device_filter_card = html.split('<div class="device-filter-card">', 1)[1].split('<div id="deviceTables"', 1)[0]
        self.assertIn('id="deviceFilters"', device_filter_card)
        self.assertIn('id="deviceJump"', device_filter_card)
        self.assertLess(device_filter_card.index('id="deviceFilters"'), device_filter_card.index('id="deviceJump"'))
        self.assertIn(".device-filter-row", css)
        self.assertIn("flex-wrap: nowrap", css)
        self.assertIn("overflow-x: auto", css)
        self.assertIn("justify-content: flex-end", css)
        self.assertIn("margin-left: auto", css)
        self.assertIn("min-width: 0", css)
        for group_name in ("风光柴", "氢储能", "电储能"):
            self.assertIn(group_name, script)
        self.assertLess(script.index('"电储能"'), script.index('"氢储能"'))
        self.assertLess(script.index('"储能PCS"'), script.index('"电制氢"'))
        for name in ("柴发", "风机", "光伏", "储能PCS", "储能电池组", "电制氢", "储氢罐", "燃料电池"):
            self.assertIn(name, script)

    def test_planning_page_has_planning_parameters_tab_and_summary_panel(self):
        html = (WEB_ROOT / "planning.html").read_text(encoding="utf-8")
        script = (WEB_ROOT / "assets" / "planning.js").read_text(encoding="utf-8")
        css = (WEB_ROOT / "assets" / "planning.css").read_text(encoding="utf-8")

        self.assertIn('data-tab="planning"', html)
        self.assertIn('id="planningTab"', html)
        self.assertIn('id="planningParametersTable"', html)
        planning_tab_html = html.split('<section id="planningTab"', 1)[1].split('<section id="limitsTab"', 1)[0]
        self.assertNotIn('class="panel-heading"', planning_tab_html)
        self.assertNotIn("<h1>规划参数</h1>", planning_tab_html)
        self.assertNotIn("参数随当前方案保存到 XLSX 文件。", html)
        self.assertLess(html.index('data-tab="devices"'), html.index('data-tab="planning"'))
        self.assertLess(html.index('data-tab="planning"'), html.index('data-tab="limits"'))
        self.assertIn('data-summary-tab="planning"', html)
        self.assertIn('data-summary-panel="planning"', html)
        self.assertIn('id="planningSummary"', html)
        self.assertIn("planningParameterSpecs", script)
        self.assertIn("planningParameterGroups", script)
        self.assertIn("activePlanningParameterGroup", script)
        self.assertIn("renderPlanningParameterTabs", script)
        self.assertIn("selectPlanningParameterGroup", script)
        self.assertIn("planningGroupToggle", script)
        self.assertIn("isPlanningGroupEnabled", script)
        self.assertIn('"storage_balance_mode"', script)
        self.assertIn('"电储能平衡模式"', script)
        self.assertIn('"daily", "日内平衡"', script)
        self.assertIn('"weekly", "周内平衡"', script)
        self.assertIn('"monthly", "月度平衡"', script)
        self.assertIn('"annual", "年度平衡"', script)
        self.assertIn('"none", "不闭环"', script)
        self.assertNotIn("bindPlanningParameterResizeHandles", script)
        self.assertNotIn("planning-parameter-resize-handle", script)
        self.assertIn("renderPlanningParameterGroupTable", script)
        self.assertIn("planning-parameter-name-col", script)
        self.assertIn("planning-parameter-value-col", script)
        self.assertIn("planning-parameter-range-col", script)
        self.assertIn("renderPlanningParameters", script)
        self.assertIn("renderPlanningParameterSummaryTable", script)
        self.assertIn('grid-template-columns: minmax(0, 1fr)', css)
        self.assertNotIn('grid-template-columns: repeat(3, minmax(220px, 1fr))', css)
        self.assertIn("collectPlanningParameterWarnings", script)
        self.assertIn("planning_parameters", script)
        self.assertIn(".planning-parameters-card", css)
        self.assertIn("#planningTab #planningParametersTable", css)
        planning_parameters_table_css = css.split("#planningTab #planningParametersTable", 1)[1].split("}", 1)[0]
        self.assertIn("border: 0 !important", planning_parameters_table_css)
        self.assertIn("background: transparent !important", planning_parameters_table_css)
        self.assertIn(".planning-parameter-grid", css)
        self.assertIn(".planning-parameter-tabs", css)
        planning_parameter_tabs_css = css.split(".planning-parameter-tabs", 1)[1].split("}", 1)[0]
        self.assertIn("padding: 0", planning_parameter_tabs_css)
        self.assertIn("border: 0 !important", planning_parameter_tabs_css)
        self.assertIn("background: transparent !important", planning_parameter_tabs_css)
        self.assertIn(".planning-parameter-tab.active", css)
        self.assertIn(".planning-parameter-panel", css)
        self.assertIn(".planning-parameter-group", css)
        self.assertIn(".planning-parameter-switch", css)
        self.assertIn(".planning-parameter-group.disabled", css)
        self.assertNotIn(".planning-parameter-resize-handle", css)
        self.assertIn("table-layout: fixed", css)
        self.assertIn("height: 100%", css)
        self.assertIn(".planning-parameter-name-col", css)
        self.assertIn(".planning-parameter-value-col", css)
        self.assertIn(".planning-parameter-range-col", css)
        self.assertIn("overflow-wrap: anywhere", css)
        for label in (
            "常规参数",
            "扰动后安全参数",
            "频率安全参数",
            "柴油价格(万元/吨)",
            "柴发开机持续工作小时数下限",
            "柴发关机持续工作小时数下限",
            "工作模式",
            "全年运行",
            "度夏运行",
            "冬季开始月份",
            "冬季开始日期",
            "冬季结束月份",
            "冬季结束日期",
            "绿色电量占比下限(0.0-1.0)",
            "规划求解时间上限(分钟)",
            "初始电储SOC(0.0-1.0)",
            "初始氢储SOC(0.0-1.0)",
            "储能是否参与调频",
            "是否考虑扰动后平衡约束",
            "负荷向上扰动系数(0.0-0.5)",
            "负荷向下扰动系数(0.0-0.5)",
            "新能源向下扰动系数(0.0-0.5)",
            "是否考虑频率安全约束",
            "额定频率(Hz)",
            "频率最低点下限(Hz)",
            "频率最高点上限(Hz)",
            "频率下限安全裕度(Hz)",
            "频率上限安全裕度(Hz)",
            "负荷频率系数D",
            "RoCoF上限(Hz/s)",
            "稳态频率下限(Hz)",
            "稳态频率上限(Hz)",
            "频率等效调速时间常数T(s)",
            "频率Nadir评估时长(s)",
            "Nadir线性化每轴采样点数",
            "Nadir线性化区间比例",
            "网络同步系数基值",
            "网络同步系数斜率",
            "网络同步系数基准负荷(kW)",
            "是否考虑新能源N-1",
            "是否考虑新能源扰动",
            "是否考虑负荷扰动",
        ):
            self.assertIn(label, script)
        self.assertLess(script.index('"frequency_security_constraint_enabled"'), script.index('"nominal_frequency_hz"'))
        self.assertLess(script.index('"diesel_price"'), script.index('"diesel_minimum_on_hours"'))
        self.assertLess(script.index('"diesel_minimum_on_hours"'), script.index('"diesel_minimum_off_hours"'))
        self.assertLess(script.index('"diesel_minimum_off_hours"'), script.index('"operation_mode"'))
        self.assertLess(script.index('"operation_mode"'), script.index('"winter_start_month"'))
        self.assertLess(script.index('"winter_end_day"'), script.index('"green_power_ratio_lower"'))
        self.assertLess(script.index('"nominal_frequency_hz"'), script.index('"frequency_nadir_lower_hz"'))
        self.assertLess(script.index('"frequency_governor_time_constant_s"'), script.index('"frequency_nadir_evaluation_duration_s"'))
        self.assertLess(script.index('"nadir_linearization_interval_ratio"'), script.index('"network_synchronization_coefficient_base"'))
        self.assertLess(script.index('"network_synchronization_reference_load_kw"'), script.index('"storage_frequency_regulation_enabled"'))
        self.assertNotIn("频率下限扰动功率(kW)", script)
        self.assertNotIn("频率上限扰动功率(kW)", script)
        self.assertNotIn('"frequency_lower_disturbance_kw"', script)
        self.assertNotIn('"frequency_upper_disturbance_kw"', script)
        self.assertIn("Nadir线性化每轴采样点数必须为正整数", script)
        self.assertIn("频率最低点下限(Hz)不能大于额定频率(Hz)", script)
        self.assertIn("频率最高点上限(Hz)不能小于额定频率(Hz)", script)
        self.assertIn("稳态频率上限(Hz)不能小于稳态频率下限(Hz)", script)
        self.assertNotIn('["storage_charge_efficiency", "充电效率(0.0-1.0)", "number"', script)
        self.assertNotIn('["storage_discharge_efficiency", "放电效率(0.0-1.0)", "number"', script)
        self.assertNotIn("设计使用年限(年)", script)

    def test_planning_boolean_parameters_use_yes_no_selects(self):
        script = (WEB_ROOT / "assets" / "planning.js").read_text(encoding="utf-8")
        css = (WEB_ROOT / "assets" / "planning.css").read_text(encoding="utf-8")

        self.assertIn("planning-bool-select", script)
        self.assertIn('<option value="1"', script)
        self.assertIn('<option value="0"', script)
        self.assertNotIn('<option value="true"', script)
        self.assertNotIn('<option value="false"', script)
        self.assertIn(">是</option>", script)
        self.assertIn(">否</option>", script)
        self.assertIn("numericBooleanPlanningValue", script)
        self.assertIn('input.type === "checkbox"', script)
        self.assertIn('input.tagName === "SELECT"', script)
        self.assertNotIn('type="checkbox" data-planning-key', script)
        self.assertIn(".planning-bool-select", css)

    def test_planning_save_has_parameter_alarm_validation(self):
        script = (WEB_ROOT / "assets" / "planning.js").read_text(encoding="utf-8")
        i18n_script = (WEB_ROOT / "assets" / "i18n.js").read_text(encoding="utf-8")

        self.assertIn("collectSaveWarnings", script)
        self.assertIn("bindPlanningParameterInputs", script)
        self.assertIn("syncPlanningParameterInputs", script)
        self.assertIn("syncPlanningParameterInput(input)", script)
        self.assertIn('document.addEventListener("change", onPlanningParameterInputEvent)', script)
        self.assertIn("参数校验未通过", script)
        self.assertIn("参数存在警告", script)
        self.assertIn("是否继续保存", script)
        self.assertIn("blockingWarnings", script)
        self.assertIn("advisoryWarnings", script)
        self.assertIn("参数保存成功", script)
        self.assertIn("保存参数失败：", script)
        self.assertIn("buildSchemeSavePayload", script)
        self.assertIn("timeSeriesDirty", script)
        self.assertIn("delete payload.time_series", script)
        self.assertIn("时序数据未正确加载，无法保存曲线", script)
        self.assertIn('"参数保存成功": "Parameters saved successfully"', i18n_script)
        self.assertIn('"保存参数失败：": "Failed to save parameters: "', i18n_script)
        self.assertIn("数量下限(台)", script)
        self.assertIn("数量上限(台)", script)
        self.assertIn("数量上限不能小于数量下限", script)
        self.assertIn("collectDuplicateNumericDeviceWarnings", script)
        self.assertIn("normalizeDeviceNumericSignatureValue", script)
        self.assertIn("从第2列起所有数值相同", script)
        self.assertIn('level: "warning"', script)
        self.assertNotIn("频率安全上限不能小于频率安全下限", script)
        self.assertIn("规划求解时间上限(分钟)", script)
        self.assertIn("defaultValue: 60", script)
        self.assertIn("max: 1440", script)
        self.assertIn('spec[0] === "hydrogen_tanks" ? 0.001 : 0.01', script)
        self.assertNotIn("规划负荷系数(0.1-10.0)", script)
        self.assertNotIn("设计容量上限不能小于下限", script)

    def test_planning_device_fields_have_numeric_validation_rules(self):
        script = (WEB_ROOT / "assets" / "planning.js").read_text(encoding="utf-8")

        self.assertIn("deviceFieldRules", script)
        self.assertIn("integer: true", script)
        self.assertIn("positive: true", script)
        self.assertIn("nonNegative: true", script)
        self.assertIn("deviceInputAttributes", script)
        self.assertIn('type="number"', script)
        self.assertIn('min="0"', script)
        self.assertIn('min="1"', script)
        self.assertIn('step="1"', script)
        self.assertIn('step="any"', script)
        self.assertIn('inputmode="numeric"', script)
        self.assertIn('inputmode="decimal"', script)
        self.assertIn('pattern="[0-9]*"', script)
        for field in (
            "quantity_lower",
            "quantity_upper",
            "design_life_years",
            "cost",
            "capacity",
            "power_capacity",
            "storage_charge_efficiency",
            "storage_discharge_efficiency",
            "storage_equivalent_inertia_constant_h",
            "storage_equivalent_primary_frequency_coefficient_k",
            "storage_equivalent_damping_coefficient_d",
            "battery_capacity",
            "hydrogen_tank_capacity",
            "electric_to_hydrogen_efficiency",
            "hydrogen_to_electric_efficiency",
            "fuel_rate",
            "inertia_constant_h",
            "primary_frequency_coefficient_k",
            "damping_coefficient_d",
            "governor_time_constant_t",
            "power_lower",
            "cut_in_wind_speed",
            "rated_wind_speed",
            "cut_out_wind_speed",
            "is_grid_forming",
            "soc_upper",
            "soc_lower",
            "self_discharge_rate",
        ):
            self.assertIn(field, script)
        for message in (
            "数量上下限必须为非负整数",
            "设计年限(年）必须为正整数",
            "成本(万元/台)必须为非负浮点数",
            "容量(kW)必须为正实数",
            "容量(kWh)必须为正实数",
            "容量(Nm3)必须为正实数",
            "电-氢效率(Nm3/kWh)必须为正实数",
            "氢-电效率(kWh/Nm3)必须为正实数",
            "油耗率(kg/kWh)必须为正实数",
            "惯量常数H(s)必须在0到20.0之间",
            "一次调频系数K必须在0到10.0之间",
            "阻尼系数D必须在0到20.0之间",
            "调速时间常数T(s)必须在0.0001到20.0之间",
            "等效惯量常数H(s)必须在0到20.0之间",
            "等效一次调频系数K必须在0到10.0之间",
            "等效阻尼系数D必须在0到20.0之间",
            "功率下限(kW)必须为非负实数",
            "切入风速(m/s)必须为非负实数",
            "额定风速(m/s)必须为正实数",
            "切出风速(m/s)必须为非负实数",
            "是否构网必须为0或1",
            "充电效率(0.0-1.0)必须在0到1之间，且必须大于0",
            "放电效率(0.0-1.0)必须在0到1之间，且必须大于0",
            "SOC上限(0.0-1.0)必须在0到1之间",
            "SOC下限(0.0-1.0)必须在0到1之间",
            "自损耗率(0-1%/天)必须在0到0.01之间",
        ):
            self.assertIn(message, script)
        self.assertIn("自损耗率(0-1%/天)", script)
        self.assertIn("充电效率(0.0-1.0)", script)
        self.assertIn("放电效率(0.0-1.0)", script)
        self.assertIn("dieselGeneratorDefaultValues", script)
        for default_pair in (
            "cost: 200",
            "capacity: 300",
            "power_upper: 250",
            "power_lower: 80",
            "fuel_rate: 0.28",
            "inertia_constant_h: 3.5",
            "primary_frequency_coefficient_k: 0.4",
            "damping_coefficient_d: 0.01",
            "governor_time_constant_t: 0.6",
            "quantity_lower: 1",
            "quantity_upper: 5",
            "design_life_years: 20",
        ):
            self.assertIn(default_pair, script)
        self.assertIn('spec[0] === "diesel_generators"', script)
        self.assertIn("windTurbineDefaultValues", script)
        for default_pair in (
            "cost: 400",
            "capacity: 100",
            "cut_in_wind_speed: 3",
            "rated_wind_speed: 10",
            "cut_out_wind_speed: 25",
            "quantity_lower: 1",
            "quantity_upper: 5",
            "design_life_years: 20",
        ):
            self.assertIn(default_pair, script)
        self.assertIn('spec[0] === "wind_turbines"', script)
        self.assertIn("photovoltaicDefaultValues", script)
        for default_pair in (
            "cost: 200",
            "capacity: 100",
            "quantity_lower: 1",
            "quantity_upper: 5",
            "design_life_years: 20",
        ):
            self.assertIn(default_pair, script)
        self.assertIn('spec[0] === "photovoltaics"', script)
        self.assertIn("storagePcsDefaultValues", script)
        for default_pair in (
            "cost: 30",
            "power_capacity: 100",
            "quantity_lower: 1",
            "quantity_upper: 5",
            "design_life_years: 20",
        ):
            self.assertIn(default_pair, script)
        self.assertIn('spec[0] === "storage_pcs"', script)
        self.assertIn("storageBatteryPackDefaultValues", script)
        for default_pair in (
            "cost: 100",
            "battery_capacity: 200",
            "quantity_lower: 1",
            "quantity_upper: 5",
            "design_life_years: 20",
        ):
            self.assertIn(default_pair, script)
        self.assertIn('spec[0] === "storage_battery_packs"', script)
        self.assertIn("hydrogenElectrolyzerDefaultValues", script)
        for default_pair in (
            "cost: 400",
            "power_capacity: 70",
            "power_lower: 30",
            "electric_to_hydrogen_efficiency: 0.2",
            "quantity_lower: 1",
            "quantity_upper: 5",
            "design_life_years: 20",
        ):
            self.assertIn(default_pair, script)
        self.assertIn('spec[0] === "hydrogen_electrolyzers"', script)
        self.assertIn("hydrogenTankDefaultValues", script)
        for default_pair in (
            "cost: 100",
            "hydrogen_tank_capacity: 4000",
            "quantity_lower: 1",
            "quantity_upper: 5",
            "design_life_years: 20",
        ):
            self.assertIn(default_pair, script)
        self.assertIn('spec[0] === "hydrogen_tanks"', script)
        self.assertIn("fuelCellDefaultValues", script)
        for default_pair in (
            "cost: 200",
            "power_capacity: 500",
            "hydrogen_to_electric_efficiency: 1.5",
            "quantity_upper: 5",
            "design_life_years: 20",
        ):
            self.assertIn(default_pair, script)
        self.assertIn('spec[0] === "fuel_cells"', script)
        self.assertIn('["hydrogen_tanks", "储氢罐", ["name", "cost", "hydrogen_tank_capacity", "soc_upper", "soc_lower", "self_discharge_rate"', script)
        self.assertIn('spec[0] === "hydrogen_tanks" && field === "soc_upper"', script)
        self.assertIn("return 0.85", script)
        self.assertIn('spec[0] === "hydrogen_tanks" && field === "soc_lower"', script)
        self.assertIn("return 0.15", script)
        self.assertIn("collectHydrogenInitialSocWarning", script)
        self.assertIn("初始氢储SOC(0.0-1.0)必须位于储氢罐SOC范围", script)
        self.assertIn("collectStorageInitialSocWarning", script)
        self.assertIn("初始电储SOC(0.0-1.0)必须位于储能电池组SOC范围", script)

    def test_planning_device_tables_use_compact_cell_editing(self):
        html = (WEB_ROOT / "planning.html").read_text(encoding="utf-8")
        script = (WEB_ROOT / "assets" / "planning.js").read_text(encoding="utf-8")
        css = (WEB_ROOT / "assets" / "planning.css").read_text(encoding="utf-8")
        device_css = css.split("/* Compact editable device tables. */", 1)[1]

        self.assertIn("device-data-table", script)
        self.assertIn("device-parameter-table", script)
        self.assertIn("device-input", script)
        self.assertIn("device-cell", script)
        self.assertIn("device-cell-display", script)
        self.assertIn("device-heading-label", script)
        self.assertIn("deviceHeadingLabelHtml", script)
        self.assertIn("deviceTwoLineHeadingHtml", script)
        self.assertIn("deviceTableColumnCount", script)
        self.assertIn("deviceEmptyColumnsHtml", script)
        self.assertIn("sharedColumnCount", script)
        self.assertIn("device-empty-heading", script)
        self.assertIn("device-empty-cell", script)
        self.assertIn("<br>", script)
        self.assertIn("device-row", script)
        self.assertIn("deviceColumnClass", script)
        self.assertIn("device-sticky-col", script)
        self.assertIn("onDeviceRowContextMenu", script)
        self.assertIn("onDeviceCellPointerDown", script)
        self.assertIn("enterDeviceCellEdit", script)
        self.assertIn("exitDeviceCellEdit", script)
        self.assertIn("selectDeviceRow", script)
        self.assertIn("deviceRowContextMenu", script)
        self.assertIn("data-device-context-action", script)
        self.assertIn("data-device-cell-edit", script)
        self.assertIn('readonly="readonly"', script)
        self.assertIn('tabindex="-1"', script)
        self.assertIn("deleteDeviceRowByPosition", script)
        self.assertIn('addEventListener("contextmenu", onDeviceRowContextMenu)', script)
        self.assertIn('addEventListener("pointerdown", onDeviceCellPointerDown)', script)
        self.assertIn('addEventListener("blur", onDeviceInputBlur)', script)
        self.assertNotIn("device-action-col", script)
        self.assertNotIn("device-action-cell", script)
        self.assertIn("#devicesTab .device-data-table", css)
        self.assertIn(".device-parameter-table", css)
        self.assertIn(".device-input", css)
        self.assertIn(".device-cell-display", css)
        self.assertIn(".device-heading-label", css)
        self.assertIn(".device-empty-heading", css)
        self.assertIn(".device-empty-cell", css)
        self.assertIn(".device-cell.editing", css)
        self.assertIn(".device-row.selected .device-cell", css)
        self.assertIn(".device-input[readonly]", css)
        self.assertIn(".device-sticky-col", css)
        self.assertIn(".device-sticky-1", css)
        self.assertIn(".device-sticky-2", css)
        self.assertIn(".device-sticky-3", css)
        self.assertIn(".device-context-menu", css)
        self.assertIn("#devicesTab .device-card", device_css)
        self.assertIn("border: 0 !important", device_css)
        self.assertIn("background: transparent !important", device_css)
        self.assertIn("box-shadow: none !important", device_css)
        self.assertIn("overflow-x: hidden", device_css)
        self.assertIn("width: 100%", device_css)
        self.assertIn("min-width: 0", device_css)
        self.assertIn("table-layout: fixed", device_css)
        self.assertIn("white-space: normal", device_css)
        self.assertIn("overflow-wrap: anywhere", device_css)
        self.assertIn("-webkit-line-clamp: 2", device_css)
        self.assertIn("position: static", device_css)
        self.assertNotIn("width: max-content", device_css)
        self.assertNotIn("min-width: 156px", device_css)
        self.assertIn("border: 0", css)
        self.assertIn("outline: 0", css)
        self.assertIn("assets/planning.css?v=20260601-time-table-compact-3dp", html)
        self.assertIn("assets/planning.js?v=20260601-time-table-compact-3dp", html)

    def test_planning_scheme_actions_validate_duplicates_and_delete_current_scheme(self):
        script = (WEB_ROOT / "assets" / "planning.js").read_text(encoding="utf-8")

        self.assertIn("schemeNameExists", script)
        self.assertIn("normalizeSchemeName", script)
        self.assertIn("\\s\\u0000-\\u001f", script)
        self.assertIn("方案名称已存在", script)
        self.assertIn("是否覆盖", script)
        self.assertIn("payload.overwrite = true", script)
        self.assertIn("deleteScheme", script)
        self.assertIn("DELETE", script)
        self.assertIn("确认删除方案", script)
        self.assertIn("selectNextSchemeAfterDelete", script)

    def test_planning_overview_page_has_statistics_histograms_and_candidate_device_list(self):
        html = (WEB_ROOT / "planning.html").read_text(encoding="utf-8")
        script = (WEB_ROOT / "assets" / "planning.js").read_text(encoding="utf-8")

        self.assertIn("方案概览", html)
        self.assertNotIn("方案汇总", html)
        self.assertNotIn('id="schemeOverview"', html)
        self.assertNotIn("overviewHost", script)
        self.assertNotIn('id="summaryStats"', html)
        self.assertNotIn("时序统计量", html)
        self.assertIn('id="summaryCharts"', html)
        self.assertIn('id="quantitySummary"', html)
        self.assertIn("待选设备列表", html)
        self.assertIn('class="summary-tabs"', html)
        self.assertIn('data-summary-tab="charts"', html)
        self.assertIn('data-summary-tab="devices"', html)
        self.assertIn('data-summary-tab="planning"', html)
        self.assertIn('data-summary-panel="charts"', html)
        self.assertIn('data-summary-panel="devices"', html)
        self.assertIn('data-summary-panel="planning"', html)
        self.assertNotIn("设计容量约束", html)
        self.assertIn("bindSummaryTabs", script)
        self.assertIn("data-summary-panel", script)
        self.assertIn("renderSchemeSummary", script)
        self.assertNotIn("renderStatsTable", script)
        self.assertIn("renderCandidateDeviceTable", script)
        self.assertIn("capacityValue", script)
        self.assertIn("calculateSeriesStats", script)
        self.assertIn("buildHistogram", script)
        for name in ("风速", "太阳辐照", "环境温度", "负荷", "最大值", "最小值", "平均值", "数量下限(台)", "数量上限(台)"):
            self.assertIn(name, script)
        self.assertNotIn("formatFixed2", script)
        self.assertIn("formatInteger", script)
        self.assertNotIn(">频数</text>", script)
        self.assertIn("yAxis", script)
        self.assertIn("formatInteger(count)", script)
        self.assertIn("formatInteger(bin.count)", script)
        self.assertIn("histogram-bar", script)
        self.assertIn("data-bin-range", script)
        self.assertIn("data-bin-count", script)
        self.assertIn("onHistogramMouseMove", script)
        self.assertIn("横坐标", script)
        self.assertIn("纵坐标", script)

    def test_planning_overview_page_scrolls_when_content_overflows(self):
        css = (WEB_ROOT / "assets" / "planning.css").read_text(encoding="utf-8")

        self.assertIn(".summary-page", css)
        self.assertIn(".summary-switcher", css)
        self.assertIn(".summary-tab-panel.active", css)
        summary_switcher_css = css.split("#limitsTab .summary-section.summary-switcher", 1)[1].split("}", 1)[0]
        self.assertIn("padding: 0", summary_switcher_css)
        self.assertIn("margin-bottom: 0", summary_switcher_css)
        self.assertIn("border: 0 !important", summary_switcher_css)
        self.assertIn("background: transparent !important", summary_switcher_css)
        self.assertIn("box-shadow: none !important", summary_switcher_css)
        self.assertIn("flex: 1 1 auto", css)
        self.assertIn("overflow: auto", css)

    def test_planning_overview_charts_and_tables_adapt_to_available_height(self):
        css = (WEB_ROOT / "assets" / "planning.css").read_text(encoding="utf-8")
        script = (WEB_ROOT / "assets" / "planning.js").read_text(encoding="utf-8")
        summary_layout_script = script.split("function applyAdaptiveSummaryLayout()", 1)[1].split("function bindTimeResizeHandle()", 1)[0]

        self.assertIn("--summary-panel-height", css)
        self.assertIn("--summary-table-height", css)
        self.assertIn("--summary-histogram-grid-height", css)
        self.assertIn("--summary-histogram-svg-height", css)
        self.assertIn("height: var(--summary-table-height", css)
        self.assertIn("max-height: none", css)
        self.assertNotIn("min(50vh, 560px)", css)
        self.assertIn("height: var(--summary-histogram-grid-height", css)
        self.assertNotIn("min(52vh, 620px)", css)
        self.assertIn("grid-template-rows: repeat(2, minmax(0, 1fr))", css)
        self.assertIn("flex-direction: column", css)
        self.assertIn("applyAdaptiveSummaryLayout", script)
        self.assertIn("summaryTabs", script)
        self.assertIn("summary-histogram-grid-height", script)
        self.assertIn("summary-table-height", script)
        self.assertNotIn("Math.min(560", summary_layout_script)
        self.assertNotIn("Math.min(620", summary_layout_script)

    def test_planning_overview_table_rows_wrap_and_use_adaptive_height(self):
        css = (WEB_ROOT / "assets" / "planning.css").read_text(encoding="utf-8")

        self.assertIn("#quantitySummary table,", css)
        self.assertIn("#planningSummary table", css)
        self.assertIn("table-layout: fixed", css)
        self.assertIn("#quantitySummary th,", css)
        self.assertIn("#quantitySummary td,", css)
        self.assertIn("#planningSummary th,", css)
        self.assertIn("#planningSummary td", css)
        self.assertNotIn("#limitsTab .data-table table", css)
        self.assertIn("white-space: normal", css)
        self.assertIn("overflow-wrap: anywhere", css)
        self.assertIn("line-height: 1.45", css)
        self.assertIn("min-height: 48px", css)
        self.assertIn("vertical-align: top", css)

    def test_planning_layout_constrains_page_height_to_viewport(self):
        css = (WEB_ROOT / "assets" / "planning.css").read_text(encoding="utf-8")

        self.assertIn("height: 100vh", css)
        self.assertIn("overflow: hidden", css)
        self.assertIn("grid-template-rows: auto minmax(0, 1fr)", css)
        self.assertIn("#timeTab #timeTable", css)
        self.assertIn("height: var(--time-chart-height, clamp(180px, 28vh, 300px))", css)

    def test_planning_layout_adapts_chart_and_table_heights_to_viewport(self):
        css = (WEB_ROOT / "assets" / "planning.css").read_text(encoding="utf-8")
        script = (WEB_ROOT / "assets" / "planning.js").read_text(encoding="utf-8")

        self.assertIn("--panel-table-max-height", css)
        self.assertIn("--time-table-height", css)
        self.assertIn("height: var(--time-table-height", css)
        self.assertIn("max-height: var(--panel-table-max-height", css)
        self.assertIn("#timeTab.tab-panel.active", css)
        time_tab_active_css = css.split("#timeTab.tab-panel.active {", 1)[1].split("}", 1)[0]
        self.assertIn("overflow: hidden", time_tab_active_css)
        time_table_card_css = css.split("#timeTab .table-card {", 1)[1].split("}", 1)[0]
        self.assertIn("flex: 1 1 0", time_table_card_css)
        self.assertIn("syncAdaptiveLayout", script)
        self.assertIn("applyAdaptiveTimeSeriesLayout", script)
        self.assertIn("ResizeObserver", script)
        self.assertIn("timeChartManualHeight", script)
        self.assertIn("Math.round(tableHeight)", script)
        self.assertIn("Math.max(COLLAPSED_PANEL_SIZE, available - chartHeight)", script)
        self.assertNotIn("Math.min(620, Math.max(COLLAPSED_PANEL_SIZE, available - chartHeight))", script)

    def test_planning_frontend_defers_time_series_loading(self):
        script = (WEB_ROOT / "assets" / "planning.js").read_text(encoding="utf-8")

        self.assertIn("/overview", script)
        self.assertIn("/time-series", script)
        self.assertIn("ensureTimeSeriesLoaded", script)
        self.assertIn("ensureTimeSeriesForActiveTab", script)
        self.assertIn("shouldAutoLoadTimeSeries", script)
        self.assertIn("timeSeriesLoaded", script)
        self.assertIn("时序数据尚未加载", script)
        self.assertIn("进入时序数据或方案概览", script)
        self.assertIn("自动加载", script)
        self.assertNotIn("8760时序数据", script)
        self.assertNotIn("data-load-time-series", script)
        self.assertNotIn("点击加载", script)

    def test_planning_time_series_page_includes_temperature(self):
        html = (WEB_ROOT / "planning.html").read_text(encoding="utf-8")
        script = (WEB_ROOT / "assets" / "planning.js").read_text(encoding="utf-8")

        self.assertNotIn("8760点曲线板", html)
        for curve_key, label in (
            ("wind_speed", "风速"),
            ("solar_irradiance", "太阳辐照"),
            ("temperature", "环境温度"),
            ("load", "负荷"),
        ):
            self.assertIn(f'<button type="button" data-curve="{curve_key}"', html)
            self.assertIn(f">{label}</button>", html)
        self.assertNotIn('type="radio"', html)
        self.assertNotIn('name="timeCurve"', html)
        self.assertNotIn('type="checkbox" data-curve', html)
        self.assertIn('class="curve-switch-row"', html)
        self.assertIn('class="curve-button active"', html)
        self.assertIn('aria-pressed="true"', html)
        self.assertLess(html.index('class="weather-import-bar"'), html.index('class="curve-switch-row"'))
        self.assertIn("temperature", script)
        self.assertIn("环境温度", script)

    def test_planning_time_series_page_can_fetch_geocoded_weather_history(self):
        html = (WEB_ROOT / "planning.html").read_text(encoding="utf-8")
        script = (WEB_ROOT / "assets" / "planning.js").read_text(encoding="utf-8")
        css = (WEB_ROOT / "assets" / "planning.css").read_text(encoding="utf-8")

        for element_id in (
            "importTimeSeriesFile",
            "timeSeriesImportFile",
            "timeSeriesImportModal",
            "timeSeriesImportTitle",
            "openTimeSeriesImportFile",
            "timeSeriesImportHint",
            "timeSeriesImportChart",
            "timeSeriesImportResizeHandle",
            "timeSeriesImportPreview",
            "timeSeriesImportSummary",
            "confirmTimeSeriesImport",
            "closeTimeSeriesImport",
            "openCurveGenerator",
            "loadGeneratorModal",
            "curveGeneratorTabs",
            "loadGeneratorMaxLabel",
            "loadGeneratorMinLabel",
            "loadGeneratorAverageLabel",
            "loadGeneratorMode",
            "loadGeneratorMax",
            "loadGeneratorMin",
            "loadGeneratorAverage",
            "loadCurveImportFile",
            "generateLoadCurve",
            "saveLoadTemplate",
            "loadGeneratorPreview",
            "confirmLoadGenerator",
            "closeLoadGenerator",
            "weatherPlace",
            "geocodePlace",
            "openCoordinatePicker",
            "weatherLatitude",
            "weatherLongitude",
            "weatherYear",
            "fetchWeatherHistory",
            "weatherImportStatus",
            "mapPickerModal",
            "mapPickerCanvas",
            "weatherPreviewResizeHandle",
            "weatherPreviewStats",
            "closeMapPicker",
            "confirmMapPoint",
        ):
            self.assertIn(f'id="{element_id}"', html)
        self.assertIn('id="weatherLatitude" type="number" min="-90" max="90" step="0.001"', html)
        self.assertIn('id="weatherLongitude" type="number" min="-180" max="180" step="0.001"', html)
        self.assertIn('id="weatherYear" type="number" min="2001" step="1" value="2024"', html)
        self.assertIn('class="time-chart-toolbar"', html)
        time_chart_toolbar = html.split('<div class="time-chart-toolbar">', 1)[1].split('<svg id="timeChart"', 1)[0]
        self.assertLess(time_chart_toolbar.index('class="weather-import-bar"'), time_chart_toolbar.index('class="curve-switch-row"'))
        self.assertIn('accept=".csv,.xlsx"', html)
        weather_bar = html.split('<div class="weather-import-bar"', 1)[1].split("</div>", 1)[0]
        modal = html.split('<div id="mapPickerModal"', 1)[1].split('<div id="timeResizeHandle"', 1)[0]
        import_modal = html.split('<div id="timeSeriesImportModal"', 1)[1].split('<div id="loadGeneratorModal"', 1)[0]
        self.assertIn(">文件导入<", weather_bar)
        self.assertNotIn(">气象获取<", weather_bar)
        self.assertNotIn(">年份<", weather_bar)
        self.assertNotIn('id="weatherLatitude"', weather_bar)
        self.assertNotIn('id="weatherLongitude"', weather_bar)
        self.assertNotIn('id="weatherYear"', weather_bar)
        self.assertNotIn('id="fetchWeatherHistory"', weather_bar)
        self.assertNotIn('id="weatherImportStatus"', weather_bar)
        self.assertNotIn(">历史数据年<", weather_bar)
        self.assertIn(">曲线生成<", weather_bar)
        self.assertNotIn(">风速生成<", weather_bar)
        self.assertNotIn(">光照生成<", weather_bar)
        self.assertNotIn(">负荷生成<", weather_bar)
        self.assertIn(">坐标选择<", weather_bar)
        self.assertLess(weather_bar.index(">文件导入<"), weather_bar.index(">曲线生成<"))
        self.assertLess(weather_bar.index(">曲线生成<"), weather_bar.index(">坐标选择<"))
        for label in ("打开文件", "曲线预览", "风速", "太阳辐射", "环境温度", "负荷", "确认", "取消"):
            self.assertIn(label, import_modal)
        self.assertNotIn('id="cancelTimeSeriesImport"', import_modal)
        self.assertLess(import_modal.index('id="timeSeriesImportChart"'), import_modal.index('id="timeSeriesImportResizeHandle"'))
        self.assertLess(import_modal.index('id="timeSeriesImportResizeHandle"'), import_modal.index('id="timeSeriesImportPreview"'))
        for curve_key in ("wind_speed", "solar_irradiance", "temperature", "load"):
            self.assertIn(f'data-import-curve="{curve_key}"', import_modal)
        for label in ("随机曲线", "负荷最大值", "负荷最小值", "负荷平均值", "文件导入", "保存模板", "确认", "取消"):
            self.assertIn(label, html)
        load_generator_modal = html.split('<div id="loadGeneratorModal"', 1)[1].split('<div id="mapPickerModal"', 1)[0]
        self.assertIn('id="curveGeneratorTabs"', load_generator_modal)
        self.assertIn('class="curve-generator-tabs"', load_generator_modal)
        self.assertIn('data-curve-generator-target="wind_speed"', load_generator_modal)
        self.assertIn('data-curve-generator-target="solar_irradiance"', load_generator_modal)
        self.assertIn('data-curve-generator-target="load"', load_generator_modal)
        for label in ("曲线生成", "风速", "光照辐射", "负荷"):
            self.assertIn(label, load_generator_modal)
        self.assertIn('<option value="file">文件导入</option>', load_generator_modal)
        self.assertIn('id="loadCurveImportFile"', load_generator_modal)
        self.assertNotIn('id="importLoadCurveFile"', load_generator_modal)
        self.assertNotIn('id="cancelLoadGenerator"', load_generator_modal)
        self.assertNotIn(">文件导入</button>", load_generator_modal)
        for label in ("模式1", "模式2", "模式3"):
            self.assertNotIn(label, load_generator_modal)
        for label in ("高德地图", "OpenStreetMap"):
            self.assertIn(label, modal)
        self.assertNotIn("谷歌地图", modal)
        for provider in ('data-map-provider="amap"', 'data-map-provider="osm"'):
            self.assertIn(provider, modal)
        self.assertNotIn('data-map-provider="google"', modal)
        self.assertNotIn('id="weatherPlace"', weather_bar)
        self.assertNotIn('id="geocodePlace"', weather_bar)
        self.assertNotIn(">地图选点</button>", weather_bar)
        self.assertIn('id="weatherPlace"', modal)
        self.assertIn('id="geocodePlace"', modal)
        self.assertIn(">定位<", modal)
        self.assertNotIn(">获取坐标<", modal)
        self.assertIn('class="coordinate-weather-row"', modal)
        self.assertIn('id="weatherLatitude"', modal)
        self.assertIn('id="weatherLongitude"', modal)
        self.assertIn('id="weatherYear"', modal)
        self.assertIn('id="fetchWeatherHistory"', modal)
        self.assertIn('id="weatherImportStatus"', modal)
        self.assertIn(">年份<", modal)
        self.assertIn(">气象获取<", modal)
        self.assertIn('class="coordinate-control-row"', modal)
        self.assertLess(modal.index('id="weatherPlace"'), modal.index('id="weatherLatitude"'))
        self.assertLess(modal.index('id="weatherLatitude"'), modal.index('id="mapPickerCanvas"'))
        self.assertNotIn('class="map-picker-foot"', modal)
        self.assertIn('class="coordinate-map-hint"', modal)
        self.assertIn('class="weather-preview-panel"', modal)
        self.assertIn('class="weather-preview-legend"', modal)
        self.assertIn('id="weatherPreviewChart"', modal)
        self.assertIn('id="weatherPreviewResizeHandle"', modal)
        self.assertIn('id="weatherPreviewStats"', modal)
        self.assertNotIn('id="cancelMapPoint"', modal)
        self.assertIn('class="coordinate-header-actions"', modal)
        self.assertIn(">确认<", modal)
        self.assertIn(">取消<", modal)
        self.assertLess(modal.index('id="mapPickerCanvas"'), modal.index('id="weatherPreviewChart"'))
        self.assertLess(modal.index('id="mapPickerCanvas"'), modal.index('id="weatherPreviewResizeHandle"'))
        self.assertLess(modal.index('id="weatherPreviewResizeHandle"'), modal.index('id="weatherPreviewChart"'))
        self.assertLess(modal.index('id="confirmMapPoint"'), modal.index('id="closeMapPicker"'))
        self.assertLess(modal.index('id="closeMapPicker"'), modal.index('class="coordinate-control-row"'))
        for curve_key in ("wind_speed", "solar_irradiance", "temperature"):
            self.assertIn(f'data-weather-preview-curve="{curve_key}"', modal)
        self.assertIn("根据地名查找坐标", modal)
        self.assertIn("/api/planning/map-config", script)
        self.assertIn("/api/planning/geocode", script)
        self.assertIn("/api/planning/reverse-geocode", script)
        self.assertIn("/api/planning/weather-history", script)
        self.assertIn("/api/planning/time-series/import", script)
        self.assertIn("/api/planning/load-curve/generate", script)
        self.assertIn("/api/planning/load-curve/import", script)
        self.assertIn("/api/planning/time-series-curve/generate", script)
        self.assertIn("/api/planning/time-series-curve/import", script)
        self.assertIn("/api/planning/load-curve/templates", script)
        self.assertIn("selectMapProvider", script)
        self.assertNotIn("loadBaiduMapScript", script)
        self.assertNotIn("initBaiduMapPicker", script)
        self.assertNotIn("api.map.baidu.com/api?v=3.0", script)
        self.assertNotIn("window.BMap", script)
        self.assertNotIn("loadGoogleMapScript", script)
        self.assertNotIn("initGoogleMapPicker", script)
        self.assertNotIn("maps.googleapis.com/maps/api/js", script)
        self.assertNotIn("window.google", script)
        self.assertIn("importTimeSeriesFile", script)
        self.assertIn("openTimeSeriesImportModal", script)
        self.assertIn("openTimeSeriesImportFile", script)
        self.assertIn("onTimeSeriesImportFileChange", script)
        self.assertIn("renderTimeSeriesImportPreview", script)
        self.assertIn("renderTimeSeriesImportChart", script)
        self.assertIn("bindTimeSeriesImportResizeHandle", script)
        self.assertIn("toggleTimeSeriesImportCurve", script)
        self.assertIn("timeSeriesImportVisibleCurves", script)
        self.assertIn("timeSeriesImportManualChartHeight", script)
        self.assertIn("timeSeriesImportSeries", script)
        self.assertIn("confirmImportedTimeSeries", script)
        self.assertIn("cancelTimeSeriesImport", script)
        self.assertIn("pendingTimeSeriesImport", script)
        self.assertIn("导入曲线已保存到后台", script)
        self.assertIn("isTimeSeriesImportWarning", script)
        self.assertIn('setTimeSeriesImportHint(result.message || "导入文件解析成功，请确认后保存。", level)', script)
        self.assertIn('hint.classList.toggle("warning", level === "warning")', script)
        self.assertIn("#timeSeriesImportHint.warning", css)
        self.assertIn("#loadGeneratorHint.warning", css)
        self.assertIn(".curve-generator-tabs", css)
        self.assertIn(".curve-generator-tab", css)
        self.assertIn(".curve-generator-tab.active", css)
        self.assertIn("openCurveGenerator", script)
        self.assertIn("selectCurveGeneratorTarget", script)
        self.assertIn("syncCurveGeneratorTabs", script)
        self.assertIn("resetCurveGeneratorWorkingState", script)
        self.assertIn("refreshCurveGeneratorForTarget", script)
        self.assertIn("curveGeneratorSpecs", script)
        self.assertIn("curveGeneratorTarget", script)
        self.assertIn("generateLoadCurve", script)
        self.assertIn("onLoadGeneratorModeChange", script)
        self.assertIn("loadLoadGeneratorModeSource", script)
        self.assertIn("loadGeneratorSourceCurve", script)
        self.assertIn("loadGeneratorSourceName", script)
        self.assertIn("loadGeneratorPreviewSourceRows", script)
        self.assertIn("importLoadCurveFile", script)
        self.assertIn("onLoadCurveImportFileChange", script)
        self.assertIn('loadCurveImportFile.click()', script)
        self.assertIn('if (mode === "file")', script)
        self.assertIn("source_load_curve", script)
        self.assertIn("raw: true", script)
        self.assertIn("负荷文件已导入为原始曲线，请点击生成负荷曲线。", script)
        self.assertIn("原始负荷曲线已载入，请点击生成负荷曲线。", script)
        self.assertNotIn('document.getElementById("importLoadCurveFile").addEventListener', script)
        self.assertIn("loadLoadCurveTemplates", script)
        self.assertIn("saveLoadTemplate", script)
        self.assertIn("saveLoadTemplateRequest(name, rows, true)", script)
        self.assertIn("模板名称已存在", script)
        self.assertIn("renderLoadGeneratorPreview", script)
        self.assertIn("confirmGeneratedLoadCurve", script)
        self.assertIn("cancelLoadGenerator", script)
        self.assertIn("pendingLoadCurve", script)
        self.assertIn("originalLoadCurve", script)
        self.assertIn("修改前", script)
        self.assertIn("修改后", script)
        self.assertIn("applyGeneratedLoadCurve", script)
        self.assertIn("content_base64", script)
        self.assertIn("arrayBufferToBase64", script)
        self.assertIn("导入失败", script)
        self.assertIn("[spec.key]: curve[spec.key]", script)
        self.assertIn("负荷曲线已生成", script)
        self.assertIn("openCoordinatePicker", script)
        self.assertIn("initAmapTilePicker", script)
        self.assertIn("renderAmapTileLayer", script)
        self.assertIn("webrd0${server}.is.autonavi.com", script)
        self.assertIn("osmTileUrl", script)
        self.assertIn("OSM_TILE_PROVIDERS", script)
        self.assertIn("tile.openstreetmap.de", script)
        self.assertIn("a.tile.openstreetmap.fr/hot", script)
        self.assertIn("b.basemaps.cartocdn.com/light_all", script)
        self.assertIn("tile.openstreetmap.org", script)
        self.assertIn("data-fallback-srcs", script)
        self.assertIn("switchAmapTileToGlobalFallback", script)
        self.assertIn("OpenStreetMap 全球底图", script)
        self.assertIn("lngLatToWebMercatorPixel", script)
        self.assertIn("webMercatorPixelToLngLat", script)
        self.assertIn("setMapPoint", script)
        self.assertIn("reverseGeocodePoint", script)
        self.assertIn("setWeatherPlaceFromReverseGeocode", script)
        self.assertIn("mapReverseGeocodeToken", script)
        self.assertIn("syncMapPointFromInputs", script)
        self.assertIn("pendingWeatherRows", script)
        self.assertIn("weatherPreviewVisibleCurves", script)
        self.assertIn("renderWeatherPreviewChart", script)
        self.assertIn("renderWeatherPreviewStats", script)
        self.assertIn("calculateSeriesStats(rows, key)", script)
        self.assertIn("toggleWeatherPreviewCurve", script)
        self.assertIn("applyPendingWeatherHistory", script)
        self.assertIn("bindWeatherPreviewResizeHandle", script)
        self.assertIn("weatherPreviewManualHeight", script)
        self.assertIn("--weather-preview-panel-height", script)
        self.assertIn("WEATHER_COORDINATE_STORAGE_KEY", script)
        self.assertIn("powerPlanWeatherCoordinate", script)
        self.assertIn("rememberWeatherCoordinate", script)
        self.assertIn("restoreWeatherCoordinate", script)
        self.assertIn("place: cleanPlace", script)
        self.assertIn('document.getElementById("weatherPlace").value = place', script)
        self.assertIn("localStorage.setItem(WEATHER_COORDINATE_STORAGE_KEY", script)
        self.assertIn("localStorage.getItem(WEATHER_COORDINATE_STORAGE_KEY", script)
        self.assertIn("formatCoordinate", script)
        self.assertIn("number.toFixed(3)", script)
        self.assertNotIn("toFixed(6)", script)
        self.assertIn('setMapPoint(result.latitude, result.longitude, "geocode", result)', script)
        self.assertIn("geocodeHintLabel", script)
        self.assertIn("高德定位", script)
        self.assertIn('state.mapInstance.setZoom(11)', script)
        self.assertIn("未配置${mapProviderLabel(state.mapProvider)} Key", script)
        self.assertIn("geocodePlace", script)
        self.assertIn("正在定位...", script)
        self.assertIn("正在解析地点...", script)
        self.assertIn("地点已更新", script)
        self.assertNotIn("正在获取坐标...", script)
        self.assertIn("fetchWeatherHistory", script)
        self.assertIn("validateWeatherInputs", script)
        self.assertIn("历史数据年必须", script)
        self.assertIn("rows.length !== 8760", script)
        self.assertIn("气象数据已预览，请确认后更新主页面", script)
        self.assertIn("保存气象坐标失败", (WEB_ROOT / "assets" / "i18n.js").read_text(encoding="utf-8"))
        self.assertIn("wind_speed: weather.wind_speed", script)
        self.assertIn("solar_irradiance: weather.solar_irradiance", script)
        self.assertIn("temperature: weather.temperature", script)
        self.assertNotIn("load: weather.load", script)
        self.assertIn("气象已更新", script)
        self.assertIn("纬度：", script)
        self.assertIn("经度：", script)
        self.assertIn("latitude.toFixed(3)", script)
        self.assertIn("longitude.toFixed(3)", script)
        self.assertNotIn("风速、太阳辐照和温度数据", script)
        self.assertIn(".time-chart-toolbar", css)
        time_chart_toolbar_css = css.split(".time-chart-toolbar {", 1)[1].split("}", 1)[0]
        self.assertIn("justify-content: space-between", time_chart_toolbar_css)
        self.assertIn("overflow-x: auto", time_chart_toolbar_css)
        self.assertIn(".weather-import-bar", css)
        weather_import_bar_css = css.split(".weather-import-bar {", 1)[1].split("}", 1)[0]
        self.assertIn("justify-content: flex-start", weather_import_bar_css)
        self.assertIn("flex-wrap: nowrap", weather_import_bar_css)
        self.assertIn("flex: 1 1 auto", weather_import_bar_css)
        self.assertIn("overflow-x: auto", weather_import_bar_css)
        compact_weather_input_css = css.split("#weatherLatitude,", 1)[1].split("}", 1)[0]
        self.assertIn("#weatherLongitude", compact_weather_input_css)
        self.assertIn("#weatherYear", compact_weather_input_css)
        self.assertIn("width: 94px", compact_weather_input_css)
        curve_switch_row_css = css.split(".curve-switch-row {", 1)[1].split("}", 1)[0]
        self.assertIn("justify-content: flex-end", curve_switch_row_css)
        self.assertIn("margin: 0 0 0 auto", curve_switch_row_css)
        self.assertIn(".curve-switch-row .curve-buttons", css)
        self.assertIn(".curve-switch-row #timeChartRange", css)
        curve_switch_nowrap_css = css.split(".curve-switch-row .curve-buttons,", 1)[1].split("}", 1)[0]
        self.assertIn("flex-wrap: nowrap", curve_switch_nowrap_css)
        self.assertIn(".time-series-import-dialog", css)
        self.assertIn(".time-series-import-toolbar", css)
        self.assertIn(".time-series-import-chart-panel", css)
        self.assertIn(".time-series-import-chart", css)
        self.assertIn(".time-series-import-resize-handle", css)
        self.assertIn(".time-series-import-curve-toggle", css)
        self.assertIn(".time-series-import-preview", css)
        self.assertIn(".map-provider-tabs", css)
        self.assertIn(".map-provider-tab", css)
        self.assertIn("body.modal-open", css)
        self.assertIn(".load-generator-dialog", css)
        self.assertIn(".load-generator-preview", css)
        self.assertIn("background: var(--hud-panel-strong)", css)
        self.assertIn("border-color: rgba(20, 190, 255, 0.42)", css)
        self.assertIn("linear-gradient(rgba(33, 213, 255, 0.07)", css)
        self.assertIn(".coordinate-search-row", css)
        self.assertIn(".coordinate-control-row", css)
        self.assertIn(".coordinate-weather-row", css)
        self.assertIn(".coordinate-map-hint", css)
        self.assertIn(".weather-preview-panel", css)
        self.assertIn(".weather-preview-chart", css)
        self.assertIn(".weather-preview-legend", css)
        self.assertIn(".weather-preview-resize-handle", css)
        self.assertIn(".weather-preview-stats", css)
        self.assertIn(".weather-preview-stat-item", css)
        self.assertIn(".coordinate-header-actions", css)
        self.assertIn(".weather-import-status.error", css)
        self.assertIn(".map-picker-modal", css)
        self.assertIn(".map-picker-canvas", css)
        self.assertIn('body[data-home-theme]:not([data-home-theme="default"]) .map-picker-modal', css)
        self.assertIn('body[data-home-theme]:not([data-home-theme="default"]) .map-picker-dialog', css)
        self.assertIn('body[data-home-theme]:not([data-home-theme="default"]) .load-generator-dialog', css)
        self.assertIn('body[data-home-theme]:not([data-home-theme="default"]) .time-series-import-dialog', css)
        self.assertIn('body[data-home-theme]:not([data-home-theme="default"]) .map-picker-head', css)
        self.assertIn('body[data-home-theme]:not([data-home-theme="default"]) .coordinate-control-row', css)
        self.assertIn('body[data-home-theme]:not([data-home-theme="default"]) .weather-preview-panel', css)
        self.assertIn('body[data-home-theme]:not([data-home-theme="default"]) .map-provider-tabs', css)
        self.assertIn('body[data-home-theme]:not([data-home-theme="default"]) .time-series-import-toolbar', css)
        self.assertIn('body[data-home-theme]:not([data-home-theme="default"]) .load-generator-grid', css)
        self.assertIn('body[data-home-theme]:not([data-home-theme="default"]) .map-picker-head button', css)
        self.assertIn('body[data-home-theme]:not([data-home-theme="default"]) .coordinate-header-actions button', css)
        self.assertIn('body[data-home-theme]:not([data-home-theme="default"]) .coordinate-weather-row button', css)
        self.assertIn('body[data-home-theme]:not([data-home-theme="default"]) .weather-preview-legend button', css)
        self.assertIn('body[data-home-theme]:not([data-home-theme="default"]) .load-generator-grid input', css)
        self.assertIn('body[data-home-theme]:not([data-home-theme="default"]) .coordinate-weather-row input', css)
        theme_modal_dialog_css = css.split('body[data-home-theme]:not([data-home-theme="default"]) .time-series-import-dialog {', 1)[1].split("}", 1)[0]
        self.assertIn("background: var(--theme-panel-bg)", theme_modal_dialog_css)
        self.assertIn("color: var(--theme-text)", theme_modal_dialog_css)
        theme_modal_button_css = css.split('body[data-home-theme]:not([data-home-theme="default"]) .map-picker-foot button {', 1)[1].split("}", 1)[0]
        self.assertIn("background: var(--theme-control-bg)", theme_modal_button_css)
        self.assertIn("color: var(--theme-control-text)", theme_modal_button_css)

    def test_planning_time_series_table_uses_month_tabs(self):
        html = (WEB_ROOT / "planning.html").read_text(encoding="utf-8")
        script = (WEB_ROOT / "assets" / "planning.js").read_text(encoding="utf-8")
        css = (WEB_ROOT / "assets" / "planning.css").read_text(encoding="utf-8")

        self.assertIn('id="monthTabs"', html)
        self.assertIn('class="time-table-toolbar"', html)
        self.assertNotIn('id="prevPage"', html)
        self.assertNotIn('id="nextPage"', html)
        self.assertNotIn("<h2>小时级数据</h2>", html)
        toolbar = html.split('<div class="time-table-toolbar">', 1)[1].split('<div id="timeTable"', 1)[0]
        self.assertLess(toolbar.index('id="monthTabs"'), toolbar.index('id="pageInfo"'))
        self.assertIn(".time-table-toolbar", css)
        self.assertIn("flex-wrap: nowrap", css)
        self.assertIn("justify-content: flex-start", css)
        self.assertIn("margin-left: auto", css)
        self.assertIn("text-align: right", css)
        self.assertIn("monthRanges", script)
        self.assertIn("renderMonthTabs", script)
        self.assertIn("1月", script)
        self.assertIn("12月", script)
        self.assertIn("timeCellHtml", script)
        self.assertIn("time-cell-display", script)
        self.assertIn("time-cell-input", script)
        self.assertIn("enterTimeCellEdit", script)
        self.assertIn("exitTimeCellEdit", script)
        self.assertIn("onTimeCellPointerDown", script)
        self.assertIn("onTimeInputFocusOut", script)
        self.assertIn("finalizeTimeInput", script)
        self.assertIn("timeSeriesValueKeys", script)
        self.assertIn("normalizeTimeSeriesRows", script)
        self.assertIn("normalizeTimeSeriesCellValue", script)
        self.assertIn("formatTimeSeriesCellValue", script)
        self.assertIn("roundTimeSeriesValue", script)
        self.assertIn(".toFixed(3)", script)
        self.assertIn("normalizeTimeSeriesRows(payload.time_series)", script)
        self.assertIn('readonly="readonly"', script)
        self.assertIn('tabindex="-1"', script)
        self.assertIn(".time-cell-display", css)
        self.assertIn(".time-cell-input", css)
        self.assertIn("display: none", css)
        self.assertIn(".time-cell.editing", css)
        self.assertIn(".time-cell.editing .time-cell-input", css)

    def test_planning_time_series_chart_height_has_resizable_splitter(self):
        html = (WEB_ROOT / "planning.html").read_text(encoding="utf-8")
        script = (WEB_ROOT / "assets" / "planning.js").read_text(encoding="utf-8")
        css = (WEB_ROOT / "assets" / "planning.css").read_text(encoding="utf-8")

        self.assertIn('id="timeResizeHandle"', html)
        self.assertIn('role="separator"', html)
        self.assertIn("调整时序图高度", html)
        self.assertLess(html.index('id="timeChart"'), html.index('id="timeResizeHandle"'))
        self.assertLess(html.index('id="timeResizeHandle"'), html.index('id="timeTable"'))
        self.assertIn("bindTimeResizeHandle", script)
        self.assertIn("pointerdown", script)
        self.assertIn("--time-chart-height", script)
        self.assertIn("svg.clientHeight", script)
        self.assertIn(".time-resize-handle", css)
        self.assertIn("cursor: row-resize", css)
        self.assertIn("height: var(--time-chart-height", css)

    def test_planning_time_series_input_refreshes_visible_chart(self):
        script = (WEB_ROOT / "assets" / "planning.js").read_text(encoding="utf-8")
        css = (WEB_ROOT / "assets" / "planning.css").read_text(encoding="utf-8")

        self.assertIn("function onTimeInput", script)
        self.assertIn("renderChart();", script)
        self.assertIn("selectedCurveSpec", script)
        self.assertIn('[data-curve][aria-pressed="true"]', script)
        self.assertIn("selectCurve", script)
        self.assertNotIn("时间（月）", script)
        self.assertIn("monthRanges", script)
        self.assertIn("yTicks", script)
        self.assertIn('stroke="${color}"', script)
        self.assertIn("onChartMouseMove", script)
        self.assertIn("hideChartCursor", script)
        self.assertIn('id="chartCursor"', script)
        self.assertIn('id="chartCursorLine"', script)
        self.assertIn('id="chartCursorPoint"', script)
        self.assertIn("mousemove", script)
        self.assertIn("mouseleave", script)
        self.assertIn(".chart-cursor", css)
        self.assertIn("positionFloatingTipInRect", script)
        self.assertIn("bounds.right - tipWidth - margin", script)
        self.assertIn("bounds.bottom - tipHeight - margin", script)
        self.assertIn("tip.offsetWidth", script)
        self.assertIn("parentRect.left", script)
        self.assertIn(".chart-tip", css)
        self.assertIn("position: absolute", css)
        self.assertIn("z-index: 40", css)
        self.assertIn("white-space: nowrap", css)

    def test_planning_time_series_chart_supports_drag_value_editing(self):
        script = (WEB_ROOT / "assets" / "planning.js").read_text(encoding="utf-8")
        css = (WEB_ROOT / "assets" / "planning.css").read_text(encoding="utf-8")
        i18n_script = (WEB_ROOT / "assets" / "i18n.js").read_text(encoding="utf-8")

        self.assertIn("chartDrag", script)
        self.assertIn('timeChart.addEventListener("pointerdown", startChartValueDrag)', script)
        self.assertIn('window.addEventListener("pointermove", onChartValueDragMove)', script)
        self.assertIn("function startChartValueDrag", script)
        self.assertIn("function onChartValueDragMove", script)
        self.assertIn("function endChartValueDrag", script)
        self.assertIn("function applyChartValueEdit", script)
        self.assertIn("function chartValueFromPointer", script)
        self.assertIn("function interpolatedCurveEditPoints", script)
        self.assertIn("interpolatedCurveEditPoints(state.chartDrag?.lastPoint, point)", script)
        self.assertIn("state.chartDrag.lastPoint = point", script)
        self.assertIn("setPointerCapture", script)
        self.assertIn("state.payload.time_series[absoluteIndex][meta.curveKey]", script)
        self.assertIn('setWeatherImportStatus("曲线已修改，请保存方案", "ok")', script)
        self.assertIn("renderChart();", script)
        self.assertIn("renderTimeTable();", script)
        self.assertIn("renderLimitSummary();", script)
        self.assertIn("renderSummary();", script)
        self.assertIn("cursor: crosshair", css)
        self.assertIn("touch-action: none", css)
        self.assertIn(".time-chart.editing", css)
        self.assertIn("Curve updated. Please save the scenario", i18n_script)

    def test_planning_load_generator_preview_supports_drag_value_editing(self):
        script = (WEB_ROOT / "assets" / "planning.js").read_text(encoding="utf-8")
        css = (WEB_ROOT / "assets" / "planning.css").read_text(encoding="utf-8")
        i18n_script = (WEB_ROOT / "assets" / "i18n.js").read_text(encoding="utf-8")

        self.assertIn("loadPreviewMeta", script)
        self.assertIn("loadPreviewDrag", script)
        self.assertIn('loadGeneratorPreview.addEventListener("pointerdown", startLoadPreviewValueDrag)', script)
        self.assertIn('window.addEventListener("pointermove", onLoadPreviewValueDragMove)', script)
        self.assertIn("function startLoadPreviewValueDrag", script)
        self.assertIn("function onLoadPreviewValueDragMove", script)
        self.assertIn("function endLoadPreviewValueDrag", script)
        self.assertIn("function applyLoadPreviewValueEdit", script)
        self.assertIn("function loadPreviewValueFromPointer", script)
        self.assertIn("interpolatedCurveEditPoints(state.loadPreviewDrag?.lastPoint, point)", script)
        self.assertIn("state.loadPreviewDrag.lastPoint = point", script)
        self.assertIn("state.pendingLoadCurve[pointIndex][spec.key]", script)
        self.assertIn("spec.adjustedMessage", script)
        self.assertIn("loadGeneratorPreview.classList.add(\"editing\")", script)
        self.assertIn(".load-generator-preview.editing", css)
        self.assertIn("Load curve adjusted. Please review the preview and confirm.", i18n_script)

    def test_planning_time_series_import_preview_supports_chart_and_table_value_editing(self):
        script = (WEB_ROOT / "assets" / "planning.js").read_text(encoding="utf-8")
        css = (WEB_ROOT / "assets" / "planning.css").read_text(encoding="utf-8")
        i18n_script = (WEB_ROOT / "assets" / "i18n.js").read_text(encoding="utf-8")

        self.assertIn("timeSeriesImportChartMeta", script)
        self.assertIn("timeSeriesImportDrag", script)
        self.assertIn('timeSeriesImportChart.addEventListener("pointerdown", startTimeSeriesImportValueDrag)', script)
        self.assertIn('window.addEventListener("pointermove", onTimeSeriesImportValueDragMove)', script)
        self.assertIn("function startTimeSeriesImportValueDrag", script)
        self.assertIn("function onTimeSeriesImportValueDragMove", script)
        self.assertIn("function endTimeSeriesImportValueDrag", script)
        self.assertIn("function applyTimeSeriesImportValueEdit", script)
        self.assertIn("function timeSeriesImportValueFromPointer", script)
        self.assertIn("interpolatedCurveEditPoints(state.timeSeriesImportDrag?.lastPoint, point)", script)
        self.assertIn("state.timeSeriesImportDrag.lastPoint = point", script)
        self.assertIn("function onTimeSeriesImportInput", script)
        self.assertIn("function updateTimeSeriesImportCell", script)
        self.assertIn("formatTimeSeriesCellValue(key, row[key])", script)
        self.assertIn("formatTimeSeriesCellValue(key, nextValue)", script)
        self.assertIn("formatTimeSeriesCellValue(key, value)", script)
        self.assertIn("data-time-series-import-index", script)
        self.assertIn("data-time-series-import-key", script)
        self.assertIn('type="number"', script)
        self.assertIn('step="any"', script)
        self.assertIn('inputmode="decimal"', script)
        self.assertIn("state.pendingTimeSeriesImport[pointIndex][curveKey]", script)
        self.assertIn("state.pendingTimeSeriesImport[index][key]", script)
        self.assertIn("renderTimeSeriesImportChart(state.pendingTimeSeriesImport || [])", script)
        self.assertIn('setTimeSeriesImportHint("导入曲线已调整，请确认后保存。", "ok")', script)
        self.assertIn("timeSeriesImportChart.classList.add(\"editing\")", script)
        self.assertIn("timeSeriesImportChartMeta = null", script)
        self.assertIn("scalesByKey", script)
        self.assertIn(".time-series-import-chart.editing", css)
        self.assertIn("Imported curves adjusted. Please confirm to save.", i18n_script)

    def test_planning_device_fields_follow_latest_parameter_names(self):
        script = (WEB_ROOT / "assets" / "planning.js").read_text(encoding="utf-8")

        self.assertNotIn("design_capacity_lower", script)
        self.assertNotIn("design_capacity_upper", script)
        self.assertNotIn("generation_efficiency", script)
        self.assertNotIn("发电效率(0-1.0)", script)
        self.assertIn("氢-电效率(kWh/Nm3)", script)
        self.assertIn("电-氢效率(Nm3/kWh)", script)
        self.assertIn("切入风速(m/s)", script)
        self.assertIn("额定风速(m/s)", script)
        self.assertIn("切出风速(m/s)", script)
        self.assertIn("成本(万元/台)", script)
        self.assertIn("油耗率(kg/kWh)", script)
        self.assertIn("惯量常数H(s)", script)
        self.assertIn("一次调频系数K", script)
        self.assertIn("阻尼系数D", script)
        self.assertIn("调速时间常数T(s)", script)
        self.assertIn("等效惯量常数H(s)", script)
        self.assertIn("等效一次调频系数K", script)
        self.assertIn("等效阻尼系数D", script)
        self.assertIn("功率上限(kW)", script)
        self.assertIn("功率下限(kW)", script)
        self.assertIn('capacity: "容量(kW)"', script)
        self.assertIn('power_capacity: "容量(kW)"', script)
        self.assertIn('battery_capacity: "容量(kWh)"', script)
        self.assertIn('hydrogen_tank_capacity: "容量(Nm3)"', script)
        self.assertNotIn('capacity: "功率容量(kW)"', script)
        self.assertNotIn('power_capacity: "功率容量"', script)
        self.assertNotIn('battery_capacity: "电池容量"', script)
        self.assertNotIn('hydrogen_tank_capacity: "氢储容量(Nm3)"', script)
        self.assertNotIn('hydrogen_tank_capacity: "储氢罐容量"', script)

    def test_planning_frontend_device_fields_match_backend_sheet_specs(self):
        script = (WEB_ROOT / "assets" / "planning.js").read_text(encoding="utf-8")
        frontend_specs = {}
        for match in re.finditer(r'^\s*\["([^"]+)",\s*"[^"]+",\s*(\[[^\]]*\])\],?\s*$', script, re.MULTILINE):
            key = match.group(1)
            if key in server.planning_store.SHEET_SPECS:
                frontend_specs[key] = ast.literal_eval(match.group(2))

        for key, (_, backend_fields) in server.planning_store.SHEET_SPECS.items():
            if key in {"time_series", "planning_parameters"}:
                continue
            self.assertEqual(frontend_specs.get(key), backend_fields, key)

    def test_planning_frontend_parameters_match_backend_sheet_specs(self):
        script = (WEB_ROOT / "assets" / "planning.js").read_text(encoding="utf-8")
        parameter_block = script.split("const planningParameterSpecs = [", 1)[1].split("];", 1)[0]
        frontend_parameters = [
            match.group(1)
            for match in re.finditer(r'^\s*\["([^"]+)"', parameter_block, re.MULTILINE)
        ]

        self.assertEqual(frontend_parameters, server.planning_store.SHEET_SPECS["planning_parameters"][1])

    def test_planning_frontend_exposes_modeling_interface_and_solver_capabilities(self):
        script = (WEB_ROOT / "assets" / "planning.js").read_text(encoding="utf-8")
        html = (WEB_ROOT / "planning.html").read_text(encoding="utf-8")

        self.assertIn('["modeling_interface", "建模接口方式", "select"', script)
        self.assertIn('["cvxpy", "CVXPY通用接口"]', script)
        self.assertIn('["native", "优化求解器原生接口"]', script)
        self.assertIn('["copt", "COPT"]', script)
        self.assertIn('["mindopt", "MindOpt"]', script)
        self.assertIn('/api/planning/solver-capabilities', script)
        self.assertIn("assets/planning.js?v=20260817-modeling-interface", html)

    def test_planning_solver_capabilities_endpoint(self):
        with patch.object(server.milp_solver, "solver_capabilities", return_value={"solvers": {"scipy": {"native": True, "cvxpy": True}}}):
            status, headers, body = server.handle_planning_api_path("/api/planning/solver-capabilities", "GET", b"")

        self.assertEqual(status, 200)
        self.assertEqual(json.loads(body.decode("utf-8"))["solvers"]["scipy"]["native"], True)

    def test_planning_device_cost_columns_follow_name(self):
        script = (WEB_ROOT / "assets" / "planning.js").read_text(encoding="utf-8")

        for line in script.splitlines():
            if line.strip().startswith("[") and "quantity_upper" in line and "planningParameterSpecs" not in line:
                fields = [item.strip().strip('"') for item in line.split("[", 2)[2].split("]", 1)[0].split(",")]
                self.assertEqual(fields[0], "name")
                self.assertEqual(fields[1], "cost")

    def test_planning_device_tables_include_design_life_column(self):
        script = (WEB_ROOT / "assets" / "planning.js").read_text(encoding="utf-8")

        device_spec_lines = [
            line
            for line in script.splitlines()
            if line.strip().startswith("[") and "quantity_upper" in line and "planningParameterSpecs" not in line
        ]
        self.assertEqual(len(device_spec_lines), 8)
        for line in device_spec_lines:
            self.assertIn("design_life_years", line)
            fields = line.split("[", 2)[2].split("]", 1)[0]
            self.assertEqual(fields.rsplit('"', 2)[1], "design_life_years")
        self.assertIn("design_life_years: \"设计年限(年）\"", script)
        self.assertIn("design_life_years: 20", script)

    def test_planning_hydrogen_electrolyzer_table_has_power_lower(self):
        script = (WEB_ROOT / "assets" / "planning.js").read_text(encoding="utf-8")

        electrolyzer_line = next(
            line for line in script.splitlines() if line.strip().startswith('["hydrogen_electrolyzers"')
        )
        self.assertIn("power_lower", electrolyzer_line)
        self.assertLess(electrolyzer_line.index("cost"), electrolyzer_line.index("power_capacity"))
        self.assertLess(electrolyzer_line.index("power_capacity"), electrolyzer_line.index("power_lower"))
        self.assertIn('power_lower: "功率下限(kW)"', script)

    def test_static_path_resolves_index(self):
        resolved = server.resolve_static_path("/")

        self.assertEqual(resolved.name, "index.html")
        self.assertTrue(resolved.exists())

    def test_planning_assets_are_cache_busted_and_static_assets_are_browser_cacheable(self):
        html = (WEB_ROOT / "planning.html").read_text(encoding="utf-8")

        self.assertIn("assets/planning.css?v=", html)
        self.assertIn("assets/planning.js?v=", html)
        self.assertEqual(server.resolve_static_path("/assets/planning.js?v=test").name, "planning.js")
        server_text = (WEB_ROOT / "server.py").read_text(encoding="utf-8")
        self.assertIn("STATIC_BROWSER_CACHE_SUFFIXES", server_text)
        self.assertIn('".css", ".js"', server_text)
        self.assertIn('".png"', server_text)
        self.assertIn("STATIC_ASSET_CACHE_CONTROL", server_text)

    def test_interactive_pages_restore_last_page_state(self):
        page_state_script = (WEB_ROOT / "assets" / "page_state.js").read_text(encoding="utf-8")
        self.assertIn("PowerPlanPageState", page_state_script)
        self.assertIn("localStorage", page_state_script)
        self.assertIn("read", page_state_script)
        self.assertIn("write", page_state_script)
        self.assertIn("patch", page_state_script)
        self.assertIn("isFiniteNumber", page_state_script)
        self.assertIn("number", page_state_script)

        interactive_pages = (
            "planning.html",
            "optimize.html",
            "evaluation.html",
            "frequency.html",
            "comparison.html",
            "tasks.html",
        )
        for page_name in interactive_pages:
            with self.subTest(page=page_name):
                html = (WEB_ROOT / page_name).read_text(encoding="utf-8")
                self.assertIn("assets/page_state.js?v=", html)

        page_scripts = {
            "planning.js": (
                "PLANNING_PAGE_STATE_KEY",
                "restorePlanningPageState",
                "rememberPlanningPageState",
                "activeTab",
                "summaryTab",
                "currentScheme",
                "timeChartRange",
                "month",
                "activePlanningParameterGroup",
                "visibleDevices",
                "loadGeneratorMode",
            ),
            "optimize.js": (
                "OPTIMIZATION_PAGE_STATE_KEY",
                "restoreOptimizationPageState",
                "rememberOptimizationPageState",
                "currentScheme",
                "activeResultTab",
                "optimizationSchemeRailHeight",
                "axisRanges",
            ),
            "evaluation.js": (
                "EVALUATION_PAGE_STATE_KEY",
                "restoreEvaluationPageState",
                "rememberEvaluationPageState",
                "currentScheme",
                "selectedResultFile",
                "activeResultTab",
                "collapsedSchemes",
                "evaluationSchemeRailHeight",
                "evaluationResultRailWidth",
                "axisRanges",
            ),
            "frequency.js": (
                "FREQUENCY_PAGE_STATE_KEY",
                "restoreFrequencyPageState",
                "rememberFrequencyPageState",
                "currentScheme",
                "selectedResultFile",
                "activeResultTab",
                "collapsedSchemes",
                "frequencyTimeSelection",
                "schemeRailHeight",
                "axisRanges",
            ),
            "comparison.js": (
                "COMPARISON_PAGE_STATE_KEY",
                "restoreComparisonPageState",
                "rememberComparisonPageState",
                "tabs",
                "activeTabId",
                "tableHeight",
                "tableColumnWidths",
                "axisRanges",
                "selectedCurves",
            ),
            "tasks.js": (
                "TASKS_PAGE_STATE_KEY",
                "restoreTasksPageState",
                "rememberTasksPageState",
                "activeTaskType",
                "evaluationSchemeFilter",
                "frequencySchemeFilter",
            ),
        }
        for script_name, tokens in page_scripts.items():
            with self.subTest(script=script_name):
                script = (WEB_ROOT / "assets" / script_name).read_text(encoding="utf-8")
                self.assertIn("PowerPlanPageState", script)
                for token in tokens:
                    self.assertIn(token, script)

    def test_page_state_ignores_empty_numeric_layout_values(self):
        page_state_script = (WEB_ROOT / "assets" / "page_state.js").read_text(encoding="utf-8")
        self.assertIn("value !== null", page_state_script)
        self.assertIn('value !== ""', page_state_script)

        for script_name in ("planning.js", "optimize.js", "evaluation.js", "frequency.js", "comparison.js"):
            with self.subTest(script=script_name):
                script = (WEB_ROOT / "assets" / script_name).read_text(encoding="utf-8")
                self.assertIn("PowerPlanPageState?.number", script)
                self.assertNotIn("Number.isFinite(Number(saved[key]))", script)
                self.assertNotIn("Number.isFinite(Number(saved.schemeRailManualHeight))", script)
                self.assertNotIn("Number.isFinite(Number(saved.tableHeight))", script)

    def test_axis_range_state_ignores_empty_values(self):
        for script_name in ("optimize.js", "evaluation.js", "frequency.js", "comparison.js", "result_curves.js"):
            with self.subTest(script=script_name):
                script = (WEB_ROOT / "assets" / script_name).read_text(encoding="utf-8")
                self.assertIn("storedAxisNumber", script)
                self.assertNotIn("= Number(range.min)", script)
                self.assertNotIn("= Number(range.max)", script)
                self.assertNotIn('next.min = Number.isFinite(value) ? value : ""', script)
                self.assertNotIn('next.max = Number.isFinite(value) ? value : ""', script)

    def test_planning_import_curve_visibility_is_not_reset_on_modal_open(self):
        script = (WEB_ROOT / "assets" / "planning.js").read_text(encoding="utf-8")
        self.assertIn("timeSeriesImportVisibleCurves", script)
        self.assertNotIn("state.timeSeriesImportVisibleCurves = new Set(timeSeriesImportSeries.map", script)

    def test_planning_curve_generator_and_weather_input_state_are_remembered(self):
        script = (WEB_ROOT / "assets" / "planning.js").read_text(encoding="utf-8")

        self.assertIn("rememberPlanningPageState({ curveGeneratorTarget: state.curveGeneratorTarget })", script)
        self.assertIn("rememberWeatherInputsFromFields", script)
        self.assertIn('document.getElementById("weatherPlace").addEventListener("input", rememberWeatherInputsFromFields)', script)
        self.assertIn('document.getElementById("weatherYear").addEventListener("input", rememberWeatherInputsFromFields)', script)
        self.assertIn('document.getElementById("weatherLatitude").addEventListener("input", rememberWeatherInputsFromFields)', script)
        self.assertIn('document.getElementById("weatherLongitude").addEventListener("input", rememberWeatherInputsFromFields)', script)

    def test_tasks_scheme_filter_keeps_restored_selection_when_tasks_change(self):
        html = (WEB_ROOT / "tasks.html").read_text(encoding="utf-8")
        script = (WEB_ROOT / "assets" / "tasks.js").read_text(encoding="utf-8")

        self.assertIn("assets/tasks.js?v=20260531-task-filter-state", html)
        self.assertIn('const selectedScheme = String(taskState[stateKey] || "").trim();', script)
        self.assertIn("if (selectedScheme && !schemeNames.includes(selectedScheme)) schemeNames.push(selectedScheme);", script)
        self.assertIn("select.value = selectedScheme;", script)
        self.assertNotIn('taskState[stateKey] = "";', script)
        self.assertNotIn('rememberTasksPageState({ [stateKey]: "" });', script)

    def test_result_curve_viewer_persists_internal_selection_state(self):
        curve_script = (WEB_ROOT / "assets" / "result_curves.js").read_text(encoding="utf-8")
        self.assertIn("stateKey", curve_script)
        self.assertIn("restoreResultCurveViewerState", curve_script)
        self.assertIn("rememberResultCurveViewerState", curve_script)
        self.assertIn("selectedCurvesByGroup", curve_script)
        self.assertIn("activeGroup", curve_script)
        self.assertIn("curveRangeFilter", curve_script)
        self.assertIn("annualViewMode", curve_script)

        optimize_script = (WEB_ROOT / "assets" / "optimize.js").read_text(encoding="utf-8")
        evaluation_script = (WEB_ROOT / "assets" / "evaluation.js").read_text(encoding="utf-8")
        comparison_script = (WEB_ROOT / "assets" / "comparison.js").read_text(encoding="utf-8")
        optimize_html = (WEB_ROOT / "optimize.html").read_text(encoding="utf-8")
        evaluation_html = (WEB_ROOT / "evaluation.html").read_text(encoding="utf-8")
        comparison_html = (WEB_ROOT / "comparison.html").read_text(encoding="utf-8")

        self.assertIn('stateKey: "optimization-result-curves"', optimize_script)
        self.assertIn('stateKey: "evaluation-result-curves"', evaluation_script)
        self.assertIn('stateKey: "comparison-result-curves"', comparison_script)
        self.assertIn("assets/result_curves.js?v=20260530-page-state3", optimize_html)
        self.assertIn("assets/result_curves.js?v=20260530-page-state3", evaluation_html)
        self.assertIn("assets/result_curves.js?v=20260530-page-state3", comparison_html)

    def test_redirect_response_disables_cache_and_varies_on_cookie(self):
        status, headers, body = server._redirect_response("/login.html?next=%2Ftasks.html")

        self.assertEqual(status, 302)
        self.assertEqual(headers["Location"], "/login.html?next=%2Ftasks.html")
        self.assertEqual(headers["Content-Type"], "text/plain; charset=utf-8")
        self.assertEqual(headers["Cache-Control"], "no-store, no-cache, max-age=0, must-revalidate")
        self.assertEqual(headers["Pragma"], "no-cache")
        self.assertEqual(headers["Expires"], "0")
        self.assertEqual(headers["Vary"], "Cookie")
        self.assertEqual(body, b"")

    def test_static_headers_disable_cache_for_html_and_cache_versioned_assets(self):
        html_headers = server._static_headers(WEB_ROOT / "tasks.html", authenticated_html=True)
        js_headers = server._static_headers(WEB_ROOT / "assets" / "tasks.js")

        self.assertEqual(html_headers["Cache-Control"], "no-store, no-cache, max-age=0, must-revalidate")
        self.assertEqual(html_headers["Pragma"], "no-cache")
        self.assertEqual(html_headers["Expires"], "0")
        self.assertEqual(html_headers["Vary"], "Cookie")
        self.assertTrue(html_headers["Content-Type"].startswith("text/html"))
        self.assertEqual(js_headers["Cache-Control"], "public, max-age=86400, stale-while-revalidate=3600")
        self.assertIn("ETag", js_headers)
        self.assertIn("Last-Modified", js_headers)
        self.assertNotIn("Pragma", js_headers)
        self.assertIn("javascript", js_headers["Content-Type"])

    def test_static_validator_headers_support_not_modified_checks(self):
        headers = server._static_headers(WEB_ROOT / "assets" / "planning.js")

        self.assertTrue(server._static_request_not_modified({"If-None-Match": headers["ETag"]}, headers))
        self.assertTrue(server._static_request_not_modified({"If-Modified-Since": headers["Last-Modified"]}, headers))
        self.assertFalse(server._static_request_not_modified({"If-None-Match": 'W/"stale"'}, headers))

    def test_gzip_response_body_if_supported_compresses_large_text_payloads(self):
        headers = {"Content-Type": "application/javascript"}
        body = b"const value = 1;\n" * 500

        compressed_headers, compressed_body = server.gzip_response_body_if_supported(
            {"Accept-Encoding": "gzip, deflate"},
            headers,
            body,
        )

        self.assertEqual(compressed_headers["Content-Encoding"], "gzip")
        self.assertEqual(compressed_headers["Vary"], "Accept-Encoding")
        self.assertLess(len(compressed_body), len(body))

    def test_gzip_response_body_if_supported_keeps_small_payloads_plain(self):
        headers = {"Content-Type": "application/json; charset=utf-8"}
        body = b'{"ok":true}'

        next_headers, next_body = server.gzip_response_body_if_supported({"Accept-Encoding": "gzip"}, headers, body)

        self.assertIs(next_headers, headers)
        self.assertEqual(next_body, body)
        self.assertNotIn("Content-Encoding", next_headers)

    def test_static_headers_keep_regular_data_files_publicly_cacheable(self):
        headers = server._static_headers(WEB_ROOT / "data" / "load_curve_templates.csv")

        self.assertEqual(headers["Cache-Control"], "public, max-age=3600")
        self.assertTrue(headers["Content-Type"])
        self.assertIn("ETag", headers)
        self.assertIn("Last-Modified", headers)
        self.assertNotIn("Vary", headers)

    def test_static_path_rejects_directory_traversal(self):
        with self.assertRaises(ValueError):
            server.resolve_static_path("/../README.md")


if __name__ == "__main__":
    unittest.main()
