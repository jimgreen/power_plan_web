#!/usr/bin/env python3
"""Static web server and JSON API for the power_plan dashboard."""

from __future__ import annotations

import argparse
import base64
import binascii
from http.cookies import SimpleCookie
import csv
from email.utils import formatdate
import hashlib
import hmac
import json
import math
import mimetypes
import multiprocessing
import os
import queue
import re
import secrets
import sqlite3
import threading
import time
import traceback
import zlib
from copy import deepcopy
from datetime import datetime
from decimal import Decimal
from http import HTTPStatus
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from io import BytesIO, StringIO
from pathlib import Path
import sys


def _candidate_venv_python(venv_root: Path) -> Path:
    if os.name == "nt":
        return venv_root / "Scripts" / "python.exe"
    return venv_root / "bin" / "python"


def _ensure_project_virtualenv() -> None:
    if os.environ.get("POWER_PLAN_DISABLE_VENV_BOOTSTRAP") == "1":
        return
    current_prefix = Path(sys.prefix).resolve()
    script_path = Path(__file__).resolve()
    try:
        launched_path = Path(sys.argv[0]).resolve()
    except (OSError, RuntimeError):
        return
    if launched_path != script_path:
        return
    explicit_venv = os.environ.get("POWER_PLAN_VENV")
    candidate_roots: list[Path] = []
    if explicit_venv:
        explicit_path = Path(explicit_venv).expanduser()
        candidate_roots.append(explicit_path.parent.parent if explicit_path.name.startswith("python") else explicit_path)
    project_root = Path(__file__).resolve().parent.parent
    candidate_roots.extend([project_root / "venv", project_root / ".venv", project_root / "power_plan_web" / ".venv"])
    for venv_root in candidate_roots:
        python_exe = _candidate_venv_python(venv_root)
        if not python_exe.exists():
            continue
        if current_prefix == venv_root.resolve():
            return
        os.environ["VIRTUAL_ENV"] = str(venv_root)
        os.environ["PATH"] = f"{python_exe.parent}{os.pathsep}{os.environ.get('PATH', '')}"
        os.execv(str(python_exe), [str(python_exe), str(script_path), *sys.argv[1:]])


_ensure_project_virtualenv()
from contextlib import closing
from urllib.error import HTTPError, URLError
from urllib.parse import parse_qs, quote, urlencode, unquote, urlparse
from urllib.request import urlopen
from zipfile import BadZipFile

from openpyxl import Workbook, load_workbook
from openpyxl.styles import Font, PatternFill
from openpyxl.utils import get_column_letter
from openpyxl.utils.exceptions import InvalidFileException

import calculation_precheck
import estimate
import file_cache
import file_ops
import milp_solver
import plan_optimizer
import planning_store
import reliability


WEB_ROOT = Path(__file__).resolve().parent
DATA_DIR = WEB_ROOT / "data"
VENDOR_DIR = WEB_ROOT / "vendor"
LOAD_CURVE_TEMPLATE_PATH = DATA_DIR / "load_curve_templates.csv"
USER_DB_PATH = Path(os.environ.get("POWER_PLAN_USER_DB", WEB_ROOT / "power_plan_users.sqlite3"))
SESSION_COOKIE_NAME = "power_plan_session"
SESSION_MAX_AGE_SECONDS = 7 * 24 * 60 * 60
LOCAL_AUTH_BYPASS_ENABLED = os.environ.get("POWER_PLAN_LOCAL_AUTH_BYPASS", "0").strip().lower() in {"1", "true", "yes", "on"}
LOCAL_AUTH_USER = {"id": 0, "username": "local", "role": "admin", "created_at": "local"}
NASA_POWER_HOURLY_URL = "https://power.larc.nasa.gov/api/temporal/hourly/point"
AMAP_GEOCODING_URL = "https://restapi.amap.com/v3/geocode/geo"
AMAP_REVERSE_GEOCODING_URL = "https://restapi.amap.com/v3/geocode/regeo"
DEFAULT_AMAP_WEB_SERVICE_KEY = "21db26646aac8fed4620eaa36f210018"
OPEN_METEO_GEOCODING_URL = "https://geocoding-api.open-meteo.com/v1/search"
NOMINATIM_SEARCH_URL = "https://nominatim.openstreetmap.org/search"
NOMINATIM_REVERSE_URL = "https://nominatim.openstreetmap.org/reverse"
PHOTON_GEOCODING_URL = "https://photon.komoot.io/api/"
CHINESE_PLACE_ALIASES = {
    "上海": "shanghai",
    "上海市": "shanghai",
    "北京": "beijing",
    "北京市": "beijing",
    "广州": "guangzhou",
    "广州市": "guangzhou",
    "深圳": "shenzhen",
    "深圳市": "shenzhen",
    "天津": "tianjin",
    "天津市": "tianjin",
    "重庆": "chongqing",
    "重庆市": "chongqing",
    "香港": "hong kong",
    "澳门": "macau",
    "台北": "taipei",
}
OPTIMIZATION_RESULT_WORKBOOK_NAME = "opt_results.xlsx"
RESULT_WORKBOOK_RE = re.compile(r"^[A-Za-z0-9_\-\u4e00-\u9fff]+_results\.xlsx$")
RELIABILITY_PARAMETERS_FILE_NAME = "reliability_parameters.json"
RELIABILITY_RESULT_JSON_SUFFIX = "_reliability.json"
RELIABILITY_RESULT_WORKBOOK_SUFFIX = "_reliability_results.xlsx"
PLANNING_RESULT_SHEET_NAME = "规划结果"
PLANNING_RESULT_HEADERS = ["设备类型", "设计台数", "单台容量", "总容量", "单位"]
CSV_ROWS_CACHE = file_cache.FileCache("dashboard_csv_rows", max_entries=32)
RESULT_WORKBOOK_HEALTH_CACHE = file_cache.FileCache("result_workbook_health", max_entries=256)
EVALUATION_PLANNING_RESULT_CACHE = file_cache.FileCache("evaluation_planning_result", max_entries=256)
RESULT_DISPLAY_PAYLOAD_CACHE = file_cache.FileCache("result_display_payload", max_entries=128)
COMPARISON_WORKBOOK_CACHE = file_cache.FileCache("comparison_workbook", max_entries=128)
COMPARISON_CURVE_SLICE_CACHE = file_cache.FileCache("comparison_curve_slice", max_entries=512)
LOAD_CURVE_TEMPLATE_CACHE = file_cache.FileCache("load_curve_templates", max_entries=8)
STATIC_FILE_BYTES_CACHE = file_cache.FileCache("static_file_bytes", max_entries=256, copy_values=False)
TASK_RESULT_FILE_LIST_CACHE = file_cache.FileCache("task_result_file_list", max_entries=256)
FREQUENCY_CURVE_DURATION_SECONDS = 3.0
FREQUENCY_CURVE_STEP_SECONDS = 0.05
FREQUENCY_CURVE_POINT_COUNT = int(round(FREQUENCY_CURVE_DURATION_SECONDS / FREQUENCY_CURVE_STEP_SECONDS)) + 1
COMPARISON_CURVE_GROUPS = {
    "hourly": {"title": "小时级曲线", "sheet": "调度结果", "limit": 8760},
    "daily": {"title": "日级统计", "sheet": "供能日曲线", "limit": None},
    "safety": {"title": "安全日曲线", "sheet": "安全日曲线", "limit": None},
    "monthly": {"title": "月度统计", "sheet": "供能月曲线", "limit": None},
}
COMPARISON_CURVE_X_HEADERS = {
    "小时",
    "hour_index",
    "时间",
    "datetime",
    "day",
    "month",
    "日期",
    "月份",
    "nominal_frequency_hz",
}
RESULT_CURVE_FIELD_LABELS = {
    "load_energy": "负荷总电量",
    "diesel_energy": "柴发总发电量",
    "wind_energy": "风机总发电量",
    "pv_energy": "光伏总发电量",
    "storage_charge_energy": "电储能总储电量",
    "storage_discharge_energy": "电储能总放电量",
    "hydrogen_production_energy": "电制氢总用电量",
    "hydrogen_storage_increase": "氢储总增加量",
    "hydrogen_storage_decrease": "氢储总消耗量",
    "fuel_cell_energy": "燃料电池总发电量",
    "wind_available_energy": "风力最大可发电量",
    "pv_available_energy": "光伏最大可发电量",
    "renewable_available_energy": "新能源最大可发电量",
    "renewable_energy": "新能源实发电量",
    "wind_curtailed_energy": "弃风总电量",
    "pv_curtailed_energy": "弃光总电量",
    "curtailed_energy": "新能源总弃电量",
    "unmet_load_energy": "切负荷总电量",
    "renewable_ratio": "新能源占比",
    "renewable_curtailed_rate": "新能源弃电率",
    "load_up_disturbance_power": "负荷上扰动功率",
    "load_down_disturbance_power": "负荷下扰动功率",
    "renewable_down_disturbance_power": "新能源下扰动功率",
    "renewable_single_unit_power_max": "风光单机功率最大值",
    "grid_up_regulation_capacity": "电网向上调节能力",
    "grid_down_regulation_capacity": "电网向下调节能力",
    "grid_up_regulation_requirement": "电网向上调节需求",
    "grid_down_regulation_requirement": "电网向下调节需求",
    "frequency_min": "最低频率",
    "frequency_max": "最高频率",
    "frequency_nadir_est_hz": "最低频率保守估计值",
    "frequency_peak_est_hz": "最高频率保守估计值",
    "frequency_nadir_exact_hz": "最低频率解析值",
    "frequency_peak_exact_hz": "最高频率解析值",
    "steady_state_frequency_min_hz": "下限场景稳态频率",
    "steady_state_frequency_max_hz": "上限场景稳态频率",
    "rocof_hz_per_s": "初始频率变化率",
    "rocof_upper_hz_per_s": "上限场景初始频率变化率",
    "frequency_lower_margin_hz": "频率下限裕度",
    "frequency_upper_margin_hz": "频率上限裕度",
    "equivalent_inertia_m": "等效惯量M",
    "equivalent_primary_frequency_k": "等效调频系数K",
    "equivalent_damping_d": "等效阻尼系数D",
    "frequency_delta_p_mw": "频率下限扰动功率",
    "frequency_upper_delta_p_mw": "频率上限扰动功率",
    "frequency_fit_error_hz": "频率下限拟合误差",
    "frequency_upper_fit_error_hz": "频率上限拟合误差",
}
DEPRECATED_RESULT_CURVE_HEADERS = {"新能源N-1功率缺口", "renewable_n1_power_gap"}
RESULT_WORKBOOK_HEADER_TO_FIELD = {label: key for key, label in RESULT_CURVE_FIELD_LABELS.items()}
RESULT_WORKBOOK_HEADER_TO_FIELD.update(
    {
        "小时": "hour_index",
        "时间": "datetime",
        "风速": "wind_speed",
        "太阳辐射": "solar_irradiance",
        "环境温度": "temperature",
        "负荷总功率": "load",
        "柴发总功率": "diesel_power",
        "柴发开机容量": "diesel_capacity",
        "风力最大可发": "wind_available",
        "风机总功率": "wind_power",
        "光伏最大可发": "pv_available",
        "光伏总功率": "pv_power",
        "新能源总出力": "renewable_power",
        "新能源最大可发": "renewable_available",
        "电储能总功率": "storage_power",
        "构网储能总容量": "grid_storage_capacity",
        "构网储能总功率": "grid_storage_power",
        "电储电量": "storage_soc",
        "电制氢总功率": "hydrogen_production_power",
        "储氢罐氢储量": "hydrogen_storage",
        "燃料电池总功率": "fuel_cell_power",
        "弃风总功率": "wind_curtailed_power",
        "弃光总功率": "pv_curtailed_power",
        "新能源弃电总功率": "curtailed_power",
        "切负荷功率": "unmet_load",
        "柴发启停": "diesel_on",
        "制氢启停": "electrolyzer_on",
        "储能充电": "storage_charge",
        "储能放电": "storage_discharge",
        "日期": "day",
        "月份": "month",
    }
)
STATIC_NO_STORE_SUFFIXES = {".html"}
STATIC_BROWSER_CACHE_SUFFIXES = {".css", ".js", ".png", ".svg", ".ico", ".map", ".jpg", ".jpeg", ".webp", ".woff", ".woff2"}
NO_STORE_CACHE_CONTROL = "no-store, no-cache, max-age=0, must-revalidate"
STATIC_ASSET_CACHE_CONTROL = "public, max-age=86400, stale-while-revalidate=3600"
STATIC_DATA_CACHE_CONTROL = "public, max-age=3600"
RESULT_WORKBOOK_READ_ERRORS = (BadZipFile, zlib.error, OSError, EOFError, KeyError, InvalidFileException)
COMPRESSIBLE_CONTENT_PREFIXES = ("text/", "application/json", "application/javascript", "image/svg+xml")
MIN_GZIP_RESPONSE_BYTES = 2048
TIME_SERIES_IMPORT_ROW_COUNT = 8760
TIME_SERIES_IMPORT_REQUIRED_COLUMNS = {
    "wind_speed": ("风速", ["wind_speed", "wind", "风速", "风速(m/s)", "风速ms", "风速米秒", "ws10m"]),
    "solar_irradiance": (
        "太阳辐射",
        [
            "solar_irradiance",
            "solar",
            "irradiance",
            "solar radiation",
            "太阳辐射",
            "太阳辐照",
            "太阳辐照度",
            "单位面积太阳辐射",
            "单位面积太阳辐照",
            "太阳辐射(w/m2)",
            "太阳辐照(w/m2)",
            "太阳辐射(W/m^2)",
            "太阳辐照(W/m^2)",
            "allsky_sfc_sw_dwn",
        ],
    ),
    "temperature": ("室温", ["temperature", "temp", "室温", "温度", "环境温度", "气温", "温度(摄氏度)", "环境温度(摄氏度)", "t2m"]),
    "load": ("负荷", ["load", "负荷", "负荷功率", "负荷总功率", "用电负荷", "用电功率", "用电功率(kW)", "负荷(kW)", "负荷kw"]),
}
TIME_SERIES_IMPORT_OPTIONAL_COLUMNS = {
    "datetime": ["datetime", "time", "时间", "日期时间", "时刻", "小时", "小时序号", "hour", "hour_index"],
}
AMAP_WEB_SERVICE_KEY = (
    os.environ.get("POWER_PLAN_AMAP_KEY")
    or os.environ.get("AMAP_WEB_SERVICE_KEY")
    or os.environ.get("AMAP_KEY")
    or DEFAULT_AMAP_WEB_SERVICE_KEY
)
NASA_POWER_PARAMETERS = {
    "wind_speed": "WS10M",
    "solar_irradiance": "ALLSKY_SFC_SW_DWN",
    "temperature": "T2M",
}
if VENDOR_DIR.exists():
    sys.path.insert(0, str(VENDOR_DIR))
DB_CONFIG = {
    "host": os.environ.get("POWER_PLAN_DB_HOST", "127.0.0.1"),
    "port": int(os.environ.get("POWER_PLAN_DB_PORT", "3306")),
    "user": os.environ.get("POWER_PLAN_DB_USER", "root"),
    "password": os.environ.get("POWER_PLAN_DB_PASSWORD", "scadaems"),
    "database": os.environ.get("POWER_PLAN_DB_NAME", "scadaems"),
    "charset": "utf8mb4",
}


def _coerce_value(value: str) -> float | int | str:
    if isinstance(value, Decimal):
        value = float(value)
    if isinstance(value, (int, float)):
        if isinstance(value, float) and value.is_integer():
            return int(value)
        return value
    value = value.strip()
    if value == "":
        return ""
    try:
        number = float(value)
    except ValueError:
        return value
    if number.is_integer():
        return int(number)
    return number


def _metric(label: str, value: float | int | str, unit: str, status: str = "normal") -> dict:
    return {"label": label, "value": value, "unit": unit, "status": status}


def _summary(label: str, value: str, status: str = "normal") -> dict:
    return {"label": label, "value": value, "status": status}


def _now_iso() -> str:
    return datetime.now().isoformat(timespec="seconds")


def _time_to_hour(value: str) -> float:
    hour, minute = value.split(":", 1)
    return int(hour) + int(minute) / 60


def _hour_to_time(value: float) -> str:
    value = value % 24
    hour = int(value)
    minute = int(round((value - hour) * 60)) % 60
    return f"{hour:02d}:{minute:02d}"


class SimuRuntime:
    """In-memory runtime state for simulation controls."""

    def __init__(self, initial_time: str = "00:00", speed: float = 1.0, status: str = "STOPPED") -> None:
        self.cursor_hour = _time_to_hour(initial_time)
        self.speed = speed
        self.status = status
        self._last_tick = time.monotonic()
        self._lock = threading.Lock()

    def snapshot(self) -> dict:
        with self._lock:
            self._advance_locked()
            return {
                "sim_time": _hour_to_time(self.cursor_hour),
                "cursor_hour": round(self.cursor_hour, 3),
                "speed": self.speed,
                "status": self.status,
            }

    def apply(self, action: str) -> dict:
        with self._lock:
            self._advance_locked()
            if action == "start":
                self.status = "RUNNING"
            elif action == "faster":
                self.speed = min(16.0, self.speed * 2)
            elif action == "slower":
                self.speed = max(0.25, self.speed / 2)
            elif action == "stop":
                self.status = "STOPPED"
            elif action == "reset":
                self.cursor_hour = 0.0
                self.speed = 1.0
                self.status = "STOPPED"
            else:
                raise ValueError(f"unknown SIMU action: {action}")
            self._last_tick = time.monotonic()
            return {
                "sim_time": _hour_to_time(self.cursor_hour),
                "cursor_hour": round(self.cursor_hour, 3),
                "speed": self.speed,
                "status": self.status,
            }

    def _advance_locked(self) -> None:
        now = time.monotonic()
        elapsed = now - self._last_tick
        if self.status == "RUNNING":
            # One real second advances one simulated minute at speed=1.
            self.cursor_hour = (self.cursor_hour + elapsed * self.speed / 60) % 24
            self._last_tick = now


class UserStore:
    """SQLite-backed user and session store."""

    def __init__(self, db_path: Path = USER_DB_PATH) -> None:
        self.db_path = Path(db_path)
        self.db_path.parent.mkdir(parents=True, exist_ok=True)
        self._lock = threading.Lock()
        self._init_db()

    def _connect(self):
        connection = sqlite3.connect(self.db_path)
        connection.row_factory = sqlite3.Row
        return connection

    def _init_db(self) -> None:
        with closing(self._connect()) as connection:
            connection.execute(
                """
                CREATE TABLE IF NOT EXISTS users (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    username TEXT NOT NULL UNIQUE,
                    password_salt TEXT NOT NULL,
                    password_hash TEXT NOT NULL,
                    role TEXT NOT NULL CHECK(role IN ('admin', 'user')),
                    created_at TEXT NOT NULL
                )
                """
            )
            connection.execute(
                """
                CREATE TABLE IF NOT EXISTS sessions (
                    token TEXT PRIMARY KEY,
                    user_id INTEGER NOT NULL,
                    created_at TEXT NOT NULL,
                    expires_at REAL NOT NULL,
                    FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE
                )
                """
            )
            connection.commit()

    def user_count(self) -> int:
        with closing(self._connect()) as connection:
            return int(connection.execute("SELECT COUNT(*) FROM users").fetchone()[0])

    def create_user(self, username: str, password: str, role: str | None = None) -> dict:
        clean_username = self._clean_username(username)
        if len(password or "") < 6:
            raise ValueError("密码长度不能少于6位")
        with self._lock:
            user_role = role or ("admin" if self.user_count() == 0 else "user")
            if user_role not in {"admin", "user"}:
                raise ValueError("用户角色不合法")
            salt, password_hash = self._hash_password(password)
            try:
                with closing(self._connect()) as connection:
                    cursor = connection.execute(
                        """
                        INSERT INTO users (username, password_salt, password_hash, role, created_at)
                        VALUES (?, ?, ?, ?, ?)
                        """,
                        (clean_username, salt, password_hash, user_role, _now_iso()),
                    )
                    connection.commit()
                    user_id = int(cursor.lastrowid)
            except sqlite3.IntegrityError as exc:
                raise ValueError("用户名已存在") from exc
        user = self.get_user_by_id(user_id)
        if not user:
            raise ValueError("用户创建失败")
        return user

    def authenticate(self, username: str, password: str) -> dict:
        clean_username = self._clean_username(username)
        with closing(self._connect()) as connection:
            row = connection.execute(
                "SELECT id, username, password_salt, password_hash, role, created_at FROM users WHERE username = ?",
                (clean_username,),
            ).fetchone()
        if not row or not self._verify_password(password or "", row["password_salt"], row["password_hash"]):
            raise ValueError("用户名或密码错误")
        return self._public_user(row)

    def create_session(self, user_id: int) -> str:
        token = secrets.token_urlsafe(32)
        now = time.time()
        with closing(self._connect()) as connection:
            connection.execute(
                "INSERT INTO sessions (token, user_id, created_at, expires_at) VALUES (?, ?, ?, ?)",
                (token, user_id, _now_iso(), now + SESSION_MAX_AGE_SECONDS),
            )
            connection.commit()
        return token

    def delete_session(self, token: str) -> None:
        if not token:
            return
        with closing(self._connect()) as connection:
            connection.execute("DELETE FROM sessions WHERE token = ?", (token,))
            connection.commit()

    def user_for_session(self, token: str) -> dict | None:
        if not token:
            return None
        now = time.time()
        with closing(self._connect()) as connection:
            connection.execute("DELETE FROM sessions WHERE expires_at <= ?", (now,))
            row = connection.execute(
                """
                SELECT users.id, users.username, users.role, users.created_at
                FROM sessions
                JOIN users ON users.id = sessions.user_id
                WHERE sessions.token = ? AND sessions.expires_at > ?
                """,
                (token, now),
            ).fetchone()
            connection.commit()
        return self._public_user(row) if row else None

    def list_users(self) -> list[dict]:
        with closing(self._connect()) as connection:
            rows = connection.execute(
                "SELECT id, username, role, created_at FROM users ORDER BY id"
            ).fetchall()
        return [self._public_user(row) for row in rows]

    def get_user_by_id(self, user_id: int) -> dict | None:
        with closing(self._connect()) as connection:
            row = connection.execute(
                "SELECT id, username, role, created_at FROM users WHERE id = ?",
                (user_id,),
            ).fetchone()
        return self._public_user(row) if row else None

    def get_user_by_username(self, username: str) -> dict | None:
        clean_username = self._clean_username(username)
        with closing(self._connect()) as connection:
            row = connection.execute(
                "SELECT id, username, role, created_at FROM users WHERE username = ?",
                (clean_username,),
            ).fetchone()
        return self._public_user(row) if row else None

    def update_role(self, user_id: int, role: str) -> dict:
        if role not in {"admin", "user"}:
            raise ValueError("用户角色不合法")
        with closing(self._connect()) as connection:
            connection.execute("UPDATE users SET role = ? WHERE id = ?", (role, user_id))
            connection.commit()
        user = self.get_user_by_id(user_id)
        if not user:
            raise FileNotFoundError("用户不存在")
        return user

    def delete_user(self, user_id: int, current_user_id: int | None = None) -> None:
        if current_user_id is not None and int(user_id) == int(current_user_id):
            raise ValueError("不能删除当前登录用户")
        with closing(self._connect()) as connection:
            cursor = connection.execute("DELETE FROM users WHERE id = ?", (user_id,))
            connection.commit()
        if cursor.rowcount == 0:
            raise FileNotFoundError("用户不存在")

    @staticmethod
    def _clean_username(username: str) -> str:
        clean = str(username or "").strip()
        if len(clean) < 2 or len(clean) > 32:
            raise ValueError("用户名长度应为2-32位")
        if any(ord(char) < 32 or char.isspace() for char in clean):
            raise ValueError("用户名不能包含空格或不可见字符")
        return clean

    @staticmethod
    def _hash_password(password: str, salt: str | None = None) -> tuple[str, str]:
        password_salt = salt or secrets.token_hex(16)
        digest = hashlib.pbkdf2_hmac("sha256", password.encode("utf-8"), bytes.fromhex(password_salt), 120_000)
        return password_salt, digest.hex()

    @classmethod
    def _verify_password(cls, password: str, salt: str, expected_hash: str) -> bool:
        _, actual_hash = cls._hash_password(password, salt)
        return hmac.compare_digest(actual_hash, expected_hash)

    @staticmethod
    def _public_user(row) -> dict:
        return {
            "id": int(row["id"]),
            "username": row["username"],
            "role": row["role"],
            "created_at": row["created_at"],
        }


class OptimizationRuntime:
    """In-memory runtime state for the planning optimization page."""

    def __init__(self, scheme: str = "") -> None:
        self.status = "待启动"
        self.scheme = str(scheme or "").strip()
        self.start_time = ""
        self.end_time = ""
        self.progress = 0
        self._started_monotonic = 0.0
        self._last_progress_log = -1
        self.result_file = ""
        self._metrics: list[dict] = []
        self._results: dict = {}
        self._results_exported = False
        self._logs: list[dict[str, str]] = []
        self._lock = threading.Lock()
        self._process: multiprocessing.Process | None = None
        self._event_queue: multiprocessing.Queue | None = None
        self.process_id: int | None = None
        self._stop_requested = False
        self._run_token = 0
        self._append_log_unlocked("info", "规划求解待启动")

    def snapshot(self, include_hourly_curves: bool = True) -> dict:
        with self._lock:
            self._drain_events_unlocked()
            self._reap_process_unlocked()
            return self._payload_unlocked(include_hourly_curves=include_hourly_curves)

    def task_snapshot(self) -> dict:
        """Return the lightweight state needed by the task-concurrency page.

        Task polling only needs process state and the latest log.  Avoid
        reading the generated result workbook here; those workbooks can be
        large and are displayed by the solver/evaluation pages on demand.
        """

        with self._lock:
            self._drain_events_unlocked()
            self._reap_process_unlocked()
            return self._task_payload_unlocked()

    def apply(self, action: str, scheme: str = "") -> dict:
        target_scheme = str(scheme or self.scheme or "未选择方案").strip() or "未选择方案"
        if action == "clear_logs":
            with self._lock:
                self.scheme = target_scheme
                self._logs.clear()
                return self._payload_unlocked()

        if action == "cancel_queue":
            with self._lock:
                if self.status == "运行中":
                    raise OptimizationStateError("running", f"方案“{target_scheme}”正在运行，无法退出队列")
                self.scheme = target_scheme
                self.status = "退出队列"
                self.start_time = ""
                self.end_time = ""
                self.progress = 0
                self.result_file = ""
                self.process_id = None
                self._metrics = []
                self._results = {}
                self._results_exported = False
                self._stop_requested = False
                self._terminate_process_unlocked()
                self._append_log_unlocked("info", "退出等待队列")
                return self._payload_unlocked()

        if action == "start":
            with self._lock:
                if self.status == "运行中":
                    if self.scheme == target_scheme:
                        raise OptimizationStateError("running", f"方案“{target_scheme}”正在运行，无法再次启动")
                    raise OptimizationStateError("running", f"方案“{self.scheme}”正在运行，无法启动方案“{target_scheme}”")
            try:
                scheme_payload = PLANNING_STORE.read_scheme(target_scheme)
            except FileNotFoundError:
                scheme_payload = None
            if scheme_payload:
                try:
                    calculation_precheck.validate_optimization_fast_feasibility(scheme_payload)
                except ValueError as exc:
                    self._mark_start_failure(target_scheme, str(exc))
                    raise
            with self._lock:
                if self.status == "运行中":
                    if self.scheme == target_scheme:
                        raise OptimizationStateError("running", f"方案“{target_scheme}”正在运行，无法再次启动")
                    raise OptimizationStateError("running", f"方案“{self.scheme}”正在运行，无法启动方案“{target_scheme}”")
                self.status = "运行中"
                self.scheme = target_scheme
                self.start_time = _now_text()
                self.end_time = ""
                self.progress = 0
                self._started_monotonic = time.monotonic()
                self._last_progress_log = -1
                self.result_file = ""
                self._metrics = []
                self._results = {}
                self._results_exported = False
                self._stop_requested = False
                self._terminate_process_unlocked()
                self._run_token += 1
                self._append_log_unlocked("ok", f"启动规划求解，方案：{self.scheme}")
                self._append_log_unlocked("info", "后台规划求解程序已启动")
                self._event_queue = multiprocessing.Queue()
                self._process = multiprocessing.Process(
                    target=optimization_process_worker,
                    args=(self._event_queue, target_scheme, str(PLANNING_STORE.root)),
                    daemon=True,
                )
                self._process.start()
                self.process_id = self._process.pid
                return self._payload_unlocked()

        if action == "stop":
            with self._lock:
                if self.status != "运行中" or self.scheme != target_scheme:
                    raise OptimizationStateError("not_running", f"方案“{target_scheme}”没有运行")
                self._stop_requested = True
                self._terminate_process_unlocked()
                self.status = "计算中止"
                self.end_time = _now_text()
                self._append_log_unlocked("warn", "停止规划求解")
                return self._payload_unlocked()

        raise ValueError(f"unknown optimization action: {action}")

    def _mark_start_failure(self, scheme: str, message: str) -> None:
        with self._lock:
            if self.status == "运行中":
                return
            self.scheme = scheme
            now = _now_text()
            self.start_time = now
            self.end_time = now
            self.progress = 0
            self.result_file = ""
            self._metrics = []
            self._results = {}
            self._results_exported = False
            self._stop_requested = False
            self._terminate_process_unlocked()
            self.status = "失败"
            self._append_log_unlocked("error", message)

    def _drain_events_unlocked(self) -> None:
        if not self._event_queue:
            return
        while True:
            try:
                event = self._event_queue.get_nowait()
            except queue.Empty:
                break
            if not isinstance(event, dict):
                continue
            self._handle_process_event_unlocked(event)
            if not self._event_queue:
                break
        if self.status == "运行中" and self._process and not self._process.is_alive():
            self._process.join(timeout=0)
            if self.status == "运行中":
                exit_code = self._process.exitcode
                self.status = "失败"
                self.end_time = _now_text()
                self._append_log_unlocked("error", f"规划求解进程异常退出，退出码：{exit_code}")
            self._close_event_queue_unlocked()

    def _handle_process_event_unlocked(self, event: dict) -> None:
        event_type = str(event.get("type") or "log")
        if event_type == "log":
            self._append_optimizer_event_unlocked(event)
            return
        if event_type == "done":
            if self.status != "运行中" or self._stop_requested:
                return
            self.progress = 100
            self._metrics = event.get("metrics") if isinstance(event.get("metrics"), list) else []
            self._results = event.get("results") if isinstance(event.get("results"), dict) else {}
            completed_end_time = _now_text()
            result_path = export_optimization_results_workbook(
                self._payload_unlocked(
                    read_workbook=False,
                    status_override="已完成",
                    end_time_override=completed_end_time,
                )
            )
            self.result_file = str(result_path)
            self._results_exported = True
            self.status = "已完成"
            self.end_time = completed_end_time
            self._append_log_unlocked("ok", f"优化结果已写入：{result_path.name}")
            self._join_finished_process_unlocked()
            self._close_event_queue_unlocked()
            return
        if event_type == "timeout" or (event_type == "error" and calculation_timeout_message(event.get("message"))):
            if self.status != "运行中":
                return
            self.status = "超时"
            self.end_time = _now_text()
            self._append_log_unlocked("error", str(event.get("message") or "规划求解达到最大用时，计算超时"))
            self._join_finished_process_unlocked()
            self._close_event_queue_unlocked()
            return
        if event_type == "error":
            if self.status != "运行中":
                return
            self.status = "失败"
            self.end_time = _now_text()
            self._append_log_unlocked("error", str(event.get("message") or "规划求解失败"))
            self._join_finished_process_unlocked()
            self._close_event_queue_unlocked()

    def _payload_unlocked(
        self,
        include_hourly_curves: bool = True,
        read_workbook: bool = True,
        status_override: str | None = None,
        end_time_override: str | None = None,
    ) -> dict:
        result_path = optimization_result_workbook_path(self.scheme)
        workbook_payload = (
            read_result_workbook_display_payload_for_response(result_path, include_hourly_curves=include_hourly_curves)
            if read_workbook and self.status != "运行中"
            else None
        )
        if workbook_payload:
            self.result_file = str(result_path)
        status = self.status if status_override is None else status_override
        end_time = self.end_time if end_time_override is None else end_time_override
        return {
            "status": status,
            "scheme": self.scheme,
            "start_time": self.start_time,
            "end_time": end_time,
            "progress": self.progress,
            "result_file": self.result_file,
            "process_id": self.process_id or "",
            "elapsed_seconds": elapsed_seconds_from_times(self.start_time, end_time),
            "metrics": merge_runtime_metrics(
                self._metrics_unlocked(status_override=status, end_time_override=end_time),
                workbook_payload.get("metrics", []) if workbook_payload else [],
            ),
            "results": workbook_payload.get("results", {}) if workbook_payload else (self._results if self._results else self._default_results_unlocked()),
            "logs": list(self._logs),
        }

    def _runtime_payload_unlocked(self) -> dict:
        return {
            "status": self.status,
            "scheme": self.scheme,
            "start_time": self.start_time,
            "end_time": self.end_time,
            "progress": self.progress,
            "result_file": self.result_file,
            "process_id": self.process_id or "",
            "elapsed_seconds": elapsed_seconds_from_times(self.start_time, self.end_time),
            "metrics": self._metrics_unlocked(),
            "results": self._results if self._results else self._default_results_unlocked(),
            "logs": list(self._logs),
        }

    def _task_payload_unlocked(self) -> dict:
        result_path = optimization_result_workbook_path(self.scheme)
        result_file = self.result_file or (str(result_path) if result_path.exists() else "")
        return {
            "status": self.status,
            "scheme": self.scheme,
            "start_time": self.start_time,
            "end_time": self.end_time,
            "progress": self.progress,
            "result_file": result_file,
            "process_id": self.process_id or "",
            "elapsed_seconds": elapsed_seconds_from_times(self.start_time, self.end_time),
            "logs": list(self._logs),
        }

    def _metrics_unlocked(self, status_override: str | None = None, end_time_override: str | None = None) -> list[dict]:
        status = self.status if status_override is None else status_override
        end_time = self.end_time if end_time_override is None else end_time_override
        base = [
            {"label": "状态", "value": status, "unit": ""},
            {"label": "开始", "value": self.start_time or "-", "unit": ""},
            {"label": "完成", "value": end_time or "-", "unit": ""},
        ]
        existing_labels = {item["label"] for item in base}
        for metric in self._metrics:
            if not isinstance(metric, dict):
                continue
            label = metric.get("label", "")
            if label in existing_labels:
                continue
            base.append(metric)
            existing_labels.add(label)
        if "度电成本" not in existing_labels:
            base.append({"label": "度电成本", "value": "-", "unit": "元"})
        if "绿电占比" not in existing_labels:
            base.append({"label": "绿电占比", "value": "-", "unit": "%"})
        return base

    @staticmethod
    def _default_results_unlocked() -> dict:
        return {
            "overview_tables": [
                {"title": "规划结果", "rows": []},
                {"title": "规划年指标", "rows": []},
            ],
            "overview_disks": [],
            "overview": [],
            "green": [],
            "green_table": [],
            "safety": [],
            "safety_table": [],
            "curves": {"green_daily": [], "green_monthly": [], "green_hourly": [], "safety_daily": []},
        }

    def _append_optimizer_event_unlocked(self, event: dict) -> None:
        level = str(event.get("level") or "info")
        message = str(event.get("message") or "")
        progress = event.get("progress")
        if progress is not None:
            try:
                self.progress = max(self.progress, min(100, max(0, int(progress))))
            except (TypeError, ValueError):
                pass
        if message:
            self._append_log_unlocked(level, message)

    def _results_unlocked(self) -> dict:
        cost = round(max(0.42, 0.78 - self.progress * 0.002), 3)
        green_ratio = round(min(92.0, 52.0 + self.progress * 0.34), 1)
        reserve_margin = round(max(12.0, 28.0 - self.progress * 0.05), 1)
        frequency_margin = round(1.08 + min(0.12, self.progress * 0.001), 3)
        diesel_energy = round(max(480, 1580 - self.progress * 7.2), 1)
        wind_energy = round(2260 + self.progress * 8.6, 1)
        pv_energy = round(1880 + self.progress * 7.4, 1)
        storage_energy = round(720 + self.progress * 3.5, 1)
        hydrogen_production = round(148 + self.progress * 1.1, 1)
        fuel_cell_energy = round(420 + self.progress * 2.6, 1)
        construction_cost = 2180.0
        operation_cost = round(980 - self.progress * 1.4, 1)
        renewable_energy = round(wind_energy + pv_energy + fuel_cell_energy, 1)
        load_energy_kwh = 6840.0 * 1000
        diesel_energy_kwh = diesel_energy * 1000
        wind_energy_kwh = wind_energy * 1000
        pv_energy_kwh = pv_energy * 1000
        storage_energy_kwh = storage_energy * 1000
        storage_charge_energy_kwh = round(storage_energy_kwh / 0.9, 1)
        fuel_cell_energy_kwh = fuel_cell_energy * 1000
        hydrogen_production_nm3 = hydrogen_production * 10000
        hydrogen_production_energy_kwh = hydrogen_production_nm3 / 0.2
        seed_green_daily = self._green_daily_curve_unlocked(
            load_energy_kwh,
            diesel_energy_kwh,
            wind_energy_kwh,
            pv_energy_kwh,
            fuel_cell_energy_kwh,
            hydrogen_production_energy_kwh,
            storage_energy_kwh,
            storage_charge_energy_kwh,
        )
        green_hourly = self._green_hourly_curve_unlocked(seed_green_daily)
        green_daily = estimate.aggregate_daily(green_hourly)
        green_monthly = estimate.aggregate_monthly(green_daily)
        energy_totals = self._energy_totals_from_daily_unlocked(green_daily)
        energy_totals["hydrogen_production"] = hydrogen_production_nm3
        energy_totals["diesel_consumption"] = round(diesel_energy * 0.24, 1)
        annual_energy_rows = estimate.annual_energy_rows(
            energy_totals,
            energy_totals["renewable_ratio"],
            energy_totals["renewable_curtailed_rate"],
        )
        safety_daily = self._safety_daily_frequency_curve_unlocked()
        highest_frequency = max(point["frequency_max"] for point in safety_daily)
        lowest_frequency = min(point["frequency_min"] for point in safety_daily)
        upward_disturbance = round(max(860.0, 1480.0 - self.progress * 5.8), 1)
        downward_disturbance = round(max(720.0, 1310.0 - self.progress * 5.1), 1)
        frequency_risk_hours = max(0, int(round(68 - self.progress * 0.55)))
        planning_rows = [
            {"设备类型": "柴发", "设计台数": 2, "单台容量": 320, "总容量": 640, "单位": "kW"},
            {"设备类型": "风机", "设计台数": 6, "单台容量": 120, "总容量": 720, "单位": "kW"},
            {"设备类型": "光伏", "设计台数": 18, "单台容量": 55, "总容量": 990, "单位": "kW"},
            {"设备类型": "储能", "设计台数": 4, "单台容量": 250, "总容量": 1000, "单位": "kWh"},
            {"设备类型": "电制氢", "设计台数": 2, "单台容量": 180, "总容量": 360, "单位": "kW"},
            {"设备类型": "储氢罐", "设计台数": 3, "单台容量": 420, "总容量": 1260, "单位": "Nm3"},
            {"设备类型": "燃料电池", "设计台数": 2, "单台容量": 160, "总容量": 320, "单位": "kW"},
        ]
        return {
            "overview_tables": [
                {
                    "title": "规划结果",
                    "rows": planning_rows,
                },
                {
                    "title": "规划年指标",
                    "rows": [
                        {"指标": "柴发总容量", "数值": 640, "单位": "kW"},
                        {"指标": "风电总容量", "数值": 720, "单位": "kW"},
                        {"指标": "光伏总容量", "数值": 990, "单位": "kW"},
                        {"指标": "氢能总容量", "数值": 320, "单位": "kW"},
                        {"指标": "储能总容量", "数值": 1000, "单位": "kWh"},
                        {"指标": "负荷总电量", "数值": 6840.0, "单位": "MWh"},
                        {"指标": "柴发总电量", "数值": diesel_energy, "单位": "MWh"},
                        {"指标": "风能总电量", "数值": wind_energy, "单位": "MWh"},
                        {"指标": "光伏总电量", "数值": pv_energy, "单位": "MWh"},
                        {"指标": "弃电量", "数值": round(420 - self.progress * 2.1, 1), "单位": "MWh"},
                        {"指标": "储能总电量", "数值": storage_energy, "单位": "MWh"},
                        {"指标": "制氢总量", "数值": hydrogen_production, "单位": "万Nm3"},
                        {"指标": "燃料电池发电量", "数值": fuel_cell_energy, "单位": "MWh"},
                        *annual_energy_rows,
                        {"指标": "总成本", "数值": round(3280 - self.progress * 3.2, 1), "单位": "万元"},
                        {"指标": "绿电占比", "数值": green_ratio, "单位": "%"},
                        {"指标": "频率风险点", "数值": max(0, 6 - self.progress // 18), "单位": "个"},
                    ],
                },
            ],
            "overview_disks": [
                {
                    "title": "成本构成",
                    "left_label": "运行成本",
                    "left_value": operation_cost,
                    "right_label": "建设成本",
                    "right_value": construction_cost,
                    "unit": "万元",
                },
                plan_optimizer.capacity_composition_disk(planning_rows),
                {
                    "title": "电量构成",
                    "left_label": "柴发电量",
                    "left_value": diesel_energy,
                    "right_label": "新能源电量",
                    "right_value": renewable_energy,
                    "unit": "MWh",
                },
            ],
            "overview": [
                {"指标": "度电成本", "数值": cost, "单位": "元", "说明": "基于当前候选方案的综合成本估计"},
                {"指标": "绿电占比", "数值": green_ratio, "单位": "%", "说明": "按柴油年发电量占负荷年用电量的比例反推"},
                {"指标": "优化进度", "数值": self.progress, "单位": "%", "说明": self.status},
            ],
            "green": [
                {"指标": "风电消纳率", "数值": round(min(98.0, 82.0 + self.progress * 0.08), 1), "单位": "%", "说明": "风机出力消纳水平"},
                {"指标": "光伏消纳率", "数值": round(min(98.5, 84.0 + self.progress * 0.07), 1), "单位": "%", "说明": "光伏出力消纳水平"},
                {"指标": "弃电率", "数值": round(max(1.0, 9.0 - self.progress * 0.04), 1), "单位": "%", "说明": "新能源未利用电量占比"},
            ],
            "green_table": [
                *annual_energy_rows,
                {"指标": "负荷总电量", "数值": round(load_energy_kwh, 1), "单位": "kWh"},
                {"指标": "柴发总电量", "数值": round(diesel_energy_kwh, 1), "单位": "kWh"},
                {"指标": "风机总发电量", "数值": round(wind_energy_kwh, 1), "单位": "kWh"},
                {"指标": "光伏总发电量", "数值": round(pv_energy_kwh, 1), "单位": "kWh"},
                {"指标": "电储总发电量", "数值": round(storage_energy_kwh, 1), "单位": "kWh"},
                {"指标": "氢储总发电量", "数值": round(fuel_cell_energy_kwh, 1), "单位": "kWh"},
                {"指标": "新能源弃电率", "数值": energy_totals["renewable_curtailed_rate"], "单位": "%"},
                {"指标": "柴油消耗", "数值": round(diesel_energy * 0.24, 1), "单位": "吨"},
                {"指标": "制氢总量", "数值": round(hydrogen_production_nm3, 1), "单位": "Nm3"},
            ],
            "safety": [
                {"指标": "备用裕度", "数值": reserve_margin, "单位": "%", "说明": "负荷扰动后的可用备用"},
                {"指标": "频率安全裕度", "数值": frequency_margin, "单位": "p.u.", "说明": "频率约束裕度"},
                {"指标": "N-1校核", "数值": "通过" if self.progress >= 35 else "计算中", "单位": "", "说明": "新能源 N-1 约束校核"},
            ],
            "safety_table": [
                {"指标": "向上扰动最大量", "数值": upward_disturbance, "单位": "kW"},
                {"指标": "向下扰动最大量", "数值": downward_disturbance, "单位": "kW"},
                {"指标": "最高频率", "数值": highest_frequency, "单位": "Hz"},
                {"指标": "最低频率", "数值": lowest_frequency, "单位": "Hz"},
                {"指标": "频率安全风险小时数", "数值": frequency_risk_hours, "单位": "h"},
            ],
            "curves": {
                "overview": [
                    {"label": "20%", "value": round(cost + 0.08, 3)},
                    {"label": "40%", "value": round(cost + 0.05, 3)},
                    {"label": "60%", "value": round(cost + 0.03, 3)},
                    {"label": "80%", "value": round(cost + 0.01, 3)},
                    {"label": "当前", "value": cost},
                ],
                "green": [
                    {"label": "风电", "value": round(min(98.0, 82.0 + self.progress * 0.08), 1)},
                    {"label": "光伏", "value": round(min(98.5, 84.0 + self.progress * 0.07), 1)},
                    {"label": "氢储", "value": round(min(90.0, 66.0 + self.progress * 0.09), 1)},
                    {"label": "总体", "value": green_ratio},
                ],
                "green_daily": green_daily,
                "green_monthly": green_monthly,
                "green_hourly": green_hourly,
                "safety": [
                    {"label": "备用", "value": reserve_margin},
                    {"label": "频率", "value": round(frequency_margin * 10, 2)},
                    {"label": "N-1", "value": 100 if self.progress >= 35 else max(10, self.progress)},
                ],
                "safety_daily": safety_daily,
            },
        }

    @staticmethod
    def _energy_totals_from_daily_unlocked(daily_rows: list[dict[str, float | int]]) -> dict[str, float]:
        totals = {
            field: round(sum(float(row.get(field, 0.0) or 0.0) for row in daily_rows), 4)
            for field in estimate.ENERGY_AGGREGATE_FIELDS
        }
        estimate.add_energy_ratios(totals)
        return totals

    def _green_daily_curve_unlocked(
        self,
        load_total: float,
        diesel_total: float,
        wind_total: float,
        pv_total: float,
        hydrogen_total: float,
        hydrogen_production_total: float,
        storage_discharge_total: float,
        storage_charge_total: float,
    ) -> list[dict[str, float | int]]:
        days = list(range(1, 366))

        def seasonal(day: int, phase_shift: float, amplitude: float, second: float = 0.0) -> float:
            phase = 2 * math.pi * (day - 1) / 365
            return max(0.05, 1 + amplitude * math.sin(phase + phase_shift) + second * math.sin(phase * 2 + phase_shift / 2))

        raw = {
            "load_energy": [seasonal(day, -0.25, 0.12, 0.05) for day in days],
            "diesel_energy": [seasonal(day, 0.8, 0.18, 0.03) for day in days],
            "wind_energy": [seasonal(day, 1.25, 0.22, 0.05) for day in days],
            "pv_energy": [seasonal(day, -1.15, 0.34, 0.04) for day in days],
            "hydrogen_energy": [seasonal(day, 0.35, 0.16, 0.03) for day in days],
            "hydrogen_production_energy": [seasonal(day, -0.75, 0.2, 0.04) for day in days],
            "storage_discharge_energy": [seasonal(day, 0.1, 0.11, 0.06) for day in days],
            "storage_charge_energy": [seasonal(day, -0.9, 0.18, 0.05) for day in days],
        }
        totals = {
            "load_energy": load_total,
            "diesel_energy": diesel_total,
            "wind_energy": wind_total,
            "pv_energy": pv_total,
            "hydrogen_energy": hydrogen_total,
            "hydrogen_production_energy": hydrogen_production_total,
            "storage_discharge_energy": storage_discharge_total,
            "storage_charge_energy": storage_charge_total,
        }
        scaled: dict[str, list[float]] = {}
        for key, values in raw.items():
            raw_total = sum(values) or 1
            scaled[key] = [round(totals[key] * value / raw_total, 1) for value in values]

        return [
            {
                "day": day,
                "diesel_energy": scaled["diesel_energy"][index],
                "wind_energy": scaled["wind_energy"][index],
                "pv_energy": scaled["pv_energy"][index],
                "hydrogen_energy": scaled["hydrogen_energy"][index],
                "storage_discharge_energy": scaled["storage_discharge_energy"][index],
                "load_energy": scaled["load_energy"][index],
                "hydrogen_production_energy": scaled["hydrogen_production_energy"][index],
                "storage_charge_energy": scaled["storage_charge_energy"][index],
            }
            for index, day in enumerate(days)
        ]

    def _safety_daily_frequency_curve_unlocked(self) -> list[dict[str, float | int]]:
        days = list(range(1, 366))
        max_deviation_base = max(0.07, 0.22 - self.progress * 0.0012)
        min_deviation_base = max(0.06, 0.2 - self.progress * 0.001)
        points = []
        for day in days:
            phase = 2 * math.pi * (day - 1) / 365
            frequency_max = 50 + max_deviation_base + 0.035 * math.sin(phase + 0.7) + 0.012 * math.sin(phase * 2.4)
            frequency_min = 50 - min_deviation_base - 0.032 * math.cos(phase - 0.25) - 0.01 * math.sin(phase * 2.1)
            points.append(
                {
                    "day": day,
                    "frequency_max": round(frequency_max, 3),
                    "frequency_min": round(frequency_min, 3),
                }
            )
        return points

    def _green_hourly_curve_unlocked(self, daily_points: list[dict[str, float | int]]) -> list[dict[str, float | int | str]]:
        rows: list[dict[str, float | int | str]] = []
        hour_index = 1
        for day_point in daily_points:
            day = int(day_point.get("day") or 1)
            for hour in range(24):
                load_shape = 0.85 + 0.28 * math.sin((hour - 7) / 24 * 2 * math.pi) + 0.12 * math.sin((hour - 18) / 24 * 4 * math.pi)
                wind_shape = 0.9 + 0.18 * math.sin((hour + 3) / 24 * 2 * math.pi)
                daylight = max(0.0, math.sin((hour - 6) / 12 * math.pi))
                storage_discharge_shape = 1.0 + 0.35 * math.sin((hour - 17) / 24 * 2 * math.pi)
                storage_charge_shape = 0.35 + daylight

                load = max(0.0, float(day_point.get("load_energy", 0.0)) * load_shape / 24)
                wind_power = max(0.0, float(day_point.get("wind_energy", 0.0)) * wind_shape / 24)
                pv_power = max(0.0, float(day_point.get("pv_energy", 0.0)) * daylight / 7.6)
                storage_discharge = max(0.0, float(day_point.get("storage_discharge_energy", 0.0)) * storage_discharge_shape / 24)
                storage_charge = max(0.0, float(day_point.get("storage_charge_energy", 0.0)) * storage_charge_shape / 24)
                diesel_power = max(0.0, float(day_point.get("diesel_energy", 0.0)) * (1.0 + 0.16 * math.sin((hour + 9) / 24 * 2 * math.pi)) / 24)
                renewable_surplus = max(0.0, wind_power + pv_power + storage_discharge + diesel_power - load - storage_charge)
                curtailed_power = min(renewable_surplus, max(0.0, float(day_point.get("wind_energy", 0.0)) + float(day_point.get("pv_energy", 0.0))) / 24)
                unmet_load = max(0.0, load + storage_charge - wind_power - pv_power - storage_discharge - diesel_power)
                storage_soc = max(0.0, min(100.0, 52.0 + 30.0 * math.sin((day + hour / 24) / 365 * 2 * math.pi) + 12.0 * math.sin((hour - 5) / 24 * 2 * math.pi)))
                wind_speed = max(0.0, 7.5 + 2.2 * math.sin((day + hour / 24) / 365 * 2 * math.pi) + 1.1 * math.sin((hour + 2) / 24 * 2 * math.pi))
                solar_irradiance = max(0.0, daylight * (780.0 + 120.0 * math.sin((day - 80) / 365 * 2 * math.pi)))
                temperature = -5.0 + 18.0 * math.sin((day - 80) / 365 * 2 * math.pi) + 4.0 * math.sin((hour - 14) / 24 * 2 * math.pi)
                wind_available = wind_power + curtailed_power * 0.55
                pv_available = pv_power + curtailed_power * 0.45
                wind_curtailed_power = max(0.0, wind_available - wind_power)
                pv_curtailed_power = max(0.0, pv_available - pv_power)
                renewable_available = wind_available + pv_available
                renewable_power = wind_power + pv_power
                renewable_curtailed_power = wind_curtailed_power + pv_curtailed_power
                rows.append(
                    {
                        "hour_index": hour_index,
                        "datetime": f"D{day:03d} H{hour + 1:02d}",
                        "wind_speed": round(wind_speed, 4),
                        "solar_irradiance": round(solar_irradiance, 4),
                        "temperature": round(temperature, 4),
                        "load": round(load, 4),
                        "diesel_power": round(diesel_power, 4),
                        "wind_available": round(wind_available, 4),
                        "wind_power": round(wind_power, 4),
                        "pv_available": round(pv_available, 4),
                        "pv_power": round(pv_power, 4),
                        "renewable_available": round(renewable_available, 4),
                        "renewable_ratio": round(estimate.green_power_ratio_from_diesel_load(diesel_power, load), 4),
                        "storage_power": round(storage_discharge - storage_charge, 4),
                        "storage_soc": round(storage_soc, 4),
                        "hydrogen_production_power": 0,
                        "hydrogen_storage": 0,
                        "fuel_cell_power": 0,
                        "wind_curtailed_power": round(wind_curtailed_power, 4),
                        "pv_curtailed_power": round(pv_curtailed_power, 4),
                        "curtailed_power": round(renewable_curtailed_power, 4),
                        "renewable_curtailed_rate": round(renewable_curtailed_power / renewable_available * 100 if renewable_available else 0.0, 4),
                        "unmet_load": round(unmet_load, 4),
                        "diesel_on": 1 if diesel_power > 0 else 0,
                        "electrolyzer_on": 0,
                        "storage_charge": round(storage_charge, 4),
                        "storage_discharge": round(storage_discharge, 4),
                    }
                )
                hour_index += 1
        return rows[:8760]

    def _append_log_unlocked(self, level: str, message: str) -> None:
        self._logs.append({"time": _now_text(), "level": level, "message": message})
        if len(self._logs) > 2000:
            del self._logs[:-2000]

    def _terminate_process_unlocked(self) -> None:
        if self._process and self._process.is_alive():
            self._process.terminate()
            self._process.join(timeout=2)
            if self._process.is_alive():
                self._process.kill()
                self._process.join(timeout=2)
        if self._process and not self._process.is_alive():
            self._process.join(timeout=0)
        self._close_event_queue_unlocked()

    def _join_finished_process_unlocked(self) -> None:
        if self._process and not self._process.is_alive():
            self._process.join(timeout=0)

    def _reap_process_unlocked(self) -> None:
        if self._process and self.status != "运行中" and not self._process.is_alive():
            self._process.join(timeout=0)

    def _close_event_queue_unlocked(self) -> None:
        if not self._event_queue:
            return
        try:
            self._event_queue.close()
            self._event_queue.join_thread()
        except Exception:
            pass
        self._event_queue = None

    def _export_results_once_unlocked(self) -> None:
        if self._results_exported:
            return
        result_path = export_optimization_results_workbook(self._payload_unlocked(read_workbook=False))
        self.result_file = str(result_path)
        self._results_exported = True


def _now_text() -> str:
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")


def elapsed_seconds_from_times(start_time: str, end_time: str = "") -> int:
    if not start_time:
        return 0
    try:
        start = datetime.strptime(start_time, "%Y-%m-%d %H:%M:%S")
        end = datetime.strptime(end_time, "%Y-%m-%d %H:%M:%S") if end_time else datetime.now()
    except ValueError:
        return 0
    return max(0, int((end - start).total_seconds()))


class OptimizationStateError(RuntimeError):
    """Raised when optimization start/stop violates the current runtime state."""

    def __init__(self, code: str, message: str) -> None:
        super().__init__(message)
        self.code = code


def calculation_timeout_message(message: object) -> bool:
    """Return True when a worker error message represents a solver time-limit stop."""

    return milp_solver.is_timeout_text(message)


def optimization_process_worker(event_queue, scheme: str, planning_root: str = "") -> None:
    """Run planning optimization in a child process and report serializable events."""
    try:
        event_queue.put({"type": "log", "level": "info", "message": "读取方案参数和8760时序数据"})
        store = planning_store.PlanningStore(root=Path(planning_root)) if planning_root else PLANNING_STORE
        scheme_payload = store.read_scheme(scheme)
        result = plan_optimizer.run_optimization(
            scheme_payload,
            log=lambda event: event_queue.put({"type": "log", **dict(event or {})}),
        )
        event_queue.put(
            {
                "type": "done",
                "metrics": result.get("metrics") if isinstance(result.get("metrics"), list) else [],
                "results": result.get("results") if isinstance(result.get("results"), dict) else {},
            }
        )
    except milp_solver.CalculationTimeoutError as exc:
        event_queue.put({"type": "timeout", "message": f"规划求解超时：{exc}", "traceback": traceback.format_exc()})
    except Exception as exc:
        event_queue.put({"type": "error", "message": f"规划求解失败：{exc}", "traceback": traceback.format_exc()})


def evaluation_process_worker(event_queue, scheme: str, filename: str, planning_root: str = "") -> None:
    """Run fixed-plan evaluation in a child process and report serializable events."""
    try:
        event_queue.put({"type": "log", "level": "info", "message": "读取方案参数和当前规划结果"})
        store = planning_store.PlanningStore(root=Path(planning_root)) if planning_root else PLANNING_STORE
        scheme_payload = store.read_scheme(scheme)
        planning_rows = read_evaluation_planning_result_rows_with_store(store, scheme, filename)
        if not planning_rows:
            raise ValueError("当前结果文件缺少规划结果")
        result = estimate.run_estimation(
            scheme_payload,
            planning_rows,
            log=lambda event: event_queue.put({"type": "log", **dict(event or {})}),
        )
        event_queue.put(
            {
                "type": "done",
                "metrics": result.get("metrics") if isinstance(result.get("metrics"), list) else [],
                "results": result.get("results") if isinstance(result.get("results"), dict) else {},
                "dispatch_rows": result.get("dispatch_rows") if isinstance(result.get("dispatch_rows"), list) else [],
            }
        )
    except milp_solver.CalculationTimeoutError as exc:
        event_queue.put({"type": "timeout", "message": f"方案评估超时：{exc}", "traceback": traceback.format_exc()})
    except Exception as exc:
        event_queue.put({"type": "error", "message": f"方案评估失败：{exc}", "traceback": traceback.format_exc()})


def export_optimization_results_workbook(payload: dict) -> Path:
    scheme = str(payload.get("scheme") or "未选择方案")
    result_path = optimization_result_workbook_path(scheme)
    result_path.parent.mkdir(parents=True, exist_ok=True)
    save_result_workbook(build_optimization_results_workbook(payload, include_curves=False), result_path, "结果文件")
    save_result_workbook(build_result_curves_workbook(payload), result_curves_workbook_path(result_path), "曲线结果文件")
    return result_path


def optimization_result_workbook_path(scheme: str) -> Path:
    return PLANNING_STORE.scheme_dir(str(scheme or "未选择方案")) / OPTIMIZATION_RESULT_WORKBOOK_NAME


def export_evaluation_results_workbook(payload: dict, dispatch_rows: list[dict]) -> Path:
    scheme = str(payload.get("scheme") or "未选择方案")
    filename = str(payload.get("result_filename") or "").strip()
    result_path = evaluation_result_path(scheme, filename)
    if result_path.name == OPTIMIZATION_RESULT_WORKBOOK_NAME:
        raise ValueError("默认结果文件不允许修改")
    result_path.parent.mkdir(parents=True, exist_ok=True)
    save_result_workbook(build_optimization_results_workbook(payload, include_curves=False), result_path, "结果文件")
    save_result_workbook(build_result_curves_workbook(payload), result_curves_workbook_path(result_path), "曲线结果文件")
    return result_path


def replace_result_workbook_with_retry(source: Path, target: Path, attempts: int = 20, delay_seconds: float = 0.1) -> None:
    file_ops.retry_file_operation(
        lambda: source.replace(target),
        f"结果文件被占用，无法保存：{target.name}。请关闭正在打开该文件的 Excel 或预览窗口后重试。",
        attempts=attempts,
        delay_seconds=delay_seconds,
    )
    file_cache.invalidate_path(source)
    file_cache.invalidate_path(target)


def save_result_workbook(workbook: Workbook, result_path: Path, label: str) -> None:
    tmp_path = result_path.with_name(f".{result_path.name}.tmp")
    try:
        file_ops.save_workbook_with_retry(workbook, tmp_path, label)
    finally:
        workbook.close()
    replace_result_workbook_with_retry(tmp_path, result_path)


def result_curves_workbook_path(result_path: Path) -> Path:
    stem = result_path.stem
    if stem.endswith("_results"):
        stem = stem[: -len("_results")]
    return result_path.with_name(f"{stem}_curves.xlsx")


def build_optimization_results_workbook(payload: dict, include_curves: bool = False) -> Workbook:
    workbook = Workbook()
    workbook.remove(workbook.active)
    results = payload.get("results") if isinstance(payload.get("results"), dict) else {}
    metrics = payload.get("metrics") if isinstance(payload.get("metrics"), list) else []
    logs = payload.get("logs") if isinstance(payload.get("logs"), list) else []

    append_rows_sheet(
        workbook,
        "总体指标",
        [
            {"指标": "方案", "数值": payload.get("scheme", ""), "单位": ""},
            {"指标": "状态", "数值": payload.get("status", ""), "单位": ""},
            {"指标": "进度", "数值": payload.get("progress", ""), "单位": "%"},
            *[
                {
                    "指标": metric.get("label", ""),
                    "数值": metric.get("value", ""),
                    "单位": metric.get("unit", ""),
                }
                for metric in metrics
            ],
        ],
        ["指标", "数值", "单位"],
    )

    for table in results.get("overview_tables", []):
        append_rows_sheet(workbook, str(table.get("title", "结果表")), table.get("rows", []))
    append_rows_sheet(workbook, "供能分析", results.get("green_table", []))
    append_rows_sheet(workbook, "安全评估", results.get("safety_table", []))
    if include_curves:
        append_result_curve_sheets(workbook, payload)
    append_rows_sheet(workbook, "运行日志", logs, ["time", "level", "message"], {"time": "时间", "level": "级别", "message": "消息"})
    return workbook


def build_result_curves_workbook(payload: dict) -> Workbook:
    workbook = Workbook()
    workbook.remove(workbook.active)
    append_result_curve_sheets(workbook, payload)
    return workbook


def append_result_curve_sheets(workbook: Workbook, payload: dict) -> None:
    results = payload.get("results") if isinstance(payload.get("results"), dict) else {}
    curves = results.get("curves") if isinstance(results.get("curves"), dict) else {}
    append_rows_sheet(workbook, "供能日曲线", curves.get("green_daily", []))
    append_rows_sheet(workbook, "供能月曲线", curves.get("green_monthly", []))
    append_rows_sheet(workbook, "安全日曲线", curves.get("safety_daily", []))
    append_dispatch_rows_sheet(workbook, curves.get("green_hourly", []))


def append_dispatch_rows_sheet(workbook: Workbook, dispatch_rows: list[dict]) -> None:
    append_rows_sheet(
        workbook,
        "调度结果",
        dispatch_rows,
        [
            "hour_index",
            "datetime",
            "wind_speed",
            "solar_irradiance",
            "temperature",
            "load",
            "diesel_power",
            "diesel_capacity",
            "wind_available",
            "wind_power",
            "pv_available",
            "pv_power",
            "renewable_power",
            "renewable_available",
            "renewable_ratio",
            "renewable_curtailed_rate",
            "storage_power",
            "grid_storage_capacity",
            "grid_storage_power",
            "storage_soc",
            "hydrogen_production_power",
            "hydrogen_storage",
            "fuel_cell_power",
            "wind_curtailed_power",
            "pv_curtailed_power",
            "curtailed_power",
            "unmet_load",
            "load_up_disturbance_power",
            "load_down_disturbance_power",
            "renewable_down_disturbance_power",
            "renewable_single_unit_power_max",
            "grid_up_regulation_capacity",
            "grid_down_regulation_capacity",
            "grid_up_regulation_requirement",
            "grid_down_regulation_requirement",
            "frequency_min",
            "frequency_max",
            "frequency_nadir_est_hz",
            "frequency_peak_est_hz",
            "frequency_nadir_exact_hz",
            "frequency_peak_exact_hz",
            "steady_state_frequency_min_hz",
            "steady_state_frequency_max_hz",
            "rocof_hz_per_s",
            "rocof_upper_hz_per_s",
            "frequency_lower_margin_hz",
            "frequency_upper_margin_hz",
            "equivalent_inertia_m",
            "equivalent_primary_frequency_k",
            "equivalent_damping_d",
            "frequency_delta_p_mw",
            "frequency_upper_delta_p_mw",
            "frequency_fit_error_hz",
            "frequency_upper_fit_error_hz",
            "diesel_on",
            "electrolyzer_on",
            "storage_charge",
            "storage_discharge",
        ],
        {
            "hour_index": "小时",
            "datetime": "时间",
            "wind_speed": "风速",
            "solar_irradiance": "太阳辐射",
            "temperature": "环境温度",
            "load": "负荷总功率",
            "diesel_power": "柴发总功率",
            "diesel_capacity": "柴发开机容量",
            "wind_available": "风力最大可发",
            "wind_power": "风机总功率",
            "pv_available": "光伏最大可发",
            "pv_power": "光伏总功率",
            "renewable_power": "新能源总出力",
            "renewable_available": "新能源最大可发",
            "renewable_ratio": "新能源占比",
            "renewable_curtailed_rate": "新能源弃电率",
            "storage_power": "电储能总功率",
            "grid_storage_capacity": "构网储能总容量",
            "grid_storage_power": "构网储能总功率",
            "storage_soc": "电储电量",
            "hydrogen_production_power": "电制氢总功率",
            "hydrogen_storage": "储氢罐氢储量",
            "fuel_cell_power": "燃料电池总功率",
            "wind_curtailed_power": "弃风总功率",
            "pv_curtailed_power": "弃光总功率",
            "curtailed_power": "新能源弃电总功率",
            "unmet_load": "切负荷功率",
            "load_up_disturbance_power": "负荷上扰动功率",
            "load_down_disturbance_power": "负荷下扰动功率",
            "renewable_down_disturbance_power": "新能源下扰动功率",
            "renewable_single_unit_power_max": "风光单机功率最大值",
            "grid_up_regulation_capacity": "电网向上调节能力",
            "grid_down_regulation_capacity": "电网向下调节能力",
            "grid_up_regulation_requirement": "电网向上调节需求",
            "grid_down_regulation_requirement": "电网向下调节需求",
            "frequency_min": "最低频率",
            "frequency_max": "最高频率",
            "frequency_nadir_est_hz": "最低频率保守估计值",
            "frequency_peak_est_hz": "最高频率保守估计值",
            "frequency_nadir_exact_hz": "最低频率解析值",
            "frequency_peak_exact_hz": "最高频率解析值",
            "steady_state_frequency_min_hz": "下限场景稳态频率",
            "steady_state_frequency_max_hz": "上限场景稳态频率",
            "rocof_hz_per_s": "初始频率变化率",
            "rocof_upper_hz_per_s": "上限场景初始频率变化率",
            "frequency_lower_margin_hz": "频率下限裕度",
            "frequency_upper_margin_hz": "频率上限裕度",
            "equivalent_inertia_m": "等效惯量M",
            "equivalent_primary_frequency_k": "等效调频系数K",
            "equivalent_damping_d": "等效阻尼系数D",
            "frequency_delta_p_mw": "频率下限扰动功率",
            "frequency_upper_delta_p_mw": "频率上限扰动功率",
            "frequency_fit_error_hz": "频率下限拟合误差",
            "frequency_upper_fit_error_hz": "频率上限拟合误差",
            "diesel_on": "柴发启停",
            "electrolyzer_on": "制氢启停",
            "storage_charge": "储能充电",
            "storage_discharge": "储能放电",
        },
    )


def append_rows_sheet(
    workbook: Workbook,
    title: str,
    rows: list[dict],
    headers: list[str] | None = None,
    labels: dict[str, str] | None = None,
) -> None:
    sheet = workbook.create_sheet(safe_sheet_title(title, workbook.sheetnames))
    normalized_rows = rows if isinstance(rows, list) else []
    header_keys = headers or keys_from_rows(normalized_rows)
    if not header_keys:
        header_keys = ["内容"]
    header_labels = labels or {}
    sheet.append([header_labels.get(key, key) for key in header_keys])
    for row in normalized_rows:
        if isinstance(row, dict):
            sheet.append([row.get(key, "") for key in header_keys])
    style_result_sheet(sheet)


def keys_from_rows(rows: list[dict]) -> list[str]:
    keys: list[str] = []
    for row in rows:
        if not isinstance(row, dict):
            continue
        for key in row:
            if key not in keys:
                keys.append(key)
    return keys


def safe_sheet_title(title: str, existing: list[str]) -> str:
    clean = re.sub(r"[\[\]:*?/\\]", "_", str(title or "Sheet")).strip() or "Sheet"
    clean = clean[:31]
    if clean not in existing:
        return clean
    for index in range(2, 1000):
        suffix = f"_{index}"
        candidate = f"{clean[:31 - len(suffix)]}{suffix}"
        if candidate not in existing:
            return candidate
    raise ValueError("无法生成唯一工作表名称")


def style_result_sheet(sheet) -> None:
    header_fill = PatternFill(fill_type="solid", fgColor="D9EAF7")
    for cell in sheet[1]:
        cell.font = Font(bold=True)
        cell.fill = header_fill
    sheet.freeze_panes = "A2"
    for column_cells in sheet.columns:
        max_length = max(len(str(cell.value)) if cell.value is not None else 0 for cell in column_cells)
        sheet.column_dimensions[get_column_letter(column_cells[0].column)].width = min(max(max_length + 2, 10), 28)


def is_reliability_export_workbook(path: Path | str) -> bool:
    """Return True for reliability-only XLSX exports.

    Reliability exports deliberately end in ``_results.xlsx`` so they are
    convenient to recognize outside the application.  They are not planning
    result workbooks, however, and must never be offered as inputs to another
    optimization/evaluation run.
    """

    return Path(path).name.endswith(RELIABILITY_RESULT_WORKBOOK_SUFFIX)


def list_evaluation_result_files(scheme: str) -> list[dict]:
    folder = PLANNING_STORE.scheme_dir(scheme)
    if not folder.exists():
        raise FileNotFoundError(f"方案不存在: {scheme}")
    files = []
    for path in sorted(folder.glob("*_results.xlsx"), key=lambda item: item.name):
        if path.is_file() and RESULT_WORKBOOK_RE.fullmatch(path.name) and not is_reliability_export_workbook(path):
            item = {"name": path.name, "modified_at": path.stat().st_mtime, "readable": True, "message": ""}
            error_message = result_workbook_error_message(path)
            if error_message:
                item["readable"] = False
                item["message"] = error_message
            files.append(item)
    return files


def list_evaluation_result_files_for_tasks(scheme: str) -> list[dict]:
    folder = PLANNING_STORE.scheme_dir(scheme)
    if not folder.exists():
        raise FileNotFoundError(f"方案不存在: {scheme}")
    return TASK_RESULT_FILE_LIST_CACHE.get(folder, list_evaluation_result_files_for_tasks_uncached, variant="task_results")


def list_evaluation_result_files_for_tasks_uncached(folder: Path) -> list[dict]:
    files = []
    for path in sorted(folder.glob("*_results.xlsx"), key=lambda item: item.name):
        if path.is_file() and RESULT_WORKBOOK_RE.fullmatch(path.name) and not is_reliability_export_workbook(path):
            files.append({"name": path.name, "modified_at": path.stat().st_mtime, "readable": True, "message": ""})
    return files


def selected_evaluation_result_filename(scheme: str, filename: str = "") -> str:
    files = list_evaluation_result_files(scheme)
    names = [item["name"] for item in files]
    readable_names = [item["name"] for item in files if item.get("readable", True)]
    selected = str(filename or "").strip()
    if selected and selected in names:
        return selected
    return readable_names[0] if readable_names else ""


def result_workbook_error_message(path: Path) -> str:
    return RESULT_WORKBOOK_HEALTH_CACHE.get(path, result_workbook_error_message_uncached)


def result_workbook_error_message_uncached(path: Path) -> str:
    try:
        workbook = load_workbook(path, read_only=True, data_only=True)
        workbook.close()
        return ""
    except RESULT_WORKBOOK_READ_ERRORS:
        return "结果文件无法读取，可能已损坏或格式不正确"


def evaluation_result_path(scheme: str, filename: str) -> Path:
    return evaluation_result_path_with_store(PLANNING_STORE, scheme, filename)


def evaluation_result_path_with_store(store: planning_store.PlanningStore, scheme: str, filename: str) -> Path:
    name = str(filename or "").strip()
    if not RESULT_WORKBOOK_RE.fullmatch(name):
        raise ValueError("结果文件名必须符合 xxxx_results.xlsx")
    folder = store.scheme_dir(scheme)
    path = (folder / name).resolve()
    if folder not in path.parents or path.parent != folder:
        raise ValueError("结果文件路径越界")
    return path


def read_evaluation_planning_result_rows(scheme: str, filename: str) -> list[dict]:
    return read_evaluation_planning_result_rows_with_store(PLANNING_STORE, scheme, filename)


def read_evaluation_planning_result_rows_with_store(
    store: planning_store.PlanningStore,
    scheme: str,
    filename: str,
) -> list[dict]:
    if not filename:
        return []
    result_path = evaluation_result_path_with_store(store, scheme, filename)
    if not result_path.exists():
        return []
    return EVALUATION_PLANNING_RESULT_CACHE.get(
        result_path,
        read_evaluation_planning_result_rows_from_path,
        variant="planning_result_rows",
    )


def read_evaluation_planning_result_rows_from_path(result_path: Path) -> list[dict]:
    try:
        workbook = load_workbook(result_path, read_only=True, data_only=True)
    except RESULT_WORKBOOK_READ_ERRORS as exc:
        raise ValueError(f"结果文件无法读取: {result_path.name}") from exc
    try:
        if PLANNING_RESULT_SHEET_NAME not in workbook.sheetnames:
            return []
        sheet = workbook[PLANNING_RESULT_SHEET_NAME]
        rows = list(sheet.iter_rows(values_only=True))
        if not rows:
            return []
        headers = [str(value or "").strip() for value in rows[0]]
        result_rows = []
        for row in rows[1:]:
            item = {}
            for index, header in enumerate(headers):
                if header:
                    item[header] = row[index] if index < len(row) else ""
            if any(value not in (None, "") for value in item.values()):
                normalize_planning_result_total_capacity(item)
                result_rows.append(item)
        return result_rows
    finally:
        workbook.close()


def normalize_planning_result_total_capacity(row: dict) -> dict:
    if "设计台数" in row and "单台容量" in row:
        row["总容量"] = round(
            estimate.numeric(row.get("设计台数"), 0.0) * estimate.numeric(row.get("单台容量"), 0.0),
            4,
        )
    return row


def read_evaluation_planning_result_rows_for_response(scheme: str, filename: str) -> list[dict]:
    try:
        return read_evaluation_planning_result_rows(scheme, filename)
    except ValueError:
        return []


def handle_comparison_data_api_path(
    path: str,
    query: str = "",
    current_user: dict | None = None,
) -> tuple[int, dict[str, str], bytes]:
    if path != "/api/comparison/data":
        return _json_response({"error": "not_found", "path": path}, HTTPStatus.NOT_FOUND)
    try:
        query_params = parse_qs(query)
        items_text = query_params.get("items", ["[]"])[0]
        mode = str(query_params.get("mode", ["full"])[0] or "full").strip().lower()
        curve_group = str(query_params.get("group", ["hourly"])[0] or "hourly").strip().lower()
        curve_names = parse_comparison_curve_names(query_params.get("curves", ["[]"])[0])
        include_hourly_curves = mode not in {"summary", "tables", "light"}
        items = json.loads(items_text or "[]")
        if not isinstance(items, list):
            raise ValueError("对比项必须为列表")
        for item in items:
            if isinstance(item, dict) and str(item.get("scheme", "")).strip():
                ensure_planning_scheme_access(str(item.get("scheme", "")), current_user)
        if mode in {"curve", "curves"}:
            return _json_response(build_comparison_curve_payload(items[:8], curve_group, curve_names[:32]))
        return _json_response(build_comparison_payload(items[:8], include_hourly_curves=include_hourly_curves))
    except FileNotFoundError as exc:
        return _json_response({"error": "not_found", "message": str(exc)}, HTTPStatus.NOT_FOUND)
    except (ValueError, json.JSONDecodeError) as exc:
        return _json_response({"error": "bad_request", "message": str(exc)}, HTTPStatus.BAD_REQUEST)


def parse_comparison_curve_names(text: str) -> list[str]:
    try:
        parsed = json.loads(text or "[]")
    except json.JSONDecodeError:
        parsed = [item.strip() for item in str(text or "").split(",")]
    if not isinstance(parsed, list):
        raise ValueError("曲线名称必须为列表")
    names = []
    for item in parsed:
        name = str(item or "").strip()
        if name and name not in names:
            names.append(name)
    return names


def build_comparison_payload(items: list[dict], include_hourly_curves: bool = True) -> dict:
    selected_items: list[dict] = []
    capacity_tables: list[list[dict]] = []
    energy_tables: list[list[dict]] = []
    safety_tables: list[list[dict]] = []
    annual_tables: list[list[dict]] = []
    curve_groups = empty_comparison_curve_groups()

    for index, item in enumerate(items):
        if not isinstance(item, dict):
            continue
        scheme = str(item.get("scheme", "")).strip()
        filename = str(item.get("filename", "")).strip()
        if not scheme or not filename:
            continue
        result_path = evaluation_result_path(scheme, filename)
        if not result_path.exists():
            raise FileNotFoundError(f"结果文件不存在: {result_path.name}")
        result_display_name = result_display_name_from_filename(filename)
        label = f"{scheme} / {result_display_name}"
        workbook_data = read_comparison_workbook(result_path, include_hourly_curves=include_hourly_curves)
        selected_items.append(
            {
                "id": f"item-{index + 1}",
                "scheme": scheme,
                "filename": filename,
                "result_display_name": result_display_name,
                "label": label,
            }
        )
        capacity_tables.append(workbook_data["capacity"])
        energy_tables.append(workbook_data["energy"])
        safety_tables.append(workbook_data["safety"])
        annual_tables.append(workbook_data["annual"])
        append_comparison_curve_groups(curve_groups, workbook_data["curve_groups"], label, scheme, filename)

    hourly_group = curve_groups["hourly"]

    return {
        "items": selected_items,
        "tables": {
            "capacity": merge_comparison_rows(capacity_tables, selected_items, "设备类型"),
            "energy": merge_comparison_rows(energy_tables, selected_items, "指标"),
            "safety": merge_comparison_rows(safety_tables, selected_items, "指标"),
        },
        "curve_groups": curve_groups,
        "annual_table": merge_comparison_rows(annual_tables, selected_items, "指标"),
        "curves": hourly_group["curves"],
        "series": hourly_group["series"],
    }


def build_comparison_curve_payload(items: list[dict], curve_group: str, curve_names: list[str]) -> dict:
    if curve_group not in COMPARISON_CURVE_GROUPS:
        raise ValueError("曲线类型不合法")
    selected_items: list[dict] = []
    curve_groups = empty_comparison_curve_groups()
    selected_curve_names = [str(name or "").strip() for name in curve_names if str(name or "").strip()]
    if not selected_curve_names:
        hourly_group = curve_groups["hourly"]
        return {
            "items": [],
            "tables": {"capacity": [], "energy": [], "safety": []},
            "curve_groups": curve_groups,
            "annual_table": [],
            "curves": hourly_group["curves"],
            "series": hourly_group["series"],
        }

    for index, item in enumerate(items):
        if not isinstance(item, dict):
            continue
        scheme = str(item.get("scheme", "")).strip()
        filename = str(item.get("filename", "")).strip()
        if not scheme or not filename:
            continue
        result_path = evaluation_result_path(scheme, filename)
        if not result_path.exists():
            raise FileNotFoundError(f"结果文件不存在: {result_path.name}")
        result_display_name = result_display_name_from_filename(filename)
        label = f"{scheme} / {result_display_name}"
        selected_items.append(
            {
                "id": f"item-{index + 1}",
                "scheme": scheme,
                "filename": filename,
                "result_display_name": result_display_name,
                "label": label,
            }
        )
        source_group = read_comparison_curve_group(result_path, curve_group, selected_curve_names)
        append_comparison_curve_groups(curve_groups, {curve_group: source_group}, label, scheme, filename)

    hourly_group = curve_groups["hourly"]
    return {
        "items": selected_items,
        "tables": {"capacity": [], "energy": [], "safety": []},
        "curve_groups": curve_groups,
        "annual_table": [],
        "curves": hourly_group["curves"],
        "series": hourly_group["series"],
    }


def read_comparison_workbook(path: Path, include_hourly_curves: bool = True) -> dict:
    if include_hourly_curves:
        ensure_split_result_workbook(path)
    curve_signature = result_curve_file_signature(path)
    return COMPARISON_WORKBOOK_CACHE.get(
        path,
        lambda resolved: read_comparison_workbook_uncached(resolved, include_hourly_curves=include_hourly_curves),
        variant=("comparison", bool(include_hourly_curves), curve_signature),
    )


def read_comparison_workbook_uncached(path: Path, include_hourly_curves: bool = True) -> dict:
    try:
        workbook = load_workbook(path, read_only=True, data_only=True)
    except RESULT_WORKBOOK_READ_ERRORS as exc:
        raise ValueError(f"结果文件无法读取: {path.name}") from exc
    try:
        curve_groups = read_comparison_curve_groups(path, fallback_workbook=workbook, include_hourly_curves=include_hourly_curves)
        return {
            "capacity": read_named_sheet_rows(workbook, "规划结果"),
            "energy": read_named_sheet_rows(workbook, "供能分析"),
            "safety": read_named_sheet_rows(workbook, "安全评估"),
            "annual": read_annual_comparison_rows(workbook),
            "curve_groups": curve_groups,
            "curves": curve_groups["hourly"],
        }
    finally:
        workbook.close()


def read_comparison_curve_group(path: Path, curve_group: str, curve_names: list[str]) -> dict[str, list[dict]]:
    ensure_split_result_workbook(path)
    selected_names = tuple(name for name in curve_names if name)
    curve_signature = result_curve_file_signature(path)
    return COMPARISON_CURVE_SLICE_CACHE.get(
        path,
        lambda resolved: read_comparison_curve_group_uncached(resolved, curve_group, selected_names),
        variant=("comparison_curve_slice", curve_group, selected_names, curve_signature),
    )


def read_comparison_curve_group_uncached(path: Path, curve_group: str, curve_names: tuple[str, ...]) -> dict[str, list[dict]]:
    config = COMPARISON_CURVE_GROUPS.get(curve_group)
    if not config:
        raise ValueError("曲线类型不合法")
    source_path = result_curves_workbook_path(path) if result_curves_workbook_path(path).exists() else path
    try:
        workbook = load_workbook(source_path, read_only=True, data_only=True)
    except RESULT_WORKBOOK_READ_ERRORS as exc:
        raise ValueError(f"曲线结果文件无法读取: {source_path.name}") from exc
    try:
        return read_curve_sheet(workbook, config["sheet"], config["limit"], selected_names=set(curve_names))
    finally:
        workbook.close()


def result_curve_file_signature(path: Path):
    curve_path = result_curves_workbook_path(path)
    if not curve_path.exists():
        return None
    try:
        return file_cache.file_signature(curve_path)
    except OSError:
        return None


def ensure_split_result_workbook(path: Path) -> None:
    curve_path = result_curves_workbook_path(path)
    if not path.exists():
        return
    curve_sheet_names = ["供能日曲线", "供能月曲线", "安全日曲线", "调度结果"]
    try:
        workbook = load_workbook(path, read_only=True, data_only=True)
    except RESULT_WORKBOOK_READ_ERRORS:
        return
    try:
        if not any(sheet_name in workbook.sheetnames for sheet_name in curve_sheet_names):
            return
        curve_payload = (
            {"results": {"curves": read_result_curves_from_workbook(workbook, include_hourly_curves=True)}}
            if not curve_path.exists()
            else None
        )
    finally:
        workbook.close()
    if curve_payload is not None:
        save_result_workbook(build_result_curves_workbook(curve_payload), curve_path, "曲线结果文件")
    try:
        workbook = load_workbook(path)
    except RESULT_WORKBOOK_READ_ERRORS:
        return
    tmp_path = path.with_name(f".{path.name}.tmp")
    try:
        for sheet_name in curve_sheet_names:
            if sheet_name in workbook.sheetnames:
                workbook.remove(workbook[sheet_name])
        file_ops.save_workbook_with_retry(workbook, tmp_path, "结果文件")
    finally:
        workbook.close()
    replace_result_workbook_with_retry(tmp_path, path)


def read_comparison_curve_groups(path: Path, fallback_workbook=None, include_hourly_curves: bool = True) -> dict:
    curve_path = result_curves_workbook_path(path)
    if curve_path.exists():
        try:
            curve_workbook = load_workbook(curve_path, read_only=True, data_only=True)
        except RESULT_WORKBOOK_READ_ERRORS as exc:
            raise ValueError(f"曲线结果文件无法读取: {curve_path.name}") from exc
        try:
            return read_comparison_curve_groups_from_workbook(curve_workbook, include_hourly_curves=include_hourly_curves)
        finally:
            curve_workbook.close()
    if fallback_workbook is not None:
        return read_comparison_curve_groups_from_workbook(fallback_workbook, include_hourly_curves=include_hourly_curves)
    return empty_comparison_curve_groups()


def read_comparison_curve_groups_from_workbook(workbook, include_hourly_curves: bool = True) -> dict:
    curve_groups = {}
    for key, config in COMPARISON_CURVE_GROUPS.items():
        if key == "hourly" and not include_hourly_curves:
            curve_groups[key] = read_curve_sheet_headers(workbook, config["sheet"])
        else:
            curve_groups[key] = read_curve_sheet(workbook, config["sheet"], config["limit"])
    return curve_groups


def read_result_workbook_display_payload_for_response(path: Path, include_hourly_curves: bool = True) -> dict | None:
    try:
        if not path.exists():
            return None
        return read_result_workbook_display_payload(path, include_hourly_curves=include_hourly_curves)
    except (ValueError, FileNotFoundError):
        return None


def read_result_workbook_display_payload(path: Path, include_hourly_curves: bool = True) -> dict:
    if include_hourly_curves:
        ensure_split_result_workbook(path)
    curve_signature = result_curve_file_signature(path)
    return RESULT_DISPLAY_PAYLOAD_CACHE.get(
        path,
        lambda resolved: read_result_workbook_display_payload_uncached(resolved, include_hourly_curves=include_hourly_curves),
        variant=("display", bool(include_hourly_curves), curve_signature),
    )


def read_result_workbook_display_payload_uncached(path: Path, include_hourly_curves: bool = True) -> dict:
    try:
        workbook = load_workbook(path, read_only=True, data_only=True)
    except RESULT_WORKBOOK_READ_ERRORS as exc:
        raise ValueError(f"结果文件无法读取: {path.name}") from exc
    try:
        planning_rows = read_named_sheet_rows(workbook, "规划结果")
        for row in planning_rows:
            normalize_planning_result_total_capacity(row)
        annual_rows = read_named_sheet_rows(workbook, "规划年指标")
        curve_rows = read_result_curve_workbook_payload(path, fallback_workbook=workbook, include_hourly_curves=include_hourly_curves)
        return {
            "metrics": read_result_workbook_metrics(workbook),
            "results": {
                "overview_tables": [
                    {"title": "规划结果", "rows": planning_rows},
                    {"title": "规划年指标", "rows": annual_rows},
                ],
                "overview_disks": build_overview_composition_from_workbook(workbook),
                "overview": [],
                "green": [],
                "green_table": read_named_sheet_rows(workbook, "供能分析"),
                "safety": [],
                "safety_table": read_named_sheet_rows(workbook, "安全评估"),
                "curves": curve_rows,
            },
        }
    finally:
        workbook.close()


def read_result_curve_workbook_payload(path: Path, fallback_workbook=None, include_hourly_curves: bool = True) -> dict:
    curve_path = result_curves_workbook_path(path)
    if curve_path.exists():
        try:
            curve_workbook = load_workbook(curve_path, read_only=True, data_only=True)
        except RESULT_WORKBOOK_READ_ERRORS as exc:
            raise ValueError(f"曲线结果文件无法读取: {curve_path.name}") from exc
        try:
            return read_result_curves_from_workbook(curve_workbook, include_hourly_curves=include_hourly_curves)
        finally:
            curve_workbook.close()
    if fallback_workbook is not None:
        return read_result_curves_from_workbook(fallback_workbook, include_hourly_curves=include_hourly_curves)
    return {"green_daily": [], "green_monthly": [], "green_hourly": [], "safety_daily": []}


def read_result_curves_from_workbook(workbook, include_hourly_curves: bool = True) -> dict:
    return {
        "green_daily": read_workbook_rows_with_field_map(workbook, "供能日曲线", limit=365),
        "green_monthly": read_workbook_rows_with_field_map(workbook, "供能月曲线", limit=12),
        "green_hourly": (
            read_workbook_rows_with_field_map(workbook, "调度结果", limit=8760)
            if include_hourly_curves
            else []
        ),
        "safety_daily": read_workbook_rows_with_field_map(workbook, "安全日曲线", limit=365),
    }


def read_result_workbook_metrics(workbook) -> list[dict]:
    rows = read_named_sheet_rows(workbook, "总体指标")
    metrics = []
    for row in rows:
        label = str(row.get("指标", "")).strip()
        if not label or label in {"方案", "当前状态", "状态", "进度"}:
            continue
        unit = row.get("单位", "")
        if label == "度电成本":
            unit = "元"
        metrics.append({"label": label, "value": row.get("数值", ""), "unit": unit})
    return metrics


EVALUATION_REPORT_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
EVALUATION_REPORT_CJK_FONT_PATHS = [
    "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
    "/usr/share/fonts/opentype/noto/NotoSansCJK-Bold.ttc",
    "/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf",
]
EVALUATION_REPORT_CJK_FONT_NAMES = [
    "Noto Sans CJK SC",
    "Noto Sans CJK JP",
    "Noto Sans CJK TC",
    "Droid Sans Fallback",
    "SimHei",
    "Microsoft YaHei",
    "Arial Unicode MS",
    "DejaVu Sans",
]
EVALUATION_REPORT_DEVICE_SECTIONS = [
    ("diesel_generators", "柴发类型数"),
    ("wind_turbines", "风机类型数"),
    ("photovoltaics", "光伏类型数"),
    ("storage_pcs", "储能PCS类型数"),
    ("storage_battery_packs", "储能电池组类型数"),
    ("hydrogen_electrolyzers", "电解槽类型数"),
    ("hydrogen_tanks", "储氢罐类型数"),
    ("fuel_cells", "燃料电池类型数"),
]
EVALUATION_REPORT_GREEN_SERIES = [
    ("load_energy", "负荷"),
    ("diesel_energy", "柴发"),
    ("wind_energy", "风机"),
    ("pv_energy", "光伏"),
    ("renewable_energy", "新能源"),
    ("storage_discharge_energy", "储能放电"),
    ("fuel_cell_energy", "燃料电池"),
]
EVALUATION_REPORT_SAFETY_SERIES = [
    ("frequency_min", "最低频率"),
    ("frequency_max", "最高频率"),
]


def build_evaluation_report_docx(scheme: str, filename: str) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    clean_scheme = planning_store.validate_scheme_name(scheme)
    selected = resolve_evaluation_report_result_filename(clean_scheme, filename)
    if not selected:
        raise FileNotFoundError(f"结果文件不存在: {filename or OPTIMIZATION_RESULT_WORKBOOK_NAME}")
    result_path = evaluation_result_path(clean_scheme, selected)
    if not result_path.exists():
        raise FileNotFoundError(f"结果文件不存在: {result_path.name}")

    payload = read_result_workbook_display_payload(result_path, include_hourly_curves=False)
    overview = read_evaluation_report_scheme_overview(clean_scheme)
    meta = PLANNING_STORE.read_scheme_meta(clean_scheme)
    results = payload.get("results") if isinstance(payload.get("results"), dict) else {}

    document = Document()
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Microsoft YaHei"
    normal_style.font.size = Pt(10)

    title = document.add_heading("方案结果报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = document.add_paragraph(f"{clean_scheme} / {result_display_name_from_filename(selected) or selected}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_heading("基本信息", level=1)
    report_add_key_value_table(
        document,
        evaluation_report_basic_rows(clean_scheme, selected, result_path, overview, meta),
    )

    document.add_heading("关键指标", level=1)
    metrics = payload.get("metrics") if isinstance(payload.get("metrics"), list) else []
    if metrics:
        report_add_rows_table(
            document,
            [{"指标": item.get("label", ""), "数值": item.get("value", ""), "单位": item.get("unit", "")} for item in metrics],
            headers=["指标", "数值", "单位"],
        )
    else:
        document.add_paragraph("暂无关键指标。")

    document.add_heading("结果表格", level=1)
    for table in results.get("overview_tables", []):
        rows = table.get("rows", []) if isinstance(table, dict) else []
        document.add_heading(str(table.get("title") or "结果表格"), level=2)
        report_add_rows_table(document, rows, empty_text="暂无数据。")

    green_table = results.get("green_table", [])
    if green_table:
        document.add_heading("经济性指标", level=2)
        report_add_rows_table(document, green_table, empty_text="暂无经济性指标。")
    safety_table = results.get("safety_table", [])
    if safety_table:
        document.add_heading("安全性指标", level=2)
        report_add_rows_table(document, safety_table, empty_text="暂无安全性指标。")

    document.add_heading("图表展示", level=1)
    charts_added = 0
    charts_added += report_add_overview_composition_charts(document, results.get("overview_disks", []))
    curves = results.get("curves") if isinstance(results.get("curves"), dict) else {}
    if report_add_line_chart(
        document,
        "供能日曲线",
        curves.get("green_daily", []),
        "day",
        EVALUATION_REPORT_GREEN_SERIES,
        "电量",
    ):
        charts_added += 1
    if report_add_line_chart(
        document,
        "供能月曲线",
        curves.get("green_monthly", []),
        "month",
        EVALUATION_REPORT_GREEN_SERIES,
        "电量",
    ):
        charts_added += 1
    if report_add_line_chart(
        document,
        "频率安全日曲线",
        curves.get("safety_daily", []),
        "day",
        EVALUATION_REPORT_SAFETY_SERIES,
        "频率/裕度",
    ):
        charts_added += 1
    if charts_added == 0:
        document.add_paragraph("暂无可展示图表。")

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def resolve_evaluation_report_result_filename(scheme: str, filename: str = "") -> str:
    requested = str(filename or "").strip()
    selected = selected_evaluation_result_filename(scheme, requested)
    if requested and selected != requested:
        raise FileNotFoundError(f"结果文件不存在: {requested}")
    return selected


def read_evaluation_report_scheme_overview(scheme: str) -> dict:
    try:
        return PLANNING_STORE.read_scheme_overview(scheme)
    except (FileNotFoundError, ValueError):
        return {"scheme": scheme}


def evaluation_report_basic_rows(
    scheme: str,
    filename: str,
    result_path: Path,
    overview: dict,
    meta: dict,
) -> list[tuple[str, object]]:
    rows: list[tuple[str, object]] = [
        ("方案名称", overview.get("scheme") or scheme),
        ("结果名称", result_display_name_from_filename(filename) or filename),
        ("结果文件", filename),
        ("报告生成时间", datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
        ("结果更新时间", datetime.fromtimestamp(result_path.stat().st_mtime).strftime("%Y-%m-%d %H:%M:%S")),
        ("方案归属", meta.get("owner_username") or "未设置"),
        ("创建人", meta.get("created_by") or meta.get("owner_username") or "未设置"),
        ("方案创建时间", meta.get("created_at") or ""),
        ("时序数据行数", overview.get("time_series_count", "")),
    ]
    for key, label in EVALUATION_REPORT_DEVICE_SECTIONS:
        value = overview.get(key)
        if isinstance(value, list):
            rows.append((label, len(value)))
    return rows


def report_add_key_value_table(document, rows: list[tuple[str, object]]) -> None:
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in rows:
        row = table.add_row().cells
        row[0].text = str(label)
        row[1].text = report_display_value(value)


def report_add_rows_table(
    document,
    rows: list[dict],
    headers: list[str] | None = None,
    empty_text: str = "暂无数据。",
    max_rows: int = 120,
) -> None:
    normalized_rows = [row for row in rows if isinstance(row, dict)] if isinstance(rows, list) else []
    if not normalized_rows:
        document.add_paragraph(empty_text)
        return
    table_headers = headers or report_headers_from_rows(normalized_rows)
    if not table_headers:
        document.add_paragraph(empty_text)
        return
    table = document.add_table(rows=1, cols=len(table_headers))
    table.style = "Table Grid"
    for index, header in enumerate(table_headers):
        table.rows[0].cells[index].text = str(header)
    for row in normalized_rows[:max_rows]:
        cells = table.add_row().cells
        for index, header in enumerate(table_headers):
            cells[index].text = report_display_value(row.get(header, ""))
    if len(normalized_rows) > max_rows:
        document.add_paragraph(f"仅展示前 {max_rows} 行，共 {len(normalized_rows)} 行。")


def report_headers_from_rows(rows: list[dict]) -> list[str]:
    headers: list[str] = []
    for row in rows:
        if not isinstance(row, dict):
            continue
        for key in row:
            clean = str(key or "").strip()
            if clean and clean not in headers:
                headers.append(clean)
    return headers


def report_display_value(value: object) -> str:
    if value is None:
        return ""
    if isinstance(value, datetime):
        return value.strftime("%Y-%m-%d %H:%M:%S")
    if isinstance(value, Decimal):
        value = float(value)
    if isinstance(value, float):
        if math.isnan(value) or math.isinf(value):
            return ""
        text = f"{value:.6f}".rstrip("0").rstrip(".")
        return text or "0"
    return str(value)


def report_add_overview_composition_charts(document, disks: list[dict]) -> int:
    added = 0
    for disk in disks if isinstance(disks, list) else []:
        segments = report_overview_disk_segments(disk)
        if not segments:
            continue
        title = str(disk.get("title") or "构成").strip() or "构成"
        if report_add_bar_chart(document, title, segments):
            added += 1
    return added


def report_overview_disk_segments(disk: dict) -> list[dict]:
    if not isinstance(disk, dict):
        return []
    segments: list[dict] = []
    raw_segments = disk.get("segments")
    if isinstance(raw_segments, list):
        for segment in raw_segments:
            if not isinstance(segment, dict):
                continue
            value = _numeric_or_none(segment.get("value"))
            label = str(segment.get("label") or "").strip()
            if value is not None and value > 0 and label:
                segments.append({"label": label, "value": value, "unit": segment.get("unit") or disk.get("unit") or ""})
    else:
        for label_key, value_key in (("left_label", "left_value"), ("right_label", "right_value")):
            value = _numeric_or_none(disk.get(value_key))
            label = str(disk.get(label_key) or "").strip()
            if value is not None and value > 0 and label:
                segments.append({"label": label, "value": value, "unit": disk.get("unit") or ""})
    return segments


def report_add_bar_chart(document, title: str, segments: list[dict]) -> bool:
    plt = report_pyplot()
    labels = [str(segment.get("label") or "") for segment in segments]
    values = [_numeric_or_none(segment.get("value")) or 0 for segment in segments]
    if not labels or not any(value > 0 for value in values):
        return False
    unit = str(next((segment.get("unit") for segment in segments if segment.get("unit")), "") or "")
    fig, ax = plt.subplots(figsize=(6.4, 3.2))
    bars = ax.bar(labels, values, color=["#0d5c59", "#d98a2b", "#4a7a9f", "#8a9b3f", "#7d5fb2", "#a85d5d"])
    ax.set_title(title)
    if unit:
        ax.set_ylabel(unit)
    ax.tick_params(axis="x", rotation=20)
    for bar, value in zip(bars, values):
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height(), report_display_value(value), ha="center", va="bottom", fontsize=8)
    ax.grid(axis="y", alpha=0.2)
    fig.tight_layout()
    report_add_figure(document, fig)
    return True


def report_add_line_chart(
    document,
    title: str,
    rows: list[dict],
    x_key: str,
    series_specs: list[tuple[str, str]],
    y_label: str,
) -> bool:
    normalized_rows = [row for row in rows if isinstance(row, dict)] if isinstance(rows, list) else []
    if not normalized_rows:
        return False
    plt = report_pyplot()
    fig, ax = plt.subplots(figsize=(6.4, 3.2))
    plotted = False
    x_values = [
        _numeric_or_none(row.get(x_key)) if _numeric_or_none(row.get(x_key)) is not None else index + 1
        for index, row in enumerate(normalized_rows)
    ]
    for key, label in series_specs:
        y_values = [_numeric_or_none(row.get(key)) for row in normalized_rows]
        if not any(value is not None for value in y_values):
            continue
        plotted = True
        ax.plot(
            x_values,
            [value if value is not None else float("nan") for value in y_values],
            linewidth=1.6,
            label=label,
        )
    if not plotted:
        plt.close(fig)
        return False
    ax.set_title(title)
    ax.set_xlabel("月份" if x_key == "month" else "天")
    ax.set_ylabel(y_label)
    ax.grid(alpha=0.25)
    ax.legend(loc="best", fontsize=8)
    fig.tight_layout()
    report_add_figure(document, fig)
    return True


def report_pyplot():
    import matplotlib

    matplotlib.use("Agg", force=True)
    from matplotlib import font_manager
    import matplotlib.pyplot as plt

    for font_path in EVALUATION_REPORT_CJK_FONT_PATHS:
        path = Path(font_path)
        if path.exists():
            try:
                font_manager.fontManager.addfont(str(path))
            except RuntimeError:
                pass
    available_fonts = {font.name for font in font_manager.fontManager.ttflist}
    selected_fonts = [name for name in EVALUATION_REPORT_CJK_FONT_NAMES if name in available_fonts]
    plt.rcParams["font.sans-serif"] = selected_fonts or EVALUATION_REPORT_CJK_FONT_NAMES
    plt.rcParams["axes.unicode_minus"] = False
    return plt


def report_add_figure(document, figure) -> None:
    from docx.shared import Inches

    stream = BytesIO()
    try:
        figure.savefig(stream, format="png", dpi=150, bbox_inches="tight")
        stream.seek(0)
        document.add_picture(stream, width=Inches(6.4))
    finally:
        report_pyplot().close(figure)


def merge_runtime_metrics(runtime_metrics: list[dict], workbook_metrics: list[dict]) -> list[dict]:
    if not workbook_metrics:
        return runtime_metrics
    primary_labels = {"状态", "开始", "完成"}
    primary = [item for item in runtime_metrics if isinstance(item, dict) and str(item.get("label", "")) in primary_labels]
    secondary = [item for item in runtime_metrics if isinstance(item, dict) and str(item.get("label", "")) not in primary_labels]
    merged = list(primary)
    existing = {str(item.get("label", "")) for item in merged if isinstance(item, dict)}
    for metric in workbook_metrics:
        if not isinstance(metric, dict):
            continue
        label = str(metric.get("label", "")).strip()
        if not label or label in existing:
            continue
        merged.append(metric)
        existing.add(label)
    for metric in secondary:
        label = str(metric.get("label", "")).strip()
        if label and label not in existing:
            merged.append(metric)
            existing.add(label)
    return merged


def read_workbook_rows_with_field_map(workbook, sheet_name: str, limit: int | None = None) -> list[dict]:
    if sheet_name not in workbook.sheetnames:
        return []
    sheet = workbook[sheet_name]
    rows_iter = sheet.iter_rows(values_only=True)
    raw_headers = [str(value or "").strip() for value in next(rows_iter, [])]
    headers = [result_workbook_header_to_field(header) for header in raw_headers]
    rows = []
    for row_index, row in enumerate(rows_iter, start=1):
        if limit is not None and row_index > limit:
            break
        item = {}
        for index, header in enumerate(headers):
            if not header:
                continue
            value = row[index] if index < len(row) else ""
            if value not in (None, ""):
                item[header] = value
        if item:
            rows.append(item)
    return rows


def result_workbook_header_to_field(header: str) -> str:
    clean = str(header or "").strip()
    if clean in DEPRECATED_RESULT_CURVE_HEADERS:
        return ""
    return RESULT_WORKBOOK_HEADER_TO_FIELD.get(clean, clean)


def build_overview_composition_from_workbook(workbook) -> list[dict]:
    annual_rows = read_annual_comparison_rows(workbook)
    energy_rows = read_named_sheet_rows(workbook, "供能分析")
    planning_rows = read_named_sheet_rows(workbook, PLANNING_RESULT_SHEET_NAME)
    metrics = rows_by_metric([*annual_rows, *energy_rows])
    construction_cost = metric_value(metrics, "年均建设成本")
    diesel_cost = metric_value(metrics, "年柴油成本")
    if construction_cost == 0 and diesel_cost == 0:
        total_cost = metric_value(metrics, "年总成本") or metric_value(metrics, "总成本")
        if total_cost:
            construction_cost = total_cost
    diesel_energy = metric_value(metrics, "柴发总发电量") or metric_value(metrics, "柴发总电量")
    green_energy = metric_value(metrics, "绿电年发电量") or metric_value(metrics, "新能源实发电量")
    if green_energy == 0:
        green_energy = (
            metric_value(metrics, "风机总发电量")
            + metric_value(metrics, "光伏总发电量")
            + metric_value(metrics, "电储能总放电量")
            + metric_value(metrics, "燃料电池总发电量")
        )
    return [
        {
            "title": "成本构成",
            "left_label": "年柴油成本",
            "left_value": diesel_cost,
            "right_label": "年均建设成本",
            "right_value": construction_cost,
            "unit": "万元",
        },
        capacity_composition_disk_from_workbook_rows(planning_rows, metrics),
        {
            "title": "电量构成",
            "left_label": "柴发电量",
            "left_value": diesel_energy,
            "right_label": "绿电电量",
            "right_value": green_energy,
            "unit": energy_unit(metrics, ["柴发总发电量", "柴发总电量", "绿电年发电量", "新能源实发电量"]) or "kWh",
        },
    ]


def capacity_composition_disk_from_workbook_rows(planning_rows: list[dict], metrics: dict[str, dict]) -> dict:
    capacity_disk = plan_optimizer.capacity_composition_disk(planning_rows)
    if any(metric_number(segment.get("value")) > 0 for segment in capacity_disk["segments"]):
        return capacity_disk
    return plan_optimizer.capacity_composition_disk_from_values(
        diesel_capacity=metric_value(metrics, "柴发总容量"),
        wind_capacity=metric_value(metrics, "风电总容量"),
        pv_capacity=metric_value(metrics, "光伏总容量"),
        storage_energy_capacity=metric_value(metrics, "储能总容量"),
        fuel_cell_power_capacity=metric_value(metrics, "氢能总容量"),
    )


def metric_number(value) -> float:
    parsed = _numeric_or_none(value)
    return float(parsed) if parsed is not None else 0.0


def rows_by_metric(rows: list[dict]) -> dict[str, dict]:
    result = {}
    for row in rows:
        if not isinstance(row, dict):
            continue
        name = str(row.get("指标", "")).strip()
        if name and name not in result:
            result[name] = row
    return result


def metric_value(metrics: dict[str, dict], name: str) -> float:
    row = metrics.get(name) or {}
    value = _numeric_or_none(row.get("数值"))
    return float(value) if value is not None else 0.0


def energy_unit(metrics: dict[str, dict], names: list[str]) -> str:
    for name in names:
        unit = str((metrics.get(name) or {}).get("单位", "")).strip()
        if unit:
            return unit
    return ""


def empty_comparison_curve_groups() -> dict:
    return {
        key: {"title": config["title"], "curves": [], "series": {}}
        for key, config in COMPARISON_CURVE_GROUPS.items()
    }


def append_comparison_curve_groups(target_groups: dict, source_groups: dict, label: str, scheme: str, filename: str) -> None:
    for key, config in COMPARISON_CURVE_GROUPS.items():
        target_group = target_groups.setdefault(key, {"title": config["title"], "curves": [], "series": {}})
        source_group = source_groups.get(key, {}) if isinstance(source_groups, dict) else {}
        for name, points in source_group.items():
            if name not in target_group["curves"]:
                target_group["curves"].append(name)
            if points:
                target_group["series"].setdefault(name, []).append(
                    {
                        "label": label,
                        "scheme": scheme,
                        "filename": filename,
                        "points": points,
                    }
                )


def read_named_sheet_rows(workbook, sheet_name: str) -> list[dict]:
    if sheet_name not in workbook.sheetnames:
        return []
    sheet = workbook[sheet_name]
    rows = list(sheet.iter_rows(values_only=True))
    if not rows:
        return []
    headers = [str(value or "").strip() for value in rows[0]]
    result_rows = []
    for row in rows[1:]:
        item = {}
        for index, header in enumerate(headers):
            if header:
                item[header] = row[index] if index < len(row) else ""
        if any(value not in (None, "") for value in item.values()):
            result_rows.append(item)
    return result_rows


def read_annual_comparison_rows(workbook) -> list[dict]:
    rows = read_named_sheet_rows(workbook, "规划年指标")
    if rows:
        return rows
    return read_named_sheet_rows(workbook, "供能分析")


def read_dispatch_curves(workbook) -> dict[str, list[dict]]:
    return read_curve_sheet(workbook, "调度结果", 8760)


def read_curve_sheet_headers(workbook, sheet_name: str) -> dict[str, list[dict]]:
    if sheet_name not in workbook.sheetnames:
        return {}
    sheet = workbook[sheet_name]
    rows_iter = sheet.iter_rows(values_only=True)
    headers = [str(value or "").strip() for value in next(rows_iter, [])]
    curves: dict[str, list[dict]] = {}
    for header in headers:
        display_name = result_curve_display_name(header)
        if header and header not in COMPARISON_CURVE_X_HEADERS and display_name:
            curves[display_name] = []
    return curves


def read_curve_sheet(
    workbook,
    sheet_name: str,
    limit: int | None = None,
    selected_names: set[str] | None = None,
) -> dict[str, list[dict]]:
    curves = read_curve_sheet_headers(workbook, sheet_name)
    if not curves:
        return {}
    if selected_names is not None:
        wanted = {str(name or "").strip() for name in selected_names if str(name or "").strip()}
        curves = {name: points for name, points in curves.items() if name in wanted}
        if not curves:
            return {}
    sheet = workbook[sheet_name]
    rows_iter = sheet.iter_rows(values_only=True)
    headers = [str(value or "").strip() for value in next(rows_iter, [])]
    for row_index, row in enumerate(rows_iter, start=1):
        if limit is not None and row_index > limit:
            break
        x_value = row[0] if len(row) > 0 and row[0] not in (None, "") else row_index
        for column_index, header in enumerate(headers):
            display_name = result_curve_display_name(header)
            if display_name not in curves:
                continue
            value = row[column_index] if column_index < len(row) else None
            number = _numeric_or_none(value)
            if number is not None:
                curves[display_name].append({"x": x_value, "y": number})
    return curves


def result_curve_display_name(header: str) -> str:
    clean = str(header or "").strip()
    if clean in DEPRECATED_RESULT_CURVE_HEADERS:
        return ""
    return RESULT_CURVE_FIELD_LABELS.get(clean, clean)


def merge_comparison_rows(tables: list[list[dict]], items: list[dict], key_field: str) -> list[dict]:
    merged: dict[str, dict] = {}
    order: list[str] = []
    for item, rows in zip(items, tables):
        value_field = "总容量" if key_field == "设备类型" else "数值"
        unit_field = "单位"
        label = item["label"]
        for row in rows:
            if not isinstance(row, dict):
                continue
            key = str(row.get(key_field, "")).strip()
            if not key:
                continue
            if key not in merged:
                merged[key] = {key_field: key, "单位": row.get(unit_field, "")}
                order.append(key)
            if not merged[key].get("单位") and row.get(unit_field, ""):
                merged[key]["单位"] = row.get(unit_field, "")
            merged[key][label] = row.get(value_field, "")
    return [merged[key] for key in order]


def result_display_name_from_filename(filename: str) -> str:
    return re.sub(r"_results\.xlsx$", "", str(filename or ""))


def _numeric_or_none(value):
    if isinstance(value, (int, float)) and not isinstance(value, bool):
        if isinstance(value, float) and (math.isnan(value) or math.isinf(value)):
            return None
        return value
    try:
        text = str(value).strip()
        if not text:
            return None
        return float(text)
    except (TypeError, ValueError):
        return None


def write_evaluation_planning_counts(scheme: str, filename: str, planning_rows: list[dict]) -> list[dict]:
    result_path = evaluation_result_path(scheme, filename)
    if not result_path.exists():
        raise FileNotFoundError(f"结果文件不存在: {result_path.name}")
    ensure_split_result_workbook(result_path)

    counts_by_device: dict[str, object] = {}
    for row in planning_rows if isinstance(planning_rows, list) else []:
        if not isinstance(row, dict):
            continue
        device_type = str(row.get("设备类型", "")).strip()
        if not device_type or "设计台数" not in row:
            continue
        counts_by_device[device_type] = row.get("设计台数")

    workbook = load_workbook(result_path)
    try:
        if PLANNING_RESULT_SHEET_NAME not in workbook.sheetnames:
            sheet = workbook.create_sheet(PLANNING_RESULT_SHEET_NAME)
            sheet.append(PLANNING_RESULT_HEADERS)
        else:
            sheet = workbook[PLANNING_RESULT_SHEET_NAME]
            if sheet.max_row < 1:
                sheet.append(PLANNING_RESULT_HEADERS)

        headers = [str(cell.value or "").strip() for cell in sheet[1]]
        if "设备类型" not in headers or "设计台数" not in headers:
            raise ValueError("规划结果工作表缺少设备类型或设计台数列")
        device_column = headers.index("设备类型") + 1
        count_column = headers.index("设计台数") + 1
        unit_capacity_column = headers.index("单台容量") + 1 if "单台容量" in headers else None
        total_capacity_column = headers.index("总容量") + 1 if "总容量" in headers else None

        for row_index in range(2, sheet.max_row + 1):
            device_type = str(sheet.cell(row=row_index, column=device_column).value or "").strip()
            if device_type in counts_by_device:
                count = normalize_planning_count(counts_by_device[device_type])
                sheet.cell(row=row_index, column=count_column).value = count
                if total_capacity_column and unit_capacity_column and count != "":
                    unit_capacity = estimate.numeric(sheet.cell(row=row_index, column=unit_capacity_column).value, 0.0)
                    sheet.cell(row=row_index, column=total_capacity_column).value = round(count * unit_capacity, 4)

        tmp_path = result_path.with_name(f".{result_path.name}.tmp")
        file_ops.save_workbook_with_retry(workbook, tmp_path, "结果文件")
    finally:
        workbook.close()
    replace_result_workbook_with_retry(tmp_path, result_path)
    return read_evaluation_planning_result_rows(scheme, filename)


def normalize_planning_count(value):
    if value in (None, ""):
        return ""
    text = str(value).strip()
    if not re.fullmatch(r"\d+", text):
        raise ValueError("设计台数必须为非负整数")
    return int(text)


def evaluation_result_filename_from_name(name: str) -> str:
    clean_name = str(name or "").strip()
    if not clean_name:
        raise ValueError("结果名称不能为空")
    if clean_name.endswith(".xlsx"):
        raise ValueError("结果名称只需输入 xxxx，不要输入扩展名")
    return f"{clean_name}_results.xlsx"


def save_evaluation_result_workbook(scheme: str, filename: str) -> Path:
    result_path = evaluation_result_path(scheme, filename or OPTIMIZATION_RESULT_WORKBOOK_NAME)
    if not result_path.exists():
        raise FileNotFoundError(f"结果文件不存在: {result_path.name}")
    ensure_split_result_workbook(result_path)
    return result_path


def rename_evaluation_result_workbook(scheme: str, filename: str, target_name: str) -> Path:
    source_path = evaluation_result_path(scheme, filename or OPTIMIZATION_RESULT_WORKBOOK_NAME)
    if source_path.name == OPTIMIZATION_RESULT_WORKBOOK_NAME:
        raise ValueError("默认结果文件不允许重命名")
    if not source_path.exists():
        raise FileNotFoundError(f"结果文件不存在: {source_path.name}")
    ensure_split_result_workbook(source_path)
    source_error = result_workbook_error_message(source_path)
    if source_error:
        raise ValueError(f"重命名失败，当前结果文件无法读取: {source_path.name}")

    target_filename = evaluation_result_filename_from_name(target_name)
    target_path = evaluation_result_path(scheme, target_filename)
    if target_path.name == OPTIMIZATION_RESULT_WORKBOOK_NAME:
        raise ValueError("默认结果文件不允许作为重命名目标")
    if target_path == source_path:
        return source_path

    source_curve_path = result_curves_workbook_path(source_path)
    target_curve_path = result_curves_workbook_path(target_path)
    if target_path.exists():
        raise FileExistsError(f"重命名失败，结果文件已存在: {target_path.name}")
    if target_curve_path.exists():
        raise FileExistsError(f"重命名失败，曲线结果文件已存在: {target_curve_path.name}")

    file_ops.retry_file_operation(
        lambda: source_path.rename(target_path),
        f"结果文件被占用，无法重命名：{source_path.name}。请关闭正在打开该文件的 Excel 或预览窗口后重试。",
    )
    file_cache.invalidate_path(source_path)
    file_cache.invalidate_path(target_path)
    file_cache.invalidate_path(source_path.parent)
    if source_curve_path.exists():
        file_ops.retry_file_operation(
            lambda: source_curve_path.rename(target_curve_path),
            f"曲线结果文件被占用，无法重命名：{source_curve_path.name}。请关闭正在打开该文件的 Excel 或预览窗口后重试。",
        )
        file_cache.invalidate_path(source_curve_path)
        file_cache.invalidate_path(target_curve_path)
    file_cache.invalidate_path(target_path.parent)
    return target_path


def handle_evaluation_results_api_path(
    path: str,
    method: str = "GET",
    body: bytes = b"",
    query: str = "",
    current_user: dict | None = None,
) -> tuple[int, dict[str, str], bytes]:
    if path not in {"/api/evaluation/results", "/api/evaluation/report"}:
        return _json_response({"error": "not_found", "path": path}, HTTPStatus.NOT_FOUND)
    try:
        if path == "/api/evaluation/report":
            if method != "GET":
                return _json_response({"error": "not_found", "path": path}, HTTPStatus.NOT_FOUND)
            query_params = parse_qs(query)
            scheme = query_params.get("scheme", [""])[0]
            filename = query_params.get("filename", [""])[0]
            scheme = ensure_planning_scheme_access(scheme, current_user)
            selected = resolve_evaluation_report_result_filename(scheme, filename)
            if not selected:
                raise FileNotFoundError("结果文件不存在")
            report_body = build_evaluation_report_docx(scheme, selected)
            display_name = result_display_name_from_filename(selected) or "结果"
            return _download_response(
                report_body,
                f"{scheme}_{display_name}_报告.docx",
                EVALUATION_REPORT_CONTENT_TYPE,
            )

        if method == "GET":
            query_params = parse_qs(query)
            scheme = query_params.get("scheme", [""])[0]
            filename = query_params.get("filename", [""])[0]
            ensure_planning_scheme_access(scheme, current_user)
            light = query_params.get("light", [""])[0] in {"1", "true", "yes"}
            selected = selected_evaluation_result_filename(scheme, filename)
            if light:
                return _json_response(
                    {
                        "selected": selected,
                        "results": list_evaluation_result_files(scheme),
                    }
                )
            return _json_response(
                {
                    "selected": selected,
                    "results": list_evaluation_result_files(scheme),
                    "planning_result_rows": read_evaluation_planning_result_rows_for_response(scheme, selected),
                }
            )

        if method != "POST":
            return _json_response({"error": "not_found", "path": path}, HTTPStatus.NOT_FOUND)

        payload = _read_json_body(body)
        scheme = str(payload.get("scheme", ""))
        action = str(payload.get("action", ""))
        filename = str(payload.get("filename", ""))
        ensure_planning_scheme_manage_access(scheme, current_user)

        if action == "delete":
            result_path = evaluation_result_path(scheme, filename)
            if result_path.name == OPTIMIZATION_RESULT_WORKBOOK_NAME:
                raise ValueError("默认结果文件不允许删除")
            if result_path.exists():
                file_ops.delete_file_with_retry(result_path, "结果文件")
            curve_path = result_curves_workbook_path(result_path)
            if curve_path.exists():
                file_ops.delete_file_with_retry(curve_path, "曲线结果文件")
            selected = selected_evaluation_result_filename(scheme)
            return _json_response(
                {
                    "selected": selected,
                    "results": list_evaluation_result_files(scheme),
                    "planning_result_rows": read_evaluation_planning_result_rows_for_response(scheme, selected),
                }
            )

        if action == "copy":
            source_path = evaluation_result_path(scheme, filename)
            if not source_path.exists():
                raise FileNotFoundError(f"结果文件不存在: {source_path.name}")
            ensure_split_result_workbook(source_path)
            source_error = result_workbook_error_message(source_path)
            if source_error:
                raise ValueError(f"复制失败，当前结果文件无法读取: {source_path.name}")
            target_filename = evaluation_result_filename_from_name(str(payload.get("target_name", "")))
            target_path = evaluation_result_path(scheme, target_filename)
            if target_path.exists():
                if not result_workbook_error_message(target_path):
                    return _json_response(
                        {"error": "exists", "message": f"复制失败，结果文件已存在: {target_path.name}"},
                        HTTPStatus.CONFLICT,
                    )
                if target_path.name == OPTIMIZATION_RESULT_WORKBOOK_NAME:
                    raise ValueError("默认结果文件不允许覆盖")
            file_ops.copy_file_with_retry(source_path, target_path, "结果文件")
            return _json_response(
                {
                    "selected": target_path.name,
                    "results": list_evaluation_result_files(scheme),
                    "planning_result_rows": read_evaluation_planning_result_rows(scheme, target_path.name),
                }
            )

        if action == "rename":
            target_path = rename_evaluation_result_workbook(scheme, filename, str(payload.get("target_name", "")))
            return _json_response(
                {
                    "selected": target_path.name,
                    "results": list_evaluation_result_files(scheme),
                    "planning_result_rows": read_evaluation_planning_result_rows(scheme, target_path.name),
                }
            )

        if action == "save":
            if (filename or OPTIMIZATION_RESULT_WORKBOOK_NAME) == OPTIMIZATION_RESULT_WORKBOOK_NAME:
                raise ValueError("默认结果文件不允许修改")
            planning_rows = payload.get("planning_result_rows")
            if isinstance(planning_rows, list):
                planning_result_rows = write_evaluation_planning_counts(
                    scheme,
                    filename or OPTIMIZATION_RESULT_WORKBOOK_NAME,
                    planning_rows,
                )
                selected = filename or OPTIMIZATION_RESULT_WORKBOOK_NAME
            else:
                result_path = save_evaluation_result_workbook(scheme, filename or OPTIMIZATION_RESULT_WORKBOOK_NAME)
                selected = result_path.name
                planning_result_rows = read_evaluation_planning_result_rows(scheme, selected)
            return _json_response(
                {
                    "selected": selected,
                    "results": list_evaluation_result_files(scheme),
                    "planning_result_rows": planning_result_rows,
                }
            )

        raise ValueError("未知结果文件操作")
    except FileNotFoundError as exc:
        return _json_response({"error": "not_found", "message": str(exc)}, HTTPStatus.NOT_FOUND)
    except FileExistsError as exc:
        return _json_response({"error": "exists", "message": str(exc)}, HTTPStatus.CONFLICT)
    except PermissionError as exc:
        return _json_response({"error": "file_locked", "message": str(exc)}, HTTPStatus.CONFLICT)
    except ValueError as exc:
        return _json_response({"error": "bad_request", "message": str(exc)}, HTTPStatus.BAD_REQUEST)


class OptimizationRuntimeManager:
    """Holds independent optimization runtimes for multiple schemes."""

    def __init__(self) -> None:
        self._runtimes: dict[str, OptimizationRuntime] = {}
        self._lock = threading.Lock()

    def snapshot(self, scheme: str = "", include_hourly_curves: bool = True) -> dict:
        runtime = self._runtime_for_scheme(scheme)
        payload = runtime.snapshot(include_hourly_curves=include_hourly_curves)
        payload["running_schemes"] = self.running_schemes()
        append_task_control_state(payload, "optimization", self._scheme_name(scheme), OPTIMIZATION_RESULT_WORKBOOK_NAME)
        return payload

    def apply(self, action: str, scheme: str = "") -> dict:
        runtime = self._runtime_for_scheme(scheme)
        payload = runtime.apply(action, scheme=self._scheme_name(scheme))
        payload["running_schemes"] = self.running_schemes()
        return payload

    def running_schemes(self) -> list[str]:
        with self._lock:
            runtimes = list(self._runtimes.items())
        running = []
        for scheme, runtime in runtimes:
            if runtime.status == "运行中":
                running.append(scheme)
        return running

    def runtimes(self) -> dict[str, OptimizationRuntime]:
        with self._lock:
            return dict(self._runtimes)

    def _runtime_for_scheme(self, scheme: str = "") -> OptimizationRuntime:
        name = self._scheme_name(scheme)
        with self._lock:
            if name not in self._runtimes:
                self._runtimes[name] = OptimizationRuntime(scheme=name)
            return self._runtimes[name]

    @staticmethod
    def _scheme_name(scheme: str = "") -> str:
        return str(scheme or "未选择方案").strip() or "未选择方案"


class EvaluationRuntime:
    """Independent runtime state for fixed-plan evaluation dispatch."""

    def __init__(self, scheme: str = "") -> None:
        self.status = "待启动"
        self.scheme = str(scheme or "").strip()
        self.start_time = ""
        self.end_time = ""
        self.progress = 0
        self.result_filename = ""
        self.result_file = ""
        self._metrics: list[dict] = []
        self._results: dict = {}
        self._logs: list[dict[str, str]] = []
        self._lock = threading.Lock()
        self._process: multiprocessing.Process | None = None
        self._event_queue: multiprocessing.Queue | None = None
        self.process_id: int | None = None
        self._stop_requested = False
        self._run_token = 0
        self._append_log_unlocked("info", "方案评估待启动")

    def snapshot(self, include_hourly_curves: bool = True) -> dict:
        with self._lock:
            self._drain_events_unlocked()
            self._reap_process_unlocked()
            return self._payload_unlocked(include_hourly_curves=include_hourly_curves)

    def task_snapshot(self) -> dict:
        """Return task-list state without opening the result workbook."""

        with self._lock:
            self._drain_events_unlocked()
            self._reap_process_unlocked()
            return self._task_payload_unlocked()

    def apply(self, action: str, scheme: str = "", filename: str = "") -> dict:
        target_scheme = str(scheme or self.scheme or "未选择方案").strip() or "未选择方案"
        if action == "clear_logs":
            with self._lock:
                self.scheme = target_scheme
                if filename:
                    self.result_filename = str(filename or "").strip()
                self._logs.clear()
                return self._payload_unlocked()

        if action == "cancel_queue":
            target_filename = selected_evaluation_result_filename(target_scheme, filename)
            with self._lock:
                if self.status == "运行中":
                    raise OptimizationStateError("running", f"方案“{target_scheme}”正在评估，无法退出队列")
                self.scheme = target_scheme
                self.status = "退出队列"
                self.start_time = ""
                self.end_time = ""
                self.progress = 0
                self.result_filename = target_filename
                try:
                    self.result_file = str(evaluation_result_path(target_scheme, target_filename)) if target_filename else ""
                except ValueError:
                    self.result_file = ""
                self.process_id = None
                self._metrics = []
                self._results = {}
                self._stop_requested = False
                self._terminate_process_unlocked()
                self._append_log_unlocked("info", "退出等待队列")
                return self._payload_unlocked()

        if action == "start":
            target_filename = selected_evaluation_result_filename(target_scheme, filename)
            if not target_filename:
                raise ValueError("请先选择结果文件")
            target_path = evaluation_result_path(target_scheme, target_filename)
            if not target_path.exists():
                raise FileNotFoundError(f"结果文件不存在: {target_path.name}")
            if target_path.name == OPTIMIZATION_RESULT_WORKBOOK_NAME:
                raise ValueError("默认结果文件不允许修改")
            with self._lock:
                if self.status == "运行中":
                    if self.scheme == target_scheme:
                        raise OptimizationStateError("running", f"方案“{target_scheme}”正在评估，无法再次启动")
                    raise OptimizationStateError("running", f"方案“{self.scheme}”正在评估，无法启动方案“{target_scheme}”")
            scheme_payload = PLANNING_STORE.read_scheme(target_scheme)
            planning_rows = read_evaluation_planning_result_rows(target_scheme, target_filename)
            if not planning_rows:
                raise ValueError("当前结果文件缺少规划结果")
            try:
                calculation_precheck.validate_evaluation_fast_feasibility(scheme_payload, planning_rows)
            except ValueError as exc:
                self._mark_start_failure(target_scheme, target_filename, str(exc), str(target_path))
                raise

            with self._lock:
                if self.status == "运行中":
                    if self.scheme == target_scheme:
                        raise OptimizationStateError("running", f"方案“{target_scheme}”正在评估，无法再次启动")
                    raise OptimizationStateError("running", f"方案“{self.scheme}”正在评估，无法启动方案“{target_scheme}”")
                self.status = "运行中"
                self.scheme = target_scheme
                self.start_time = _now_text()
                self.end_time = ""
                self.progress = 0
                self.result_filename = target_filename
                self.result_file = str(target_path)
                self._metrics = []
                self._results = {}
                self._stop_requested = False
                self._terminate_process_unlocked()
                self._run_token += 1
                self._append_log_unlocked("ok", f"启动方案评估，方案：{self.scheme}，结果：{self.result_filename}")
                self._append_log_unlocked("info", "后台评估程序已启动")
                self._event_queue = multiprocessing.Queue()
                self._process = multiprocessing.Process(
                    target=evaluation_process_worker,
                    args=(self._event_queue, target_scheme, target_filename, str(PLANNING_STORE.root)),
                    daemon=True,
                )
                self._process.start()
                self.process_id = self._process.pid
                return self._payload_unlocked()

        if action == "stop":
            with self._lock:
                if self.status != "运行中" or self.scheme != target_scheme:
                    raise OptimizationStateError("not_running", f"方案“{target_scheme}”没有运行")
                self._stop_requested = True
                self._terminate_process_unlocked()
                self.status = "计算中止"
                self.end_time = _now_text()
                self._append_log_unlocked("warn", "停止方案评估")
                return self._payload_unlocked()

        raise ValueError(f"unknown evaluation action: {action}")

    def _mark_start_failure(self, scheme: str, filename: str, message: str, result_file: str = "") -> None:
        with self._lock:
            if self.status == "运行中":
                return
            self.scheme = scheme
            self.result_filename = filename
            self.result_file = result_file
            now = _now_text()
            self.start_time = now
            self.end_time = now
            self.progress = 0
            self._metrics = []
            self._results = {}
            self._stop_requested = False
            self._terminate_process_unlocked()
            self.status = "失败"
            self._append_log_unlocked("error", message)

    def _drain_events_unlocked(self) -> None:
        if not self._event_queue:
            return
        while True:
            try:
                event = self._event_queue.get_nowait()
            except queue.Empty:
                break
            if not isinstance(event, dict):
                continue
            self._handle_process_event_unlocked(event)
            if not self._event_queue:
                break
        if self.status == "运行中" and self._process and not self._process.is_alive():
            self._process.join(timeout=0)
            if self.status == "运行中":
                exit_code = self._process.exitcode
                self.status = "失败"
                self.end_time = _now_text()
                self._append_log_unlocked("error", f"方案评估进程异常退出，退出码：{exit_code}")
            self._close_event_queue_unlocked()

    def _handle_process_event_unlocked(self, event: dict) -> None:
        event_type = str(event.get("type") or "log")
        if event_type == "log":
            self._append_estimate_event_unlocked(event)
            return
        if event_type == "done":
            if self.status != "运行中" or self._stop_requested:
                return
            self.progress = 100
            self._metrics = event.get("metrics") if isinstance(event.get("metrics"), list) else []
            self._results = event.get("results") if isinstance(event.get("results"), dict) else {}
            dispatch_rows = event.get("dispatch_rows") if isinstance(event.get("dispatch_rows"), list) else []
            completed_end_time = _now_text()
            result_path = export_evaluation_results_workbook(
                self._payload_unlocked(
                    read_workbook=False,
                    status_override="已完成",
                    end_time_override=completed_end_time,
                ),
                dispatch_rows,
            )
            self.result_file = str(result_path)
            self.status = "已完成"
            self.end_time = completed_end_time
            self._append_log_unlocked("ok", f"评估结果已写入：{result_path.name}")
            self._join_finished_process_unlocked()
            self._close_event_queue_unlocked()
            return
        if event_type == "timeout" or (event_type == "error" and calculation_timeout_message(event.get("message"))):
            if self.status != "运行中":
                return
            self.status = "超时"
            self.end_time = _now_text()
            self._append_log_unlocked("error", str(event.get("message") or "方案评估达到最大用时，计算超时"))
            self._join_finished_process_unlocked()
            self._close_event_queue_unlocked()
            return
        if event_type == "error":
            if self.status != "运行中":
                return
            self.status = "失败"
            self.end_time = _now_text()
            self._append_log_unlocked("error", str(event.get("message") or "方案评估失败"))
            self._join_finished_process_unlocked()
            self._close_event_queue_unlocked()

    def _append_estimate_event_unlocked(self, event: dict) -> None:
        level = str(event.get("level") or "info")
        message = str(event.get("message") or "")
        progress = event.get("progress")
        if progress is not None:
            try:
                self.progress = max(self.progress, min(100, max(0, int(progress))))
            except (TypeError, ValueError):
                pass
        if message:
            self._append_log_unlocked(level, message)

    def _payload_unlocked(
        self,
        include_hourly_curves: bool = True,
        read_workbook: bool = True,
        status_override: str | None = None,
        end_time_override: str | None = None,
    ) -> dict:
        workbook_payload = None
        if read_workbook and self.status != "运行中" and self.result_filename:
            try:
                workbook_payload = read_result_workbook_display_payload_for_response(
                    evaluation_result_path(self.scheme, self.result_filename),
                    include_hourly_curves=include_hourly_curves,
                )
            except ValueError:
                workbook_payload = None
        status = self.status if status_override is None else status_override
        end_time = self.end_time if end_time_override is None else end_time_override
        return {
            "status": status,
            "scheme": self.scheme,
            "start_time": self.start_time,
            "end_time": end_time,
            "progress": self.progress,
            "result_filename": self.result_filename,
            "result_file": self.result_file,
            "process_id": self.process_id or "",
            "elapsed_seconds": elapsed_seconds_from_times(self.start_time, end_time),
            "metrics": merge_runtime_metrics(
                self._metrics_unlocked(status_override=status, end_time_override=end_time),
                workbook_payload.get("metrics", []) if workbook_payload else [],
            ),
            "results": workbook_payload.get("results", {}) if workbook_payload else (self._results if self._results else self._default_results_unlocked()),
            "logs": list(self._logs),
        }

    def _task_payload_unlocked(self) -> dict:
        result_file = self.result_file
        if not result_file and self.result_filename:
            try:
                result_path = evaluation_result_path(self.scheme, self.result_filename)
                if result_path.exists():
                    result_file = str(result_path)
            except ValueError:
                result_file = ""
        return {
            "status": self.status,
            "scheme": self.scheme,
            "start_time": self.start_time,
            "end_time": self.end_time,
            "progress": self.progress,
            "result_filename": self.result_filename,
            "result_file": result_file,
            "process_id": self.process_id or "",
            "elapsed_seconds": elapsed_seconds_from_times(self.start_time, self.end_time),
            "logs": list(self._logs),
        }

    def _metrics_unlocked(self, status_override: str | None = None, end_time_override: str | None = None) -> list[dict]:
        status = self.status if status_override is None else status_override
        end_time = self.end_time if end_time_override is None else end_time_override
        base = [
            {"label": "状态", "value": status, "unit": ""},
            {"label": "开始", "value": self.start_time or "-", "unit": ""},
            {"label": "完成", "value": end_time or "-", "unit": ""},
        ]
        existing_labels = {item["label"] for item in base}
        for metric in self._metrics:
            if not isinstance(metric, dict):
                continue
            label = metric.get("label", "")
            if label in existing_labels:
                continue
            base.append(metric)
            existing_labels.add(label)
        if "度电成本" not in existing_labels:
            base.append({"label": "度电成本", "value": "-", "unit": "元"})
            existing_labels.add("度电成本")
        if "绿电占比" not in existing_labels:
            base.append({"label": "绿电占比", "value": "-", "unit": "%"})
            existing_labels.add("绿电占比")
        return base

    @staticmethod
    def _default_results_unlocked() -> dict:
        return {
            "overview_tables": [
                {"title": "规划结果", "rows": []},
                {"title": "规划年指标", "rows": []},
            ],
            "overview_disks": [],
            "green_table": [],
            "safety_table": [],
            "curves": {"green_daily": [], "green_monthly": [], "green_hourly": [], "safety_daily": []},
        }

    def _append_log_unlocked(self, level: str, message: str) -> None:
        self._logs.append({"time": _now_text(), "level": level, "message": message})
        if len(self._logs) > 2000:
            del self._logs[:-2000]

    def _terminate_process_unlocked(self) -> None:
        if self._process and self._process.is_alive():
            self._process.terminate()
            self._process.join(timeout=2)
            if self._process.is_alive():
                self._process.kill()
                self._process.join(timeout=2)
        if self._process and not self._process.is_alive():
            self._process.join(timeout=0)
        self._close_event_queue_unlocked()

    def _join_finished_process_unlocked(self) -> None:
        if self._process and not self._process.is_alive():
            self._process.join(timeout=0)

    def _reap_process_unlocked(self) -> None:
        if self._process and self.status != "运行中" and not self._process.is_alive():
            self._process.join(timeout=0)

    def _close_event_queue_unlocked(self) -> None:
        if not self._event_queue:
            return
        try:
            self._event_queue.close()
            self._event_queue.join_thread()
        except Exception:
            pass
        self._event_queue = None


class EvaluationRuntimeManager:
    """Holds independent evaluation runtimes for multiple schemes."""

    def __init__(self) -> None:
        self._runtimes: dict[str, EvaluationRuntime] = {}
        self._lock = threading.Lock()

    def snapshot(self, scheme: str = "", filename: str = "", include_hourly_curves: bool = True) -> dict:
        runtime = self._runtime_for_result(scheme, filename)
        payload = runtime.snapshot(include_hourly_curves=include_hourly_curves)
        payload["running_schemes"] = self.running_schemes()
        append_task_control_state(
            payload,
            "evaluation",
            self._scheme_name(scheme),
            payload.get("result_filename") or self._result_filename(self._scheme_name(scheme), filename),
        )
        return payload

    def apply(self, action: str, scheme: str = "", filename: str = "") -> dict:
        runtime = self._runtime_for_result(scheme, filename)
        payload = runtime.apply(action, scheme=self._scheme_name(scheme), filename=filename)
        payload["running_schemes"] = self.running_schemes()
        return payload

    def running_schemes(self) -> list[str]:
        with self._lock:
            runtimes = list(self._runtimes.items())
        running = []
        for key, runtime in runtimes:
            if runtime.status == "运行中":
                running.append(key.split("\0", 1)[0])
        return sorted(set(running))

    def runtimes(self) -> dict[str, EvaluationRuntime]:
        with self._lock:
            return dict(self._runtimes)

    def _runtime_for_result(self, scheme: str = "", filename: str = "") -> EvaluationRuntime:
        scheme_name = self._scheme_name(scheme)
        result_filename = self._result_filename(scheme_name, filename)
        key = f"{scheme_name}\0{result_filename}"
        with self._lock:
            if key not in self._runtimes:
                runtime = EvaluationRuntime(scheme=scheme_name)
                runtime.result_filename = result_filename
                if result_filename:
                    try:
                        runtime.result_file = str(evaluation_result_path(scheme_name, result_filename))
                    except ValueError:
                        runtime.result_file = ""
                self._runtimes[key] = runtime
            return self._runtimes[key]

    @staticmethod
    def _result_filename(scheme: str, filename: str = "") -> str:
        selected = str(filename or "").strip()
        if selected:
            return selected
        try:
            return selected_evaluation_result_filename(scheme)
        except (FileNotFoundError, ValueError):
            return ""

    def _runtime_for_scheme(self, scheme: str = "") -> EvaluationRuntime:
        runtime = self._runtime_for_result(scheme)
        return runtime

    @staticmethod
    def _scheme_name(scheme: str = "") -> str:
        return str(scheme or "未选择方案").strip() or "未选择方案"


class FrequencyEvaluationRuntime:
    """Lightweight runtime for frequency-only checks against a result workbook."""

    def __init__(self, scheme: str = "") -> None:
        self.status = "待启动"
        self.scheme = str(scheme or "").strip()
        self.result_filename = ""
        self.result_file = ""
        self.frequency_result_file = ""
        self.start_time = ""
        self.end_time = ""
        self.process_id: int | None = None
        self._metrics: list[dict] = []
        self._summary: list[dict] = []
        self._frequency_table: list[dict] = []
        self._frequency_8760_table: list[dict] = []
        self._curves: dict = {"safety_daily": []}
        self._logs: list[dict[str, str]] = []
        self._lock = threading.Lock()
        self._append_log_unlocked("info", "频率计算待启动")

    def snapshot(self, include_hourly_curves: bool = True) -> dict:
        with self._lock:
            if self.status != "运行中" and self.result_filename:
                self._load_workbook_payload_unlocked()
            return self._payload_unlocked()

    def task_snapshot(self) -> dict:
        with self._lock:
            return self._task_payload_unlocked()

    def apply(self, action: str, scheme: str = "", filename: str = "") -> dict:
        target_scheme = str(scheme or self.scheme or "未选择方案").strip() or "未选择方案"
        if action == "cancel_queue":
            target_filename = selected_evaluation_result_filename(target_scheme, filename)
            with self._lock:
                self.scheme = target_scheme
                self.result_filename = target_filename
                self.result_file = self._safe_result_file_unlocked(target_scheme, target_filename)
                self.frequency_result_file = self._safe_frequency_result_file_unlocked(target_scheme, target_filename)
                self.status = "退出队列"
                self.start_time = ""
                self.end_time = ""
                self.process_id = None
                self._append_log_unlocked("info", "退出等待队列")
                return self._payload_unlocked()

        if action == "start":
            target_filename = selected_evaluation_result_filename(target_scheme, filename)
            if not target_filename:
                raise ValueError("请先选择结果文件")
            target_path = evaluation_result_path(target_scheme, target_filename)
            if not target_path.exists():
                raise FileNotFoundError(f"结果文件不存在: {target_path.name}")
            with self._lock:
                if self.status == "运行中":
                    raise OptimizationStateError("running", f"方案“{target_scheme}”正在进行频率计算，无法再次启动")
                self.status = "运行中"
                self.scheme = target_scheme
                self.result_filename = target_filename
                self.result_file = str(target_path)
                self.frequency_result_file = ""
                self.start_time = _now_text()
                self.end_time = ""
                self.process_id = None
                self._metrics = []
                self._summary = []
                self._frequency_table = []
                self._frequency_8760_table = []
                self._curves = {"safety_daily": []}
                self._append_log_unlocked("ok", f"启动频率计算，方案：{self.scheme}，结果：{self.result_filename}")
            try:
                self._append_log_threadsafe("info", "读取方案参数")
                scheme_payload = PLANNING_STORE.read_scheme(target_scheme)
                payload = build_frequency_evaluation_payload(
                    target_path,
                    scheme=target_scheme,
                    scheme_payload=scheme_payload,
                    generate_curves=True,
                    log_callback=self._append_log_threadsafe,
                )
            except Exception as exc:
                with self._lock:
                    self.status = "失败"
                    self.end_time = _now_text()
                    self._append_log_unlocked("error", f"频率计算失败：{exc}")
                    return self._payload_unlocked()
                raise
            with self._lock:
                self._metrics = payload["metrics"]
                self._summary = payload["summary"]
                self._frequency_table = payload["frequency_table"]
                self._frequency_8760_table = payload.get("frequency_8760_table", [])
                self._curves = payload["curves"]
                self.frequency_result_file = payload.get("frequency_result_file", "")
                self.status = "已完成"
                self.end_time = _now_text()
                self._append_log_unlocked("ok", "频率计算完成")
                return self._payload_unlocked()

        if action == "stop":
            with self._lock:
                if self.status != "运行中":
                    raise OptimizationStateError("not_running", f"方案“{target_scheme}”没有运行中的频率计算")
                self.status = "计算中止"
                self.end_time = _now_text()
                self._append_log_unlocked("warn", "停止频率计算")
                return self._payload_unlocked()

        raise ValueError(f"unknown frequency action: {action}")

    def _payload_unlocked(self) -> dict:
        return {
            "status": self.status,
            "scheme": self.scheme,
            "result_filename": self.result_filename,
            "result_file": self.result_file,
            "frequency_result_file": self.frequency_result_file,
            "start_time": self.start_time,
            "end_time": self.end_time,
            "process_id": self.process_id or "",
            "elapsed_seconds": elapsed_seconds_from_times(self.start_time, self.end_time),
            "metrics": self._metrics_unlocked(),
            "summary": list(self._summary),
            "frequency_table": list(self._frequency_table),
            "frequency_8760_table": list(self._frequency_8760_table),
            "curves": dict(self._curves),
            "logs": list(self._logs),
        }

    def _task_payload_unlocked(self) -> dict:
        return {
            "status": self.status,
            "scheme": self.scheme,
            "result_filename": self.result_filename,
            "result_file": self.result_file,
            "frequency_result_file": self.frequency_result_file,
            "start_time": self.start_time,
            "end_time": self.end_time,
            "process_id": self.process_id or "",
            "elapsed_seconds": elapsed_seconds_from_times(self.start_time, self.end_time),
            "logs": list(self._logs),
        }

    def _metrics_unlocked(self) -> list[dict]:
        base = [
            {"label": "状态", "value": self.status, "unit": ""},
            {"label": "开始", "value": self.start_time or "-", "unit": ""},
            {"label": "完成", "value": self.end_time or "-", "unit": ""},
        ]
        labels = {item["label"] for item in base}
        for metric in self._metrics:
            if isinstance(metric, dict) and metric.get("label") not in labels:
                base.append(metric)
                labels.add(metric.get("label"))
        for label, unit in [("最低频率", "Hz"), ("最高频率", "Hz"), ("频率安全风险小时数", "h")]:
            if label not in labels:
                base.append({"label": label, "value": "-", "unit": unit})
        return base

    def _load_workbook_payload_unlocked(self) -> None:
        try:
            path = evaluation_result_path(self.scheme, self.result_filename)
            if not path.exists():
                return
            payload = build_frequency_evaluation_payload(path, scheme=self.scheme)
        except Exception:
            return
        self.result_file = str(path)
        self._metrics = payload["metrics"]
        self._summary = payload["summary"]
        self._frequency_table = payload["frequency_table"]
        self._frequency_8760_table = payload.get("frequency_8760_table", [])
        self._curves = payload["curves"]
        self.frequency_result_file = payload.get("frequency_result_file", self.frequency_result_file)

    def _safe_result_file_unlocked(self, scheme: str, filename: str) -> str:
        try:
            return str(evaluation_result_path(scheme, filename)) if filename else ""
        except ValueError:
            return ""

    def _safe_frequency_result_file_unlocked(self, scheme: str, filename: str) -> str:
        try:
            path = frequency_curve_result_path(scheme, filename)
            return str(path) if path.exists() else ""
        except ValueError:
            return ""

    def _append_log_unlocked(self, level: str, message: str) -> None:
        self._logs.append({"time": _now_text(), "level": level, "message": message})
        if len(self._logs) > 1000:
            del self._logs[:-1000]

    def _append_log_threadsafe(self, level: str, message: str) -> None:
        with self._lock:
            self._append_log_unlocked(level, message)


class FrequencyEvaluationRuntimeManager:
    def __init__(self) -> None:
        self._runtimes: dict[str, FrequencyEvaluationRuntime] = {}
        self._lock = threading.Lock()

    def snapshot(self, scheme: str = "", filename: str = "", include_hourly_curves: bool = True) -> dict:
        runtime = self._runtime_for_result(scheme, filename)
        payload = runtime.snapshot(include_hourly_curves=include_hourly_curves)
        append_task_control_state(
            payload,
            "frequency",
            self._scheme_name(scheme),
            payload.get("result_filename") or self._result_filename(self._scheme_name(scheme), filename),
        )
        return payload

    def apply(self, action: str, scheme: str = "", filename: str = "") -> dict:
        return self._runtime_for_result(scheme, filename).apply(action, scheme=self._scheme_name(scheme), filename=filename)

    def runtimes(self) -> dict[str, FrequencyEvaluationRuntime]:
        with self._lock:
            return dict(self._runtimes)

    def _runtime_for_result(self, scheme: str = "", filename: str = "") -> FrequencyEvaluationRuntime:
        scheme_name = self._scheme_name(scheme)
        result_filename = self._result_filename(scheme_name, filename)
        key = f"{scheme_name}\0{result_filename}"
        with self._lock:
            if key not in self._runtimes:
                runtime = FrequencyEvaluationRuntime(scheme=scheme_name)
                runtime.result_filename = result_filename
                runtime.result_file = runtime._safe_result_file_unlocked(scheme_name, result_filename)
                self._runtimes[key] = runtime
            return self._runtimes[key]

    @staticmethod
    def _result_filename(scheme: str, filename: str = "") -> str:
        selected = str(filename or "").strip()
        if selected:
            return selected
        try:
            return selected_evaluation_result_filename(scheme)
        except (FileNotFoundError, ValueError):
            return ""

    @staticmethod
    def _scheme_name(scheme: str = "") -> str:
        return str(scheme or "未选择方案").strip() or "未选择方案"


RELIABILITY_DEVICE_DEFAULTS: dict[str, dict] = {
    "wind": {
        "device_type": "wind",
        "label": "风机",
        "unit_count": 0,
        "unit_capacity_kw": 0.0,
        "forced_outage_rate": 0.08,
        "mttr_hours": 72.0,
        "extreme_cold_capacity_factor": 0.75,
        "capex_wan_per_unit": 0.0,
        "fixed_om_rate": 0.025,
        "design_life_years": 20.0,
    },
    "pv": {
        "device_type": "pv",
        "label": "光伏",
        "unit_count": 0,
        "unit_capacity_kw": 0.0,
        "forced_outage_rate": 0.02,
        "mttr_hours": 24.0,
        "extreme_cold_capacity_factor": 0.70,
        "capex_wan_per_unit": 0.0,
        "fixed_om_rate": 0.015,
        "design_life_years": 20.0,
    },
    "storage": {
        "device_type": "storage",
        "label": "储能",
        "unit_count": 0,
        "unit_capacity_kw": 0.0,
        "unit_capacity_kwh": 0.0,
        "forced_outage_rate": 0.03,
        "mttr_hours": 48.0,
        "battery_forced_outage_rate": 0.02,
        "battery_mttr_hours": 96.0,
        "extreme_cold_capacity_factor": 0.70,
        "capex_wan_per_unit": 0.0,
        "fixed_om_rate": 0.015,
        "design_life_years": 10.0,
    },
    "diesel": {
        "device_type": "diesel",
        "label": "柴油发电机",
        "unit_count": 0,
        "unit_capacity_kw": 0.0,
        "forced_outage_rate": 0.05,
        "mttr_hours": 36.0,
        "extreme_cold_capacity_factor": 0.85,
        "capex_wan_per_unit": 0.0,
        "fixed_om_rate": 0.04,
        "design_life_years": 20.0,
        "startup_failure_rate": 0.02,
        "variable_om_yuan_per_kwh": 0.15,
    },
}
RELIABILITY_DEVICE_TYPE_ALIASES = {
    "wind": "wind",
    "wind_turbine": "wind",
    "wind_turbines": "wind",
    "pv": "pv",
    "solar": "pv",
    "photovoltaic": "pv",
    "photovoltaics": "pv",
    "storage": "storage",
    "battery": "storage",
    "ess": "storage",
    "pcs": "storage",
    "diesel": "diesel",
    "generator": "diesel",
    "diesel_generator": "diesel",
    "diesel_generators": "diesel",
}
RELIABILITY_SCHEME_DEVICE_KEYS = {
    "diesel_generators": "diesel",
    "wind_turbines": "wind",
    "photovoltaics": "pv",
    "storage_pcs": "storage",
    "storage_battery_packs": "storage",
}


def reliability_parameters_path_with_store(store: planning_store.PlanningStore, scheme: str) -> Path:
    folder = store.scheme_dir(str(scheme or "").strip())
    return folder / RELIABILITY_PARAMETERS_FILE_NAME


def reliability_parameters_path(scheme: str) -> Path:
    return reliability_parameters_path_with_store(PLANNING_STORE, scheme)


def reliability_result_json_path_with_store(
    store: planning_store.PlanningStore,
    scheme: str,
    filename: str = "",
) -> Path:
    folder = store.scheme_dir(str(scheme or "").strip())
    source_stem = Path(str(filename or OPTIMIZATION_RESULT_WORKBOOK_NAME)).stem
    return folder / f"{source_stem}{RELIABILITY_RESULT_JSON_SUFFIX}"


def reliability_result_json_path(scheme: str, filename: str = "") -> Path:
    return reliability_result_json_path_with_store(PLANNING_STORE, scheme, filename)


def reliability_result_workbook_path_with_store(
    store: planning_store.PlanningStore,
    scheme: str,
    filename: str = "",
) -> Path:
    folder = store.scheme_dir(str(scheme or "").strip())
    source_stem = Path(str(filename or OPTIMIZATION_RESULT_WORKBOOK_NAME)).stem
    return folder / f"{source_stem}{RELIABILITY_RESULT_WORKBOOK_SUFFIX}"


def reliability_result_workbook_path(scheme: str, filename: str = "") -> Path:
    return reliability_result_workbook_path_with_store(PLANNING_STORE, scheme, filename)


def normalize_reliability_source_filename(scheme: str, filename: str = "", *, require_exists: bool = False) -> str:
    scheme_name = str(scheme or "").strip()
    if not scheme_name:
        raise ValueError("请选择规划方案")
    PLANNING_STORE.read_scheme(scheme_name)
    selected = str(filename or "").strip()
    if selected:
        if not RESULT_WORKBOOK_RE.fullmatch(selected) or is_reliability_export_workbook(selected):
            raise ValueError("可靠性评估源文件必须是规划或方案评估结果工作簿")
        source_path = evaluation_result_path(scheme_name, selected)
        if require_exists and not source_path.exists():
            raise FileNotFoundError(f"结果文件不存在: {selected}")
        return selected
    try:
        return selected_evaluation_result_filename(scheme_name)
    except FileNotFoundError:
        return ""


def _reliability_number(
    value: object,
    default: float,
    *,
    field_name: str,
    minimum: float | None = None,
    maximum: float | None = None,
) -> float:
    if value in (None, ""):
        number = float(default)
    else:
        try:
            number = float(value)
        except (TypeError, ValueError) as exc:
            raise ValueError(f"{field_name}必须是数字") from exc
    if not math.isfinite(number):
        raise ValueError(f"{field_name}必须是有限数字")
    if minimum is not None and number < minimum:
        raise ValueError(f"{field_name}不能小于{minimum}")
    if maximum is not None and number > maximum:
        raise ValueError(f"{field_name}不能大于{maximum}")
    return number


def _reliability_ratio(value: object, default: float, *, field_name: str) -> float:
    number = _reliability_number(value, default, field_name=field_name)
    if 1.0 < number <= 100.0:
        number /= 100.0
    if number < 0.0 or number > 1.0:
        raise ValueError(f"{field_name}必须位于0到1之间")
    return number


def _reliability_device_type(row: dict) -> str:
    raw = str(row.get("device_type") or row.get("type") or row.get("device") or "").strip().lower()
    return RELIABILITY_DEVICE_TYPE_ALIASES.get(raw, raw)


def _weighted_device_value(rows: list[dict], field: str, count_field: str = "quantity_upper") -> float:
    weighted = 0.0
    total_count = 0
    fallback = None
    for row in rows:
        value = row.get(field)
        if value in (None, ""):
            continue
        numeric_value = _reliability_number(value, 0.0, field_name=field, minimum=0.0)
        if fallback is None:
            fallback = numeric_value
        count = max(0, int(round(_reliability_number(row.get(count_field), 0.0, field_name=count_field))))
        if count:
            weighted += numeric_value * count
            total_count += count
    if total_count:
        return weighted / total_count
    return float(fallback or 0.0)


def _minimum_positive_device_value(rows: list[dict], field: str, fallback: float) -> float:
    values = [
        _reliability_number(row.get(field), 0.0, field_name=field, minimum=0.0)
        for row in rows
        if row.get(field) not in (None, "")
    ]
    positive = [value for value in values if value > 0.0]
    return min(positive) if positive else fallback


def build_default_reliability_parameters(
    scheme: str,
    filename: str = "",
    *,
    store: planning_store.PlanningStore | None = None,
) -> tuple[dict, list[str], str]:
    active_store = store or PLANNING_STORE
    scheme_name = str(scheme or "").strip()
    scheme_payload = active_store.read_scheme(scheme_name)
    selected = str(filename or "").strip()
    if selected:
        if not RESULT_WORKBOOK_RE.fullmatch(selected) or is_reliability_export_workbook(selected):
            raise ValueError("可靠性评估源文件必须是规划或方案评估结果工作簿")
        planning_rows = read_evaluation_planning_result_rows_with_store(active_store, scheme_name, selected)
    else:
        planning_rows = []
    fixed_payload = estimate.fixed_quantity_payload(scheme_payload, planning_rows) if planning_rows else deepcopy(scheme_payload)

    parameters = {
        "mode": "both",
        "simulation_years": int(reliability.DEFAULT_CONFIG.get("simulation_years", 100)),
        "random_seed": int(reliability.DEFAULT_CONFIG.get("seed", 20260712)),
        "critical_load_ratio": 0.60,
        "reserve_duration_hours": 24.0,
        "dispatch_policy": "reliability_first",
        "assumption_source": "demo_scenario_assumption",
        "devices": [],
    }
    for device_type in ("wind", "pv", "storage", "diesel"):
        device = deepcopy(RELIABILITY_DEVICE_DEFAULTS[device_type])
        if device_type == "wind":
            rows = [row for row in fixed_payload.get("wind_turbines", []) if isinstance(row, dict)]
            device["unit_count"] = sum(max(0, int(round(float(row.get("quantity_upper") or 0)))) for row in rows)
            device["unit_capacity_kw"] = _weighted_device_value(rows, "capacity")
        elif device_type == "pv":
            rows = [row for row in fixed_payload.get("photovoltaics", []) if isinstance(row, dict)]
            device["unit_count"] = sum(max(0, int(round(float(row.get("quantity_upper") or 0)))) for row in rows)
            device["unit_capacity_kw"] = _weighted_device_value(rows, "capacity")
        elif device_type == "diesel":
            rows = [row for row in fixed_payload.get("diesel_generators", []) if isinstance(row, dict)]
            device["unit_count"] = sum(max(0, int(round(float(row.get("quantity_upper") or 0)))) for row in rows)
            device["unit_capacity_kw"] = _weighted_device_value(rows, "power_upper") or _weighted_device_value(rows, "capacity")
        else:
            pcs_rows = [row for row in fixed_payload.get("storage_pcs", []) if isinstance(row, dict)]
            battery_rows = [row for row in fixed_payload.get("storage_battery_packs", []) if isinstance(row, dict)]
            pcs_count = sum(max(0, int(round(float(row.get("quantity_upper") or 0)))) for row in pcs_rows)
            battery_count = sum(max(0, int(round(float(row.get("quantity_upper") or 0)))) for row in battery_rows)
            # The reliability parameter table labels this row as storage PCS
            # and displays unit_capacity_kw.  Use the installed PCS count for
            # that traceability field; the engine still reads PCS and battery
            # quantities independently from the selected planning workbook.
            device["unit_count"] = pcs_count if pcs_count > 0 else battery_count
            device["unit_capacity_kw"] = _weighted_device_value(pcs_rows, "power_capacity")
            device["unit_capacity_kwh"] = _weighted_device_value(battery_rows, "battery_capacity")
            pcs_cost = _weighted_device_value(pcs_rows, "cost")
            battery_cost = _weighted_device_value(battery_rows, "cost")
            device["capex_wan_per_unit"] = pcs_cost + battery_cost
            device["design_life_years"] = min(
                _minimum_positive_device_value(pcs_rows, "design_life_years", device["design_life_years"]),
                _minimum_positive_device_value(battery_rows, "design_life_years", device["design_life_years"]),
            )
            parameters["devices"].append(device)
            continue
        device["capex_wan_per_unit"] = _weighted_device_value(rows, "cost")
        device["design_life_years"] = _minimum_positive_device_value(
            rows,
            "design_life_years",
            float(device["design_life_years"]),
        )
        parameters["devices"].append(device)

    warnings = [
        "当前FOR、MTTR、极寒降额与成本默认值为演示场景假设，不是采购报价或厂商承诺值。",
        "critical_load_ratio和reserve_duration_hours当前仅作为决策假设留档，可靠性引擎仍按完整小时负荷进行供需充裕度计算。",
        "startup_failure_rate及运维成本字段已保存用于追溯，当前两状态故障模型尚未单独使用这些字段。",
    ]
    if planning_rows:
        warnings.append(f"设备台数已从规划结果工作簿 {selected} 读取并固定。")
    else:
        warnings.append("未选择可读规划结果工作簿，设备台数来自方案参数中的固定上下限或上限。")
    return parameters, warnings, selected


def normalize_reliability_parameters(parameters: dict | None, defaults: dict | None = None) -> dict:
    source = deepcopy(parameters) if isinstance(parameters, dict) else {}
    base = deepcopy(defaults) if isinstance(defaults, dict) else {
        "mode": "both",
        "simulation_years": int(reliability.DEFAULT_CONFIG.get("simulation_years", 100)),
        "random_seed": int(reliability.DEFAULT_CONFIG.get("seed", 20260712)),
        "critical_load_ratio": 0.60,
        "reserve_duration_hours": 24.0,
        "dispatch_policy": "reliability_first",
        "assumption_source": "demo_scenario_assumption",
        "devices": [deepcopy(RELIABILITY_DEVICE_DEFAULTS[key]) for key in ("wind", "pv", "storage", "diesel")],
    }
    merged = deepcopy(base)
    for key, value in source.items():
        if key != "devices":
            merged[key] = deepcopy(value)

    mode = str(merged.get("mode") or merged.get("assessment_mode") or "both").strip().lower()
    mode_aliases = {
        "both": "both",
        "combined": "both",
        "deterministic": "deterministic",
        "n-1": "deterministic",
        "n1": "deterministic",
        "monte_carlo": "monte_carlo",
        "monte-carlo": "monte_carlo",
        "probabilistic": "monte_carlo",
    }
    if mode not in mode_aliases:
        raise ValueError("可靠性评估模式仅支持both、deterministic或monte_carlo")
    merged["mode"] = mode_aliases[mode]
    years_value = merged.get("simulation_years", merged.get("monte_carlo_years", 100))
    years = _reliability_number(years_value, 100.0, field_name="蒙特卡洛年数", minimum=1.0, maximum=10000.0)
    merged["simulation_years"] = int(round(years))
    seed_value = merged.get("random_seed", merged.get("seed", reliability.DEFAULT_CONFIG.get("seed", 20260712)))
    seed = _reliability_number(seed_value, 20260712.0, field_name="随机种子", minimum=0.0, maximum=2147483647.0)
    merged["random_seed"] = int(round(seed))
    merged["critical_load_ratio"] = _reliability_ratio(
        merged.get("critical_load_ratio", 0.60),
        0.60,
        field_name="关键负荷比例",
    )
    merged["reserve_duration_hours"] = _reliability_number(
        merged.get("reserve_duration_hours", 24.0),
        24.0,
        field_name="备用持续时间",
        minimum=0.0,
        maximum=720.0,
    )
    merged["dispatch_policy"] = str(merged.get("dispatch_policy") or "reliability_first").strip() or "reliability_first"
    merged["assumption_source"] = str(merged.get("assumption_source") or "demo_scenario_assumption").strip()

    base_rows = base.get("devices") if isinstance(base.get("devices"), list) else []
    base_by_type = {
        _reliability_device_type(row): deepcopy(row)
        for row in base_rows
        if isinstance(row, dict) and _reliability_device_type(row)
    }
    incoming = source.get("devices", merged.get("devices"))
    if isinstance(incoming, dict):
        incoming_rows = [
            {"device_type": device_type, **(deepcopy(value) if isinstance(value, dict) else {})}
            for device_type, value in incoming.items()
        ]
    elif isinstance(incoming, list):
        incoming_rows = [deepcopy(row) for row in incoming if isinstance(row, dict)]
    else:
        incoming_rows = []
    incoming_by_type = {
        _reliability_device_type(row): row
        for row in incoming_rows
        if _reliability_device_type(row)
    }

    ordered_types = ["wind", "pv", "storage", "diesel"]
    ordered_types.extend(
        device_type
        for device_type in incoming_by_type
        if device_type not in ordered_types
    )
    normalized_devices = []
    for device_type in ordered_types:
        row = deepcopy(base_by_type.get(device_type) or RELIABILITY_DEVICE_DEFAULTS.get(device_type) or {
            "device_type": device_type,
            "label": device_type,
        })
        row.update(deepcopy(incoming_by_type.get(device_type) or {}))
        row["device_type"] = device_type
        row["label"] = str(row.get("label") or RELIABILITY_DEVICE_DEFAULTS.get(device_type, {}).get("label") or device_type)
        row["unit_count"] = int(round(_reliability_number(
            row.get("unit_count", row.get("count", 0)),
            0.0,
            field_name=f"{row['label']}台数",
            minimum=0.0,
            maximum=100000.0,
        )))
        for capacity_field in ("unit_capacity", "unit_capacity_kw", "unit_capacity_kwh"):
            if row.get(capacity_field) not in (None, ""):
                row[capacity_field] = _reliability_number(
                    row.get(capacity_field),
                    0.0,
                    field_name=f"{row['label']}{capacity_field}",
                    minimum=0.0,
                )
        row["forced_outage_rate"] = _reliability_ratio(
            row.get("forced_outage_rate", row.get("for", 0.0)),
            0.0,
            field_name=f"{row['label']} FOR",
        )
        row["mttr_hours"] = _reliability_number(
            row.get("mttr_hours", row.get("mttr", 0.0)),
            0.0,
            field_name=f"{row['label']} MTTR",
            minimum=0.0,
            maximum=8760.0,
        )
        if 0.0 < row["forced_outage_rate"] < 1.0 and row["mttr_hours"] <= 0.0:
            raise ValueError(f"{row['label']}设置非零FOR时必须同时提供大于0的MTTR")
        if device_type == "storage":
            row["battery_forced_outage_rate"] = _reliability_ratio(
                row.get("battery_forced_outage_rate", row.get("forced_outage_rate", 0.0)),
                row["forced_outage_rate"],
                field_name="储能电池组 FOR",
            )
            row["battery_mttr_hours"] = _reliability_number(
                row.get("battery_mttr_hours", row.get("mttr_hours", 0.0)),
                row["mttr_hours"],
                field_name="储能电池组 MTTR",
                minimum=0.0,
                maximum=8760.0,
            )
            if 0.0 < row["battery_forced_outage_rate"] < 1.0 and row["battery_mttr_hours"] <= 0.0:
                raise ValueError("储能电池组设置非零FOR时必须同时提供大于0的MTTR")
        for ratio_field in ("extreme_cold_capacity_factor", "fixed_om_rate", "startup_failure_rate"):
            if row.get(ratio_field) not in (None, ""):
                row[ratio_field] = _reliability_ratio(
                    row.get(ratio_field),
                    0.0,
                    field_name=f"{row['label']}{ratio_field}",
                )
        for numeric_field in ("capex_wan_per_unit", "variable_om_yuan_per_kwh", "design_life_years"):
            if row.get(numeric_field) not in (None, ""):
                row[numeric_field] = _reliability_number(
                    row.get(numeric_field),
                    0.0,
                    field_name=f"{row['label']}{numeric_field}",
                    minimum=0.0,
                )
        normalized_devices.append(row)
    merged["devices"] = normalized_devices
    return merged


def read_reliability_parameters(
    scheme: str,
    filename: str = "",
    *,
    store: planning_store.PlanningStore | None = None,
) -> tuple[dict, list[str], str]:
    active_store = store or PLANNING_STORE
    defaults, warnings, selected = build_default_reliability_parameters(scheme, filename, store=active_store)
    path = reliability_parameters_path_with_store(active_store, scheme)
    if path.exists():
        try:
            saved = json.loads(path.read_text(encoding="utf-8"))
        except (OSError, json.JSONDecodeError) as exc:
            raise ValueError(f"可靠性参数文件无法读取: {path.name}") from exc
        parameters = normalize_reliability_parameters(saved, defaults)
    else:
        parameters = normalize_reliability_parameters({}, defaults)

    if selected:
        fresh_by_type = {
            row["device_type"]: row
            for row in defaults.get("devices", [])
            if isinstance(row, dict) and row.get("device_type")
        }
        for row in parameters.get("devices", []):
            source_row = fresh_by_type.get(row.get("device_type"))
            if source_row:
                row["unit_count"] = source_row.get("unit_count", row.get("unit_count", 0))
                row["unit_count_source"] = "planning_result_workbook"
        parameters["source_result_filename"] = selected
    return parameters, warnings, selected


def write_reliability_parameters(
    scheme: str,
    parameters: dict,
    filename: str = "",
) -> tuple[dict, list[str], str]:
    defaults, warnings, selected = build_default_reliability_parameters(scheme, filename)
    normalized = normalize_reliability_parameters(parameters, defaults)
    if selected:
        default_by_type = {
            row["device_type"]: row
            for row in defaults.get("devices", [])
            if isinstance(row, dict) and row.get("device_type")
        }
        for row in normalized.get("devices", []):
            source_row = default_by_type.get(row.get("device_type"))
            if source_row:
                row["unit_count"] = source_row.get("unit_count", row.get("unit_count", 0))
                row["unit_count_source"] = "planning_result_workbook"
        normalized["source_result_filename"] = selected
    path = reliability_parameters_path(scheme)
    path.parent.mkdir(parents=True, exist_ok=True)
    tmp_path = path.with_name(f".{path.name}.tmp")
    tmp_path.write_text(json.dumps(normalized, ensure_ascii=False, indent=2), encoding="utf-8")
    os.replace(tmp_path, path)
    file_cache.invalidate_path(path)
    return normalized, warnings, selected


def prepare_reliability_scheme_payload(
    scheme_payload: dict,
    parameters: dict,
    *,
    has_planning_result: bool,
) -> dict:
    payload = deepcopy(scheme_payload)
    parameter_by_type = {
        _reliability_device_type(row): row
        for row in parameters.get("devices", [])
        if isinstance(row, dict) and _reliability_device_type(row)
    }
    for source_key, device_type in RELIABILITY_SCHEME_DEVICE_KEYS.items():
        rows = payload.get(source_key)
        device_parameters = parameter_by_type.get(device_type)
        if not isinstance(rows, list) or not isinstance(device_parameters, dict):
            continue
        for row in rows:
            if not isinstance(row, dict):
                continue
            applied_parameters = deepcopy(device_parameters)
            if source_key == "storage_battery_packs":
                applied_parameters["forced_outage_rate"] = device_parameters.get(
                    "battery_forced_outage_rate",
                    device_parameters.get("forced_outage_rate", 0.0),
                )
                applied_parameters["mttr_hours"] = device_parameters.get(
                    "battery_mttr_hours",
                    device_parameters.get("mttr_hours", 0.0),
                )
            row["forced_outage_rate"] = applied_parameters.get("forced_outage_rate", 0.0)
            row["mttr_hours"] = applied_parameters.get("mttr_hours", 0.0)
            row["reliability"] = applied_parameters
            if not has_planning_result and len(rows) == 1:
                row["installed_quantity"] = device_parameters.get("unit_count", row.get("quantity_upper", 0))
            capacity_kw = device_parameters.get("unit_capacity_kw", device_parameters.get("unit_capacity"))
            if capacity_kw not in (None, ""):
                if source_key == "diesel_generators":
                    row["capacity"] = capacity_kw
                    row["power_upper"] = capacity_kw
                elif source_key in {"wind_turbines", "photovoltaics"}:
                    row["capacity"] = capacity_kw
                elif source_key == "storage_pcs":
                    row["power_capacity"] = capacity_kw
            capacity_kwh = device_parameters.get("unit_capacity_kwh")
            if source_key == "storage_battery_packs" and capacity_kwh not in (None, ""):
                row["battery_capacity"] = capacity_kwh
            derating = device_parameters.get(
                "extreme_cold_capacity_factor",
                device_parameters.get("cold_derating_factor"),
            )
            if derating not in (None, "") and source_key in {"wind_turbines", "photovoltaics"}:
                row["output_derating_factor"] = derating
            if device_parameters.get("capex_wan_per_unit") not in (None, ""):
                row["cost"] = device_parameters.get("capex_wan_per_unit")
            if device_parameters.get("design_life_years") not in (None, ""):
                row["design_life_years"] = device_parameters.get("design_life_years")
    return payload


def reliability_engine_config(parameters: dict) -> dict:
    policy = str(parameters.get("dispatch_policy") or "reliability_first").strip().lower()
    policy_map = {
        "reliability_first": "renewable_storage_diesel",
        "renewable_first": "renewable_storage_diesel",
        "economic": "renewable_storage_diesel",
        "proportional": "renewable_storage_diesel",
        "renewable_storage_diesel": "renewable_storage_diesel",
        "diesel_first": "renewable_diesel_storage",
        "renewable_diesel_storage": "renewable_diesel_storage",
    }
    mode = str(parameters.get("mode") or "both")
    return {
        "seed": int(parameters.get("random_seed", reliability.DEFAULT_CONFIG.get("seed", 20260712))),
        "simulation_years": int(parameters.get("simulation_years", reliability.DEFAULT_CONFIG.get("simulation_years", 100))),
        "confidence_level": float(parameters.get("confidence_level", 0.95)),
        "initial_storage_soc_ratio": float(parameters.get("initial_storage_soc_ratio", 0.5)),
        "include_annual_samples": True,
        "include_device_contributions": True,
        "run_n_minus_one": mode != "monte_carlo",
        "dispatch_policy": policy_map.get(policy, "renewable_storage_diesel"),
    }


def _deterministic_reliability_result(case: dict) -> dict:
    n_minus_one = reliability.run_n_minus_one(case)
    scenarios = n_minus_one.get("scenarios", [])
    base_case = n_minus_one.get("base_case", {})
    return {
        "status": "completed",
        "schema_version": reliability.SCHEMA_VERSION,
        "input": reliability.input_summary(case),
        "method": {
            "availability_model": "deterministic_full_horizon_single_unit_outage",
            "dispatch_model": "hourly_priority_dispatch",
            "simulation_years": 0,
            "hours_per_year": int(case.get("source_hours", 0)),
        },
        "summary": {
            "simulated_years": 0,
            "eens_kwh_per_year": None,
            "lole_hours_per_year": None,
            "lolp": None,
            "lpsp": None,
            "energy_supply_reliability": base_case.get("energy_supply_reliability"),
            "time_supply_availability": base_case.get("time_supply_availability"),
            "p95_ens_kwh_per_year": None,
            "max_deficit_kw": max((float(row.get("max_deficit_kw", 0.0)) for row in scenarios), default=0.0),
            "longest_consecutive_outage_hours": max(
                (float(row.get("longest_consecutive_outage_hours", 0.0)) for row in scenarios),
                default=0.0,
            ),
        },
        "confidence_intervals": {},
        "n_minus_one": n_minus_one,
        "device_contributions": [],
        "warnings": [
            "本次仅运行确定性N-1压力测试，未生成LOLP、LOLE、EENS等概率年期望指标。",
        ],
    }


def _result_metric_number(mapping: dict, key: str, default=None):
    value = mapping.get(key, default) if isinstance(mapping, dict) else default
    if value in (None, ""):
        return default
    try:
        number = float(value)
    except (TypeError, ValueError):
        return default
    return number if math.isfinite(number) else default


def normalize_reliability_result(
    raw_result: dict,
    *,
    scheme: str,
    filename: str,
    parameters: dict,
    logs: list[dict] | None = None,
) -> dict:
    result = deepcopy(raw_result) if isinstance(raw_result, dict) else {}
    summary = result.get("summary") if isinstance(result.get("summary"), dict) else {}
    confidence = result.get("confidence_intervals") if isinstance(result.get("confidence_intervals"), dict) else {}
    eens_ci = confidence.get("eens_kwh_per_year") if isinstance(confidence.get("eens_kwh_per_year"), dict) else {}
    n_minus_one = result.get("n_minus_one") if isinstance(result.get("n_minus_one"), dict) else {}
    raw_scenarios = n_minus_one.get("scenarios") if isinstance(n_minus_one.get("scenarios"), list) else []
    scenarios = []
    for item in raw_scenarios:
        if not isinstance(item, dict):
            continue
        row = deepcopy(item)
        row.update(
            {
                "id": item.get("scenario_id"),
                "name": item.get("scenario_name") or item.get("scenario_id"),
                "outage_device": item.get("device_name"),
                "duration_hours": item.get("longest_consecutive_outage_hours"),
                "max_unmet_load_kw": item.get("max_deficit_kw"),
                "unserved_energy_kwh": item.get("ens_kwh"),
                "critical_load_supply_rate": item.get("energy_supply_reliability"),
                "pass": item.get("passed"),
            }
        )
        scenarios.append(row)
    n1_pass_rate = (
        sum(1 for row in scenarios if row.get("passed")) / len(scenarios)
        if scenarios
        else None
    )
    metrics = {
        "eens_kwh_per_year": summary.get("eens_kwh_per_year"),
        "lole_hours_per_year": summary.get("lole_hours_per_year"),
        "lolp": summary.get("lolp"),
        "lpsp": summary.get("lpsp"),
        "supply_reliability": summary.get("energy_supply_reliability"),
        "availability": summary.get("time_supply_availability"),
        "n1_pass_rate": n1_pass_rate,
        "p95_eens_kwh": summary.get("p95_ens_kwh_per_year"),
        "max_unmet_load_kw": summary.get("max_deficit_kw"),
        "max_outage_duration_hours": summary.get("longest_consecutive_outage_hours"),
        "simulation_years": summary.get("simulated_years", parameters.get("simulation_years")),
        "confidence_level": confidence.get("confidence_level"),
        "eens_ci_lower_kwh": eens_ci.get("lower"),
        "eens_ci_upper_kwh": eens_ci.get("upper"),
    }
    annual_samples = result.get("annual_samples") if isinstance(result.get("annual_samples"), list) else []
    annual_distribution = []
    for sample in annual_samples:
        if not isinstance(sample, dict):
            continue
        row = deepcopy(sample)
        row["eens_kwh"] = sample.get("ens_kwh_per_year")
        row["lole_hours"] = sample.get("lole_hours_per_year")
        annual_distribution.append(row)
    normalized_contributions = []
    for item in result.get("device_contributions", []) if isinstance(result.get("device_contributions"), list) else []:
        if not isinstance(item, dict):
            continue
        row = deepcopy(item)
        row["label"] = item.get("device_name") or item.get("device_type_label") or item.get("device_id")
        row["eens_kwh"] = item.get("marginal_eens_reduction_kwh_per_year")
        row["share"] = item.get("normalized_contribution_share")
        normalized_contributions.append(row)
    assumption_log = [
        {
            "name": "参数来源",
            "value": parameters.get("assumption_source", "demo_scenario_assumption"),
            "note": "场景假设应与现场统计、厂商承诺和采购报价区分。",
        },
        {"name": "评估模式", "value": parameters.get("mode", "both")},
        {"name": "蒙特卡洛年数", "value": parameters.get("simulation_years"), "unit": "年"},
        {"name": "随机种子", "value": parameters.get("random_seed")},
        {"name": "关键负荷比例", "value": parameters.get("critical_load_ratio"), "unit": "p.u.", "note": "当前仅留档，计算仍按完整负荷。"},
        {"name": "备用持续时间", "value": parameters.get("reserve_duration_hours"), "unit": "h", "note": "当前仅留档。"},
        {"name": "源规划结果", "value": filename or "方案参数固定台数"},
        {"name": "故障独立性", "value": "设备独立两状态模型", "note": "共同原因故障与极端天气相关故障尚未单独建模。"},
    ]
    for device in parameters.get("devices", []):
        if not isinstance(device, dict):
            continue
        label = str(device.get("label") or device.get("device_type") or "设备")
        unit_count = int(round(_result_metric_number(device, "unit_count", 0) or 0))
        unit_capacity_kw = _result_metric_number(device, "unit_capacity_kw", 0.0) or 0.0
        forced_outage_rate = _result_metric_number(device, "forced_outage_rate", 0.0) or 0.0
        mttr_hours = _result_metric_number(device, "mttr_hours", 0.0) or 0.0
        value_parts = [
            f"{unit_count}台 × {unit_capacity_kw:g} kW",
            f"FOR {forced_outage_rate * 100:g}%",
            f"MTTR {mttr_hours:g} h",
        ]
        if device.get("device_type") == "storage":
            battery_for = _result_metric_number(device, "battery_forced_outage_rate", forced_outage_rate)
            battery_mttr = _result_metric_number(device, "battery_mttr_hours", mttr_hours)
            value_parts.extend(
                [
                    f"电池组FOR {(battery_for or 0.0) * 100:g}%",
                    f"电池组MTTR {battery_mttr or 0.0:g} h",
                ]
            )
        cold_factor = _result_metric_number(device, "extreme_cold_capacity_factor", 1.0)
        capex = _result_metric_number(device, "capex_wan_per_unit", 0.0)
        fixed_om = _result_metric_number(device, "fixed_om_rate", 0.0)
        design_life = _result_metric_number(device, "design_life_years", 0.0)
        note_parts = [
            f"极寒/系统可用系数 {(cold_factor or 0.0) * 100:g}%",
            f"CAPEX {capex or 0.0:g} 万元/单元",
            f"固定运维 {(fixed_om or 0.0) * 100:g}%/年",
            f"设计寿命 {design_life or 0.0:g} 年",
        ]
        if device.get("device_type") == "diesel":
            startup_failure = _result_metric_number(device, "startup_failure_rate", 0.0)
            variable_om = _result_metric_number(device, "variable_om_yuan_per_kwh", 0.0)
            note_parts.extend(
                [
                    f"启动失败率 {(startup_failure or 0.0) * 100:g}%",
                    f"变动运维 {variable_om or 0.0:g} 元/kWh",
                ]
            )
        assumption_log.append(
            {
                "name": f"{label}可靠性参数",
                "value": "；".join(value_parts),
                "note": "；".join(note_parts) + "。成本为场景假设、非采购报价；未被内核使用的字段仍保留用于追溯。",
            }
        )
    warnings = list(result.get("warnings") or [])
    warnings.extend(
        [
            "critical_load_ratio、reserve_duration_hours、startup_failure_rate及成本字段当前作为追溯假设保存，尚未改变充裕度调度。",
            "概率结果基于设备独立两状态故障模型，未覆盖共同原因故障、燃油断供和人员无法维修等极地风险。",
        ]
    )
    result.update(
        {
            "status": "completed",
            "scheme": scheme,
            "source_result_filename": filename,
            "generated_at": _now_text(),
            "parameters": deepcopy(parameters),
            "metrics": metrics,
            "reliability_metrics": metrics,
            "n1_scenarios": scenarios,
            "annual_distribution": annual_distribution,
            "device_contributions": normalized_contributions,
            "assumption_log": assumption_log,
            "warnings": list(dict.fromkeys(str(item) for item in warnings if str(item).strip())),
            "logs": deepcopy(logs or []),
        }
    )
    return result


def reliability_process_worker(
    event_queue,
    scheme: str,
    filename: str,
    parameters: dict,
    planning_root: str = "",
) -> None:
    worker_logs: list[dict] = []

    def emit(level: str, message: str, progress: int | None = None) -> None:
        item = {"time": _now_text(), "level": level, "message": message}
        worker_logs.append(item)
        event = {"type": "log", **item}
        if progress is not None:
            event["progress"] = int(progress)
        event_queue.put(event)

    try:
        emit("info", "读取方案参数和规划结果台数", 5)
        store = planning_store.PlanningStore(root=Path(planning_root)) if planning_root else PLANNING_STORE
        scheme_payload = store.read_scheme(scheme)
        planning_rows = read_evaluation_planning_result_rows_with_store(store, scheme, filename) if filename else []
        emit("info", f"规划台数来源：{filename or '方案参数固定上下限'}", 15)
        prepared_payload = prepare_reliability_scheme_payload(
            scheme_payload,
            parameters,
            has_planning_result=bool(planning_rows),
        )
        engine_config = reliability_engine_config(parameters)
        emit("info", "构建设备两状态故障与小时级供需充裕度模型", 25)
        if str(parameters.get("mode") or "both") == "deterministic":
            case = reliability.build_reliability_case(prepared_payload, planning_rows or None, engine_config)
            emit("info", "执行确定性N-1压力测试", 40)
            raw_result = _deterministic_reliability_result(case)
        else:
            emit("info", "执行序贯蒙特卡洛与N-1可靠性评估", 35)
            raw_result = reliability.run_reliability_assessment(
                prepared_payload,
                planning_rows or None,
                engine_config,
            )
        emit("ok", "可靠性计算完成，准备写入独立结果文件", 90)
        normalized = normalize_reliability_result(
            raw_result,
            scheme=scheme,
            filename=filename,
            parameters=parameters,
            logs=worker_logs,
        )
        event_queue.put({"type": "done", "result": normalized})
    except Exception as exc:
        event_queue.put(
            {
                "type": "error",
                "message": f"可靠性评估失败：{exc}",
                "traceback": traceback.format_exc(),
            }
        )


def _reliability_excel_value(value):
    if isinstance(value, (dict, list, tuple)):
        return json.dumps(value, ensure_ascii=False, separators=(",", ":"))
    return value


def _style_reliability_sheet(sheet) -> None:
    if sheet.max_row >= 1:
        for cell in sheet[1]:
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill("solid", fgColor="1F4E78")
    sheet.freeze_panes = "A2"
    for column_cells in sheet.columns:
        max_length = max(len(str(cell.value)) if cell.value is not None else 0 for cell in column_cells)
        sheet.column_dimensions[get_column_letter(column_cells[0].column)].width = min(max(max_length + 2, 12), 48)


def _append_reliability_dict_sheet(workbook: Workbook, title: str, rows: list[dict], columns: list[tuple[str, str]]) -> None:
    sheet = workbook.create_sheet(title)
    sheet.append([label for _, label in columns])
    for row in rows:
        sheet.append([_reliability_excel_value(row.get(key)) for key, _ in columns])
    _style_reliability_sheet(sheet)


def build_reliability_results_workbook(result: dict) -> Workbook:
    workbook = Workbook()
    summary_sheet = workbook.active
    summary_sheet.title = "摘要"
    summary_sheet.append(["指标", "数值", "单位", "说明"])
    metric_specs = [
        ("EENS", result.get("metrics", {}).get("eens_kwh_per_year"), "kWh/a", "年期望未供电量"),
        ("LOLE", result.get("metrics", {}).get("lole_hours_per_year"), "h/a", "年期望失负荷小时数"),
        ("LOLP", result.get("metrics", {}).get("lolp"), "p.u.", "任一小时失负荷概率"),
        ("LPSP", result.get("metrics", {}).get("lpsp"), "p.u.", "未供电量占负荷电量比例"),
        ("供电可靠率", result.get("metrics", {}).get("supply_reliability"), "p.u.", "按电量统计"),
        ("N-1通过率", result.get("metrics", {}).get("n1_pass_rate"), "p.u.", "单设备停运场景通过比例"),
        ("P95 EENS", result.get("metrics", {}).get("p95_eens_kwh"), "kWh/a", "年度分布95分位"),
        ("最大失供功率", result.get("metrics", {}).get("max_unmet_load_kw"), "kW", "模拟样本最大功率缺口"),
        ("最大失供持续时间", result.get("metrics", {}).get("max_outage_duration_hours"), "h", "最长连续失负荷时段"),
        ("模拟年数", result.get("metrics", {}).get("simulation_years"), "年", "序贯蒙特卡洛样本数"),
    ]
    for row in metric_specs:
        summary_sheet.append(list(row))
    _style_reliability_sheet(summary_sheet)

    _append_reliability_dict_sheet(
        workbook,
        "N-1场景",
        result.get("n1_scenarios", []),
        [
            ("name", "场景"),
            ("outage_device", "停运设备"),
            ("device_type_label", "设备类型"),
            ("installed_units", "安装台数"),
            ("removed_units", "停运台数"),
            ("max_unmet_load_kw", "最大失供功率(kW)"),
            ("unserved_energy_kwh", "未供电量(kWh)"),
            ("lole_hours", "失负荷时长(h)"),
            ("longest_consecutive_outage_hours", "最长连续失供(h)"),
            ("pass", "是否通过"),
        ],
    )
    _append_reliability_dict_sheet(
        workbook,
        "年度样本",
        result.get("annual_distribution", []),
        [
            ("year", "模拟年"),
            ("seed", "随机种子"),
            ("eens_kwh", "EENS(kWh/a)"),
            ("lole_hours", "LOLE(h/a)"),
            ("lolp", "LOLP"),
            ("lpsp", "LPSP"),
            ("max_deficit_kw", "最大失供功率(kW)"),
            ("longest_consecutive_outage_hours", "最长连续失供(h)"),
        ],
    )
    _append_reliability_dict_sheet(
        workbook,
        "设备贡献",
        result.get("device_contributions", []),
        [
            ("label", "设备"),
            ("device_type_label", "设备类型"),
            ("installed_units", "安装台数"),
            ("input_forced_outage_rate", "输入FOR"),
            ("input_mttr_hours", "输入MTTR(h)"),
            ("observed_unavailability", "模拟不可用率"),
            ("eens_kwh", "边际EENS贡献(kWh/a)"),
            ("share", "归一化贡献占比"),
        ],
    )
    assumption_rows = []
    for row in result.get("assumption_log", []):
        if isinstance(row, dict):
            assumption_rows.append(
                {
                    "record_type": "假设",
                    "name": row.get("name"),
                    "value": row.get("value"),
                    "unit": row.get("unit"),
                    "note": row.get("note"),
                }
            )
    for warning in result.get("warnings", []):
        assumption_rows.append({"record_type": "边界", "name": "警告", "note": warning})
    for log in result.get("logs", []):
        if isinstance(log, dict):
            assumption_rows.append(
                {
                    "record_type": "日志",
                    "name": log.get("level"),
                    "value": log.get("time"),
                    "note": log.get("message"),
                }
            )
    _append_reliability_dict_sheet(
        workbook,
        "假设与日志",
        assumption_rows,
        [
            ("record_type", "记录类型"),
            ("name", "名称/级别"),
            ("value", "数值/时间"),
            ("unit", "单位"),
            ("note", "说明"),
        ],
    )
    return workbook


def save_reliability_result_json(path: Path, result: dict) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    tmp_path = path.with_name(f".{path.name}.tmp")
    tmp_path.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
    os.replace(tmp_path, path)
    file_cache.invalidate_path(path)


def read_reliability_result_json(path: Path) -> dict | None:
    if not path.exists():
        return None
    try:
        payload = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError) as exc:
        raise ValueError(f"可靠性结果文件无法读取: {path.name}") from exc
    if not isinstance(payload, dict):
        raise ValueError(f"可靠性结果文件格式不正确: {path.name}")
    return payload


class ReliabilityRuntime:
    """Independent background runtime for one scheme/result reliability case."""

    def __init__(self, scheme: str = "", filename: str = "") -> None:
        self.status = "idle"
        self.scheme = str(scheme or "").strip()
        self.result_filename = str(filename or "").strip()
        self.start_time = ""
        self.end_time = ""
        self.progress = 0
        self.result_file = ""
        self.export_file = ""
        self.process_id: int | None = None
        self._result: dict = {}
        self._logs: list[dict] = []
        self._lock = threading.Lock()
        self._process: multiprocessing.Process | None = None
        self._event_queue: multiprocessing.Queue | None = None
        self._stop_requested = False
        self._append_log_unlocked("info", "可靠性评估待启动")

    def snapshot(self, *, light: bool = False) -> dict:
        with self._lock:
            self._drain_events_unlocked()
            self._reap_process_unlocked()
            return self._payload_unlocked(light=light)

    def apply(self, action: str, *, scheme: str, filename: str, parameters: dict | None = None) -> dict:
        target_scheme = str(scheme or self.scheme or "").strip()
        target_filename = str(filename or self.result_filename or "").strip()
        if action == "clear_logs":
            with self._lock:
                self.scheme = target_scheme
                self.result_filename = target_filename
                self._logs.clear()
                return self._payload_unlocked()
        if action == "cancel_queue":
            with self._lock:
                if self.status == "running":
                    raise OptimizationStateError("running", "可靠性评估正在运行，无法退出队列")
                self.status = "idle"
                self.progress = 0
                self._append_log_unlocked("info", "已退出可靠性等待队列")
                return self._payload_unlocked()
        if action in {"start", "queue"}:
            with self._lock:
                self._drain_events_unlocked()
                if self.status == "running":
                    raise OptimizationStateError("running", f"方案“{target_scheme}”的可靠性评估正在运行")
            if not target_scheme:
                raise ValueError("请选择规划方案")
            PLANNING_STORE.read_scheme(target_scheme)
            if target_filename:
                source_path = evaluation_result_path(target_scheme, target_filename)
                if not source_path.exists():
                    raise FileNotFoundError(f"结果文件不存在: {target_filename}")
            if parameters is None:
                target_parameters, _, _ = read_reliability_parameters(target_scheme, target_filename)
            else:
                defaults, _, _ = build_default_reliability_parameters(target_scheme, target_filename)
                target_parameters = normalize_reliability_parameters(parameters, defaults)
                if target_filename:
                    default_by_type = {
                        row["device_type"]: row
                        for row in defaults.get("devices", [])
                        if isinstance(row, dict) and row.get("device_type")
                    }
                    for row in target_parameters.get("devices", []):
                        source_row = default_by_type.get(row.get("device_type"))
                        if source_row:
                            row["unit_count"] = source_row.get("unit_count", row.get("unit_count", 0))
                            row["unit_count_source"] = "planning_result_workbook"
            with self._lock:
                if self.status == "running":
                    raise OptimizationStateError("running", f"方案“{target_scheme}”的可靠性评估正在运行")
                self._terminate_process_unlocked()
                self.scheme = target_scheme
                self.result_filename = target_filename
                self.status = "running"
                self.start_time = _now_text()
                self.end_time = ""
                self.progress = 0
                self.result_file = ""
                self.export_file = ""
                self._result = {}
                self._stop_requested = False
                self._logs.clear()
                if action == "queue":
                    self._append_log_unlocked("info", "可靠性队列当前按独立后台任务立即启动")
                self._append_log_unlocked("ok", f"启动可靠性评估：{target_scheme} / {target_filename or '方案参数'}")
                self._event_queue = multiprocessing.Queue()
                self._process = multiprocessing.Process(
                    target=reliability_process_worker,
                    args=(
                        self._event_queue,
                        target_scheme,
                        target_filename,
                        target_parameters,
                        str(PLANNING_STORE.root),
                    ),
                    daemon=True,
                )
                self._process.start()
                self.process_id = self._process.pid
                return self._payload_unlocked(light=True)
        if action == "stop":
            with self._lock:
                self._drain_events_unlocked()
                if self.status != "running":
                    raise OptimizationStateError("not_running", "当前可靠性评估没有运行")
                self._stop_requested = True
                self._terminate_process_unlocked()
                self.status = "stopped"
                self.end_time = _now_text()
                self._append_log_unlocked("warn", "可靠性评估已停止")
                return self._payload_unlocked(light=True)
        raise ValueError(f"unknown reliability action: {action}")

    def _payload_unlocked(self, *, light: bool = False) -> dict:
        result_path = reliability_result_json_path(self.scheme, self.result_filename) if self.scheme else None
        export_path = reliability_result_workbook_path(self.scheme, self.result_filename) if self.scheme else None
        if result_path and result_path.exists() and self.status == "idle":
            self.status = "completed"
            self.progress = 100
            self.result_file = str(result_path)
            self.export_file = str(export_path) if export_path and export_path.exists() else ""
            if not self.end_time:
                self.end_time = datetime.fromtimestamp(result_path.stat().st_mtime).strftime("%Y-%m-%d %H:%M:%S")
        if not light and not self._result and result_path and result_path.exists() and self.status != "running":
            stored = read_reliability_result_json(result_path)
            if stored:
                self._result = stored
                self.result_file = str(result_path)
                self.export_file = str(export_path) if export_path and export_path.exists() else ""
                self.end_time = str(stored.get("generated_at") or self.end_time)
        payload = {
            "status": self.status,
            "scheme": self.scheme,
            "filename": self.result_filename,
            "result_filename": self.result_filename,
            "start_time": self.start_time,
            "end_time": self.end_time,
            "progress": self.progress,
            "elapsed_seconds": elapsed_seconds_from_times(self.start_time, self.end_time),
            "process_id": self.process_id or "",
            "result_file": self.result_file or (str(result_path) if result_path and result_path.exists() else ""),
            "export_file": self.export_file or (str(export_path) if export_path and export_path.exists() else ""),
            "logs": list(self._logs),
        }
        if not light:
            payload["result"] = deepcopy(self._result) if self._result else None
        return payload

    def _drain_events_unlocked(self) -> None:
        if not self._event_queue:
            return
        while True:
            try:
                event = self._event_queue.get_nowait()
            except queue.Empty:
                break
            if isinstance(event, dict):
                self._handle_event_unlocked(event)
            if not self._event_queue:
                break

    def _handle_event_unlocked(self, event: dict) -> None:
        event_type = str(event.get("type") or "log")
        if event_type == "log":
            if event.get("progress") not in (None, ""):
                self.progress = max(self.progress, min(99, int(event.get("progress"))))
            self._append_log_unlocked(str(event.get("level") or "info"), str(event.get("message") or ""))
            return
        if event_type == "done":
            if self.status != "running" or self._stop_requested:
                return
            result = event.get("result") if isinstance(event.get("result"), dict) else {}
            self.progress = 96
            self._append_log_unlocked("info", "写入独立可靠性JSON与Excel结果")
            result_path = reliability_result_json_path(self.scheme, self.result_filename)
            export_path = reliability_result_workbook_path(self.scheme, self.result_filename)
            try:
                result["logs"] = list(self._logs)
                result.setdefault("artifacts", {})
                result["artifacts"]["json"] = str(result_path)
                try:
                    save_result_workbook(build_reliability_results_workbook(result), export_path, "可靠性结果文件")
                    self.export_file = str(export_path)
                    result["artifacts"]["xlsx"] = str(export_path)
                    self._append_log_unlocked("ok", f"可靠性Excel已写入：{export_path.name}")
                except Exception as exc:
                    message = f"可靠性Excel导出失败，JSON结果仍将保留：{exc}"
                    self._append_log_unlocked("warn", message)
                    result.setdefault("warnings", []).append(message)
                self._append_log_unlocked("ok", f"可靠性JSON已写入：{result_path.name}")
                result["logs"] = list(self._logs)
                save_reliability_result_json(result_path, result)
            except Exception as exc:
                self.status = "failed"
                self.end_time = _now_text()
                self._append_log_unlocked("error", f"可靠性结果保存失败：{exc}")
                self._join_finished_process_unlocked()
                self._close_event_queue_unlocked()
                return
            self._result = result
            self.result_file = str(result_path)
            self.progress = 100
            self.status = "completed"
            self.end_time = _now_text()
            self._join_finished_process_unlocked()
            self._close_event_queue_unlocked()
            return
        if event_type == "error":
            if self.status != "running":
                return
            self.status = "failed"
            self.end_time = _now_text()
            self._append_log_unlocked("error", str(event.get("message") or "可靠性评估失败"))
            self._join_finished_process_unlocked()
            self._close_event_queue_unlocked()

    def _reap_process_unlocked(self) -> None:
        if not self._process or self._process.is_alive():
            return
        self._process.join(timeout=0)
        self._drain_events_unlocked()
        if self.status == "running":
            self.status = "failed"
            self.end_time = _now_text()
            self._append_log_unlocked("error", f"可靠性评估进程异常退出，退出码：{self._process.exitcode}")
        self._close_event_queue_unlocked()
        self._process = None

    def _terminate_process_unlocked(self) -> None:
        process = self._process
        if process and process.is_alive():
            process.terminate()
            process.join(timeout=2)
            if process.is_alive() and hasattr(process, "kill"):
                process.kill()
                process.join(timeout=1)
        self._process = None
        self.process_id = None
        self._close_event_queue_unlocked()

    def _join_finished_process_unlocked(self) -> None:
        if self._process:
            self._process.join(timeout=0.2)
        self._process = None

    def _close_event_queue_unlocked(self) -> None:
        event_queue = self._event_queue
        self._event_queue = None
        if event_queue and hasattr(event_queue, "close"):
            try:
                event_queue.close()
            except (OSError, ValueError):
                pass

    def _append_log_unlocked(self, level: str, message: str) -> None:
        self._logs.append({"time": _now_text(), "level": level, "message": message})
        if len(self._logs) > 1000:
            del self._logs[:-1000]


class ReliabilityRuntimeManager:
    def __init__(self) -> None:
        self._runtimes: dict[str, ReliabilityRuntime] = {}
        self._lock = threading.Lock()

    def snapshot(self, scheme: str = "", filename: str = "", *, light: bool = False) -> dict:
        scheme_name = str(scheme or "").strip()
        selected = normalize_reliability_source_filename(scheme_name, filename, require_exists=False)
        return self._runtime(scheme_name, selected).snapshot(light=light)

    def apply(
        self,
        action: str,
        *,
        scheme: str = "",
        filename: str = "",
        parameters: dict | None = None,
    ) -> dict:
        scheme_name = str(scheme or "").strip()
        selected = normalize_reliability_source_filename(
            scheme_name,
            filename,
            require_exists=bool(str(filename or "").strip()),
        )
        return self._runtime(scheme_name, selected).apply(
            action,
            scheme=scheme_name,
            filename=selected,
            parameters=parameters,
        )

    def runtimes(self) -> dict[str, ReliabilityRuntime]:
        with self._lock:
            return dict(self._runtimes)

    def _runtime(self, scheme: str, filename: str) -> ReliabilityRuntime:
        key = f"{scheme}\0{filename}"
        with self._lock:
            if key not in self._runtimes:
                self._runtimes[key] = ReliabilityRuntime(scheme, filename)
            return self._runtimes[key]


def handle_reliability_api_path(
    path: str,
    method: str,
    body: bytes = b"",
    query: str = "",
    current_user: dict | None = None,
) -> tuple[int, dict[str, str], bytes]:
    query_params = parse_qs(query)
    try:
        if path == "/api/reliability/parameters":
            if method == "GET":
                scheme = query_params.get("scheme", [""])[0]
                filename = query_params.get("filename", [""])[0]
                ensure_planning_scheme_access(scheme, current_user)
                selected = normalize_reliability_source_filename(scheme, filename, require_exists=False)
                parameters, warnings, selected = read_reliability_parameters(scheme, selected)
                return _json_response(
                    {
                        "scheme": scheme,
                        "selected": selected,
                        "parameters": parameters,
                        "assumption_warnings": warnings,
                    }
                )
            if method == "PUT":
                payload = json.loads(body.decode("utf-8") or "{}")
                scheme = str(payload.get("scheme") or "").strip()
                filename = str(payload.get("filename") or "").strip()
                ensure_planning_scheme_manage_access(scheme, current_user)
                selected = normalize_reliability_source_filename(
                    scheme,
                    filename,
                    require_exists=bool(filename),
                )
                parameters = payload.get("parameters")
                if not isinstance(parameters, dict):
                    raise ValueError("parameters必须是对象")
                saved, warnings, selected = write_reliability_parameters(scheme, parameters, selected)
                return _json_response(
                    {
                        "ok": True,
                        "scheme": scheme,
                        "selected": selected,
                        "parameters": saved,
                        "assumption_warnings": warnings,
                    }
                )
        if path == "/api/reliability/results" and method == "GET":
            scheme = query_params.get("scheme", [""])[0]
            filename = query_params.get("filename", [""])[0]
            ensure_planning_scheme_access(scheme, current_user)
            selected = normalize_reliability_source_filename(scheme, filename, require_exists=False)
            result_files = list_evaluation_result_files(scheme)
            result_path = reliability_result_json_path(scheme, selected)
            export_path = reliability_result_workbook_path(scheme, selected)
            result = read_reliability_result_json(result_path)
            return _json_response(
                {
                    "scheme": scheme,
                    "selected": selected,
                    "results": result_files,
                    "result_file": result_path.name if result_path.exists() else "",
                    "export_file": export_path.name if export_path.exists() else "",
                    "result": result,
                }
            )
        if path == "/api/reliability/status" and method == "GET":
            scheme = query_params.get("scheme", [""])[0]
            if scheme:
                ensure_planning_scheme_access(scheme, current_user)
            filename = query_params.get("filename", [""])[0]
            light = truthy_json_value(query_params.get("light", ["0"])[0])
            return _json_response(RELIABILITY_RUNTIME.snapshot(scheme, filename, light=light))
    except FileNotFoundError as exc:
        return _json_response({"error": "not_found", "message": str(exc)}, HTTPStatus.NOT_FOUND)
    except (ValueError, json.JSONDecodeError) as exc:
        return _json_response({"error": "bad_request", "message": str(exc)}, HTTPStatus.BAD_REQUEST)
    return _json_response({"error": "not_found", "path": path}, HTTPStatus.NOT_FOUND)


def build_frequency_evaluation_payload(
    path: Path,
    scheme: str = "",
    scheme_payload: dict | None = None,
    generate_curves: bool = False,
    log_callback=None,
) -> dict:
    emit_frequency_log(log_callback, "info", f"读取结果文件：{path.name}")
    workbook_payload = read_result_workbook_display_payload(path, include_hourly_curves=False)
    results = workbook_payload.get("results", {})
    safety_table = [dict(row) for row in results.get("safety_table", []) if isinstance(row, dict)]
    safety_daily = [dict(row) for row in results.get("curves", {}).get("safety_daily", []) if isinstance(row, dict)]
    emit_frequency_log(log_callback, "info", f"已读取频率安全指标：{len(safety_table)}行，日曲线：{len(safety_daily)}点")
    frequency_result_file = ""
    curve_row_count = 0
    frequency_8760_table: list[dict] = []
    if scheme_payload is None and scheme and generate_curves:
        try:
            scheme_payload = PLANNING_STORE.read_scheme(scheme)
        except Exception:
            scheme_payload = None
    if scheme and scheme_payload is not None and generate_curves:
        emit_frequency_log(log_callback, "info", "读取调度结果8760点")
        dispatch_rows = read_frequency_dispatch_rows(path)
        if dispatch_rows:
            emit_frequency_log(log_callback, "info", f"开始生成{len(dispatch_rows)}个时刻的频率仿真曲线")
            frequency_results = []
            for index, result in enumerate(iter_frequency_result_rows(scheme_payload, dispatch_rows), start=1):
                frequency_results.append(result)
                if index == 1 or index % 1000 == 0 or index == len(dispatch_rows):
                    emit_frequency_log(log_callback, "info", f"频率曲线生成进度：{index}/{len(dispatch_rows)}")
            emit_frequency_log(log_callback, "info", "整理频率8760点指标")
            frequency_8760_table = frequency_8760_display_rows_from_results(frequency_results, scheme_payload)
            emit_frequency_log(log_callback, "info", "写入频率计算结果文件")
            frequency_result_path = export_frequency_curve_workbook(
                scheme,
                path.name,
                scheme_payload,
                dispatch_rows,
                frequency_results=frequency_results,
                log_callback=log_callback,
            )
            frequency_result_file = str(frequency_result_path)
            curve_row_count = len(dispatch_rows) * 2
            emit_frequency_log(log_callback, "ok", f"频率计算结果文件已生成：{frequency_result_path.name}，曲线{curve_row_count}条")
    elif scheme:
        existing_curve_path = frequency_curve_result_path(scheme, path.name)
        if existing_curve_path.exists():
            emit_frequency_log(log_callback, "info", f"读取已有频率曲线文件：{existing_curve_path.name}")
            frequency_result_file = str(existing_curve_path)
            frequency_8760_table = read_frequency_8760_display_rows(existing_curve_path, scheme_payload)
            curve_row_count = len(frequency_8760_table) * 2 if frequency_8760_table else TIME_SERIES_IMPORT_ROW_COUNT * 2
    frequency_metrics = frequency_metrics_from_rows(safety_table, safety_daily)
    if frequency_result_file:
        frequency_metrics.append({"label": "频率曲线文件", "value": Path(frequency_result_file).name, "unit": ""})
        frequency_metrics.append({"label": "频率曲线条数", "value": curve_row_count, "unit": "条"})
    return {
        "metrics": frequency_metrics,
        "summary": frequency_summary_rows(frequency_metrics),
        "frequency_table": frequency_detail_rows(safety_table, frequency_metrics),
        "frequency_8760_table": frequency_8760_table,
        "curves": {"safety_daily": safety_daily},
        "frequency_result_file": frequency_result_file,
    }


def emit_frequency_log(log_callback, level: str, message: str) -> None:
    if callable(log_callback):
        log_callback(level, message)


def read_frequency_dispatch_rows(path: Path) -> list[dict]:
    ensure_split_result_workbook(path)
    dispatch_path = result_curves_workbook_path(path)
    source_path = dispatch_path if dispatch_path.exists() else path
    try:
        workbook = load_workbook(source_path, read_only=True, data_only=True)
    except RESULT_WORKBOOK_READ_ERRORS as exc:
        raise ValueError(f"结果文件无法读取: {source_path.name}") from exc
    try:
        rows = read_workbook_rows_with_field_map(workbook, "调度结果", limit=TIME_SERIES_IMPORT_ROW_COUNT)
    finally:
        workbook.close()
    if len(rows) != TIME_SERIES_IMPORT_ROW_COUNT:
        raise ValueError(f"频率计算需要调度结果包含{TIME_SERIES_IMPORT_ROW_COUNT}点，当前为{len(rows)}点")
    return rows


def frequency_curve_result_path(scheme: str, filename: str) -> Path:
    source_name = evaluation_result_filename_from_name(Path(str(filename or "")).stem.replace("_results", "") or "frequency")
    source_stem = Path(source_name).stem.replace("_results", "")
    return PLANNING_STORE.scheme_dir(str(scheme or "未选择方案")) / f"{source_stem}_frequency_curves.xlsx"


def export_frequency_curve_workbook(
    scheme: str,
    filename: str,
    scheme_payload: dict,
    dispatch_rows: list[dict],
    frequency_results: list[dict] | None = None,
    log_callback=None,
) -> Path:
    result_path = frequency_curve_result_path(scheme, filename)
    result_path.parent.mkdir(parents=True, exist_ok=True)
    workbook = Workbook(write_only=True)
    summary_sheet = workbook.create_sheet("频率8760结果")
    curve_sheet = workbook.create_sheet("频率曲线")
    summary_headers = frequency_8760_result_headers()
    curve_headers = frequency_curve_headers()
    summary_sheet.append(summary_headers)
    curve_sheet.append(curve_headers)
    results = frequency_results if frequency_results is not None else iter_frequency_result_rows(scheme_payload, dispatch_rows)
    total_results = len(frequency_results) if isinstance(frequency_results, list) else len(dispatch_rows)
    for index, result in enumerate(results, start=1):
        summary_sheet.append([result["summary"].get(header, "") for header in summary_headers])
        for row in result["curves"]:
            curve_sheet.append([row.get(header, "") for header in curve_headers])
        if index == 1 or index % 1000 == 0 or index == total_results:
            emit_frequency_log(log_callback, "info", f"频率结果写入进度：{index}/{total_results}")
    tmp_path = result_path.with_name(f".{result_path.name}.tmp")
    try:
        emit_frequency_log(log_callback, "info", "保存频率曲线Excel文件")
        file_ops.save_workbook_with_retry(workbook, tmp_path, "频率曲线文件")
    finally:
        workbook.close()
    emit_frequency_log(log_callback, "info", "替换频率曲线结果文件")
    replace_result_workbook_with_retry(tmp_path, result_path)
    return result_path


def frequency_8760_result_headers() -> list[str]:
    return [
        "hour_index",
        "datetime",
        "grid_model",
        "diesel_capacity",
        "diesel_on",
        "diesel_power",
        "load",
        "renewable_power",
        "grid_storage_capacity",
        "storage_power",
        "storage_charge",
        "storage_discharge",
        "equivalent_inertia_m",
        "load_response_d",
        "equivalent_primary_frequency_k",
        "max_up_disturbance_mw",
        "max_down_disturbance_mw",
        "source_frequency_max_hz",
        "source_frequency_min_hz",
        "calculated_max_frequency_hz",
        "calculated_min_frequency_hz",
    ]


def frequency_curve_headers() -> list[str]:
    time_headers = [frequency_curve_point_header(index) for index in range(FREQUENCY_CURVE_POINT_COUNT)]
    return [
        "hour_index",
        "datetime",
        "curve_type",
        "grid_model",
        "diesel_capacity",
        "diesel_on",
        "diesel_power",
        "load",
        "renewable_power",
        "grid_storage_capacity",
        "storage_power",
        "storage_charge",
        "storage_discharge",
        "equivalent_inertia_m",
        "load_response_d",
        "equivalent_primary_frequency_k",
        "max_up_disturbance_mw",
        "max_down_disturbance_mw",
        "source_frequency_max_hz",
        "source_frequency_min_hz",
        "calculated_max_frequency_hz",
        "calculated_min_frequency_hz",
        "disturbance_mw",
        *time_headers,
    ]


def frequency_curve_point_header(index: int) -> str:
    return f"f_{index * FREQUENCY_CURVE_STEP_SECONDS:.2f}s"


def iter_frequency_curve_rows(scheme_payload: dict, dispatch_rows: list[dict]):
    for result in iter_frequency_result_rows(scheme_payload, dispatch_rows):
        yield from result["curves"]


def frequency_diesel_unit_capacity(scheme_payload: dict) -> float | None:
    try:
        diesel_devices = plan_optimizer.normalized_device_rows(scheme_payload).get("diesel_generators", [])
    except Exception:
        return None
    capacities = [float(device.get("power_upper") or device.get("capacity") or 0.0) for device in diesel_devices]
    capacities = [capacity for capacity in capacities if capacity > 0]
    return capacities[0] if len(capacities) == 1 else None


def frequency_grid_storage_capacity(scheme_payload: dict) -> float | None:
    try:
        storage_devices = plan_optimizer.normalized_device_rows(scheme_payload).get("storage_pcs", [])
    except Exception:
        return None
    capacity = sum(
        float(device.get("capacity") or 0.0) * float(device.get("quantity_upper") or 0.0)
        for device in storage_devices
        if device.get("is_grid_forming")
    )
    return capacity if capacity > 0 else None


def frequency_8760_display_rows_from_results(results: list[dict], scheme_payload: dict | None = None) -> list[dict]:
    summaries = [dict(result.get("summary", {})) for result in results if isinstance(result, dict)]
    return frequency_8760_display_rows_from_summaries(summaries, scheme_payload)


def frequency_8760_display_rows_from_summaries(summaries: list[dict], scheme_payload: dict | None = None) -> list[dict]:
    diesel_unit_capacity = frequency_diesel_unit_capacity(scheme_payload or {})
    rows: list[dict] = []
    for summary in summaries:
        diesel_capacity = frequency_number(summary.get("diesel_capacity"))
        if diesel_capacity is None and diesel_unit_capacity is not None:
            diesel_capacity = (frequency_number(summary.get("diesel_on")) or 0.0) * diesel_unit_capacity
        rows.append(
            {
                "小时": summary.get("hour_index", ""),
                "时间": summary.get("datetime", ""),
                "柴发开机容量": round(diesel_capacity, 4) if diesel_capacity is not None else "",
                "向上最大扰动": summary.get("max_up_disturbance_mw", ""),
                "向下最大扰动": summary.get("max_down_disturbance_mw", ""),
                "优化频率最大值": summary.get("source_frequency_max_hz", ""),
                "优化频率最小值": summary.get("source_frequency_min_hz", ""),
                "仿真频率最大值": summary.get("calculated_max_frequency_hz", ""),
                "仿真频率最小值": summary.get("calculated_min_frequency_hz", ""),
            }
        )
    return rows


def read_frequency_8760_display_rows(path: Path, scheme_payload: dict | None = None) -> list[dict]:
    try:
        workbook = load_workbook(path, read_only=True, data_only=True)
    except RESULT_WORKBOOK_READ_ERRORS:
        return []
    try:
        if "频率8760结果" not in workbook.sheetnames:
            return []
        sheet = workbook["频率8760结果"]
        raw_rows = list(sheet.iter_rows(values_only=True))
    finally:
        workbook.close()
    if not raw_rows:
        return []
    headers = [str(value or "").strip() for value in raw_rows[0]]
    summaries = [
        {headers[index]: value for index, value in enumerate(row) if index < len(headers) and headers[index]}
        for row in raw_rows[1:]
    ]
    return frequency_8760_display_rows_from_summaries(summaries, scheme_payload)


def read_frequency_time_curve_payload(
    path: Path,
    *,
    hour_index: str = "",
    month: str = "",
    day: str = "",
    hour: str = "",
) -> dict:
    if not path.exists():
        raise FileNotFoundError(f"频率曲线文件不存在: {path.name}")
    try:
        workbook = load_workbook(path, read_only=True, data_only=True)
    except RESULT_WORKBOOK_READ_ERRORS as exc:
        raise ValueError(f"频率曲线文件无法读取: {path.name}") from exc
    try:
        if "频率8760结果" not in workbook.sheetnames or "频率曲线" not in workbook.sheetnames:
            raise ValueError("频率曲线文件缺少频率8760结果或频率曲线工作表")
        summaries = worksheet_dict_rows(workbook["频率8760结果"])
        target_summary = select_frequency_time_summary(summaries, hour_index=hour_index, month=month, day=day, hour=hour)
        target_hour = int(frequency_number(target_summary.get("hour_index")) or 0)
        if target_hour <= 0:
            raise ValueError("频率曲线文件中的小时序号无效")
        curve_rows = [
            row
            for row in worksheet_dict_rows(workbook["频率曲线"])
            if int(frequency_number(row.get("hour_index")) or 0) == target_hour
        ]
    finally:
        workbook.close()
    if not curve_rows:
        raise ValueError(f"未找到第{target_hour}小时的频率曲线")
    high_curve = next((row for row in curve_rows if str(row.get("curve_type", "")).strip() == "最高频率曲线"), None)
    low_curve = next((row for row in curve_rows if str(row.get("curve_type", "")).strip() == "最低频率曲线"), None)
    if high_curve is None or low_curve is None:
        raise ValueError(f"第{target_hour}小时的高频曲线或低频曲线缺失")
    selection = frequency_time_selection_payload(target_summary)
    return {
        "selection": selection,
        "summary_table": frequency_time_summary_table(target_summary),
        "curves": {
            "high": frequency_curve_points_from_row(high_curve),
            "low": frequency_curve_points_from_row(low_curve),
        },
    }


def worksheet_dict_rows(sheet) -> list[dict]:
    raw_rows = list(sheet.iter_rows(values_only=True))
    if not raw_rows:
        return []
    headers = [str(value or "").strip() for value in raw_rows[0]]
    return [
        {headers[index]: value for index, value in enumerate(row) if index < len(headers) and headers[index]}
        for row in raw_rows[1:]
    ]


def select_frequency_time_summary(
    summaries: list[dict],
    *,
    hour_index: str = "",
    month: str = "",
    day: str = "",
    hour: str = "",
) -> dict:
    if not summaries:
        raise ValueError("频率8760结果为空")
    requested_hour = frequency_int_or_none(hour_index)
    if requested_hour is not None:
        for summary in summaries:
            if int(frequency_number(summary.get("hour_index")) or 0) == requested_hour:
                return summary
        raise ValueError(f"未找到第{requested_hour}小时的频率结果")
    requested_parts = {
        "month": frequency_int_or_none(month),
        "day": frequency_int_or_none(day),
        "hour": frequency_int_or_none(hour),
    }
    if all(value is not None for value in requested_parts.values()):
        for summary in summaries:
            parts = frequency_datetime_parts(summary.get("datetime"))
            if parts and all(parts[key] == requested_parts[key] for key in requested_parts):
                return summary
        raise ValueError(f"未找到{month}月{day}日{hour}时的频率结果")
    return summaries[0]


def frequency_int_or_none(value) -> int | None:
    text = str(value or "").strip()
    if not text:
        return None
    try:
        return int(float(text))
    except (TypeError, ValueError):
        return None


def frequency_datetime_parts(value) -> dict[str, int] | None:
    text = str(value or "").strip()
    if not text:
        return None
    for fmt in ("%Y-%m-%d %H:%M:%S", "%Y-%m-%d %H:%M", "%Y/%m/%d %H:%M:%S", "%Y/%m/%d %H:%M"):
        try:
            dt = datetime.strptime(text, fmt)
            return {"year": dt.year, "month": dt.month, "day": dt.day, "hour": dt.hour}
        except ValueError:
            continue
    match = re.search(r"(\d{4})[-/](\d{1,2})[-/](\d{1,2})\s+(\d{1,2})", text)
    if not match:
        return None
    return {
        "year": int(match.group(1)),
        "month": int(match.group(2)),
        "day": int(match.group(3)),
        "hour": int(match.group(4)),
    }


def frequency_time_selection_payload(summary: dict) -> dict:
    parts = frequency_datetime_parts(summary.get("datetime")) or {}
    return {
        "hour_index": summary.get("hour_index", ""),
        "datetime": summary.get("datetime", ""),
        "year": parts.get("year", ""),
        "month": parts.get("month", ""),
        "day": parts.get("day", ""),
        "hour": parts.get("hour", ""),
    }


def frequency_time_summary_table(summary: dict) -> list[dict]:
    rows = [
        ("柴发开机总容量", "diesel_capacity", "kW"),
        ("柴发总功率", "diesel_power", "kW"),
        ("负荷总功率", "load", "kW"),
        ("新能源总出力", "renewable_power", "kW"),
        ("向上最大扰动", "max_up_disturbance_mw", "MW"),
        ("向下最大扰动", "max_down_disturbance_mw", "MW"),
        ("优化频率最大值", "source_frequency_max_hz", "Hz"),
        ("优化频率最小值", "source_frequency_min_hz", "Hz"),
        ("仿真频率最大值", "calculated_max_frequency_hz", "Hz"),
        ("仿真频率最小值", "calculated_min_frequency_hz", "Hz"),
    ]
    return [
        {"指标": label, "数值": frequency_summary_value(summary, key), "单位": unit}
        for label, key, unit in rows
    ]


def frequency_summary_value(summary: dict, key: str):
    value = summary.get(key, "")
    numeric = frequency_number(value)
    return round(numeric, 5) if numeric is not None else value


def frequency_curve_points_from_row(row: dict) -> list[dict]:
    points: list[dict] = []
    for index in range(FREQUENCY_CURVE_POINT_COUNT):
        header = frequency_curve_point_header(index)
        value = frequency_number(row.get(header))
        if value is not None:
            points.append({"time": round(index * FREQUENCY_CURVE_STEP_SECONDS, 2), "frequency": round(value, 5)})
    return points


def iter_frequency_result_rows(scheme_payload: dict, dispatch_rows: list[dict]):
    planning_parameters = frequency_planning_parameters(scheme_payload)
    diesel_unit_capacity = frequency_diesel_unit_capacity(scheme_payload)
    default_grid_storage_capacity = frequency_grid_storage_capacity(scheme_payload)
    loads = [frequency_number(row.get("load")) or 0.0 for row in dispatch_rows]
    nominal_frequency = min(65.0, max(45.0, plan_optimizer.numeric(planning_parameters.get("nominal_frequency_hz"), plan_optimizer.NOMINAL_FREQUENCY_HZ)))
    governor_t = max(0.0, plan_optimizer.numeric(planning_parameters.get("frequency_governor_time_constant_s"), 0.6))
    if governor_t <= plan_optimizer.FREQUENCY_EPS:
        governor_t = 0.6
    load_ref = max(loads) if loads else 0.0
    configured_ref = plan_optimizer.numeric(planning_parameters.get("network_synchronization_reference_load_kw"), 0.0)
    if configured_ref > 0:
        load_ref = configured_ref
    load_ref = max(load_ref, plan_optimizer.FREQUENCY_EPS)
    k_base = plan_optimizer.numeric(planning_parameters.get("network_synchronization_coefficient_base"), 1.0)
    k_slope = plan_optimizer.numeric(planning_parameters.get("network_synchronization_coefficient_slope"), 0.0)
    for row_index, row in enumerate(dispatch_rows, start=1):
        hour_index = int(frequency_number(row.get("hour_index")) or row_index)
        load = frequency_number(row.get("load")) or 0.0
        wind_power = frequency_number(row.get("wind_power")) or 0.0
        pv_power = frequency_number(row.get("pv_power")) or 0.0
        net_ratio = (wind_power + pv_power - load) / load_ref
        grid_model = k_base + k_slope * net_ratio
        m_eq = frequency_number(row.get("equivalent_inertia_m")) or 0.0
        k_eq = frequency_number(row.get("equivalent_primary_frequency_k")) or 0.0
        d_eq = frequency_number(row.get("equivalent_damping_d")) or 0.0
        diesel_on = row.get("diesel_on", "")
        diesel_capacity = frequency_number(row.get("diesel_capacity"))
        if diesel_capacity is None and diesel_unit_capacity is not None:
            diesel_capacity = (frequency_number(diesel_on) or 0.0) * diesel_unit_capacity
        renewable_power = frequency_number(row.get("renewable_power"))
        if renewable_power is None:
            renewable_power = (frequency_number(row.get("wind_power")) or 0.0) + (frequency_number(row.get("pv_power")) or 0.0)
        grid_storage_capacity = frequency_number(row.get("grid_storage_capacity"))
        if grid_storage_capacity is None:
            grid_storage_capacity = default_grid_storage_capacity
        base = {
            "hour_index": hour_index,
            "datetime": row.get("datetime", ""),
            "grid_model": round(grid_model, 4),
            "diesel_capacity": round(diesel_capacity, 4) if diesel_capacity is not None else "",
            "diesel_on": diesel_on,
            "diesel_power": round(frequency_number(row.get("diesel_power")) or 0.0, 4),
            "load": round(load, 4),
            "renewable_power": round(renewable_power, 4) if renewable_power is not None else "",
            "grid_storage_capacity": round(grid_storage_capacity, 4) if grid_storage_capacity is not None else "",
            "storage_power": round(frequency_number(row.get("storage_power")) or 0.0, 4),
            "storage_charge": round(frequency_number(row.get("storage_charge")) or 0.0, 4),
            "storage_discharge": round(frequency_number(row.get("storage_discharge")) or 0.0, 4),
            "equivalent_inertia_m": round(m_eq, 4),
            "load_response_d": round(d_eq, 4),
            "equivalent_primary_frequency_k": round(k_eq, 4),
        }
        lower_disturbance = frequency_number(row.get("frequency_delta_p_mw")) or 0.0
        upper_disturbance = frequency_number(row.get("frequency_upper_delta_p_mw")) or 0.0
        lower_points = frequency_response_curve_points(m_eq, k_eq, d_eq, governor_t, lower_disturbance, nominal_frequency)
        upper_points = frequency_response_curve_points(m_eq, k_eq, d_eq, governor_t, upper_disturbance, nominal_frequency)
        calculated_min_frequency = round(min(lower_points), 5) if lower_points else ""
        calculated_max_frequency = round(max(upper_points), 5) if upper_points else ""
        source_frequency_max = frequency_number(row.get("frequency_max"))
        source_frequency_min = frequency_number(row.get("frequency_min"))
        summary = dict(base)
        summary.update(
            {
                "max_up_disturbance_mw": round(lower_disturbance, 4),
                "max_down_disturbance_mw": round(upper_disturbance, 4),
                "source_frequency_max_hz": round(source_frequency_max, 5) if source_frequency_max is not None else "",
                "source_frequency_min_hz": round(source_frequency_min, 5) if source_frequency_min is not None else "",
                "calculated_max_frequency_hz": calculated_max_frequency,
                "calculated_min_frequency_hz": calculated_min_frequency,
            }
        )
        yield {
            "summary": summary,
            "curves": [
                frequency_curve_output_row(
                    summary,
                    "最低频率曲线",
                    lower_disturbance,
                    lower_points,
                ),
                frequency_curve_output_row(
                    summary,
                    "最高频率曲线",
                    upper_disturbance,
                    upper_points,
                ),
            ],
        }


def frequency_planning_parameters(scheme_payload: dict) -> dict:
    rows = scheme_payload.get("planning_parameters") if isinstance(scheme_payload, dict) else {}
    if isinstance(rows, dict):
        return rows
    if isinstance(rows, list) and rows and isinstance(rows[0], dict):
        return rows[0]
    return {}


def frequency_curve_output_row(base: dict, curve_type: str, disturbance_mw: float, points: list[float]) -> dict:
    row = dict(base)
    row["curve_type"] = curve_type
    row["disturbance_mw"] = round(disturbance_mw, 4)
    for index, value in enumerate(points):
        row[frequency_curve_point_header(index)] = round(value, 5)
    return row


def frequency_response_curve_points(
    m_eq: float,
    k_eq: float,
    d_eq: float,
    t_d: float,
    delta_p_mw: float,
    nominal_frequency_hz: float,
) -> list[float]:
    times = [index * FREQUENCY_CURVE_STEP_SECONDS for index in range(FREQUENCY_CURVE_POINT_COUNT)]
    if m_eq <= plan_optimizer.FREQUENCY_EPS or t_d <= plan_optimizer.FREQUENCY_EPS:
        return [round(float(nominal_frequency_hz), 5) for _ in times]
    denom = d_eq + k_eq / (2.0 * math.pi)
    if denom <= plan_optimizer.FREQUENCY_EPS:
        return [round(float(nominal_frequency_hz), 5) for _ in times]
    omega_ss = -delta_p_mw / denom
    alpha, beta = plan_optimizer.second_order_coefficients(m_eq, k_eq, d_eq, t_d)
    if beta <= plan_optimizer.FREQUENCY_EPS:
        return [round(float(nominal_frequency_hz), 5) for _ in times]
    disc = alpha**2 - 4.0 * beta
    omega_values: list[float] = []
    if disc < -plan_optimizer.FREQUENCY_EPS:
        omega_n = math.sqrt(beta)
        zeta = alpha / (2.0 * omega_n)
        omega_d_sq = beta * (1.0 - zeta**2)
        if omega_d_sq <= plan_optimizer.FREQUENCY_EPS:
            return [round(float(nominal_frequency_hz), 5) for _ in times]
        omega_d = math.sqrt(omega_d_sq)
        a = -omega_ss
        b = (-delta_p_mw / m_eq + zeta * omega_n * a) / omega_d
        omega_values = [
            omega_ss
            + math.exp(-zeta * omega_n * t) * (a * math.cos(omega_d * t) + b * math.sin(omega_d * t))
            for t in times
        ]
    elif disc > plan_optimizer.FREQUENCY_EPS:
        sqrt_disc = math.sqrt(disc)
        lam1 = -0.5 * alpha + 0.5 * sqrt_disc
        lam2 = -0.5 * alpha - 0.5 * sqrt_disc
        denom_lam = lam1 - lam2
        if abs(denom_lam) <= plan_optimizer.FREQUENCY_EPS:
            return [round(float(nominal_frequency_hz), 5) for _ in times]
        c1 = (-delta_p_mw / m_eq + omega_ss * lam2) / denom_lam
        c2 = -omega_ss - c1
        omega_values = [omega_ss + c1 * math.exp(lam1 * t) + c2 * math.exp(lam2 * t) for t in times]
    else:
        lam = -0.5 * alpha
        c1 = -omega_ss
        c2 = -delta_p_mw / m_eq + lam * omega_ss
        omega_values = [omega_ss + (c1 + c2 * t) * math.exp(lam * t) for t in times]
    return [
        round(float(nominal_frequency_hz) + (value if math.isfinite(value) else 0.0) / (2.0 * math.pi), 5)
        for value in omega_values
    ]


def frequency_metrics_from_rows(safety_table: list[dict], safety_daily: list[dict]) -> list[dict]:
    values_by_label = {str(row.get("指标", "")).strip(): row for row in safety_table}
    highest = frequency_number(values_by_label.get("最高频率", {}).get("数值"))
    lowest = frequency_number(values_by_label.get("最低频率", {}).get("数值"))
    risk_hours = frequency_number(values_by_label.get("频率安全风险小时数", {}).get("数值"))
    if highest is None:
        highest_values = [frequency_number(row.get("frequency_max")) for row in safety_daily]
        highest_values = [value for value in highest_values if value is not None]
        highest = max(highest_values) if highest_values else None
    if lowest is None:
        lowest_values = [frequency_number(row.get("frequency_min")) for row in safety_daily]
        lowest_values = [value for value in lowest_values if value is not None]
        lowest = min(lowest_values) if lowest_values else None
    if risk_hours is None and highest is not None and lowest is not None:
        risk_hours = sum(
            1
            for row in safety_daily
            if (frequency_number(row.get("frequency_max")) or 50.0) > 50.5
            or (frequency_number(row.get("frequency_min")) or 50.0) < 49.5
        )
    return [
        {"label": "最低频率", "value": round(lowest, 3) if lowest is not None else "-", "unit": "Hz"},
        {"label": "最高频率", "value": round(highest, 3) if highest is not None else "-", "unit": "Hz"},
        {"label": "频率安全风险小时数", "value": int(risk_hours) if risk_hours is not None else "-", "unit": "h"},
    ]


def frequency_summary_rows(metrics: list[dict]) -> list[dict]:
    return [
        {"指标": metric.get("label", ""), "数值": metric.get("value", ""), "单位": metric.get("unit", "")}
        for metric in metrics
    ]


def frequency_detail_rows(safety_table: list[dict], metrics: list[dict]) -> list[dict]:
    selected_labels = {
        "最高频率",
        "最低频率",
        "频率安全风险小时数",
        "向上扰动最大量",
        "向下扰动最大量",
        "频率下限裕度",
        "频率上限裕度",
        "初始频率变化率",
        "上限场景初始频率变化率",
        "等效惯量M",
        "等效调频系数K",
    }
    rows = [row for row in safety_table if str(row.get("指标", "")).strip() in selected_labels]
    existing = {str(row.get("指标", "")).strip() for row in rows}
    for metric in metrics:
        label = str(metric.get("label", "")).strip()
        if label and label not in existing:
            rows.append({"指标": label, "数值": metric.get("value", ""), "单位": metric.get("unit", "")})
    return rows


def frequency_number(value: object) -> float | None:
    if value is None:
        return None
    try:
        number = float(str(value).strip().replace("%", ""))
    except (TypeError, ValueError):
        return None
    if not math.isfinite(number):
        return None
    return number


class TaskScheduler:
    """Serial queue for calculation tasks that should wait for existing jobs."""

    def __init__(self) -> None:
        self._queue: list[dict[str, str]] = []
        self._lock = threading.Lock()

    def enqueue(self, task_type_key: str, scheme: str, result: str = "") -> dict:
        item = normalized_task_item(task_type_key, scheme, result)
        with self._lock:
            if not any(same_task_item(item, existing) for existing in self._queue):
                self._queue.append(item)
            return dict(item)

    def remove(self, task_type_key: str, scheme: str, result: str = "") -> None:
        item = normalized_task_item(task_type_key, scheme, result)
        with self._lock:
            self._queue = [existing for existing in self._queue if not same_task_item(item, existing)]

    def remove_running_or_finished(self) -> None:
        with self._lock:
            self._queue = [item for item in self._queue if not is_task_running_or_finished(item)]

    def queued_items(self) -> list[dict[str, str]]:
        with self._lock:
            return [dict(item) for item in self._queue]

    def is_queued(self, task_type_key: str, scheme: str, result: str = "") -> bool:
        item = normalized_task_item(task_type_key, scheme, result)
        with self._lock:
            return any(same_task_item(item, existing) for existing in self._queue)

    def queue_position(self, task_type_key: str, scheme: str, result: str = "") -> int:
        item = normalized_task_item(task_type_key, scheme, result)
        with self._lock:
            for index, existing in enumerate(self._queue, start=1):
                if same_task_item(item, existing):
                    return index
        return 0

    def schedule_next_if_idle(self) -> None:
        with self._lock:
            if any_calculation_running_unlocked():
                return
            while self._queue:
                item = self._queue.pop(0)
                try:
                    start_task_item_unlocked(item)
                    return
                except Exception:
                    continue


def normalized_task_item(task_type_key: str, scheme: str, result: str = "") -> dict[str, str]:
    normalized_type = normalize_task_type_key(task_type_key)
    normalized_result = str(result or "").strip()
    if normalized_type == "optimization":
        normalized_result = OPTIMIZATION_RESULT_WORKBOOK_NAME
    return {
        "task_type_key": normalized_type,
        "scheme": str(scheme or "未选择方案").strip() or "未选择方案",
        "result": normalized_result,
    }


def same_task_item(left: dict, right: dict) -> bool:
    return (
        left.get("task_type_key") == right.get("task_type_key")
        and left.get("scheme") == right.get("scheme")
        and (left.get("result") or "") == (right.get("result") or "")
    )


def is_task_running_or_finished(item: dict) -> bool:
    runtime = runtime_for_task_item(item)
    return bool(runtime and runtime.status in {"运行中", "已完成"})


def runtime_for_task_item(item: dict):
    task_type_key = item.get("task_type_key")
    scheme = item.get("scheme", "")
    result = item.get("result", "")
    if task_type_key == "optimization":
        return OPTIMIZATION_RUNTIME.runtimes().get(scheme)
    if task_type_key == "evaluation":
        return EVALUATION_RUNTIME.runtimes().get(f"{scheme}\0{result}")
    if task_type_key == "frequency":
        return FREQUENCY_EVALUATION_RUNTIME.runtimes().get(f"{scheme}\0{result}")
    return None


def any_calculation_running_unlocked() -> bool:
    return any(runtime.status == "运行中" for runtime in OPTIMIZATION_RUNTIME.runtimes().values()) or any(
        runtime.status == "运行中" for runtime in EVALUATION_RUNTIME.runtimes().values()
    ) or any(
        runtime.status == "运行中" for runtime in FREQUENCY_EVALUATION_RUNTIME.runtimes().values()
    )


def start_task_item_unlocked(item: dict) -> dict:
    task_type_key = item.get("task_type_key")
    if task_type_key == "optimization":
        return OPTIMIZATION_RUNTIME.apply("start", scheme=item.get("scheme", ""))
    if task_type_key == "evaluation":
        return EVALUATION_RUNTIME.apply("start", scheme=item.get("scheme", ""), filename=item.get("result", ""))
    if task_type_key == "frequency":
        return FREQUENCY_EVALUATION_RUNTIME.apply("start", scheme=item.get("scheme", ""), filename=item.get("result", ""))
    raise ValueError("任务类型必须为规划计算、方案评估或频率计算")


def build_task_list(schedule: bool = True, current_user: dict | None = None) -> list[dict]:
    TASK_SCHEDULER.remove_running_or_finished()
    if schedule:
        TASK_SCHEDULER.schedule_next_if_idle()
    tasks: dict[str, dict] = {}
    for scheme_item in safe_list_schemes_for_tasks(current_user):
        scheme = str(scheme_item.get("name") or "").strip()
        if not scheme:
            continue
        runtime = OPTIMIZATION_RUNTIME.runtimes().get(scheme)
        state = runtime.task_snapshot() if runtime else default_task_runtime_state(scheme)
        task = task_from_runtime_state(
            "optimization",
            state,
            scheme=scheme,
            result=OPTIMIZATION_RESULT_WORKBOOK_NAME,
            queued=TASK_SCHEDULER.is_queued("optimization", scheme, OPTIMIZATION_RESULT_WORKBOOK_NAME),
            queue_position=TASK_SCHEDULER.queue_position("optimization", scheme, OPTIMIZATION_RESULT_WORKBOOK_NAME),
        )
        tasks[task["id"]] = task
        for result_item in safe_list_results_for_tasks(scheme):
            result_name = str(result_item.get("name") or "").strip()
            if not result_name:
                continue
            key = f"{scheme}\0{result_name}"
            if task_list_evaluation_result_is_eligible(result_item):
                eval_runtime = EVALUATION_RUNTIME.runtimes().get(key)
                eval_state = (
                    eval_runtime.task_snapshot()
                    if eval_runtime
                    else default_task_runtime_state(scheme, result_filename=result_name)
                )
                eval_task = task_from_runtime_state(
                    "evaluation",
                    eval_state,
                    scheme=scheme,
                    result=result_name,
                    queued=TASK_SCHEDULER.is_queued("evaluation", scheme, result_name),
                    queue_position=TASK_SCHEDULER.queue_position("evaluation", scheme, result_name),
                )
                tasks[eval_task["id"]] = eval_task
            if task_list_frequency_result_is_eligible(result_item):
                freq_runtime = FREQUENCY_EVALUATION_RUNTIME.runtimes().get(key)
                freq_state = (
                    freq_runtime.task_snapshot()
                    if freq_runtime
                    else default_task_runtime_state(scheme, result_filename=result_name)
                )
                freq_task = task_from_runtime_state(
                    "frequency",
                    freq_state,
                    scheme=scheme,
                    result=result_name,
                    queued=TASK_SCHEDULER.is_queued("frequency", scheme, result_name),
                    queue_position=TASK_SCHEDULER.queue_position("frequency", scheme, result_name),
                )
                tasks[freq_task["id"]] = freq_task

    for scheme, runtime in OPTIMIZATION_RUNTIME.runtimes().items():
        if not user_can_manage_planning_scheme(scheme, current_user):
            continue
        state = runtime.task_snapshot()
        task = task_from_runtime_state(
            "optimization",
            state,
            scheme=scheme,
            result=OPTIMIZATION_RESULT_WORKBOOK_NAME,
            queued=TASK_SCHEDULER.is_queued("optimization", scheme, OPTIMIZATION_RESULT_WORKBOOK_NAME),
            queue_position=TASK_SCHEDULER.queue_position("optimization", scheme, OPTIMIZATION_RESULT_WORKBOOK_NAME),
        )
        tasks[task["id"]] = task

    for key, runtime in EVALUATION_RUNTIME.runtimes().items():
        scheme, result = split_evaluation_runtime_key(key)
        if not user_can_manage_planning_scheme(scheme, current_user):
            continue
        queued = TASK_SCHEDULER.is_queued("evaluation", scheme, result)
        if result == OPTIMIZATION_RESULT_WORKBOOK_NAME and runtime.status != "运行中" and not queued:
            continue
        state = runtime.task_snapshot()
        task = task_from_runtime_state(
            "evaluation",
            state,
            scheme=scheme,
            result=result,
            queued=queued,
            queue_position=TASK_SCHEDULER.queue_position("evaluation", scheme, result),
        )
        tasks[task["id"]] = task

    for key, runtime in FREQUENCY_EVALUATION_RUNTIME.runtimes().items():
        scheme, result = split_evaluation_runtime_key(key)
        if not user_can_manage_planning_scheme(scheme, current_user):
            continue
        queued = TASK_SCHEDULER.is_queued("frequency", scheme, result)
        state = runtime.task_snapshot()
        task = task_from_runtime_state(
            "frequency",
            state,
            scheme=scheme,
            result=result,
            queued=queued,
            queue_position=TASK_SCHEDULER.queue_position("frequency", scheme, result),
        )
        tasks[task["id"]] = task

    return sorted((task for task in tasks.values() if task_list_item_is_visible(task)), key=task_sort_key)


def task_list_evaluation_result_is_eligible(result_item: dict) -> bool:
    result_name = str(result_item.get("name") or "").strip()
    return bool(result_name) and result_name != OPTIMIZATION_RESULT_WORKBOOK_NAME and bool(result_item.get("readable", True))


def task_list_frequency_result_is_eligible(result_item: dict) -> bool:
    result_name = str(result_item.get("name") or "").strip()
    return bool(result_name) and bool(result_item.get("readable", True))


def task_list_item_is_visible(task: dict) -> bool:
    """Hide rows that business rules made completely non-operable."""
    return bool(task.get("queued") or task.get("can_start") or task.get("can_queue") or task.get("can_stop"))


def task_sort_key(item: dict) -> tuple[int, str, str]:
    type_ranks = {"optimization": 0, "evaluation": 1, "frequency": 2}
    type_rank = type_ranks.get(item.get("task_type_key"), 9)
    return type_rank, str(item.get("scheme") or ""), str(item.get("result") or "")


def safe_list_schemes_for_tasks(current_user: dict | None = None) -> list[dict]:
    try:
        return PLANNING_STORE.list_schemes(owner_username=scheme_owner_filter_for_user(current_user))
    except Exception:
        return []


def safe_list_results_for_tasks(scheme: str) -> list[dict]:
    try:
        return list_evaluation_result_files_for_tasks(scheme)
    except Exception:
        return []


def split_evaluation_runtime_key(key: str) -> tuple[str, str]:
    if "\0" in key:
        scheme, result = key.split("\0", 1)
        return scheme, result
    return key, ""


def default_task_runtime_state(scheme: str, result_filename: str = "") -> dict:
    return {
        "status": "待启动",
        "scheme": scheme,
        "result_filename": result_filename,
        "start_time": "",
        "end_time": "",
        "process_id": "",
        "elapsed_seconds": 0,
        "logs": [],
    }


def task_from_runtime_state(
    task_type_key: str,
    state: dict,
    scheme: str,
    result: str = "",
    queued: bool = False,
    queue_position: int = 0,
) -> dict:
    runtime_status = str(state.get("status") or "待启动")
    normalized_result = result or str(state.get("result_filename") or "") or OPTIMIZATION_RESULT_WORKBOOK_NAME
    latest_log = latest_log_message(state.get("logs"))
    display_status = task_display_status(runtime_status, queued)
    process_id = state.get("process_id") or ""
    if isinstance(process_id, str) and process_id.isdigit():
        process_id = int(process_id)
    task_type_labels = {"optimization": "规划计算", "evaluation": "方案评估", "frequency": "频率计算"}
    task_type = task_type_labels.get(task_type_key, "未知任务")
    task_id = f"{task_type_key}::{scheme}::{normalized_result}"
    return {
        "id": task_id,
        "task_key": task_id,
        "task_type": task_type,
        "task_type_key": task_type_key,
        "scheme": scheme,
        "result": normalized_result,
        "status": display_status,
        "runtime_status": runtime_status,
        "queued": bool(queued),
        "queue_position": int(queue_position or 0),
        "process_id": process_id,
        "start_time": state.get("start_time") or "",
        "end_time": state.get("end_time") or "",
        "elapsed_seconds": int(state.get("elapsed_seconds") or elapsed_seconds_from_times(state.get("start_time", ""), state.get("end_time", ""))),
        "latest_log": latest_log,
        "can_start": runtime_status != "运行中",
        "can_queue": runtime_status != "运行中" and not queued,
        "can_stop": runtime_status == "运行中",
    }


def task_display_status(runtime_status: str, queued: bool = False) -> str:
    if runtime_status == "运行中":
        return "计算中"
    if queued:
        return "排队中"
    if runtime_status == "退出队列":
        return "退出队列"
    if runtime_status == "已完成":
        return "完成计算"
    if runtime_status == "计算中止":
        return "计算中止"
    if runtime_status == "超时":
        return "计算超时"
    if runtime_status == "失败":
        return "计算失败"
    return "未计算"


def latest_log_message(logs: object) -> str:
    if not isinstance(logs, list) or not logs:
        return ""
    for item in reversed(logs):
        if isinstance(item, dict) and item.get("message"):
            return str(item.get("message"))
    return ""


def build_task_control_response(
    action: str,
    task_type: str,
    scheme: str,
    result: str = "",
    current_user: dict | None = None,
) -> dict:
    task_type_key = normalize_task_type_key(task_type)
    normalized_action = normalize_task_action(action)
    if normalized_action == "cancel_queue":
        item = normalized_task_item(task_type_key, scheme, result)
        ensure_planning_scheme_manage_access(item["scheme"], current_user)
        TASK_SCHEDULER.remove(task_type_key, scheme, result)
        if task_type_key == "optimization":
            state = OPTIMIZATION_RUNTIME.apply("cancel_queue", scheme=item["scheme"])
        elif task_type_key == "evaluation":
            state = EVALUATION_RUNTIME.apply("cancel_queue", scheme=item["scheme"], filename=item["result"])
        elif task_type_key == "frequency":
            state = FREQUENCY_EVALUATION_RUNTIME.apply("cancel_queue", scheme=item["scheme"], filename=item["result"])
        else:
            state = default_task_runtime_state(item["scheme"], item["result"])
        tasks = build_task_list(current_user=current_user)
        task = task_from_runtime_state(task_type_key, state, scheme=item["scheme"], result=item["result"])
        return {"ok": True, "task": task, "tasks": tasks}
    if normalized_action == "queue":
        item = normalized_task_item(task_type_key, scheme, result)
        ensure_planning_scheme_manage_access(item["scheme"], current_user)
        item = TASK_SCHEDULER.enqueue(item["task_type_key"], item["scheme"], item["result"])
        tasks = build_task_list(schedule=False, current_user=current_user)
        task = find_task_in_list(tasks, item)
        return {"ok": True, "task": task, "tasks": tasks}
    if task_type_key == "optimization":
        ensure_planning_scheme_manage_access(scheme, current_user)
        TASK_SCHEDULER.remove(task_type_key, scheme, result)
        state = OPTIMIZATION_RUNTIME.apply(normalized_action, scheme=scheme)
        task = task_from_runtime_state("optimization", state, scheme=state.get("scheme") or scheme, result=OPTIMIZATION_RESULT_WORKBOOK_NAME)
        tasks = build_task_list(current_user=current_user)
        return {"ok": True, "task": task_from_list_or_default(tasks, task), "tasks": tasks}
    if task_type_key == "evaluation":
        ensure_planning_scheme_manage_access(scheme, current_user)
        TASK_SCHEDULER.remove(task_type_key, scheme, result)
        state = EVALUATION_RUNTIME.apply(normalized_action, scheme=scheme, filename=result)
        task = task_from_runtime_state(
            "evaluation",
            state,
            scheme=state.get("scheme") or scheme,
            result=state.get("result_filename") or result,
        )
        tasks = build_task_list(current_user=current_user)
        return {"ok": True, "task": task_from_list_or_default(tasks, task), "tasks": tasks}
    if task_type_key == "frequency":
        ensure_planning_scheme_manage_access(scheme, current_user)
        TASK_SCHEDULER.remove(task_type_key, scheme, result)
        state = FREQUENCY_EVALUATION_RUNTIME.apply(normalized_action, scheme=scheme, filename=result)
        task = task_from_runtime_state(
            "frequency",
            state,
            scheme=state.get("scheme") or scheme,
            result=state.get("result_filename") or result,
        )
        tasks = build_task_list(current_user=current_user)
        return {"ok": True, "task": task_from_list_or_default(tasks, task), "tasks": tasks}
    raise ValueError("任务类型必须为规划计算、方案评估或频率计算")


def normalize_task_action(action: str) -> str:
    text = str(action or "").strip().lower()
    if text in {"start", "start_now", "immediate", "run", "立刻启动"}:
        return "start"
    if text in {"queue", "enqueue", "排队", "加入排队"}:
        return "queue"
    if text in {"cancel_queue", "dequeue", "remove_queue", "取消排队", "移出队列", "退出队列", "退队", "离队"}:
        return "cancel_queue"
    if text in {"stop", "停止", "停止计算"}:
        return "stop"
    return text


def find_task_in_list(tasks: list[dict], item: dict, queued: bool = True) -> dict:
    for task in tasks:
        if (
            task.get("task_type_key") == item.get("task_type_key")
            and task.get("scheme") == item.get("scheme")
            and (task.get("result") or "") == (item.get("result") or "")
        ):
            return task
    return task_from_runtime_state(
        item.get("task_type_key", ""),
        default_task_runtime_state(item.get("scheme", ""), item.get("result", "")),
        scheme=item.get("scheme", ""),
        result=item.get("result", ""),
        queued=queued,
        queue_position=TASK_SCHEDULER.queue_position(item.get("task_type_key", ""), item.get("scheme", ""), item.get("result", "")),
    )


def task_from_list_or_default(tasks: list[dict], fallback: dict) -> dict:
    item = {"task_type_key": fallback.get("task_type_key"), "scheme": fallback.get("scheme"), "result": fallback.get("result")}
    for task in tasks:
        if same_task_item(item, task):
            return task
    return fallback


def task_control_state_for_item(task_type_key: str, scheme: str, result: str = "", state: dict | None = None) -> dict:
    """Return the task scheduler view used by business pages and the task page."""
    item = normalized_task_item(task_type_key, scheme, result)
    queued = TASK_SCHEDULER.is_queued(item["task_type_key"], item["scheme"], item["result"])
    queue_position = TASK_SCHEDULER.queue_position(item["task_type_key"], item["scheme"], item["result"])
    task = task_from_runtime_state(
        item["task_type_key"],
        state or default_task_runtime_state(item["scheme"], item["result"]),
        scheme=item["scheme"],
        result=item["result"],
        queued=queued,
        queue_position=queue_position,
    )
    if queued:
        task["can_queue"] = False
    if item["task_type_key"] == "evaluation":
        task["can_start"] = task["can_start"] and bool(item["result"]) and item["result"] != OPTIMIZATION_RESULT_WORKBOOK_NAME
        task["can_queue"] = task["can_queue"] and bool(item["result"]) and item["result"] != OPTIMIZATION_RESULT_WORKBOOK_NAME
    if item["task_type_key"] == "frequency":
        task["can_start"] = task["can_start"] and bool(item["result"])
        task["can_queue"] = task["can_queue"] and bool(item["result"])
    return task


def append_task_control_state(payload: dict, task_type_key: str, scheme: str, result: str = "") -> dict:
    task = task_control_state_for_item(task_type_key, scheme, result, payload)
    payload.update(
        {
            "task_status": task["status"],
            "task_runtime_status": task["runtime_status"],
            "task_type_key": task["task_type_key"],
            "task_key": task["task_key"],
            "task_result": task["result"],
            "queued": task["queued"],
            "queue_position": task["queue_position"],
            "can_start_task": task["can_start"],
            "can_queue_task": task["can_queue"],
            "can_stop_task": task["can_stop"],
            "can_cancel_queue_task": bool(task["queued"]),
        }
    )
    replace_metric_value(payload, "状态", task["status"])
    return payload


def replace_metric_value(payload: dict, label: str, value: object, unit: str = "") -> None:
    metrics = payload.get("metrics")
    if not isinstance(metrics, list):
        return
    for metric in metrics:
        if isinstance(metric, dict) and metric.get("label") == label:
            metric["value"] = value
            metric["unit"] = unit
            return
    metrics.insert(0, {"label": label, "value": value, "unit": unit})


def normalize_task_type_key(value: str) -> str:
    text = str(value or "").strip().lower()
    if text in {"optimization", "opt", "规划计算", "规划求解"}:
        return "optimization"
    if text in {"evaluation", "eval", "方案评估"}:
        return "evaluation"
    if text in {"frequency", "freq", "频率计算"}:
        return "frequency"
    return ""


def read_csv_rows_from_path(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as file:
        return [dict(row) for row in csv.DictReader(file)]


class CsvDataSource:
    """Periodically reload dashboard data from CSV files."""

    def __init__(self, data_dir: Path = DATA_DIR, reload_interval: float = 1.0) -> None:
        self.data_dir = Path(data_dir)
        self.reload_interval = reload_interval
        self._snapshot: dict | None = None
        self._loaded_at = 0.0
        self._lock = threading.Lock()

    def snapshot(self, force_reload: bool = False) -> dict:
        now = time.monotonic()
        with self._lock:
            expired = now - self._loaded_at >= self.reload_interval
            if force_reload or self._snapshot is None or expired:
                self._snapshot = self._load_snapshot()
                self._loaded_at = now
            return self._snapshot

    def _rows(self, filename: str) -> list[dict[str, str]]:
        path = self.data_dir / filename
        if not path.exists():
            return []
        return CSV_ROWS_CACHE.get(path, read_csv_rows_from_path, variant="dashboard_csv")

    def _by_page(self, rows: list[dict[str, str]], page: str) -> list[dict[str, str]]:
        return [row for row in rows if row.get("page") == page]

    def _load_snapshot(self) -> dict:
        return {
            "system": "考察站风-光-氢-储-柴联合规划系统",
            "timestamp": datetime.now().isoformat(timespec="seconds"),
            "summary": self._load_overview_summary(),
        }

    def _load_overview_summary(self) -> dict:
        return {
            row.get("key", ""): _coerce_value(row.get("value", ""))
            for row in self._rows("summary.csv")
            if row.get("key")
        }

    def _load_metrics(self, rows: list[dict[str, str]], page: str) -> list[dict]:
        return [
            _metric(
                row.get("label", ""),
                _coerce_value(row.get("value", "")),
                row.get("unit", ""),
                row.get("status", "normal"),
            )
            for row in self._by_page(rows, page)
        ]

    def _load_alarms(self, rows: list[dict[str, str]], page: str) -> list[dict]:
        return [
            {
                "time": row.get("time", ""),
                "object": row.get("object", ""),
                "message": row.get("message", ""),
                "status": row.get("status", ""),
            }
            for row in self._by_page(rows, page)
        ]

    def _load_page_summary(self, rows: list[dict[str, str]], page: str) -> list[dict]:
        return [
            _summary(row.get("label", ""), row.get("value", ""), row.get("status", "normal"))
            for row in self._by_page(rows, page)
        ]

    def _load_label_values(self, filename: str) -> list[dict]:
        return [
            {"label": row.get("label", ""), "value": _coerce_value(row.get("value", "")), "unit": row.get("unit", "")}
            for row in self._rows(filename)
        ]

    def _load_topology(self) -> list[dict]:
        return [
            {"id": row.get("id", ""), "status": row.get("status", "normal"), "value": row.get("value", "")}
            for row in self._rows("simu_topology.csv")
        ]

    def _load_simu_daily_curves(self) -> list[dict]:
        rows = self._rows("simu_daily_curves.csv")
        specs = [
            ("wind_speed", "风速", "m/s"),
            ("temperature", "温度", "℃"),
            ("solar_irradiance", "太阳辐射", "W/m²"),
            ("load", "负荷", "kW"),
        ]
        curves = []
        for key, name, unit in specs:
            points = [
                {"hour": _coerce_value(row.get("hour", "")), "value": _coerce_value(row.get(key, ""))}
                for row in rows
            ]
            curves.append({"key": key, "name": name, "unit": unit, "points": points})
        return curves

    def _load_stations(self) -> list[dict]:
        return [
            {"name": row.get("name", ""), "status": row.get("status", "normal"), "detail": row.get("detail", "")}
            for row in self._rows("scada_stations.csv")
        ]

    def _load_agc_reserve(self) -> dict:
        rows = self._rows("agc_reserve.csv")
        if not rows:
            return {"score": 0, "up": 0, "down": 0, "response": 0, "cycle": 0}
        row = rows[0]
        return {
            "score": _coerce_value(row.get("score", "")),
            "up": _coerce_value(row.get("up", "")),
            "down": _coerce_value(row.get("down", "")),
            "response": _coerce_value(row.get("response", "")),
            "cycle": _coerce_value(row.get("cycle", "")),
        }


def _mysql_connector(config: dict):
    try:
        import pymysql
    except ImportError as exc:
        raise RuntimeError("缺少 MySQL 驱动，请先安装: python -m pip install PyMySQL") from exc
    return pymysql.connect(
        host=config["host"],
        port=config["port"],
        user=config["user"],
        password=config["password"],
        database=config["database"],
        charset=config.get("charset", "utf8mb4"),
        cursorclass=pymysql.cursors.DictCursor,
        autocommit=False,
    )


class MySqlDataSource:
    """Periodically reload dashboard data from MySQL."""

    def __init__(
        self,
        config: dict | None = None,
        reload_interval: float = 1.0,
        connector_factory=_mysql_connector,
    ) -> None:
        self.config = config or DB_CONFIG
        self.reload_interval = reload_interval
        self.connector_factory = connector_factory
        self._snapshot: dict | None = None
        self._loaded_at = 0.0
        self._lock = threading.Lock()

    def snapshot(self, force_reload: bool = False) -> dict:
        now = time.monotonic()
        with self._lock:
            expired = now - self._loaded_at >= self.reload_interval
            if force_reload or self._snapshot is None or expired:
                self._snapshot = self._load_snapshot()
                self._loaded_at = now
            return self._snapshot

    def save_simu_state(self, state: dict) -> None:
        sql = """
            UPDATE simu_state
            SET sim_time = %s, speed = %s, status = %s, updated_at = CURRENT_TIMESTAMP
            WHERE id = 1
        """
        self._execute(sql, (state["sim_time"], state["speed"], state["status"]), commit=True)

    def _connect(self):
        return self.connector_factory(self.config)

    def _query(self, sql: str, params: tuple = ()) -> list[dict]:
        connection = self._connect()
        try:
            cursor = connection.cursor()
            try:
                cursor.execute(sql, params)
                rows = cursor.fetchall()
                if rows is None:
                    return []
                if isinstance(rows, dict):
                    return [rows]
                return list(rows)
            finally:
                cursor.close()
        finally:
            connection.close()

    def _query_one(self, sql: str, params: tuple = ()) -> dict | None:
        connection = self._connect()
        try:
            cursor = connection.cursor()
            try:
                cursor.execute(sql, params)
                return cursor.fetchone()
            finally:
                cursor.close()
        finally:
            connection.close()

    def _execute(self, sql: str, params: tuple = (), commit: bool = False) -> None:
        connection = self._connect()
        try:
            cursor = connection.cursor()
            try:
                cursor.execute(sql, params)
                if commit:
                    connection.commit()
            finally:
                cursor.close()
        finally:
            connection.close()

    def _load_snapshot(self) -> dict:
        return {
            "system": "考察站风-光-氢-储-柴联合规划系统",
            "timestamp": datetime.now().isoformat(timespec="seconds"),
            "summary": self._load_overview_summary(),
        }

    def _load_overview_summary(self) -> dict:
        return {
            row.get("key", ""): _coerce_value(str(row.get("value", "")))
            for row in self._query("SELECT `key`, value, unit FROM overview_summary ORDER BY display_order, id")
            if row.get("key")
        }

    def _by_page(self, rows: list[dict], page: str) -> list[dict]:
        return [row for row in rows if row.get("page") == page]

    def _load_metrics(self, rows: list[dict], page: str) -> list[dict]:
        return [
            _metric(
                row.get("label", ""),
                _coerce_value(str(row.get("value", ""))),
                row.get("unit", ""),
                row.get("status", "normal"),
            )
            for row in self._by_page(rows, page)
        ]

    def _load_alarms(self, rows: list[dict], page: str) -> list[dict]:
        return [
            {
                "time": row.get("time", ""),
                "object": row.get("object", ""),
                "message": row.get("message", ""),
                "status": row.get("status", ""),
            }
            for row in self._by_page(rows, page)
        ]

    def _load_page_summary(self, rows: list[dict], page: str) -> list[dict]:
        return [
            _summary(row.get("label", ""), row.get("value", ""), row.get("status", "normal"))
            for row in self._by_page(rows, page)
        ]

    def _load_label_values(self, sql: str) -> list[dict]:
        return [
            {
                "label": row.get("label", ""),
                "value": _coerce_value(str(row.get("value", ""))),
                "unit": row.get("unit", ""),
            }
            for row in self._query(sql)
        ]

    def _load_topology(self) -> list[dict]:
        return [
            {"id": row.get("id", ""), "status": row.get("status", "normal"), "value": row.get("value", "")}
            for row in self._query("SELECT id, status, value FROM simu_topology ORDER BY display_order, id")
        ]

    def _load_simu_daily_curves(self) -> list[dict]:
        rows = self._query(
            "SELECT hour, wind_speed, temperature, solar_irradiance, load_value FROM simu_daily_curves ORDER BY hour"
        )
        specs = [
            ("wind_speed", "风速", "m/s"),
            ("temperature", "温度", "℃"),
            ("solar_irradiance", "太阳辐射", "W/m²"),
            ("load_value", "负荷", "kW"),
        ]
        curves = []
        for key, name, unit in specs:
            points = [
                {"hour": _coerce_value(str(row.get("hour", ""))), "value": _coerce_value(str(row.get(key, "")))}
                for row in rows
            ]
            curves.append({"key": key, "name": name, "unit": unit, "points": points})
        return curves

    def _load_stations(self) -> list[dict]:
        return [
            {"name": row.get("name", ""), "status": row.get("status", "normal"), "detail": row.get("detail", "")}
            for row in self._query("SELECT name, status, detail FROM scada_stations ORDER BY display_order, id")
        ]

    def _load_agc_reserve(self) -> dict:
        row = self._query_one("SELECT score, up, down, response, cycle FROM agc_reserve ORDER BY id LIMIT 1")
        if not row:
            return {"score": 0, "up": 0, "down": 0, "response": 0, "cycle": 0}
        return {
            "score": _coerce_value(str(row.get("score", ""))),
            "up": _coerce_value(str(row.get("up", ""))),
            "down": _coerce_value(str(row.get("down", ""))),
            "response": _coerce_value(str(row.get("response", ""))),
            "cycle": _coerce_value(str(row.get("cycle", ""))),
        }


def _load_initial_simu_runtime() -> SimuRuntime:
    return SimuRuntime()


SIMU_RUNTIME = _load_initial_simu_runtime()
OPTIMIZATION_RUNTIME = OptimizationRuntimeManager()
EVALUATION_RUNTIME = EvaluationRuntimeManager()
FREQUENCY_EVALUATION_RUNTIME = FrequencyEvaluationRuntimeManager()
RELIABILITY_RUNTIME = ReliabilityRuntimeManager()
TASK_SCHEDULER = TaskScheduler()
DATA_SOURCE = CsvDataSource()
PLANNING_STORE = planning_store.PlanningStore()
USER_STORE = UserStore()


def build_snapshot(force_reload: bool = False) -> dict:
    """Build a snapshot from CSV files, reloading periodically."""
    return DATA_SOURCE.snapshot(force_reload=force_reload)


def _json_response(payload: dict, status: int = 200, extra_headers: dict[str, str] | None = None) -> tuple[int, dict[str, str], bytes]:
    body = json.dumps(payload, ensure_ascii=False, separators=(",", ":")).encode("utf-8")
    headers = {
        "Content-Type": "application/json; charset=utf-8",
        "Cache-Control": "no-store",
    }
    if extra_headers:
        headers.update(extra_headers)
    return status, headers, body


def _download_response(body: bytes, filename: str, content_type: str) -> tuple[int, dict[str, str], bytes]:
    safe_filename = re.sub(r"[^A-Za-z0-9._-]+", "_", str(filename or "").strip()) or "download"
    return (
        HTTPStatus.OK,
        {
            "Content-Type": content_type,
            "Cache-Control": "no-store",
            "Content-Disposition": f"attachment; filename=\"{safe_filename}\"; filename*=UTF-8''{quote(filename)}",
        },
        body,
    )


def response_is_compressible(headers: dict[str, str], body: bytes) -> bool:
    if len(body) < MIN_GZIP_RESPONSE_BYTES:
        return False
    if headers.get("Content-Encoding"):
        return False
    content_type = str(headers.get("Content-Type", "")).lower()
    return any(content_type.startswith(prefix) for prefix in COMPRESSIBLE_CONTENT_PREFIXES)


def gzip_response_body_if_supported(request_headers, headers: dict[str, str], body: bytes) -> tuple[dict[str, str], bytes]:
    if "gzip" not in str(request_headers.get("Accept-Encoding", "")).lower():
        return headers, body
    if not response_is_compressible(headers, body):
        return headers, body
    compressed = zlib.compressobj(level=6, wbits=16 + zlib.MAX_WBITS)
    compressed_body = compressed.compress(body) + compressed.flush()
    if len(compressed_body) >= len(body):
        return headers, body
    next_headers = dict(headers)
    next_headers["Content-Encoding"] = "gzip"
    next_headers["Vary"] = append_vary_header(next_headers.get("Vary", ""), "Accept-Encoding")
    return next_headers, compressed_body


def append_vary_header(current: str, value: str) -> str:
    names = [item.strip() for item in str(current or "").split(",") if item.strip()]
    if not any(item.lower() == value.lower() for item in names):
        names.append(value)
    return ", ".join(names)


def _no_store_headers(*, vary_cookie: bool = False) -> dict[str, str]:
    headers = {
        "Cache-Control": NO_STORE_CACHE_CONTROL,
        "Pragma": "no-cache",
        "Expires": "0",
    }
    if vary_cookie:
        headers["Vary"] = "Cookie"
    return headers


def _static_headers(path: Path, *, authenticated_html: bool = False) -> dict[str, str]:
    content_type = mimetypes.guess_type(str(path))[0] or "application/octet-stream"
    suffix = path.suffix.lower()
    headers = {"Content-Type": content_type}
    if suffix in STATIC_NO_STORE_SUFFIXES:
        headers.update(_no_store_headers(vary_cookie=authenticated_html or path.suffix.lower() == ".html"))
        return headers
    if suffix in STATIC_BROWSER_CACHE_SUFFIXES:
        headers["Cache-Control"] = STATIC_ASSET_CACHE_CONTROL
        headers.update(_static_validator_headers(path))
        return headers
    headers["Cache-Control"] = STATIC_DATA_CACHE_CONTROL
    headers.update(_static_validator_headers(path))
    return headers


def _static_validator_headers(path: Path) -> dict[str, str]:
    stat = path.stat()
    return {
        "ETag": f'W/"{stat.st_mtime_ns:x}-{stat.st_size:x}"',
        "Last-Modified": formatdate(stat.st_mtime, usegmt=True),
    }


def _static_request_not_modified(request_headers, response_headers: dict[str, str]) -> bool:
    etag = response_headers.get("ETag")
    if etag and request_headers.get("If-None-Match") == etag:
        return True
    last_modified = response_headers.get("Last-Modified")
    if last_modified and request_headers.get("If-Modified-Since") == last_modified:
        return True
    return False


def _read_json_body(body: bytes) -> dict:
    try:
        return json.loads(body.decode("utf-8") or "{}")
    except json.JSONDecodeError as exc:
        raise ValueError("请求体不是合法 JSON") from exc


def truthy_json_value(value) -> bool:
    if isinstance(value, bool):
        return value
    if isinstance(value, (int, float)):
        return float(value) != 0.0
    return str(value or "").strip().lower() in {"1", "true", "yes", "y", "是"}


def _session_cookie(token: str, max_age: int = SESSION_MAX_AGE_SECONDS) -> str:
    return f"{SESSION_COOKIE_NAME}={token}; Path=/; HttpOnly; SameSite=Lax; Max-Age={max_age}"


def _expired_session_cookie() -> str:
    return f"{SESSION_COOKIE_NAME}=; Path=/; HttpOnly; SameSite=Lax; Max-Age=0"


def _session_token_from_cookie(cookie_header: str | None) -> str:
    if not cookie_header:
        return ""
    cookie = SimpleCookie()
    try:
        cookie.load(cookie_header)
    except Exception:
        return ""
    morsel = cookie.get(SESSION_COOKIE_NAME)
    return morsel.value if morsel else ""


def _authenticated_user(token: str = "") -> dict | None:
    if LOCAL_AUTH_BYPASS_ENABLED:
        return dict(LOCAL_AUTH_USER)
    return USER_STORE.user_for_session(token)


def handle_auth_api_path(path: str, method: str, body: bytes = b"", token: str = "") -> tuple[int, dict[str, str], bytes]:
    current_user = _authenticated_user(token)
    try:
        if path == "/api/auth/me" and method == "GET":
            if not current_user:
                return _json_response({"error": "unauthorized", "message": "请先登录"}, HTTPStatus.UNAUTHORIZED)
            return _json_response({"user": current_user})
        if path == "/api/auth/register" and method == "POST":
            payload = _read_json_body(body)
            user = USER_STORE.create_user(str(payload.get("username", "")), str(payload.get("password", "")))
            session_token = USER_STORE.create_session(user["id"])
            return _json_response({"ok": True, "user": user}, extra_headers={"Set-Cookie": _session_cookie(session_token)})
        if path == "/api/auth/login" and method == "POST":
            payload = _read_json_body(body)
            user = USER_STORE.authenticate(str(payload.get("username", "")), str(payload.get("password", "")))
            session_token = USER_STORE.create_session(user["id"])
            return _json_response({"ok": True, "user": user}, extra_headers={"Set-Cookie": _session_cookie(session_token)})
        if path == "/api/auth/logout" and method == "POST":
            USER_STORE.delete_session(token)
            return _json_response({"ok": True}, extra_headers={"Set-Cookie": _expired_session_cookie()})
    except ValueError as exc:
        return _json_response({"error": "bad_request", "message": str(exc)}, HTTPStatus.BAD_REQUEST)
    return _json_response({"error": "not_found", "path": path}, HTTPStatus.NOT_FOUND)


def handle_users_api_path(path: str, method: str, body: bytes, current_user: dict | None) -> tuple[int, dict[str, str], bytes]:
    if not current_user:
        return _json_response({"error": "unauthorized", "message": "请先登录"}, HTTPStatus.UNAUTHORIZED)
    if current_user.get("role") != "admin":
        return _json_response({"error": "forbidden", "message": "需要管理员权限"}, HTTPStatus.FORBIDDEN)
    try:
        if path == "/api/users" and method == "GET":
            return _json_response({"users": USER_STORE.list_users()})
        if path.startswith("/api/users/"):
            user_id = int(path.rsplit("/", 1)[1])
            if method == "PUT":
                payload = _read_json_body(body)
                return _json_response({"user": USER_STORE.update_role(user_id, str(payload.get("role", "")))})
            if method == "DELETE":
                USER_STORE.delete_user(user_id, current_user_id=current_user["id"])
                return _json_response({"ok": True})
    except FileNotFoundError as exc:
        return _json_response({"error": "not_found", "message": str(exc)}, HTTPStatus.NOT_FOUND)
    except (TypeError, ValueError) as exc:
        return _json_response({"error": "bad_request", "message": str(exc)}, HTTPStatus.BAD_REQUEST)
    return _json_response({"error": "not_found", "path": path}, HTTPStatus.NOT_FOUND)


class WeatherHistoryError(RuntimeError):
    """Raised when historical weather data cannot be fetched or parsed."""


class GeocodingError(RuntimeError):
    """Raised when a place name cannot be resolved to coordinates."""


class LoadCurveTemplateExistsError(FileExistsError):
    """Raised when saving a load-curve template would overwrite an existing template."""


def geocode_place_name(place: str) -> dict:
    query_text = str(place or "").strip()
    if not query_text:
        raise ValueError("地名不能为空")
    errors: list[str] = []
    for candidate in geocode_query_candidates(query_text):
        providers = geocode_provider_order(candidate)
        for provider in providers:
            try:
                result = provider(candidate)
                if candidate != query_text:
                    result["place"] = query_text
                    result["display_name"] = f"{query_text} / {result.get('display_name') or candidate}"
                return result
            except GeocodingError as exc:
                errors.append(str(exc))
    raise GeocodingError("；".join(errors) or "未找到该地名对应的经纬度坐标")


def geocode_query_candidates(place: str) -> list[str]:
    """Return the original query plus an English alias for common Chinese place names."""
    query_text = str(place or "").strip()
    normalized = re.sub(r"\s+", "", query_text.lower())
    alias = CHINESE_PLACE_ALIASES.get(normalized)
    candidates = [query_text]
    if alias and alias not in candidates:
        candidates.append(alias)
    return candidates


def geocode_provider_order(place: str):
    """Prefer Amap for Chinese place names, and global providers for foreign names."""
    global_providers = [geocode_with_open_meteo, geocode_with_photon, geocode_with_nominatim]
    if not AMAP_WEB_SERVICE_KEY:
        return global_providers
    if should_prefer_global_geocoder(place):
        return [*global_providers, geocode_with_amap]
    return [geocode_with_amap, *global_providers]


def should_prefer_global_geocoder(place: str) -> bool:
    """Amap geocoding can match English names to domestic POIs, so route them globally first."""
    text = str(place or "").strip()
    has_cjk = bool(re.search(r"[\u3400-\u9fff]", text))
    has_non_cjk_letter = any(character.isalpha() and not ("\u3400" <= character <= "\u9fff") for character in text)
    return has_non_cjk_letter and not has_cjk


def geocode_with_amap(place: str) -> dict:
    query = urlencode({"address": place, "output": "JSON", "key": AMAP_WEB_SERVICE_KEY})
    url = f"{AMAP_GEOCODING_URL}?{query}"
    try:
        with urlopen_with_user_agent(url, timeout=8) as response:
            data = json.loads(response.read().decode("utf-8"))
    except HTTPError as exc:
        raise GeocodingError(f"高德地名解析接口返回错误: HTTP {exc.code}") from exc
    except (URLError, TimeoutError) as exc:
        raise GeocodingError(f"高德地名解析接口连接失败: {exc}") from exc
    except json.JSONDecodeError as exc:
        raise GeocodingError("高德地名解析接口返回内容不是合法 JSON") from exc
    if str(data.get("status")) != "1":
        raise GeocodingError(f"高德地名解析失败: {data.get('info') or data.get('infocode') or '未知错误'}")
    geocodes = data.get("geocodes", [])
    if not geocodes:
        raise GeocodingError("高德未找到该地名对应的经纬度坐标")
    first = geocodes[0]
    location = str(first.get("location", ""))
    try:
        longitude_text, latitude_text = location.split(",", 1)
        longitude = float(longitude_text)
        latitude = float(latitude_text)
    except (ValueError, AttributeError) as exc:
        raise GeocodingError("高德地名解析结果缺少有效经纬度") from exc
    display_parts = [
        first.get("formatted_address"),
        first.get("province"),
        first.get("city"),
        first.get("district"),
    ]
    display_name = "，".join(str(part) for part in display_parts if part and part != [])
    return {
        "place": place,
        "display_name": display_name or place,
        "latitude": latitude,
        "longitude": longitude,
        "source": "高德地图 Web 服务地理编码 API",
    }


def geocode_with_open_meteo(place: str) -> dict:
    query = urlencode({"name": place, "count": 1, "language": "zh", "format": "json"})
    url = f"{OPEN_METEO_GEOCODING_URL}?{query}"
    try:
        with urlopen_with_user_agent(url, timeout=8) as response:
            data = json.loads(response.read().decode("utf-8"))
    except HTTPError as exc:
        raise GeocodingError(f"Open-Meteo 地名解析接口返回错误: HTTP {exc.code}") from exc
    except (URLError, TimeoutError) as exc:
        raise GeocodingError(f"Open-Meteo 地名解析接口连接失败: {exc}") from exc
    except json.JSONDecodeError as exc:
        raise GeocodingError("Open-Meteo 地名解析接口返回内容不是合法 JSON") from exc
    results = data.get("results", []) if isinstance(data, dict) else []
    if not results:
        raise GeocodingError("Open-Meteo 未找到该地名对应的经纬度坐标")
    first = results[0]
    try:
        latitude = float(first["latitude"])
        longitude = float(first["longitude"])
    except (KeyError, TypeError, ValueError) as exc:
        raise GeocodingError("Open-Meteo 地名解析结果缺少有效经纬度") from exc
    display_parts = [
        first.get("name"),
        first.get("admin2"),
        first.get("admin1"),
        first.get("country"),
    ]
    display_name = "，".join(str(part) for part in display_parts if part)
    return {
        "place": place,
        "display_name": display_name or place,
        "latitude": latitude,
        "longitude": longitude,
        "source": "Open-Meteo Geocoding API",
    }


def geocode_with_nominatim(place: str) -> dict:
    query = urlencode({"q": place, "format": "json", "limit": 1, "accept-language": "zh-CN"})
    url = f"{NOMINATIM_SEARCH_URL}?{query}"
    try:
        with urlopen_with_user_agent(url, timeout=8) as response:
            data = json.loads(response.read().decode("utf-8"))
    except HTTPError as exc:
        raise GeocodingError(f"Nominatim 地名解析接口返回错误: HTTP {exc.code}") from exc
    except (URLError, TimeoutError) as exc:
        raise GeocodingError(f"Nominatim 地名解析接口连接失败: {exc}") from exc
    except json.JSONDecodeError as exc:
        raise GeocodingError("Nominatim 地名解析接口返回内容不是合法 JSON") from exc
    if not isinstance(data, list) or not data:
        raise GeocodingError("Nominatim 未找到该地名对应的经纬度坐标")
    first = data[0]
    try:
        latitude = float(first["lat"])
        longitude = float(first["lon"])
    except (KeyError, TypeError, ValueError) as exc:
        raise GeocodingError("Nominatim 地名解析结果缺少有效经纬度") from exc
    return {
        "place": place,
        "display_name": first.get("display_name", place),
        "latitude": latitude,
        "longitude": longitude,
        "source": "OpenStreetMap Nominatim",
    }


def geocode_with_photon(place: str) -> dict:
    query = urlencode({"q": place, "limit": 1, "lang": "en"})
    url = f"{PHOTON_GEOCODING_URL}?{query}"
    try:
        with urlopen_with_user_agent(url, timeout=8) as response:
            data = json.loads(response.read().decode("utf-8"))
    except HTTPError as exc:
        raise GeocodingError(f"Photon 地名解析接口返回错误: HTTP {exc.code}") from exc
    except (URLError, TimeoutError) as exc:
        raise GeocodingError(f"Photon 地名解析接口连接失败: {exc}") from exc
    except json.JSONDecodeError as exc:
        raise GeocodingError("Photon 地名解析接口返回内容不是合法 JSON") from exc
    features = data.get("features", []) if isinstance(data, dict) else []
    if not features:
        raise GeocodingError("Photon 未找到该地名对应的经纬度坐标")
    first = features[0]
    coordinates = first.get("geometry", {}).get("coordinates", [])
    try:
        longitude = float(coordinates[0])
        latitude = float(coordinates[1])
    except (IndexError, TypeError, ValueError) as exc:
        raise GeocodingError("Photon 地名解析结果缺少有效经纬度") from exc
    properties = first.get("properties", {})
    display_parts = [
        properties.get("name"),
        properties.get("city"),
        properties.get("state"),
        properties.get("country"),
    ]
    display_name = "，".join(str(part) for part in display_parts if part)
    return {
        "place": place,
        "display_name": display_name or place,
        "latitude": latitude,
        "longitude": longitude,
        "source": "OpenStreetMap Photon API",
    }


def reverse_geocode_coordinates(latitude: float, longitude: float) -> dict:
    latitude, longitude = validate_reverse_geocode_inputs(latitude, longitude)
    providers = [reverse_geocode_with_nominatim]
    if AMAP_WEB_SERVICE_KEY:
        providers.insert(0, reverse_geocode_with_amap)
    errors: list[str] = []
    for provider in providers:
        try:
            return provider(latitude, longitude)
        except GeocodingError as exc:
            errors.append(str(exc))
    raise GeocodingError("；".join(errors) or "未找到该坐标对应的地点")


def validate_reverse_geocode_inputs(latitude: float, longitude: float) -> tuple[float, float]:
    try:
        lat = float(latitude)
        lng = float(longitude)
    except (TypeError, ValueError) as exc:
        raise ValueError("经纬度必须为有效数值") from exc
    if lat < -90 or lat > 90:
        raise ValueError("纬度范围应为 -90 到 90")
    if lng < -180 or lng > 180:
        raise ValueError("经度范围应为 -180 到 180")
    return lat, lng


def reverse_geocode_with_amap(latitude: float, longitude: float) -> dict:
    query = urlencode(
        {
            "location": f"{longitude},{latitude}",
            "output": "JSON",
            "key": AMAP_WEB_SERVICE_KEY,
            "extensions": "base",
        }
    )
    url = f"{AMAP_REVERSE_GEOCODING_URL}?{query}"
    try:
        with urlopen_with_user_agent(url, timeout=8) as response:
            data = json.loads(response.read().decode("utf-8"))
    except HTTPError as exc:
        raise GeocodingError(f"高德坐标反查接口返回错误: HTTP {exc.code}") from exc
    except (URLError, TimeoutError) as exc:
        raise GeocodingError(f"高德坐标反查接口连接失败: {exc}") from exc
    except json.JSONDecodeError as exc:
        raise GeocodingError("高德坐标反查接口返回内容不是合法 JSON") from exc
    if str(data.get("status")) != "1":
        raise GeocodingError(f"高德坐标反查失败: {data.get('info') or data.get('infocode') or '未知错误'}")
    regeocode = data.get("regeocode", {}) if isinstance(data, dict) else {}
    if not regeocode:
        raise GeocodingError("高德未找到该坐标对应的地点")
    component = regeocode.get("addressComponent", {}) if isinstance(regeocode, dict) else {}
    display_parts = [
        regeocode.get("formatted_address"),
        component.get("province"),
        component.get("city"),
        component.get("district"),
        component.get("township"),
    ]
    display_name = "，".join(str(part) for part in display_parts if part and part != [])
    place = str(regeocode.get("formatted_address") or display_name or "").strip()
    if not place:
        raise GeocodingError("高德坐标反查结果缺少有效地点名称")
    return {
        "place": place,
        "display_name": place,
        "latitude": latitude,
        "longitude": longitude,
        "source": "高德地图 Web 服务逆地理编码 API",
    }


def reverse_geocode_with_nominatim(latitude: float, longitude: float) -> dict:
    query = urlencode({"format": "json", "lat": latitude, "lon": longitude, "accept-language": "zh-CN"})
    url = f"{NOMINATIM_REVERSE_URL}?{query}"
    try:
        with urlopen_with_user_agent(url, timeout=8) as response:
            data = json.loads(response.read().decode("utf-8"))
    except HTTPError as exc:
        raise GeocodingError(f"Nominatim 坐标反查接口返回错误: HTTP {exc.code}") from exc
    except (URLError, TimeoutError) as exc:
        raise GeocodingError(f"Nominatim 坐标反查接口连接失败: {exc}") from exc
    except json.JSONDecodeError as exc:
        raise GeocodingError("Nominatim 坐标反查接口返回内容不是合法 JSON") from exc
    if not isinstance(data, dict):
        raise GeocodingError("Nominatim 未找到该坐标对应的地点")
    display_name = str(data.get("display_name") or "").strip()
    address = data.get("address", {}) if isinstance(data.get("address"), dict) else {}
    place = display_name or ", ".join(
        str(part)
        for part in (
            address.get("road"),
            address.get("suburb"),
            address.get("city") or address.get("town") or address.get("village"),
            address.get("state"),
            address.get("country"),
        )
        if part
    )
    if not place:
        raise GeocodingError("Nominatim 坐标反查结果缺少有效地点名称")
    return {
        "place": place,
        "display_name": display_name or place,
        "latitude": latitude,
        "longitude": longitude,
        "source": "OpenStreetMap Nominatim Reverse",
    }


def urlopen_with_user_agent(url: str, timeout: int):
    from urllib.request import Request

    request = Request(url, headers={"User-Agent": "power-plan-local-web/1.0"})
    return urlopen(request, timeout=timeout)


def fetch_weather_history(latitude: float, longitude: float, year: int) -> dict:
    latitude, longitude, year = validate_weather_history_inputs(latitude, longitude, year)
    query = urlencode(
        {
            "parameters": ",".join(NASA_POWER_PARAMETERS.values()),
            "community": "RE",
            "longitude": longitude,
            "latitude": latitude,
            "start": f"{year}0101",
            "end": f"{year}1231",
            "format": "JSON",
            "time-standard": "LST",
        }
    )
    url = f"{NASA_POWER_HOURLY_URL}?{query}"
    try:
        with urlopen_with_user_agent(url, timeout=45) as response:
            data = json.loads(response.read().decode("utf-8"))
    except HTTPError as exc:
        raise WeatherHistoryError(f"历史气象数据接口返回错误: HTTP {exc.code}") from exc
    except (URLError, TimeoutError) as exc:
        raise WeatherHistoryError(f"历史气象数据接口连接失败: {exc}") from exc
    except json.JSONDecodeError as exc:
        raise WeatherHistoryError("历史气象数据接口返回内容不是合法 JSON") from exc
    rows = parse_nasa_power_hourly_response(data, year)
    return {
        "source": "NASA POWER Hourly API",
        "source_url": url,
        "latitude": latitude,
        "longitude": longitude,
        "year": year,
        "rows": rows,
    }


def validate_weather_history_inputs(latitude: float, longitude: float, year: int) -> tuple[float, float, int]:
    try:
        latitude_number = float(latitude)
        longitude_number = float(longitude)
        year_number = int(year)
    except (TypeError, ValueError) as exc:
        raise ValueError("经纬度和历史数据年必须为数值") from exc
    current_year = datetime.now().year
    if not -90 <= latitude_number <= 90:
        raise ValueError("纬度范围应为 -90 到 90")
    if not -180 <= longitude_number <= 180:
        raise ValueError("经度范围应为 -180 到 180")
    if year_number < 2001:
        raise ValueError("NASA POWER 小时历史数据年份不能早于 2001")
    if year_number >= current_year:
        raise ValueError(f"历史数据年必须小于当前年 {current_year}")
    return latitude_number, longitude_number, year_number


def parse_nasa_power_hourly_response(data: dict, year: int) -> list[dict]:
    parameters = data.get("properties", {}).get("parameter", {})
    fill_value = data.get("header", {}).get("fill_value", -999)
    missing = [api_name for api_name in NASA_POWER_PARAMETERS.values() if api_name not in parameters]
    if missing:
        raise WeatherHistoryError(f"历史气象数据缺少字段: {', '.join(missing)}")
    wind_values = parameters[NASA_POWER_PARAMETERS["wind_speed"]]
    keys = sorted(key for key in wind_values if str(key).startswith(str(year)) and str(key)[4:8] != "0229")
    rows = []
    for hour_index, key in enumerate(keys, start=1):
        row = {"hour_index": hour_index, "datetime": power_hour_key_to_datetime(str(key))}
        for field, api_name in NASA_POWER_PARAMETERS.items():
            value = parameters.get(api_name, {}).get(key)
            if value in (None, "", fill_value):
                raise WeatherHistoryError(f"历史气象数据在 {key} 缺少 {api_name}")
            try:
                number = float(value)
            except (TypeError, ValueError) as exc:
                raise WeatherHistoryError(f"历史气象数据在 {key} 的 {api_name} 不是数值") from exc
            row[field] = round(number, 4)
        rows.append(row)
    if len(rows) != 8760:
        raise WeatherHistoryError(f"历史气象数据小时数应为8760，当前为{len(rows)}")
    return rows


def power_hour_key_to_datetime(key: str) -> str:
    return f"{key[0:4]}-{key[4:6]}-{key[6:8]} {key[8:10]}:00"


def import_time_series_file(filename: str, content: bytes) -> dict:
    suffix = Path(filename or "").suffix.lower()
    if suffix == ".csv":
        headers, rows = read_time_series_csv(content)
    elif suffix == ".xlsx":
        headers, rows = read_time_series_xlsx(content)
    else:
        raise ValueError("导入文件仅支持 .csv 或 .xlsx 格式")
    return normalize_imported_time_series(headers, rows, filename)


def read_time_series_csv(content: bytes) -> tuple[list[str], list[dict[str, object]]]:
    text = decode_csv_text(content)
    reader = csv.DictReader(StringIO(text))
    headers = [str(header or "").strip() for header in (reader.fieldnames or [])]
    rows = [row for row in reader if any(str(value or "").strip() for value in row.values())]
    return headers, rows


def read_time_series_xlsx(content: bytes) -> tuple[list[str], list[dict[str, object]]]:
    workbook = load_workbook(BytesIO(content), read_only=True, data_only=True)
    try:
        sheet = workbook.active
        row_iter = sheet.iter_rows(values_only=True)
        headers = [str(value or "").strip() for value in next(row_iter, ())]
        rows = []
        for values in row_iter:
            if not any(value not in ("", None) for value in values):
                continue
            rows.append({headers[index]: value if index < len(values) else "" for index, value in enumerate(values) if index < len(headers)})
        return headers, rows
    finally:
        workbook.close()


def decode_csv_text(content: bytes) -> str:
    for encoding in ("utf-8-sig", "gb18030"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ValueError("CSV文件编码无法识别，请使用 UTF-8 或 GBK 编码")


def normalize_imported_time_series(headers: list[str], raw_rows: list[dict[str, object]], filename: str) -> dict:
    column_map = match_time_series_import_columns(headers)
    if not raw_rows:
        raise ValueError("导入失败，文件没有可用数据行")
    parsed_by_hour: dict[int, dict] = {}
    source_hour_count = min(len(raw_rows), TIME_SERIES_IMPORT_ROW_COUNT)
    for row_index, raw_row in enumerate(raw_rows[:TIME_SERIES_IMPORT_ROW_COUNT], start=1):
        target_hour = imported_hour_index(raw_row, column_map.get("datetime"), row_index)
        if target_hour < 1 or target_hour > TIME_SERIES_IMPORT_ROW_COUNT or target_hour in parsed_by_hour:
            target_hour = next_available_time_series_hour(parsed_by_hour, row_index)
        item = {"hour_index": target_hour, "datetime": imported_datetime(raw_row, column_map.get("datetime"), target_hour)}
        for key in TIME_SERIES_IMPORT_REQUIRED_COLUMNS:
            item[key] = imported_numeric_value_or_none(raw_row.get(column_map[key]))
        parsed_by_hour[target_hour] = item
    repaired_numeric_count = repair_imported_numeric_values(parsed_by_hour)
    imported_rows = fill_imported_time_series_hours(parsed_by_hour)
    missing_count = sum(1 for item in imported_rows if item.get("_filled"))
    for item in imported_rows:
        item.pop("_filled", None)
    message = f"已从{filename}导入8760行时序数据"
    if len(raw_rows) > TIME_SERIES_IMPORT_ROW_COUNT:
        message += f"，文件共有{len(raw_rows)}行，已使用前8760行"
    elif source_hour_count < TIME_SERIES_IMPORT_ROW_COUNT:
        message += f"，文件共有{source_hour_count}行，已按最后一行自动补齐{TIME_SERIES_IMPORT_ROW_COUNT - source_hour_count}行"
    if missing_count:
        message += f"，已补齐{missing_count}个缺失时点"
    if repaired_numeric_count:
        message += f"，已修复{repaired_numeric_count}个无效数值"
    return {"time_series": imported_rows, "time_series_count": len(imported_rows), "message": message}


def import_load_curve_file(filename: str, content: bytes, minimum: object, maximum: object, average: object, raw: object = False) -> dict:
    suffix = Path(filename or "").suffix.lower()
    if suffix == ".csv":
        headers, rows = read_time_series_csv(content)
    elif suffix == ".xlsx":
        headers, rows = read_time_series_xlsx(content)
    else:
        raise ValueError("导入文件仅支持 .csv 或 .xlsx 格式")
    return normalize_imported_load_curve(headers, rows, filename, minimum, maximum, average, raw)


def normalize_imported_load_curve(
    headers: list[str],
    raw_rows: list[dict[str, object]],
    filename: str,
    minimum: object,
    maximum: object,
    average: object,
    raw: object = False,
) -> dict:
    if not raw_rows:
        raise ValueError("导入失败，文件没有可用数据行")
    column_map = match_load_curve_import_columns(headers)
    values, repaired_numeric_count, missing_count, duplicate_count = normalized_imported_load_values(raw_rows, column_map)
    if truthy_json_value(raw):
        scaled_values = [round_load_value(float(value)) for value in values]
    else:
        min_value, max_value, avg_value = validate_load_curve_targets(minimum, maximum, average)
        scaled_values = scale_load_values_to_targets(values, min_value, max_value, avg_value)
    rows = [{"hour_index": index + 1, "load": value} for index, value in enumerate(scaled_values)]
    message = f"已从{filename}导入8760点负荷{'原始' if truthy_json_value(raw) else ''}曲线"
    if len(raw_rows) != TIME_SERIES_IMPORT_ROW_COUNT:
        message += f"，文件共有{len(raw_rows)}行，已自适应扩展到8760点"
    if missing_count:
        message += f"，已按相邻点自动补齐{missing_count}个缺失时点"
    if duplicate_count:
        message += f"，已合并{duplicate_count}个重复小时数据"
    if repaired_numeric_count:
        message += f"，已修复{repaired_numeric_count}个无效数值"
    return {
        "source": "file",
        "load_curve": rows,
        "load_curve_count": len(rows),
        "statistics": {
            "max": round_load_value(max(scaled_values)),
            "min": round_load_value(min(scaled_values)),
            "average": round_load_value(sum(scaled_values) / len(scaled_values)),
        },
        "message": message,
    }


def match_load_curve_import_columns(headers: list[str]) -> dict[str, str]:
    normalized_headers = [(normalize_import_header(header), header) for header in headers if str(header or "").strip()]
    load_aliases = TIME_SERIES_IMPORT_REQUIRED_COLUMNS["load"][1]
    load_column = match_time_series_header(normalized_headers, load_aliases)
    if not load_column:
        raise ValueError("导入失败，找不到对应的列：负荷")
    column_map = {"load": load_column}
    matched_time = match_time_series_header(normalized_headers, TIME_SERIES_IMPORT_OPTIONAL_COLUMNS["datetime"])
    if matched_time:
        column_map["datetime"] = matched_time
    return column_map


def normalized_imported_load_values(raw_rows: list[dict[str, object]], column_map: dict[str, str]) -> tuple[list[float | int], int, int, int]:
    load_column = column_map["load"]
    time_column = column_map.get("datetime")
    raw_values = [imported_numeric_value_or_none(row.get(load_column)) for row in raw_rows]
    grouped_by_hour: dict[int, list[float | int | None]] = {}
    if time_column:
        for raw_row, value in zip(raw_rows, raw_values):
            target_hour = imported_load_curve_hour_index(raw_row, time_column)
            if target_hour is None or target_hour < 1 or target_hour > TIME_SERIES_IMPORT_ROW_COUNT:
                continue
            grouped_by_hour.setdefault(target_hour, []).append(value)
    if grouped_by_hour:
        parsed_by_hour = {}
        duplicate_count = 0
        for hour, hour_values in grouped_by_hour.items():
            duplicate_count += max(0, len(hour_values) - 1)
            valid_values = [float(value) for value in hour_values if value is not None]
            parsed_by_hour[hour] = {
                "hour_index": hour,
                "datetime": f"H{hour:04d}",
                "load": sum(valid_values) / len(valid_values) if valid_values else None,
            }
        repaired_numeric_count = repair_imported_load_values(parsed_by_hour)
        imported_rows = fill_imported_load_curve_hours(parsed_by_hour)
        missing_count = sum(1 for item in imported_rows if item.get("_filled"))
        source_values = [row["load"] for row in imported_rows[: max(grouped_by_hour)]]
        return resample_load_values(source_values, TIME_SERIES_IMPORT_ROW_COUNT), repaired_numeric_count, missing_count, duplicate_count
    repaired_values, repaired_numeric_count = repair_load_value_sequence(raw_values)
    return resample_load_values(repaired_values, TIME_SERIES_IMPORT_ROW_COUNT), repaired_numeric_count, 0, 0


def imported_load_curve_hour_index(raw_row: dict[str, object], column: str) -> int | None:
    value = raw_row.get(column)
    if value in ("", None):
        return None
    if isinstance(value, datetime):
        return (value.timetuple().tm_yday - 1) * 24 + value.hour + 1
    if isinstance(value, (int, float)) and math.isfinite(float(value)):
        return int(value)
    text = str(value).strip()
    if not text:
        return None
    hour_match = re.fullmatch(r"[hH]\s*0*(\d{1,4})", text)
    if hour_match:
        return int(hour_match.group(1))
    numeric_match = re.fullmatch(r"0*(\d{1,4})(?:\.0+)?", text)
    if numeric_match:
        return int(numeric_match.group(1))
    datetime_match = re.search(r"\b(\d{1,2})[:：](\d{1,2})(?::\d{1,2})?\b", text)
    date_match = re.search(r"(\d{4})[-/年.](\d{1,2})[-/月.](\d{1,2})", text)
    if datetime_match and date_match:
        try:
            parsed = datetime(
                int(date_match.group(1)),
                int(date_match.group(2)),
                int(date_match.group(3)),
                int(datetime_match.group(1)),
            )
            return (parsed.timetuple().tm_yday - 1) * 24 + parsed.hour + 1
        except ValueError:
            return None
    return None


def repair_imported_load_values(parsed_by_hour: dict[int, dict]) -> int:
    hours = sorted(parsed_by_hour)
    valid_hours = [hour for hour in hours if parsed_by_hour[hour].get("load") is not None]
    if not valid_hours:
        raise ValueError("导入失败，负荷没有任何有效数值，无法用相邻点修复")
    repaired_count = 0
    previous_valid_hour = None
    next_valid_index = 0
    for hour in hours:
        if parsed_by_hour[hour].get("load") is not None:
            previous_valid_hour = hour
            if next_valid_index < len(valid_hours) and valid_hours[next_valid_index] == hour:
                next_valid_index += 1
            continue
        if previous_valid_hour is not None:
            source_hour = previous_valid_hour
        else:
            while next_valid_index < len(valid_hours) and valid_hours[next_valid_index] < hour:
                next_valid_index += 1
            source_hour = valid_hours[next_valid_index]
        parsed_by_hour[hour]["load"] = parsed_by_hour[source_hour]["load"]
        repaired_count += 1
    return repaired_count


def fill_imported_load_curve_hours(parsed_by_hour: dict[int, dict]) -> list[dict]:
    first_row = parsed_by_hour[min(parsed_by_hour)]
    previous = None
    imported_rows = []
    for hour in range(1, TIME_SERIES_IMPORT_ROW_COUNT + 1):
        if hour in parsed_by_hour:
            current = dict(parsed_by_hour[hour])
            current["hour_index"] = hour
            previous = current
            imported_rows.append(current)
            continue
        base = previous or first_row
        filled = {"hour_index": hour, "datetime": f"H{hour:04d}", "load": base["load"], "_filled": True}
        previous = filled
        imported_rows.append(filled)
    return imported_rows


def repair_load_value_sequence(values: list[float | int | None]) -> tuple[list[float | int], int]:
    valid_indexes = [index for index, value in enumerate(values) if value is not None]
    if not valid_indexes:
        raise ValueError("导入失败，负荷没有任何有效数值，无法用相邻点修复")
    repaired = list(values)
    repaired_count = 0
    previous_valid_index = None
    next_valid_index = 0
    for index, value in enumerate(repaired):
        if value is not None:
            previous_valid_index = index
            if next_valid_index < len(valid_indexes) and valid_indexes[next_valid_index] == index:
                next_valid_index += 1
            continue
        if previous_valid_index is not None:
            source_index = previous_valid_index
        else:
            source_index = valid_indexes[next_valid_index]
        repaired[index] = repaired[source_index]
        repaired_count += 1
    return [value for value in repaired if value is not None], repaired_count


def resample_load_values(values: list[float | int], target_count: int) -> list[float]:
    if len(values) == target_count:
        return [float(value) for value in values]
    if len(values) == 1:
        return [float(values[0]) for _ in range(target_count)]
    source_last = len(values) - 1
    target_last = target_count - 1
    resampled = []
    for index in range(target_count):
        source_position = (index / target_last) * source_last
        left_index = int(math.floor(source_position))
        right_index = min(source_last, left_index + 1)
        ratio = source_position - left_index
        left_value = float(values[left_index])
        right_value = float(values[right_index])
        resampled.append(left_value + (right_value - left_value) * ratio)
    return resampled


def scale_load_values_to_targets(values: list[float | int], min_value: float, max_value: float, avg_value: float) -> list[float | int]:
    if max_value == min_value:
        return [round_load_value(min_value) for _ in range(TIME_SERIES_IMPORT_ROW_COUNT)]
    raw_min = min(float(value) for value in values)
    raw_max = max(float(value) for value in values)
    raw_span = raw_max - raw_min
    if raw_span <= 0:
        raise ValueError("导入负荷曲线没有变化，无法按指定最大值和最小值缩放")
    target_mean = (avg_value - min_value) / (max_value - min_value)
    shape = [(float(value) - raw_min) / raw_span for value in values]
    adjusted_shape = adjust_shape_mean(shape, target_mean)
    return [round_load_value(min_value + (max_value - min_value) * value) for value in adjusted_shape]


def match_time_series_import_columns(headers: list[str]) -> dict[str, str]:
    normalized_headers = [(normalize_import_header(header), header) for header in headers if str(header or "").strip()]
    column_map: dict[str, str] = {}
    missing = []
    for key, (display_name, aliases) in TIME_SERIES_IMPORT_REQUIRED_COLUMNS.items():
        matched = match_time_series_header(normalized_headers, aliases)
        if matched:
            column_map[key] = matched
        else:
            missing.append(display_name)
    for key, aliases in TIME_SERIES_IMPORT_OPTIONAL_COLUMNS.items():
        matched = match_time_series_header(normalized_headers, aliases)
        if matched:
            column_map[key] = matched
    if missing:
        raise ValueError(f"导入失败，找不到对应的列：{', '.join(missing)}")
    return column_map


def match_time_series_header(normalized_headers: list[tuple[str, str]], aliases: list[str]) -> str:
    normalized_aliases = [normalize_import_header(alias) for alias in aliases]
    for alias in normalized_aliases:
        for header_norm, header in normalized_headers:
            if header_norm == alias:
                return header
    for alias in normalized_aliases:
        if len(alias) < 2:
            continue
        for header_norm, header in normalized_headers:
            if alias in header_norm or header_norm in alias:
                return header
    return ""


def normalize_import_header(value: object) -> str:
    text = str(value or "").strip().lower()
    return re.sub(r"[\s_\-^（）()［\]\[\]/\\:：,，。.%％]+", "", text)


def imported_hour_index(raw_row: dict[str, object], column: str | None, fallback_index: int) -> int:
    if not column:
        return fallback_index
    value = raw_row.get(column)
    if value in ("", None):
        return fallback_index
    if isinstance(value, datetime):
        day_of_year = value.timetuple().tm_yday
        return (day_of_year - 1) * 24 + value.hour + 1
    if isinstance(value, (int, float)) and math.isfinite(float(value)):
        return int(value)
    text = str(value).strip()
    if not text:
        return fallback_index
    hour_match = re.fullmatch(r"[hH]\s*0*(\d{1,4})", text)
    if hour_match:
        return int(hour_match.group(1))
    numeric_match = re.fullmatch(r"0*(\d{1,4})(?:\.0+)?", text)
    if numeric_match:
        return int(numeric_match.group(1))
    datetime_match = re.search(r"\b(\d{1,2})[:：](\d{1,2})(?::\d{1,2})?\b", text)
    date_match = re.search(r"(\d{4})[-/年.](\d{1,2})[-/月.](\d{1,2})", text)
    if datetime_match and date_match:
        try:
            parsed = datetime(
                int(date_match.group(1)),
                int(date_match.group(2)),
                int(date_match.group(3)),
                int(datetime_match.group(1)),
            )
            return (parsed.timetuple().tm_yday - 1) * 24 + parsed.hour + 1
        except ValueError:
            return fallback_index
    return fallback_index


def next_available_time_series_hour(parsed_by_hour: dict[int, dict], fallback_index: int) -> int:
    hour = max(1, min(TIME_SERIES_IMPORT_ROW_COUNT, fallback_index))
    while hour in parsed_by_hour and hour <= TIME_SERIES_IMPORT_ROW_COUNT:
        hour += 1
    if hour <= TIME_SERIES_IMPORT_ROW_COUNT:
        return hour
    for candidate in range(1, TIME_SERIES_IMPORT_ROW_COUNT + 1):
        if candidate not in parsed_by_hour:
            return candidate
    return TIME_SERIES_IMPORT_ROW_COUNT


def fill_imported_time_series_hours(parsed_by_hour: dict[int, dict]) -> list[dict]:
    first_row = parsed_by_hour[min(parsed_by_hour)]
    previous = None
    imported_rows = []
    for hour in range(1, TIME_SERIES_IMPORT_ROW_COUNT + 1):
        if hour in parsed_by_hour:
            current = dict(parsed_by_hour[hour])
            current["hour_index"] = hour
            current["datetime"] = imported_datetime_for_output(current.get("datetime"), hour)
            previous = current
            imported_rows.append(current)
            continue
        base = previous or first_row
        filled = {key: base[key] for key in TIME_SERIES_IMPORT_REQUIRED_COLUMNS}
        filled.update({"hour_index": hour, "datetime": f"H{hour:04d}", "_filled": True})
        previous = filled
        imported_rows.append(filled)
    return imported_rows


def imported_datetime_for_output(value: object, index: int) -> str:
    text = str(value or "").strip()
    return text or f"H{index:04d}"


def imported_datetime(raw_row: dict[str, object], column: str | None, index: int) -> str:
    if column:
        value = raw_row.get(column)
        if value not in ("", None):
            if isinstance(value, datetime):
                return value.isoformat(sep=" ", timespec="minutes")
            return str(value)
    return f"H{index:04d}"


def imported_numeric_value_or_none(value: object) -> float | int | None:
    if value in ("", None) or not str(value).strip():
        return None
    try:
        number = float(value)
    except (TypeError, ValueError):
        return None
    if not math.isfinite(number):
        return None
    return int(number) if number.is_integer() else number


def repair_imported_numeric_values(parsed_by_hour: dict[int, dict]) -> int:
    hours = sorted(parsed_by_hour)
    repaired_count = 0
    for key, (label, _) in TIME_SERIES_IMPORT_REQUIRED_COLUMNS.items():
        valid_hours = [hour for hour in hours if parsed_by_hour[hour].get(key) is not None]
        if not valid_hours:
            raise ValueError(f"导入失败，{label}没有任何有效数值，无法用相邻点修复")
        previous_valid_hour = None
        next_valid_index = 0
        for hour in hours:
            if parsed_by_hour[hour].get(key) is not None:
                previous_valid_hour = hour
                if next_valid_index < len(valid_hours) and valid_hours[next_valid_index] == hour:
                    next_valid_index += 1
                continue
            if previous_valid_hour is not None:
                source_hour = previous_valid_hour
            else:
                while next_valid_index < len(valid_hours) and valid_hours[next_valid_index] < hour:
                    next_valid_index += 1
                source_hour = valid_hours[next_valid_index]
            parsed_by_hour[hour][key] = parsed_by_hour[source_hour][key]
            repaired_count += 1
    return repaired_count


def generate_load_curve(mode: str, minimum: object, maximum: object, average: object, source_load_curve: object = None) -> dict:
    mode_key = str(mode or "random").strip() or "random"
    min_value, max_value, avg_value = validate_load_curve_targets(minimum, maximum, average)
    if max_value == min_value:
        values = [round_load_value(min_value) for _ in range(TIME_SERIES_IMPORT_ROW_COUNT)]
    else:
        shape = normalized_source_load_shape(source_load_curve) if mode_key == "file" else normalized_load_shape(mode_key)
        adjusted_shape = adjust_shape_mean(shape, (avg_value - min_value) / (max_value - min_value))
        values = [round_load_value(min_value + (max_value - min_value) * value) for value in adjusted_shape]
    rows = [{"hour_index": index + 1, "load": value} for index, value in enumerate(values)]
    return {
        "mode": mode_key,
        "load_curve": rows,
        "load_curve_count": len(rows),
        "statistics": {
            "max": round_load_value(max(values)),
            "min": round_load_value(min(values)),
            "average": round_load_value(sum(values) / len(values)),
        },
    }


CURVE_GENERATION_FIELDS = {
    "wind_speed": {
        "label": "风速",
        "row_key": "wind_speed_curve",
        "count_key": "wind_speed_curve_count",
        "average": "风速平均值",
        "minimum": "风速最小值",
        "maximum": "风速最大值",
    },
    "solar_irradiance": {
        "label": "太阳辐射",
        "row_key": "solar_irradiance_curve",
        "count_key": "solar_irradiance_curve_count",
        "average": "太阳辐射平均值",
        "minimum": "太阳辐射最小值",
        "maximum": "太阳辐射最大值",
    },
}


def generate_time_series_curve(
    curve: str,
    mode: str,
    minimum: object,
    maximum: object,
    average: object,
    source_curve: object = None,
) -> dict:
    curve_key, spec = curve_generation_spec(curve)
    mode_key = str(mode or "random").strip() or "random"
    min_value, max_value, avg_value = validate_curve_targets(curve_key, minimum, maximum, average)
    if max_value == min_value:
        values = [round_load_value(min_value) for _ in range(TIME_SERIES_IMPORT_ROW_COUNT)]
    else:
        shape = normalized_source_curve_shape(curve_key, source_curve) if mode_key == "file" else normalized_time_series_curve_shape(curve_key, mode_key)
        adjusted_shape = adjust_shape_mean(shape, (avg_value - min_value) / (max_value - min_value))
        values = [round_load_value(min_value + (max_value - min_value) * value) for value in adjusted_shape]
    rows = [{"hour_index": index + 1, curve_key: value} for index, value in enumerate(values)]
    return {
        "curve": curve_key,
        "mode": mode_key,
        "curve_data": rows,
        spec["row_key"]: rows,
        spec["count_key"]: len(rows),
        "curve_count": len(rows),
        "statistics": {
            "max": round_load_value(max(values)),
            "min": round_load_value(min(values)),
            "average": round_load_value(sum(values) / len(values)),
        },
    }


def import_time_series_curve_file(
    curve: str,
    filename: str,
    content: bytes,
    minimum: object,
    maximum: object,
    average: object,
    raw: object = False,
) -> dict:
    curve_key, spec = curve_generation_spec(curve)
    suffix = Path(filename or "").suffix.lower()
    if suffix == ".csv":
        headers, rows = read_time_series_csv(content)
    elif suffix == ".xlsx":
        headers, rows = read_time_series_xlsx(content)
    else:
        raise ValueError("导入文件仅支持 .csv 或 .xlsx 格式")
    if not rows:
        raise ValueError("导入失败，文件没有可用数据行")
    column_map = match_curve_import_columns(headers, curve_key)
    values, repaired_numeric_count, missing_count, duplicate_count = normalized_imported_curve_values(rows, column_map, curve_key)
    if truthy_json_value(raw):
        scaled_values = [round_load_value(float(value)) for value in values]
    else:
        min_value, max_value, avg_value = validate_curve_targets(curve_key, minimum, maximum, average)
        scaled_values = scale_curve_values_to_targets(curve_key, values, min_value, max_value, avg_value)
    output_rows = [{"hour_index": index + 1, curve_key: value} for index, value in enumerate(scaled_values)]
    message = f"已从{filename}导入8760点{spec['label']}{'原始' if truthy_json_value(raw) else ''}曲线"
    if len(rows) != TIME_SERIES_IMPORT_ROW_COUNT:
        message += f"，文件共有{len(rows)}行，已自适应扩展到8760点"
    if missing_count:
        message += f"，已按相邻点自动补齐{missing_count}个缺失时点"
    if duplicate_count:
        message += f"，已合并{duplicate_count}个重复小时数据"
    if repaired_numeric_count:
        message += f"，已修复{repaired_numeric_count}个无效数值"
    return {
        "source": "file",
        "curve": curve_key,
        "curve_data": output_rows,
        spec["row_key"]: output_rows,
        spec["count_key"]: len(output_rows),
        "curve_count": len(output_rows),
        "statistics": {
            "max": round_load_value(max(scaled_values)),
            "min": round_load_value(min(scaled_values)),
            "average": round_load_value(sum(scaled_values) / len(scaled_values)),
        },
        "message": message,
    }


def curve_generation_spec(curve: str) -> tuple[str, dict[str, str]]:
    curve_key = str(curve or "").strip()
    if curve_key not in CURVE_GENERATION_FIELDS:
        raise ValueError("曲线类型必须为风速或太阳辐射")
    return curve_key, CURVE_GENERATION_FIELDS[curve_key]


def validate_curve_targets(curve: str, minimum: object, maximum: object, average: object) -> tuple[float, float, float]:
    _, spec = curve_generation_spec(curve)
    min_value = load_curve_number(minimum, spec["minimum"])
    max_value = load_curve_number(maximum, spec["maximum"])
    avg_value = load_curve_number(average, spec["average"])
    if min_value < 0 or max_value < 0 or avg_value < 0:
        raise ValueError(f"{spec['maximum']}、{spec['minimum']}、{spec['average']}必须为非负数")
    if max_value < min_value:
        raise ValueError(f"{spec['maximum']}不能小于{spec['minimum']}")
    if max_value == min_value:
        if avg_value != min_value:
            raise ValueError("最大值等于最小值时，平均值必须与其相等")
        return min_value, max_value, avg_value
    low_mean = min_value + (max_value - min_value) / TIME_SERIES_IMPORT_ROW_COUNT
    high_mean = max_value - (max_value - min_value) / TIME_SERIES_IMPORT_ROW_COUNT
    if avg_value < low_mean or avg_value > high_mean:
        raise ValueError("平均值必须介于最小值和最大值之间，并能同时满足最大/最小约束")
    return min_value, max_value, avg_value


def normalized_time_series_curve_shape(curve: str, mode: str) -> list[float]:
    if mode != "random":
        raise ValueError("风速和太阳辐射生成模式目前支持随机曲线或文件导入")
    raw = deterministic_wind_shape() if curve == "wind_speed" else deterministic_solar_irradiance_shape()
    minimum = min(raw)
    maximum = max(raw)
    span = maximum - minimum
    if span <= 0:
        return [0.5 for _ in raw]
    return [(value - minimum) / span for value in raw]


def normalized_source_curve_shape(curve: str, source_curve: object) -> list[float]:
    values = normalize_curve_values(curve, source_curve)
    minimum = min(float(value) for value in values)
    maximum = max(float(value) for value in values)
    span = maximum - minimum
    if span <= 0:
        raise ValueError("文件导入原始曲线没有变化，无法按指定最大值和最小值缩放")
    return [(float(value) - minimum) / span for value in values]


def normalize_curve_values(curve: str, rows: object) -> list[float | int]:
    curve_key, spec = curve_generation_spec(curve)
    if not isinstance(rows, list) or len(rows) != TIME_SERIES_IMPORT_ROW_COUNT:
        raise ValueError(f"{spec['label']}曲线必须包含{TIME_SERIES_IMPORT_ROW_COUNT}点")
    values = []
    for index, row in enumerate(rows, start=1):
        raw_value = row.get(curve_key) if isinstance(row, dict) else row
        value = imported_numeric_value_or_none(raw_value)
        if value is None:
            raise ValueError(f"{spec['label']}曲线第{index}点不是有效数值")
        values.append(round_load_value(value))
    return values


def match_curve_import_columns(headers: list[str], curve: str) -> dict[str, str]:
    curve_key, spec = curve_generation_spec(curve)
    normalized_headers = [(normalize_import_header(header), header) for header in headers if str(header or "").strip()]
    aliases = TIME_SERIES_IMPORT_REQUIRED_COLUMNS[curve_key][1]
    value_column = match_time_series_header(normalized_headers, aliases)
    if not value_column:
        raise ValueError(f"导入失败，找不到对应的列：{spec['label']}")
    column_map = {curve_key: value_column}
    matched_time = match_time_series_header(normalized_headers, TIME_SERIES_IMPORT_OPTIONAL_COLUMNS["datetime"])
    if matched_time:
        column_map["datetime"] = matched_time
    return column_map


def normalized_imported_curve_values(
    raw_rows: list[dict[str, object]],
    column_map: dict[str, str],
    curve: str,
) -> tuple[list[float | int], int, int, int]:
    curve_column = column_map[curve]
    time_column = column_map.get("datetime")
    raw_values = [imported_numeric_value_or_none(row.get(curve_column)) for row in raw_rows]
    grouped_by_hour: dict[int, list[float | int | None]] = {}
    if time_column:
        for raw_row, value in zip(raw_rows, raw_values):
            target_hour = imported_load_curve_hour_index(raw_row, time_column)
            if target_hour is None or target_hour < 1 or target_hour > TIME_SERIES_IMPORT_ROW_COUNT:
                continue
            grouped_by_hour.setdefault(target_hour, []).append(value)
    if grouped_by_hour:
        parsed_by_hour = {}
        duplicate_count = 0
        for hour, hour_values in grouped_by_hour.items():
            duplicate_count += max(0, len(hour_values) - 1)
            valid_values = [float(value) for value in hour_values if value is not None]
            parsed_by_hour[hour] = {
                "hour_index": hour,
                "datetime": f"H{hour:04d}",
                curve: sum(valid_values) / len(valid_values) if valid_values else None,
            }
        repaired_numeric_count = repair_imported_curve_values(parsed_by_hour, curve)
        imported_rows = fill_imported_curve_hours(parsed_by_hour, curve)
        missing_count = sum(1 for item in imported_rows if item.get("_filled"))
        source_values = [row[curve] for row in imported_rows[: max(grouped_by_hour)]]
        return resample_load_values(source_values, TIME_SERIES_IMPORT_ROW_COUNT), repaired_numeric_count, missing_count, duplicate_count
    repaired_values, repaired_numeric_count = repair_curve_value_sequence(raw_values, curve)
    return resample_load_values(repaired_values, TIME_SERIES_IMPORT_ROW_COUNT), repaired_numeric_count, 0, 0


def repair_imported_curve_values(parsed_by_hour: dict[int, dict], curve: str) -> int:
    curve_key, spec = curve_generation_spec(curve)
    hours = sorted(parsed_by_hour)
    valid_hours = [hour for hour in hours if parsed_by_hour[hour].get(curve_key) is not None]
    if not valid_hours:
        raise ValueError(f"导入失败，{spec['label']}没有任何有效数值，无法用相邻点修复")
    repaired_count = 0
    previous_valid_hour = None
    next_valid_index = 0
    for hour in hours:
        if parsed_by_hour[hour].get(curve_key) is not None:
            previous_valid_hour = hour
            if next_valid_index < len(valid_hours) and valid_hours[next_valid_index] == hour:
                next_valid_index += 1
            continue
        if previous_valid_hour is not None:
            source_hour = previous_valid_hour
        else:
            while next_valid_index < len(valid_hours) and valid_hours[next_valid_index] < hour:
                next_valid_index += 1
            source_hour = valid_hours[next_valid_index]
        parsed_by_hour[hour][curve_key] = parsed_by_hour[source_hour][curve_key]
        repaired_count += 1
    return repaired_count


def fill_imported_curve_hours(parsed_by_hour: dict[int, dict], curve: str) -> list[dict]:
    first_row = parsed_by_hour[min(parsed_by_hour)]
    previous = None
    imported_rows = []
    for hour in range(1, TIME_SERIES_IMPORT_ROW_COUNT + 1):
        if hour in parsed_by_hour:
            current = dict(parsed_by_hour[hour])
            current["hour_index"] = hour
            previous = current
            imported_rows.append(current)
            continue
        base = previous or first_row
        filled = {"hour_index": hour, "datetime": f"H{hour:04d}", curve: base[curve], "_filled": True}
        previous = filled
        imported_rows.append(filled)
    return imported_rows


def repair_curve_value_sequence(values: list[float | int | None], curve: str) -> tuple[list[float | int], int]:
    _, spec = curve_generation_spec(curve)
    valid_indexes = [index for index, value in enumerate(values) if value is not None]
    if not valid_indexes:
        raise ValueError(f"导入失败，{spec['label']}没有任何有效数值，无法用相邻点修复")
    repaired = list(values)
    repaired_count = 0
    previous_valid_index = None
    next_valid_index = 0
    for index, value in enumerate(repaired):
        if value is not None:
            previous_valid_index = index
            if next_valid_index < len(valid_indexes) and valid_indexes[next_valid_index] == index:
                next_valid_index += 1
            continue
        if previous_valid_index is not None:
            source_index = previous_valid_index
        else:
            source_index = valid_indexes[next_valid_index]
        repaired[index] = repaired[source_index]
        repaired_count += 1
    return [value for value in repaired if value is not None], repaired_count


def scale_curve_values_to_targets(curve: str, values: list[float | int], min_value: float, max_value: float, avg_value: float) -> list[float | int]:
    if max_value == min_value:
        return [round_load_value(min_value) for _ in range(TIME_SERIES_IMPORT_ROW_COUNT)]
    shape = normalized_source_curve_shape(curve, [{"hour_index": index + 1, curve: value} for index, value in enumerate(values)])
    adjusted_shape = adjust_shape_mean(shape, (avg_value - min_value) / (max_value - min_value))
    return [round_load_value(min_value + (max_value - min_value) * value) for value in adjusted_shape]


def list_load_curve_templates() -> dict:
    templates = read_load_curve_templates()
    return {"templates": load_curve_template_summaries(templates)}


def save_load_curve_template(name: str, rows: object, overwrite: object = False) -> dict:
    clean_name = planning_store.sanitize_scheme_name(str(name or ""))
    if clean_name in {"", ".", ".."} or planning_store.INVALID_NAME_RE.search(clean_name) or ".." in clean_name:
        raise ValueError("模板名称不能为空，且不能包含路径或非法字符")
    values = normalize_load_curve_template_values(rows)
    templates = read_load_curve_templates()
    exists = any(template.get("name") == clean_name for template in templates)
    if exists and not truthy_json_value(overwrite):
        raise LoadCurveTemplateExistsError(f"模板名称已存在：{clean_name}")
    next_template = {
        "name": clean_name,
        "load_curve": values,
        "load_curve_count": len(values),
        "updated_at": _now_iso(),
    }
    next_templates = [template for template in templates if template.get("name") != clean_name]
    next_templates.append(next_template)
    next_templates.sort(key=lambda item: item["name"])
    write_load_curve_templates(next_templates)
    return {
        "template": load_curve_template_summary(next_template),
        "templates": load_curve_template_summaries(next_templates),
        "message": f"{'已覆盖' if exists else '已保存'}负荷模板：{clean_name}",
    }


def read_load_curve_templates() -> list[dict]:
    if not LOAD_CURVE_TEMPLATE_PATH.exists():
        return []
    return LOAD_CURVE_TEMPLATE_CACHE.get(
        LOAD_CURVE_TEMPLATE_PATH,
        read_load_curve_templates_from_path,
        variant="load_curve_templates",
    )


def read_load_curve_templates_from_path(path: Path) -> list[dict]:
    try:
        with path.open("r", encoding="utf-8-sig", newline="") as handle:
            reader = csv.DictReader(handle)
            headers = [header for header in (reader.fieldnames or []) if str(header or "").strip()]
            template_names = [header for header in headers if header != "hour_index"]
            columns = {name: [] for name in template_names}
            for row in reader:
                for name in template_names:
                    columns[name].append(row.get(name, ""))
    except OSError as exc:
        raise ValueError(f"负荷模板文件读取失败：{exc}") from exc
    normalized = []
    for name, values in columns.items():
        clean_name = planning_store.sanitize_scheme_name(str(name))
        try:
            normalized_values = normalize_load_curve_template_values(values)
        except ValueError:
            continue
        if clean_name:
            normalized.append(
                {
                    "name": clean_name,
                    "load_curve": normalized_values,
                    "load_curve_count": len(normalized_values),
                    "updated_at": "",
                }
            )
    normalized.sort(key=lambda item: item["name"])
    return normalized


def write_load_curve_templates(templates: list[dict]) -> None:
    LOAD_CURVE_TEMPLATE_PATH.parent.mkdir(parents=True, exist_ok=True)
    tmp_path = LOAD_CURVE_TEMPLATE_PATH.with_name(f".{LOAD_CURVE_TEMPLATE_PATH.name}.tmp")
    headers = ["hour_index", *[template["name"] for template in templates]]
    try:
        with tmp_path.open("w", encoding="utf-8-sig", newline="") as handle:
            writer = csv.DictWriter(handle, fieldnames=headers)
            writer.writeheader()
            for index in range(TIME_SERIES_IMPORT_ROW_COUNT):
                row = {"hour_index": index + 1}
                for template in templates:
                    row[template["name"]] = template["load_curve"][index]
                writer.writerow(row)
        file_ops.replace_file_with_retry(tmp_path, LOAD_CURVE_TEMPLATE_PATH, "负荷模板文件")
    except Exception:
        tmp_path.unlink(missing_ok=True)
        raise


def normalize_load_curve_template_values(rows: object) -> list[float | int]:
    if not isinstance(rows, list) or len(rows) != TIME_SERIES_IMPORT_ROW_COUNT:
        raise ValueError(f"负荷模板必须包含{TIME_SERIES_IMPORT_ROW_COUNT}点曲线")
    values = []
    for index, row in enumerate(rows, start=1):
        raw_value = row.get("load") if isinstance(row, dict) else row
        value = imported_numeric_value_or_none(raw_value)
        if value is None:
            raise ValueError(f"负荷模板第{index}点不是有效数值")
        values.append(round_load_value(value))
    return values


def load_curve_template_summaries(templates: list[dict]) -> list[dict]:
    return [load_curve_template_summary(template) for template in templates]


def load_curve_template_summary(template: dict) -> dict:
    values = template.get("load_curve") or []
    return {
        "name": template.get("name", ""),
        "load_curve_count": int(template.get("load_curve_count") or TIME_SERIES_IMPORT_ROW_COUNT),
        "load_curve": [{"hour_index": index + 1, "load": round_load_value(value)} for index, value in enumerate(values)],
        "updated_at": template.get("updated_at", ""),
    }


def load_curve_template_shape(mode: str) -> list[float]:
    template_name = mode.split(":", 1)[1].strip() if ":" in mode else ""
    if not template_name:
        raise ValueError("负荷模板名称不能为空")
    clean_name = planning_store.sanitize_scheme_name(template_name)
    for template in read_load_curve_templates():
        if template.get("name") == clean_name:
            values = template["load_curve"]
            minimum = min(float(value) for value in values)
            maximum = max(float(value) for value in values)
            span = maximum - minimum
            if span <= 0:
                return [0.5 for _ in values]
            return [(float(value) - minimum) / span for value in values]
    raise ValueError(f"负荷模板不存在：{clean_name}")


def validate_load_curve_targets(minimum: object, maximum: object, average: object) -> tuple[float, float, float]:
    min_value = load_curve_number(minimum, "负荷最小值")
    max_value = load_curve_number(maximum, "负荷最大值")
    avg_value = load_curve_number(average, "负荷平均值")
    if min_value < 0 or max_value < 0 or avg_value < 0:
        raise ValueError("负荷最大值、负荷最小值、负荷平均值必须为非负数")
    if max_value < min_value:
        raise ValueError("负荷最大值不能小于负荷最小值")
    if max_value == min_value:
        if avg_value != min_value:
            raise ValueError("最大值等于最小值时，平均值必须与其相等")
        return min_value, max_value, avg_value
    low_mean = min_value + (max_value - min_value) / TIME_SERIES_IMPORT_ROW_COUNT
    high_mean = max_value - (max_value - min_value) / TIME_SERIES_IMPORT_ROW_COUNT
    if avg_value < low_mean or avg_value > high_mean:
        raise ValueError("平均值必须介于最小值和最大值之间，并能同时满足最大/最小约束")
    return min_value, max_value, avg_value


def load_curve_number(value: object, label: str) -> float:
    try:
        number = float(value)
    except (TypeError, ValueError) as exc:
        raise ValueError(f"{label}必须为数值") from exc
    if not math.isfinite(number):
        raise ValueError(f"{label}必须为有效数值")
    return number


def normalized_load_shape(mode: str) -> list[float]:
    if mode.startswith("template:"):
        return load_curve_template_shape(mode)
    if mode == "random":
        raw = deterministic_random_shape()
    elif mode in {"pattern1", "pattern2", "pattern3"}:
        raw = standard_load_shape(mode)
    else:
        raise ValueError("负荷生成模式必须为随机曲线、文件导入、模式1、模式2、模式3或已保存模板")
    minimum = min(raw)
    maximum = max(raw)
    span = maximum - minimum
    if span <= 0:
        return [0.5 for _ in raw]
    return [(value - minimum) / span for value in raw]


def normalized_source_load_shape(source_load_curve: object) -> list[float]:
    values = normalize_load_curve_template_values(source_load_curve)
    minimum = min(float(value) for value in values)
    maximum = max(float(value) for value in values)
    span = maximum - minimum
    if span <= 0:
        raise ValueError("文件导入原始曲线没有变化，无法按指定最大值和最小值缩放")
    return [(float(value) - minimum) / span for value in values]


def deterministic_random_shape() -> list[float]:
    seed = 2463534242
    values = []
    previous = 0.5
    for hour in range(TIME_SERIES_IMPORT_ROW_COUNT):
        seed ^= (seed << 13) & 0xFFFFFFFF
        seed ^= seed >> 17
        seed ^= (seed << 5) & 0xFFFFFFFF
        random_value = (seed & 0xFFFFFFFF) / 0xFFFFFFFF
        previous = previous * 0.72 + random_value * 0.28
        daily = 0.5 + 0.28 * math.sin((hour % 24 - 7) / 24 * 2 * math.pi)
        values.append(previous * 0.65 + daily * 0.35 + hour * 1e-9)
    return values


def deterministic_wind_shape() -> list[float]:
    seed = 362436069
    values = []
    previous = 0.55
    for hour in range(TIME_SERIES_IMPORT_ROW_COUNT):
        seed ^= (seed << 13) & 0xFFFFFFFF
        seed ^= seed >> 17
        seed ^= (seed << 5) & 0xFFFFFFFF
        random_value = (seed & 0xFFFFFFFF) / 0xFFFFFFFF
        previous = previous * 0.86 + random_value * 0.14
        day = hour // 24
        hour_of_day = hour % 24
        season = 0.5 + 0.5 * math.cos((day - 20) / 365 * 2 * math.pi)
        diurnal = 0.5 + 0.5 * math.sin((hour_of_day - 14) / 24 * 2 * math.pi)
        values.append(0.18 + 0.54 * previous + 0.22 * season + 0.06 * diurnal + hour * 1e-9)
    return values


def deterministic_solar_irradiance_shape() -> list[float]:
    seed = 521288629
    cloud = 0.72
    values = []
    for hour in range(TIME_SERIES_IMPORT_ROW_COUNT):
        seed ^= (seed << 13) & 0xFFFFFFFF
        seed ^= seed >> 17
        seed ^= (seed << 5) & 0xFFFFFFFF
        random_value = (seed & 0xFFFFFFFF) / 0xFFFFFFFF
        cloud = min(1.0, max(0.35, cloud * 0.90 + random_value * 0.10))
        day = hour // 24
        hour_of_day = hour % 24
        season = 0.5 + 0.5 * math.sin((day - 80) / 365 * 2 * math.pi)
        daylight_hours = 4.0 + 16.0 * season
        distance_from_noon = abs(hour_of_day + 0.5 - 12.0)
        half_day = daylight_hours / 2.0
        if distance_from_noon <= half_day:
            daylight = math.cos((distance_from_noon / max(0.1, half_day)) * math.pi / 2) ** 1.35
            value = daylight * (0.28 + 0.72 * season) * (0.62 + 0.38 * cloud)
        else:
            # Keep a tiny non-zero background so the mean-scaling step remains
            # well-conditioned while still rendering night hours near zero.
            value = 0.0002 * (0.8 + 0.2 * cloud)
        values.append(value + hour * 1e-12)
    return values


def standard_load_shape(mode: str) -> list[float]:
    values = []
    for hour in range(TIME_SERIES_IMPORT_ROW_COUNT):
        day = hour // 24
        hour_of_day = hour % 24
        season = math.sin((day - 20) / 365 * 2 * math.pi)
        weekday_factor = 1.0 if day % 7 < 5 else 0.88
        morning_peak = math.exp(-((hour_of_day - 8) / 3.0) ** 2)
        evening_peak = math.exp(-((hour_of_day - 19) / 3.5) ** 2)
        noon_peak = math.exp(-((hour_of_day - 13) / 4.0) ** 2)
        night_valley = math.exp(-((hour_of_day - 3) / 4.0) ** 2)
        if mode == "pattern1":
            value = (0.56 + 0.20 * morning_peak + 0.34 * evening_peak - 0.10 * night_valley + 0.09 * season) * weekday_factor
        elif mode == "pattern2":
            value = (0.66 + 0.24 * noon_peak - 0.08 * night_valley + 0.05 * season) * (0.96 + 0.04 * weekday_factor)
        else:
            winter_summer = abs(season)
            value = 0.50 + 0.24 * morning_peak + 0.22 * noon_peak + 0.22 * evening_peak + 0.18 * winter_summer - 0.12 * night_valley
        values.append(value + hour * 1e-9)
    return values


def adjust_shape_mean(shape: list[float], target_mean: float) -> list[float]:
    if not shape:
        raise ValueError("曲线形状不能为空")
    current_mean = sum(shape) / len(shape)
    if abs(current_mean - target_mean) < 1e-10:
        return shape
    if target_mean < current_mean:
        low, high = 1.0, 2.0
        min_reachable = sum(1 for value in shape if value >= 1.0 - 1e-12) / len(shape)
        if target_mean <= min_reachable + 1e-10:
            raise ValueError("平均值过低，无法在保持最大/最小值的同时生成该形状曲线")
        while mean_power(shape, high) > target_mean:
            high *= 2
            if high > 1_000_000:
                raise ValueError("平均值过低，无法在保持最大/最小值的同时生成该形状曲线")
        for _ in range(80):
            mid = (low + high) / 2
            if mean_power(shape, mid) > target_mean:
                low = mid
            else:
                high = mid
        power = high
        return [value**power for value in shape]
    low, high = 1.0, 2.0
    max_reachable = sum(1 for value in shape if value > 1e-12) / len(shape)
    if target_mean >= max_reachable - 1e-10:
        raise ValueError("平均值过高，无法在保持最大/最小值的同时生成该形状曲线")
    while mean_inverse_power(shape, high) < target_mean:
        high *= 2
        if high > 1_000_000:
            raise ValueError("平均值过高，无法在保持最大/最小值的同时生成该形状曲线")
    for _ in range(80):
        mid = (low + high) / 2
        if mean_inverse_power(shape, mid) < target_mean:
            low = mid
        else:
            high = mid
    power = high
    return [1 - (1 - value) ** power for value in shape]


def mean_power(shape: list[float], power: float) -> float:
    return sum(value**power for value in shape) / len(shape)


def mean_inverse_power(shape: list[float], power: float) -> float:
    return sum(1 - (1 - value) ** power for value in shape) / len(shape)


def round_load_value(value: float) -> float | int:
    rounded = round(float(value), 6)
    return int(rounded) if rounded.is_integer() else rounded


def user_has_admin_scope(current_user: dict | None) -> bool:
    return current_user is None or current_user.get("role") == "admin"


def current_username(current_user: dict | None) -> str:
    return str((current_user or {}).get("username") or "").strip()


def scheme_owner_filter_for_user(current_user: dict | None) -> str | None:
    if user_has_admin_scope(current_user):
        return None
    return current_username(current_user)


def scheme_visible_filter_for_user(current_user: dict | None) -> str | None:
    if user_has_admin_scope(current_user):
        return None
    return current_username(current_user)


def new_scheme_owner_for_user(current_user: dict | None) -> str | None:
    if current_user is None:
        return None
    return current_username(current_user)


def ensure_planning_scheme_access(name: str, current_user: dict | None) -> str:
    clean = planning_store.validate_scheme_name(name)
    if user_has_admin_scope(current_user):
        return clean
    username = current_username(current_user)
    if not PLANNING_STORE.scheme_dir(clean).exists() or not username or not PLANNING_STORE.user_can_read_scheme(clean, username):
        raise FileNotFoundError(f"方案不存在: {clean}")
    return clean


def ensure_planning_scheme_manage_access(name: str, current_user: dict | None) -> str:
    clean = planning_store.validate_scheme_name(name)
    if user_has_admin_scope(current_user):
        return clean
    username = current_username(current_user)
    if not PLANNING_STORE.scheme_dir(clean).exists() or not username or PLANNING_STORE.scheme_owner_username(clean) != username:
        raise FileNotFoundError(f"方案不存在: {clean}")
    return clean


def user_can_access_planning_scheme(name: str, current_user: dict | None) -> bool:
    try:
        ensure_planning_scheme_access(name, current_user)
    except (ValueError, FileNotFoundError):
        return False
    return True


def user_can_manage_planning_scheme(name: str, current_user: dict | None) -> bool:
    try:
        ensure_planning_scheme_manage_access(name, current_user)
    except (ValueError, FileNotFoundError):
        return False
    return True


def planning_scheme_access_error_response(
    name: str,
    current_user: dict | None,
) -> tuple[int, dict[str, str], bytes] | None:
    try:
        ensure_planning_scheme_access(name, current_user)
    except FileNotFoundError as exc:
        return _json_response({"error": "not_found", "message": str(exc)}, HTTPStatus.NOT_FOUND)
    except ValueError as exc:
        return _json_response({"error": "bad_request", "message": str(exc)}, HTTPStatus.BAD_REQUEST)
    return None


def ensure_planning_copy_target_access(name: str, current_user: dict | None) -> str:
    clean = planning_store.validate_scheme_name(name)
    if PLANNING_STORE.scheme_dir(clean).exists():
        ensure_planning_scheme_manage_access(clean, current_user)
    return clean


def planning_scheme_list_item_for_user(item: dict, current_user: dict | None) -> dict:
    name = str(item.get("name") or "")
    owner = str(item.get("owner_username") or "").strip()
    username = current_username(current_user)
    shared_with = planning_store.normalize_shared_usernames(item.get("shared_with_usernames"))
    can_manage = user_can_manage_planning_scheme(name, current_user)
    if user_has_admin_scope(current_user):
        access_level = "admin"
    elif owner == username:
        access_level = "owner"
    elif username in shared_with:
        access_level = "shared"
    else:
        access_level = ""
    return {
        **item,
        "shared_with_usernames": shared_with,
        "access_level": access_level,
        "can_manage": can_manage,
        "is_shared_with_me": bool(username and username in shared_with and owner != username),
    }


def planning_scheme_share_payload(name: str) -> dict:
    clean = planning_store.validate_scheme_name(name)
    meta = PLANNING_STORE.read_scheme_meta(clean)
    return {
        "scheme": clean,
        "owner_username": str(meta.get("owner_username") or "").strip(),
        "shared_with_usernames": planning_store.normalize_shared_usernames(meta.get("shared_with_usernames")),
    }


def existing_share_target_username(username: str, current_user: dict | None) -> str:
    target = USER_STORE.get_user_by_username(username)
    if not target:
        raise FileNotFoundError(f"用户不存在: {username}")
    target_username = str(target.get("username") or "").strip()
    if target_username and target_username == current_username(current_user):
        raise ValueError("不能将方案分享给自己")
    return target_username


def handle_planning_api_path(
    path: str,
    method: str = "GET",
    body: bytes = b"",
    current_user: dict | None = None,
) -> tuple[int, dict[str, str], bytes]:
    prefix = "/api/planning/schemes"
    try:
        if path == "/api/planning/time-series/import" and method == "POST":
            payload = _read_json_body(body)
            filename = str(payload.get("filename", ""))
            content_base64 = str(payload.get("content_base64", ""))
            try:
                content = base64.b64decode(content_base64, validate=True)
            except (binascii.Error, ValueError) as exc:
                raise ValueError("导入失败，文件内容无法解析") from exc
            return _json_response(import_time_series_file(filename, content))
        if path == "/api/planning/load-curve/generate" and method == "POST":
            payload = _read_json_body(body)
            return _json_response(
                generate_load_curve(
                    str(payload.get("mode", "random")),
                    payload.get("min"),
                    payload.get("max"),
                    payload.get("average"),
                    payload.get("source_load_curve"),
                )
            )
        if path == "/api/planning/load-curve/import" and method == "POST":
            payload = _read_json_body(body)
            filename = str(payload.get("filename", ""))
            content_base64 = str(payload.get("content_base64", ""))
            try:
                content = base64.b64decode(content_base64, validate=True)
            except (binascii.Error, ValueError) as exc:
                raise ValueError("导入失败，文件内容无法解析") from exc
            return _json_response(
                import_load_curve_file(
                    filename,
                    content,
                    payload.get("min"),
                    payload.get("max"),
                    payload.get("average"),
                    payload.get("raw"),
                )
            )
        if path == "/api/planning/time-series-curve/generate" and method == "POST":
            payload = _read_json_body(body)
            return _json_response(
                generate_time_series_curve(
                    str(payload.get("curve", "")),
                    str(payload.get("mode", "random")),
                    payload.get("min"),
                    payload.get("max"),
                    payload.get("average"),
                    payload.get("source_curve"),
                )
            )
        if path == "/api/planning/time-series-curve/import" and method == "POST":
            payload = _read_json_body(body)
            filename = str(payload.get("filename", ""))
            content_base64 = str(payload.get("content_base64", ""))
            try:
                content = base64.b64decode(content_base64, validate=True)
            except (binascii.Error, ValueError) as exc:
                raise ValueError("导入失败，文件内容无法解析") from exc
            return _json_response(
                import_time_series_curve_file(
                    str(payload.get("curve", "")),
                    filename,
                    content,
                    payload.get("min"),
                    payload.get("max"),
                    payload.get("average"),
                    payload.get("raw"),
                )
            )
        if path == "/api/planning/load-curve/templates" and method == "GET":
            return _json_response(list_load_curve_templates())
        if path == "/api/planning/load-curve/templates" and method == "POST":
            payload = _read_json_body(body)
            return _json_response(
                save_load_curve_template(
                    str(payload.get("name", "")),
                    payload.get("load_curve"),
                    payload.get("overwrite"),
                )
            )
        if path == "/api/planning/map-config" and method == "GET":
            providers = [
                {"key": "amap", "label": "高德地图", "enabled": bool(AMAP_WEB_SERVICE_KEY)},
                {"key": "osm", "label": "OpenStreetMap", "enabled": True},
            ]
            preferred = next((provider["key"] for provider in providers if provider["enabled"]), "manual")
            return _json_response(
                {
                    "amap_key": AMAP_WEB_SERVICE_KEY,
                    "osm_key": "",
                    "providers": providers,
                    "preferred_provider": preferred,
                }
            )
        if path == "/api/planning/weather-history" and method == "POST":
            payload = _read_json_body(body)
            return _json_response(
                fetch_weather_history(payload.get("latitude"), payload.get("longitude"), payload.get("year"))
            )
        if path == "/api/planning/geocode" and method == "POST":
            payload = _read_json_body(body)
            return _json_response(geocode_place_name(str(payload.get("place", ""))))
        if path == "/api/planning/reverse-geocode" and method == "POST":
            payload = _read_json_body(body)
            return _json_response(reverse_geocode_coordinates(payload.get("latitude"), payload.get("longitude")))
        if path == prefix and method == "GET":
            schemes = [
                planning_scheme_list_item_for_user(item, current_user)
                for item in PLANNING_STORE.list_schemes(visible_username=scheme_visible_filter_for_user(current_user))
            ]
            return _json_response({"schemes": schemes})
        if path == prefix and method == "POST":
            payload = _read_json_body(body)
            return _json_response(
                PLANNING_STORE.create_scheme(
                    str(payload.get("name", "")),
                    owner_username=new_scheme_owner_for_user(current_user),
                )
            )
        if path == f"{prefix}/copy" and method == "POST":
            payload = _read_json_body(body)
            source = ensure_planning_scheme_access(str(payload.get("source", "")), current_user)
            target = ensure_planning_copy_target_access(str(payload.get("target", "")), current_user)
            return _json_response(
                PLANNING_STORE.copy_scheme(
                    source,
                    target,
                    overwrite=truthy_json_value(payload.get("overwrite")),
                    owner_username=new_scheme_owner_for_user(current_user),
                )
            )
        if path == f"{prefix}/import" and method == "POST":
            payload = _read_json_body(body)
            filename = str(payload.get("filename", ""))
            target_name = str(payload.get("name") or Path(filename).stem or "").strip()
            target = ensure_planning_copy_target_access(target_name, current_user)
            content_base64 = str(payload.get("content_base64", ""))
            try:
                content = base64.b64decode(content_base64, validate=True)
            except (binascii.Error, ValueError) as exc:
                raise ValueError("导入失败，方案压缩包内容无法解析") from exc
            return _json_response(
                PLANNING_STORE.import_scheme_archive(
                    content,
                    target,
                    owner_username=new_scheme_owner_for_user(current_user),
                    overwrite=truthy_json_value(payload.get("overwrite")),
                )
            )
        if path == f"{prefix}/share" and method == "POST":
            payload = _read_json_body(body)
            name = ensure_planning_scheme_manage_access(str(payload.get("scheme", "")), current_user)
            username = existing_share_target_username(str(payload.get("username", "")), current_user)
            PLANNING_STORE.share_scheme(name, username)
            return _json_response(planning_scheme_share_payload(name))
        if path == f"{prefix}/unshare" and method == "POST":
            payload = _read_json_body(body)
            name = ensure_planning_scheme_manage_access(str(payload.get("scheme", "")), current_user)
            username = str(payload.get("username", "")).strip()
            PLANNING_STORE.unshare_scheme(name, username)
            return _json_response(planning_scheme_share_payload(name))
        if path == f"{prefix}/rename" and method == "POST":
            payload = _read_json_body(body)
            source = ensure_planning_scheme_manage_access(str(payload.get("source", "")), current_user)
            return _json_response(
                PLANNING_STORE.rename_scheme(source, str(payload.get("target", "")))
            )
        if path.startswith(f"{prefix}/") and path.endswith("/shares") and method == "GET":
            name = unquote(path[len(prefix) + 1 : -len("/shares")])
            name = ensure_planning_scheme_manage_access(name, current_user)
            return _json_response(planning_scheme_share_payload(name))
        if path.startswith(f"{prefix}/") and path.endswith("/export") and method == "GET":
            name = unquote(path[len(prefix) + 1 : -len("/export")])
            name = ensure_planning_scheme_access(name, current_user)
            return _download_response(
                PLANNING_STORE.export_scheme_archive(name),
                f"{name}.zip",
                "application/zip",
            )
        if path.startswith(f"{prefix}/") and path.endswith("/overview") and method == "GET":
            name = unquote(path[len(prefix) + 1 : -len("/overview")])
            ensure_planning_scheme_access(name, current_user)
            return _json_response(PLANNING_STORE.read_scheme_overview(name))
        if path.startswith(f"{prefix}/") and path.endswith("/time-series") and method == "GET":
            name = unquote(path[len(prefix) + 1 : -len("/time-series")])
            ensure_planning_scheme_access(name, current_user)
            return _json_response(PLANNING_STORE.read_time_series(name))
        if path.startswith(f"{prefix}/"):
            name = unquote(path[len(prefix) + 1 :])
            if method == "GET":
                name = ensure_planning_scheme_access(name, current_user)
                return _json_response(PLANNING_STORE.read_scheme(name))
            if method == "PUT":
                name = ensure_planning_scheme_manage_access(name, current_user)
                payload = _read_json_body(body)
                includes_time_series = "time_series" in payload
                PLANNING_STORE.write_scheme(name, payload)
                if includes_time_series:
                    return _json_response(PLANNING_STORE.read_scheme(name))
                return _json_response(PLANNING_STORE.read_scheme_overview(name))
            if method == "DELETE":
                name = ensure_planning_scheme_manage_access(name, current_user)
                return _json_response(PLANNING_STORE.delete_scheme(name))
    except WeatherHistoryError as exc:
        return _json_response({"error": "weather_history_error", "message": str(exc)}, HTTPStatus.BAD_GATEWAY)
    except GeocodingError as exc:
        return _json_response({"error": "geocoding_error", "message": str(exc)}, HTTPStatus.BAD_GATEWAY)
    except PermissionError as exc:
        return _json_response({"error": "file_locked", "message": str(exc)}, HTTPStatus.CONFLICT)
    except LoadCurveTemplateExistsError as exc:
        return _json_response({"error": "exists", "message": str(exc)}, HTTPStatus.CONFLICT)
    except FileNotFoundError as exc:
        return _json_response({"error": "not_found", "message": str(exc)}, HTTPStatus.NOT_FOUND)
    except (ValueError, FileExistsError) as exc:
        return _json_response({"error": "bad_request", "message": str(exc)}, HTTPStatus.BAD_REQUEST)
    return _json_response({"error": "not_found", "path": path}, HTTPStatus.NOT_FOUND)


def handle_api_path(path: str, query: str = "", current_user: dict | None = None) -> tuple[int, dict[str, str], bytes]:
    parsed_api_path = urlparse(path)
    if parsed_api_path.query and not query:
        query = parsed_api_path.query
        path = parsed_api_path.path
    query_params = parse_qs(query)
    if path == "/api/tasks":
        return _json_response({"tasks": build_task_list(current_user=current_user)})
    if path.startswith("/api/planning/"):
        return handle_planning_api_path(path, "GET", b"", current_user=current_user)
    if path.startswith("/api/comparison/"):
        return handle_comparison_data_api_path(path, query, current_user=current_user)
    if path.startswith("/api/reliability/"):
        return handle_reliability_api_path(path, "GET", b"", query, current_user=current_user)
    if path == "/api/evaluation/status":
        scheme = query_params.get("scheme", [""])[0]
        if scheme:
            access_error = planning_scheme_access_error_response(scheme, current_user)
            if access_error:
                return access_error
        filename = query_params.get("filename", [""])[0]
        light = truthy_json_value(query_params.get("light", ["0"])[0])
        include_hourly_curves = not light
        return _json_response(EVALUATION_RUNTIME.snapshot(scheme=scheme, filename=filename, include_hourly_curves=include_hourly_curves))
    if path == "/api/frequency/status":
        scheme = query_params.get("scheme", [""])[0]
        if scheme:
            access_error = planning_scheme_access_error_response(scheme, current_user)
            if access_error:
                return access_error
        filename = query_params.get("filename", [""])[0]
        light = truthy_json_value(query_params.get("light", ["0"])[0])
        include_hourly_curves = not light
        return _json_response(FREQUENCY_EVALUATION_RUNTIME.snapshot(scheme=scheme, filename=filename, include_hourly_curves=include_hourly_curves))
    if path == "/api/frequency/time-curve":
        try:
            scheme = query_params.get("scheme", [""])[0]
            filename = query_params.get("filename", [""])[0]
            access_error = planning_scheme_access_error_response(scheme, current_user)
            if access_error:
                return access_error
            selected = selected_evaluation_result_filename(scheme, filename)
            hour_index = query_params.get("hour_index", [""])[0]
            month = query_params.get("month", [""])[0]
            day = query_params.get("day", [""])[0]
            hour = query_params.get("hour", [""])[0]
            payload = read_frequency_time_curve_payload(
                frequency_curve_result_path(scheme, selected),
                hour_index=hour_index,
                month=month,
                day=day,
                hour=hour,
            )
            return _json_response(payload)
        except FileNotFoundError as exc:
            return _json_response({"error": "not_found", "message": str(exc)}, HTTPStatus.NOT_FOUND)
        except ValueError as exc:
            return _json_response({"error": "bad_request", "message": str(exc)}, HTTPStatus.BAD_REQUEST)
    if path.startswith("/api/evaluation/"):
        return handle_evaluation_results_api_path(path, "GET", b"", query, current_user=current_user)
    if path == "/api/optimization/status":
        scheme = query_params.get("scheme", [""])[0]
        if scheme:
            access_error = planning_scheme_access_error_response(scheme, current_user)
            if access_error:
                return access_error
        light = truthy_json_value(query_params.get("light", ["0"])[0])
        include_hourly_curves = not light
        return _json_response(OPTIMIZATION_RUNTIME.snapshot(scheme=scheme, include_hourly_curves=include_hourly_curves))
    snapshot = build_snapshot()
    routes = {
        "/api/health": {"ok": True, "timestamp": snapshot["timestamp"]},
        "/api/overview": snapshot,
    }
    if path not in routes:
        return _json_response({"error": "not_found", "path": path}, HTTPStatus.NOT_FOUND)
    return _json_response(routes[path])


def handle_control_path(
    path: str,
    body: bytes,
    current_user: dict | None = None,
) -> tuple[int, dict[str, str], bytes]:
    if path == "/api/reliability/control":
        try:
            payload = json.loads(body.decode("utf-8") or "{}")
            action = str(payload.get("action") or "").strip()
            scheme = str(payload.get("scheme") or "").strip()
            filename = str(payload.get("filename") or "").strip()
            parameters = payload.get("parameters")
            if scheme or action != "clear_logs":
                ensure_planning_scheme_manage_access(scheme, current_user)
            if parameters is not None and not isinstance(parameters, dict):
                raise ValueError("parameters必须是对象")
            selected = normalize_reliability_source_filename(
                scheme,
                filename,
                require_exists=bool(filename),
            )
            if action in {"start", "queue"} and isinstance(parameters, dict):
                parameters, _, selected = write_reliability_parameters(scheme, parameters, selected)
            state = RELIABILITY_RUNTIME.apply(
                action,
                scheme=scheme,
                filename=selected,
                parameters=parameters,
            )
        except OptimizationStateError as exc:
            return _json_response({"error": exc.code, "message": str(exc)}, HTTPStatus.CONFLICT)
        except FileNotFoundError as exc:
            return _json_response({"error": "not_found", "message": str(exc)}, HTTPStatus.NOT_FOUND)
        except (ValueError, json.JSONDecodeError) as exc:
            return _json_response({"error": "bad_request", "message": str(exc)}, HTTPStatus.BAD_REQUEST)
        message = "可靠性评估已启动"
        if action == "queue":
            message = "可靠性队列当前按独立后台任务立即启动"
        elif action == "stop":
            message = "可靠性评估已停止"
        elif action == "clear_logs":
            message = "可靠性日志已清空"
        return _json_response({"ok": True, "message": message, "state": state})
    if path == "/api/optimization/control":
        try:
            payload = json.loads(body.decode("utf-8") or "{}")
            action = str(payload.get("action", ""))
            scheme = str(payload.get("scheme", ""))
            if scheme or action != "clear_logs":
                ensure_planning_scheme_manage_access(scheme, current_user)
            state = OPTIMIZATION_RUNTIME.apply(action, scheme=scheme)
        except OptimizationStateError as exc:
            return _json_response({"error": exc.code, "message": str(exc)}, HTTPStatus.CONFLICT)
        except FileNotFoundError as exc:
            return _json_response({"error": "not_found", "message": str(exc)}, HTTPStatus.NOT_FOUND)
        except (ValueError, json.JSONDecodeError) as exc:
            return _json_response({"error": "bad_request", "message": str(exc)}, HTTPStatus.BAD_REQUEST)
        return _json_response({"ok": True, "state": state})
    if path == "/api/evaluation/control":
        try:
            payload = json.loads(body.decode("utf-8") or "{}")
            action = str(payload.get("action", ""))
            scheme = str(payload.get("scheme", ""))
            filename = str(payload.get("filename", ""))
            if scheme or action != "clear_logs":
                ensure_planning_scheme_manage_access(scheme, current_user)
            state = EVALUATION_RUNTIME.apply(action, scheme=scheme, filename=filename)
        except OptimizationStateError as exc:
            return _json_response({"error": exc.code, "message": str(exc)}, HTTPStatus.CONFLICT)
        except FileNotFoundError as exc:
            return _json_response({"error": "not_found", "message": str(exc)}, HTTPStatus.NOT_FOUND)
        except (ValueError, json.JSONDecodeError) as exc:
            return _json_response({"error": "bad_request", "message": str(exc)}, HTTPStatus.BAD_REQUEST)
        return _json_response({"ok": True, "state": state})
    if path == "/api/frequency/control":
        try:
            payload = json.loads(body.decode("utf-8") or "{}")
            action = str(payload.get("action", ""))
            scheme = str(payload.get("scheme", ""))
            filename = str(payload.get("filename", ""))
            if scheme or action != "clear_logs":
                ensure_planning_scheme_manage_access(scheme, current_user)
            state = FREQUENCY_EVALUATION_RUNTIME.apply(action, scheme=scheme, filename=filename)
        except OptimizationStateError as exc:
            return _json_response({"error": exc.code, "message": str(exc)}, HTTPStatus.CONFLICT)
        except FileNotFoundError as exc:
            return _json_response({"error": "not_found", "message": str(exc)}, HTTPStatus.NOT_FOUND)
        except (ValueError, json.JSONDecodeError) as exc:
            return _json_response({"error": "bad_request", "message": str(exc)}, HTTPStatus.BAD_REQUEST)
        return _json_response({"ok": True, "state": state})
    if path == "/api/tasks/control":
        try:
            payload = json.loads(body.decode("utf-8") or "{}")
            action = str(payload.get("action", ""))
            task_type = str(payload.get("task_type") or payload.get("task_type_key") or "")
            scheme = str(payload.get("scheme", ""))
            result = str(payload.get("result") or payload.get("filename") or "")
            response = build_task_control_response(action, task_type, scheme, result, current_user=current_user)
        except OptimizationStateError as exc:
            return _json_response({"error": exc.code, "message": str(exc)}, HTTPStatus.CONFLICT)
        except FileNotFoundError as exc:
            return _json_response({"error": "not_found", "message": str(exc)}, HTTPStatus.NOT_FOUND)
        except (ValueError, json.JSONDecodeError) as exc:
            return _json_response({"error": "bad_request", "message": str(exc)}, HTTPStatus.BAD_REQUEST)
        return _json_response(response)
    return _json_response({"error": "not_found", "path": path}, HTTPStatus.NOT_FOUND)


def _unauthorized_response() -> tuple[int, dict[str, str], bytes]:
    return _json_response({"error": "unauthorized", "message": "请先登录"}, HTTPStatus.UNAUTHORIZED)


def _redirect_response(location: str) -> tuple[int, dict[str, str], bytes]:
    headers = {
        "Location": location,
        "Content-Type": "text/plain; charset=utf-8",
    }
    headers.update(_no_store_headers(vary_cookie=True))
    return HTTPStatus.FOUND, headers, b""


def safe_api_call(handler) -> tuple[int, dict[str, str], bytes]:
    try:
        return handler()
    except Exception as exc:
        traceback.print_exc()
        return _json_response(
            {"error": "internal_error", "message": f"后台处理失败: {exc}"},
            HTTPStatus.INTERNAL_SERVER_ERROR,
        )


def resolve_static_path(request_path: str) -> Path:
    parsed_path = unquote(urlparse(request_path).path)
    if parsed_path == "/":
        parsed_path = "/index.html"

    relative = parsed_path.lstrip("/")
    candidate = (WEB_ROOT / relative).resolve()
    if WEB_ROOT not in candidate.parents and candidate != WEB_ROOT:
        raise ValueError(f"path escapes web root: {request_path}")
    if candidate.is_dir():
        candidate = candidate / "index.html"
    return candidate


class PowerPlanHandler(BaseHTTPRequestHandler):
    server_version = "PowerPlan/1.0"

    def _current_user(self) -> dict | None:
        token = _session_token_from_cookie(self.headers.get("Cookie"))
        return _authenticated_user(token)

    def do_GET(self) -> None:
        parsed = urlparse(self.path)
        if parsed.path.startswith("/api/"):
            token = _session_token_from_cookie(self.headers.get("Cookie"))
            current_user = _authenticated_user(token)
            if parsed.path.startswith("/api/auth/"):
                status, headers, body = safe_api_call(lambda: handle_auth_api_path(parsed.path, "GET", b"", token))
                self._send(status, headers, body)
                return
            if parsed.path.startswith("/api/users"):
                status, headers, body = safe_api_call(lambda: handle_users_api_path(parsed.path, "GET", b"", current_user))
                self._send(status, headers, body)
                return
            if parsed.path != "/api/health" and not current_user:
                status, headers, body = _unauthorized_response()
                self._send(status, headers, body)
                return
            status, headers, body = safe_api_call(lambda: handle_api_path(parsed.path, parsed.query, current_user=current_user))
            self._send(status, headers, body)
            return

        try:
            path = resolve_static_path(self.path)
        except ValueError:
            self._send_text(HTTPStatus.FORBIDDEN, "Forbidden")
            return

        if not path.exists() or not path.is_file():
            self._send_text(HTTPStatus.NOT_FOUND, "Not Found")
            return

        if path.suffix == ".html":
            current_user = self._current_user()
            public_pages = {"login.html", "register.html"}
            if path.name not in public_pages and not current_user:
                next_path = parsed.path or "/index.html"
                self._send(*_redirect_response(f"/login.html?next={urlencode({'next': next_path})[5:]}"))
                return
            if path.name == "users.html" and current_user and current_user.get("role") != "admin":
                self._send_text(HTTPStatus.FORBIDDEN, "Forbidden")
                return
            if path.name in public_pages and current_user:
                self._send(*_redirect_response("/index.html"))
                return

        headers = _static_headers(path, authenticated_html=path.suffix == ".html")
        if _static_request_not_modified(self.headers, headers):
            self._send(HTTPStatus.NOT_MODIFIED, headers, b"")
            return
        body = STATIC_FILE_BYTES_CACHE.get(path, lambda resolved: resolved.read_bytes(), variant="static_bytes")
        self._send(HTTPStatus.OK, headers, body)

    def do_POST(self) -> None:
        parsed = urlparse(self.path)
        length = int(self.headers.get("Content-Length", "0"))
        body = self.rfile.read(length) if length else b"{}"
        if parsed.path.startswith("/api/"):
            token = _session_token_from_cookie(self.headers.get("Cookie"))
            current_user = _authenticated_user(token)
            if parsed.path.startswith("/api/auth/"):
                status, headers, response_body = safe_api_call(lambda: handle_auth_api_path(parsed.path, "POST", body, token))
                self._send(status, headers, response_body)
                return
            if parsed.path.startswith("/api/users"):
                status, headers, response_body = safe_api_call(lambda: handle_users_api_path(parsed.path, "POST", body, current_user))
                self._send(status, headers, response_body)
                return
            if not current_user:
                status, headers, response_body = _unauthorized_response()
                self._send(status, headers, response_body)
                return
            if parsed.path.startswith("/api/planning/"):
                status, headers, response_body = safe_api_call(
                    lambda: handle_planning_api_path(parsed.path, "POST", body, current_user=current_user)
                )
                self._send(status, headers, response_body)
                return
            if parsed.path == "/api/evaluation/control":
                status, headers, response_body = safe_api_call(
                    lambda: handle_control_path(parsed.path, body, current_user=current_user)
                )
                self._send(status, headers, response_body)
                return
            if parsed.path.startswith("/api/evaluation/"):
                status, headers, response_body = safe_api_call(
                    lambda: handle_evaluation_results_api_path(
                        parsed.path,
                        "POST",
                        body,
                        current_user=current_user,
                    )
                )
                self._send(status, headers, response_body)
                return
            status, headers, response_body = safe_api_call(lambda: handle_control_path(parsed.path, body, current_user=current_user))
            self._send(status, headers, response_body)
            return
        self._send_text(HTTPStatus.NOT_FOUND, "Not Found")

    def do_PUT(self) -> None:
        parsed = urlparse(self.path)
        length = int(self.headers.get("Content-Length", "0"))
        body = self.rfile.read(length) if length else b"{}"
        token = _session_token_from_cookie(self.headers.get("Cookie"))
        current_user = _authenticated_user(token)
        if parsed.path.startswith("/api/users"):
            status, headers, response_body = safe_api_call(lambda: handle_users_api_path(parsed.path, "PUT", body, current_user))
            self._send(status, headers, response_body)
            return
        if parsed.path.startswith("/api/") and not current_user:
            status, headers, response_body = _unauthorized_response()
            self._send(status, headers, response_body)
            return
        if parsed.path.startswith("/api/planning/"):
            status, headers, response_body = safe_api_call(
                lambda: handle_planning_api_path(parsed.path, "PUT", body, current_user=current_user)
            )
            self._send(status, headers, response_body)
            return
        if parsed.path.startswith("/api/reliability/"):
            status, headers, response_body = safe_api_call(
                lambda: handle_reliability_api_path(parsed.path, "PUT", body, parsed.query, current_user=current_user)
            )
            self._send(status, headers, response_body)
            return
        self._send_text(HTTPStatus.NOT_FOUND, "Not Found")

    def do_DELETE(self) -> None:
        parsed = urlparse(self.path)
        token = _session_token_from_cookie(self.headers.get("Cookie"))
        current_user = _authenticated_user(token)
        if parsed.path.startswith("/api/users"):
            status, headers, response_body = safe_api_call(lambda: handle_users_api_path(parsed.path, "DELETE", b"", current_user))
            self._send(status, headers, response_body)
            return
        if parsed.path.startswith("/api/") and not current_user:
            status, headers, response_body = _unauthorized_response()
            self._send(status, headers, response_body)
            return
        if parsed.path.startswith("/api/planning/"):
            status, headers, response_body = safe_api_call(
                lambda: handle_planning_api_path(parsed.path, "DELETE", b"", current_user=current_user)
            )
            self._send(status, headers, response_body)
            return
        self._send_text(HTTPStatus.NOT_FOUND, "Not Found")

    def log_message(self, format: str, *args: object) -> None:
        return

    def _send_text(self, status: int, text: str) -> None:
        self._send(status, {"Content-Type": "text/plain; charset=utf-8"}, text.encode("utf-8"))

    def _send(self, status: int, headers: dict[str, str], body: bytes) -> None:
        if int(status) not in (HTTPStatus.NOT_MODIFIED, HTTPStatus.NO_CONTENT):
            headers, body = gzip_response_body_if_supported(self.headers, headers, body)
        self.send_response(int(status))
        for key, value in headers.items():
            self.send_header(key, value)
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)


def run(host: str = "127.0.0.1", port: int = 8866) -> None:
    server = ThreadingHTTPServer((host, port), PowerPlanHandler)
    server.serve_forever()


def main() -> None:
    parser = argparse.ArgumentParser(description="Run the power plan server.")
    parser.add_argument("--host", default="127.0.0.1")
    parser.add_argument("--port", default=8866, type=int)
    args = parser.parse_args()
    run(args.host, args.port)


if __name__ == "__main__":
    main()
